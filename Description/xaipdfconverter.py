#!/usr/bin/env python3
"""
xAI PDF to DOCX Converter with Structure Preservation and Sensitive Information Detection
"""

import sys
import logging
import re
import os
import tempfile
import warnings
from pathlib import Path
from datetime import datetime
from multiprocessing import Pool, cpu_count
from contextlib import contextmanager
from io import StringIO

# Supprimer les warnings PyPDF2 concernant les PDFs mal formés (non critique)
warnings.filterwarnings("ignore", category=UserWarning, module="PyPDF2")
warnings.filterwarnings("ignore", message=".*incorrect startxref pointer.*")
warnings.filterwarnings("ignore", message=".*startxref.*")

# Try to import pdf2docx
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    print("WARNING: pdf2docx not available. Install with: pip install pdf2docx")

# Try to import PyPDF2 for text extraction
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Try to import phonenumbers for phone validation
try:
    import phonenumbers
    from phonenumbers import PhoneNumberFormat
    PHONENUMBERS_AVAILABLE = True
except ImportError:
    PHONENUMBERS_AVAILABLE = False

# Try to import openpyxl for Excel reports
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# Try to import python-docx for Word document post-processing
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    WD_ALIGN_PARAGRAPH = None

# Try to import pdfplumber for advanced table extraction
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Try to import tqdm for progress bars
try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False

# ==================== NLP LIBRARIES FOR ENTITY RECOGNITION ====================
# Try to import spaCy for Named Entity Recognition (NER)
SPACY_AVAILABLE = False
SPACY_NLP = None

try:
    import spacy
    SPACY_AVAILABLE = True
    # Try to load the English model (en_core_web_sm is the small model)
    try:
        SPACY_NLP = spacy.load("en_core_web_sm")
    except OSError:
        # Model not installed, try to load medium or large model
        try:
            SPACY_NLP = spacy.load("en_core_web_md")
        except OSError:
            try:
                SPACY_NLP = spacy.load("en_core_web_lg")
            except OSError:
                # No model available, disable spaCy
                SPACY_AVAILABLE = False
                SPACY_NLP = None
                logger.warning("spaCy models not found. Install with: python -m spacy download en_core_web_sm")
except ImportError:
    SPACY_AVAILABLE = False
    SPACY_NLP = None

# Try to import transformers (Hugging Face) for advanced NER
TRANSFORMERS_AVAILABLE = False
try:
    from transformers import pipeline
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False

# ==================== OCR LIBRARIES ====================
# Try to import OCR libraries (optional, only for scanned PDFs)
PADDLEOCR_AVAILABLE = False
PPSTRUCTURE_AVAILABLE = False

# Essayer d'importer PaddleOCR (peut échouer à cause de dépendances manquantes)
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
    # Essayer d'importer PPStructure séparément
    try:
        from paddleocr import PPStructure
        PPSTRUCTURE_AVAILABLE = True
    except (ImportError, RuntimeError, AttributeError):
        # PPStructure peut ne pas être disponible ou avoir des problèmes de dépendances
        PPSTRUCTURE_AVAILABLE = False
except (ImportError, RuntimeError, ModuleNotFoundError) as e:
    # PaddleOCR peut avoir des problèmes de dépendances (langchain, etc.)
    PADDLEOCR_AVAILABLE = False
    PPSTRUCTURE_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    import pytesseract
    from pytesseract import TesseractError
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    TesseractError = None

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Try to import OCRmyPDF (recommended for scanned PDFs)
try:
    import ocrmypdf
    OCRMYPDF_AVAILABLE = True
except ImportError:
    OCRMYPDF_AVAILABLE = False

# Determine best available OCR engine
OCR_AVAILABLE = OCRMYPDF_AVAILABLE or PADDLEOCR_AVAILABLE or EASYOCR_AVAILABLE or PYTESSERACT_AVAILABLE

if OCR_AVAILABLE and not PDF2IMAGE_AVAILABLE and not OCRMYPDF_AVAILABLE:
    print("⚠️  WARNING: pdf2image not available. Install with: pip install pdf2image")
    print("   OCR will not work without it (unless using OCRmyPDF).")
    OCR_AVAILABLE = False

# Configuration du logging (sera reconfiguré selon le mode)
logging.basicConfig(
    level=logging.WARNING,  # Par défaut, seulement warnings et erreurs
    format='[%(levelname)s] %(message)s'
)
logger = logging.getLogger(__name__)

# Supprimer les warnings PyPDF2/pikepdf dès le démarrage (avant toute utilisation)
logging.getLogger('PyPDF2').setLevel(logging.ERROR)
logging.getLogger('pikepdf').setLevel(logging.ERROR)

# Configure logging levels for external libraries
def configure_logging(verbose=False):
    """
    Configure logging levels based on verbose mode.
    
    Args:
        verbose: If True, show all INFO logs. If False, only WARNING and ERROR.
    """
    if verbose:
        # Mode verbose : tout afficher
        logging.getLogger().setLevel(logging.INFO)
        logger.setLevel(logging.INFO)
        logging.getLogger('pdf2docx').setLevel(logging.INFO)
    else:
        # Mode simple : silence tout sauf warnings/errors
        logging.getLogger().setLevel(logging.WARNING)
        logger.setLevel(logging.WARNING)
        logging.getLogger('pdf2docx').setLevel(logging.ERROR)  # Changé de WARNING à ERROR
        
        # Désactiver aussi les logs de pdfplumber si présent
        logging.getLogger('pdfplumber').setLevel(logging.ERROR)
        logging.getLogger('PIL').setLevel(logging.ERROR)  # Pillow (images)
    
    # Supprimer les warnings PyPDF2 (startxref pointer, etc.)
    logging.getLogger('PyPDF2').setLevel(logging.ERROR)
    logging.getLogger('pikepdf').setLevel(logging.ERROR)  # pikepdf est utilisé par PyPDF2


@contextmanager
def suppress_output(verbose=False):
    """
    Context manager to suppress stdout/stderr in non-verbose mode.
    
    Args:
        verbose: If True, don't suppress output. If False, suppress it.
    """
    if verbose:
        # Don't suppress in verbose mode
        yield
    else:
        # Suppress stdout and stderr
        old_stdout = sys.stdout
        old_stderr = sys.stderr
        try:
            sys.stdout = StringIO()
            sys.stderr = StringIO()
            yield
        finally:
            sys.stdout = old_stdout
            sys.stderr = old_stderr


# Fonction helper pour multiprocessing (doit être au niveau module)
def _convert_chunk(args):
    """
    Convertit un chunk de pages du PDF en DOCX.
    Fonction helper pour multiprocessing (doit être au niveau module).
    
    Args:
        args: Tuple (pdf_path, output_path, start_page, end_page)
    
    Returns:
        str: Chemin du fichier DOCX généré, ou None en cas d'erreur
    """
    pdf_path, output_path, start_page, end_page = args
    
    try:
        if not PDF2DOCX_AVAILABLE:
            logger.warning(f"pdf2docx not available for chunk {start_page}-{end_page}")
            return None
        
        # ⚠️ IMPORTANT : Désactiver les logs dans chaque processus worker
        # (multiprocessing n'hérite pas des paramètres de logging du processus parent)
        logging.getLogger('pdf2docx').setLevel(logging.ERROR)
        logging.getLogger('pdfplumber').setLevel(logging.ERROR)
        logging.getLogger('PIL').setLevel(logging.ERROR)
        
        converter = Converter(str(pdf_path))
        try:
            converter.convert(str(output_path), start=start_page, end=end_page)
            logger.info(f"✓ Converted pages {start_page}-{end_page}")
            return str(output_path)
        finally:
            converter.close()
    except Exception as e:
        logger.error(f"Error converting chunk {start_page}-{end_page}: {e}")
        return None


class SECPDFConverter:
    """
    Convertisseur PDF vers DOCX avec préservation de structure et détection d'informations sensibles.
    """
    
    def __init__(self, pdf_path, docx_path=None, verbose=False):
        """
        Initialise le convertisseur.
        
        Args:
            pdf_path: Chemin vers le fichier PDF
            docx_path: Chemin de sortie (optionnel, généré automatiquement si non fourni)
        """
        self.pdf_path = Path(pdf_path)
        
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {self.pdf_path}")
        
        if docx_path:
            self.docx_path = Path(docx_path)
        else:
            self.docx_path = self.pdf_path.with_suffix('.docx')
        
        # Créer le dossier parent si nécessaire
        self.docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.converter = None
        self.images_by_page = {}
        self.sensitive_info_by_page = {}
        self.verbose = verbose
        
        # Cache OCR pour éviter de retraiter les mêmes pages
        self.ocr_cache = {}
        self._init_ocr_cache()
    
    def _init_ocr_cache(self):
        """Initialise le cache OCR basé sur un hash du PDF."""
        import hashlib
        
        try:
            # Calculer un hash du PDF pour le cache
            with open(self.pdf_path, 'rb') as f:
                pdf_hash = hashlib.md5(f.read()).hexdigest()
            
            # Créer un dossier de cache si nécessaire
            cache_dir = Path.home() / '.xaipdf_cache'
            cache_dir.mkdir(exist_ok=True)
            
            self.cache_file = cache_dir / f"ocr_cache_{pdf_hash}.json"
            
            # Charger le cache existant
            if self.cache_file.exists():
                import json
                try:
                    with open(self.cache_file, 'r') as f:
                        self.ocr_cache = json.load(f)
                    if self.verbose:
                        logger.info(f"📦 Loaded OCR cache: {len(self.ocr_cache)} pages cached")
                except Exception:
                    self.ocr_cache = {}
            else:
                self.ocr_cache = {}
        except Exception as e:
            logger.warning(f"Could not initialize OCR cache: {e}")
            self.ocr_cache = {}
            self.cache_file = None
    
    def _save_ocr_cache(self):
        """Sauvegarde le cache OCR."""
        if not self.cache_file:
            return
        
        try:
            import json
            with open(self.cache_file, 'w') as f:
                json.dump(self.ocr_cache, f)
        except Exception as e:
            logger.warning(f"Could not save OCR cache: {e}")
    
    def _get_ocr_cache_key(self, page_num, dpi, lang):
        """Génère une clé unique pour le cache OCR d'une page."""
        return f"{page_num}_{dpi}_{lang}"
    
    def _transform_coords_to_pdf(self, ocr_data, img_width, img_height, page_num):
        """
        Transforme les coordonnées OCR (relatives à l'image) en coordonnées PDF.
        Utilise une transformation matricielle pour mapper image → PDF.
        """
        try:
            # Obtenir les dimensions de la page PDF
            if PYPDF2_AVAILABLE:
                with open(self.pdf_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    if page_num <= len(pdf_reader.pages):
                        page = pdf_reader.pages[page_num - 1]
                        # Dimensions de la page PDF en points (72 DPI = 1 point = 1/72 inch)
                        pdf_width_pt = float(page.mediabox.width)
                        pdf_height_pt = float(page.mediabox.height)
                        
                        # Calculer les facteurs de transformation
                        scale_x = pdf_width_pt / img_width
                        scale_y = pdf_height_pt / img_height
                        
                        # Transformer les coordonnées
                        if 'left' in ocr_data and ocr_data['left']:
                            ocr_data['left'] = [int(x * scale_x) for x in ocr_data['left']]
                            ocr_data['top'] = [int(y * scale_y) for y in ocr_data['top']]
                            ocr_data['width'] = [int(w * scale_x) for w in ocr_data['width']]
                            ocr_data['height'] = [int(h * scale_y) for h in ocr_data['height']]
        except Exception as e:
            if self.verbose:
                logger.warning(f"Coordinate transformation failed: {e}")
        
        return ocr_data
    
    def _detect_rotation_angle(self, img):
        """
        Détecte l'angle de rotation d'une image scannée.
        Retourne l'angle en degrés.
        """
        import numpy as np
        from PIL import ImageFilter
        
        try:
            # Convertir en niveaux de gris
            if img.mode != 'L':
                img_gray = img.convert('L')
            else:
                img_gray = img
            
            # Détecter les bords
            img_edges = img_gray.filter(ImageFilter.FIND_EDGES)
            img_array = np.array(img_edges)
            
            # Projection horizontale pour détecter l'inclinaison
            h_projection = img_array.sum(axis=1)
            
            # Trouver les pics dans la projection
            peaks = []
            for i in range(1, len(h_projection) - 1):
                if h_projection[i] > h_projection[i-1] and h_projection[i] > h_projection[i+1]:
                    peaks.append(i)
            
            if len(peaks) > 5:
                # Calculer l'angle moyen basé sur les différences entre pics
                angles = []
                for i in range(1, min(10, len(peaks))):
                    deviation = peaks[i] - peaks[0]
                    if deviation != 0:
                        angle = np.arctan(deviation / img.width) * 180 / np.pi
                        angles.append(angle)
                
                if angles:
                    avg_angle = np.mean(angles)
                    # Filtrer les angles trop petits (< 0.5 degré)
                    if abs(avg_angle) > 0.5:
                        return avg_angle
        except Exception as e:
            if self.verbose:
                logger.debug(f"Rotation detection failed: {e}")
        
        return 0.0
    
    def _add_rotation_to_ocr_data(self, ocr_data, img):
        """
        Ajoute le paramètre de rotation aux données OCR.
        Les bounding boxes deviennent: left, top, width, height, rotation
        """
        # Détecter l'angle de rotation
        rotation_angle = self._detect_rotation_angle(img)
        
        # Ajouter la rotation aux données OCR
        if 'left' in ocr_data and ocr_data['left']:
            n_boxes = len(ocr_data['left'])
            ocr_data['rotation'] = [rotation_angle] * n_boxes
        else:
            ocr_data['rotation'] = []
        
        return ocr_data
    
    def _get_pdf_page_count(self):
        """Retourne le nombre total de pages du PDF."""
        if not PYPDF2_AVAILABLE:
            return None
        try:
            with open(self.pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            logger.warning(f"Could not get page count: {e}")
            return None
    
    def is_scanned_pdf(self):
        """
        Détecte automatiquement si le PDF est scanné (image-based).
        
        Returns:
            bool: True si scanné, False si texte normal
        """
        if not PYPDF2_AVAILABLE:
            logger.warning("Cannot detect PDF type without PyPDF2")
            return False
        
        try:
            with open(self.pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Analyser les 3 premières pages
                pages_to_check = min(3, len(pdf_reader.pages))
                text_chars = 0
                
                for i in range(pages_to_check):
                    page = pdf_reader.pages[i]
                    text = page.extract_text()
                    text_chars += len(text.strip())
                
                # Moyenne de caractères par page
                avg_chars_per_page = text_chars / pages_to_check if pages_to_check > 0 else 0
                
                # Seuil : < 100 caractères/page = probablement scanné
                is_scanned = avg_chars_per_page < 100
                
                if is_scanned:
                    logger.info(f"📷 PDF detected as SCANNED (avg {avg_chars_per_page:.0f} chars/page)")
                else:
                    logger.info(f"📄 PDF detected as TEXT-BASED (avg {avg_chars_per_page:.0f} chars/page)")
                
                return is_scanned
        
        except Exception as e:
            logger.warning(f"Could not detect PDF type: {e}")
            return False
    
    def map_language_code(self, lang_code):
        """
        Convertit les codes langue ISO en codes Tesseract.
        
        Args:
            lang_code: Code ISO (en, fr, de, etc.)
        
        Returns:
            Code Tesseract (eng, fra, deu, etc.)
        """
        lang_map = {
            'en': 'eng',
            'fr': 'fra',
            'de': 'deu',
            'es': 'spa',
            'it': 'ita',
            'pt': 'por',
            'nl': 'nld',
            'pl': 'pol',
            'ru': 'rus',
            'ja': 'jpn',
            'zh': 'chi_sim',
            'ar': 'ara',
        }
        
        return lang_map.get(lang_code, lang_code)  # Retourne le code tel quel si non mappé
    
    def convert_with_ocr(self, start_page=0, end_page=None, ocr_language='eng'):
        """
        🔥 OCR avec Préservation de Structure EXACTE (comme onlineocr.net)
        Architecture: pdf2image → pytesseract (image_to_data) → python-docx
        Préserve FIDÈLEMENT le layout, les colonnes, tableaux et positions spatiales.
        
        Args:
            start_page: Page de départ (0-indexed)
            end_page: Page de fin (None = toutes les pages)
            ocr_language: Code langue ISO (eng, fra, etc.) pour Tesseract
        
        Raises:
            ImportError: Si pytesseract, pdf2image ou python-docx ne sont pas disponibles
            Exception: Si la conversion échoue
        """
        if not PYTESSERACT_AVAILABLE or not PDF2IMAGE_AVAILABLE:
            raise ImportError(
                "=" * 70 + "\n"
                "OCR DEPENDENCIES NOT AVAILABLE\n"
                "=" * 70 + "\n"
                "Required packages:\n"
                "  pip install pytesseract pdf2image pillow\n\n"
                "System dependencies:\n"
                "  macOS: brew install tesseract poppler\n"
                "  Ubuntu: sudo apt-get install tesseract-ocr poppler-utils\n"
                "=" * 70
            )
        
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        # Créer le répertoire de sortie
        self.docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from pytesseract import Output, TesseractError
            from docx import Document
            from docx.shared import Pt, Inches, Twips
            from PIL import Image, ImageEnhance
            import numpy as np
            
            # Créer le document Word
            doc = Document()
            
            # Configuration des marges (standard)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Style par défaut
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            
            total_pages = self._get_pdf_page_count()
            actual_end = end_page if end_page is not None else total_pages
            
            # DPI : 300 minimum pour documents financiers SEC (Best Practice)
            dpi = 300
            
            logger.info(f"🔄 Converting PDF to images ({dpi} DPI) for pages {start_page+1} to {actual_end}...")
            
            # 1. Conversion PDF → Images (300 DPI = qualité professionnelle pour documents financiers)
            try:
                images = convert_from_path(
                    str(self.pdf_path),
                    first_page=start_page+1,
                    last_page=actual_end,
                    dpi=dpi,  # 300 DPI = qualité professionnelle (Best Practice pour documents financiers)
                    fmt='PNG',
                    thread_count=4  # Paralléliser la conversion d'images
                )
            except Exception as e:
                error_msg = str(e).lower()
                if 'poppler' in error_msg or 'pdfinfo' in error_msg:
                    raise Exception(
                        "❌ POPPLER NOT FOUND\n"
                        "Poppler is required to convert PDF to images.\n"
                        "Install:\n"
                        "  macOS: brew install poppler\n"
                        "  Ubuntu: sudo apt-get install poppler-utils\n"
                        "  Windows: Download from poppler.freedesktop.org"
                    )
                else:
                    raise Exception(f"Error converting PDF to images: {e}")
            
            lang = self.map_language_code(ocr_language)
            
            # 2. Méthode OCR PROFESSIONNELLE : PaddleOCR (PPStructure) > OCRmyPDF > Tesseract
            # Ordre de priorité : PaddleOCR (meilleur) > OCRmyPDF > Tesseract (fallback)
            use_advanced_ocr = False
            
            # Essayer PaddleOCR avec PPStructure (meilleure qualité)
            if PADDLEOCR_AVAILABLE:
                try:
                    logger.info("🏆 Using PaddleOCR with PPStructure (best quality)...")
                    self._convert_with_paddleocr_ppstructure(doc, images, start_page, actual_end, ocr_language, dpi)
                    use_advanced_ocr = True
                except Exception as e:
                    logger.warning(f"⚠️  PaddleOCR failed: {e}, trying fallback...")
                    use_advanced_ocr = False
            
            # Essayer OCRmyPDF si PaddleOCR n'est pas disponible
            if not use_advanced_ocr and OCRMYPDF_AVAILABLE and PDF2DOCX_AVAILABLE:
                try:
                    logger.info("🔄 Using OCRmyPDF (Sandwich method - industry standard)...")
                    self._convert_with_ocrmypdf_sandwich(doc, start_page, actual_end, ocr_language)
                    use_advanced_ocr = True
                except Exception as e:
                    # Si ocrmypdf échoue (dépendances manquantes), utiliser méthode manuelle
                    if 'unpaper' in str(e).lower() or 'missing' in str(e).lower():
                        logger.warning("⚠️  OCRmyPDF dependencies missing (unpaper), using manual method...")
                        use_advanced_ocr = False
                    else:
                        raise
            
            # Fallback : Tesseract avec méthode manuelle
            if not use_advanced_ocr:
                logger.info("📸 Using Tesseract OCR (manual method)...")
                for i, img in enumerate(images):
                    page_num = start_page + i + 1
                    logger.info(f"  📸 Processing page {page_num}/{actual_end} with OCR...")
                    
                    if i > 0:
                        doc.add_page_break()
                    
                    # Amélioration de l'image pour meilleur OCR
                    img_enhanced = self._enhance_image_for_ocr(img)
                    
                    # Vérifier le cache OCR
                    cache_key = self._get_ocr_cache_key(page_num, dpi, lang)
                    if cache_key in self.ocr_cache:
                        logger.info(f"  📦 Using cached OCR for page {page_num}...")
                        ocr_data = self.ocr_cache[cache_key]
                    else:
                        # OCR avec extraction de coordonnées (hOCR pour meilleure précision)
                        try:
                            # Utiliser hOCR pour obtenir les bounding boxes précises
                            hocr_output = pytesseract.image_to_pdf_or_hocr(
                                img_enhanced,
                                lang=lang,
                                extension='hocr',
                                config='--psm 1 --oem 3'
                            )
                            # Parser hOCR pour extraire les coordonnées (avec rotation)
                            ocr_data = self._parse_hocr(hocr_output, img.width, img.height, img_enhanced)
                        except Exception as e:
                            # Fallback vers image_to_data si hOCR échoue
                            try:
                                ocr_data_raw = pytesseract.image_to_data(
                                    img_enhanced,
                                    lang=lang,
                                    output_type=Output.DICT,
                                    config='--psm 1 --oem 3'
                                )
                                # Ajouter la détection de rotation
                                ocr_data = self._add_rotation_to_ocr_data(ocr_data_raw, img_enhanced)
                            except Exception as e2:
                                error_msg = str(e2).lower()
                                if 'tessdata' in error_msg or 'traineddata' in error_msg:
                                    raise Exception(
                                        "❌ TESSERACT LANGUAGE DATA NOT FOUND\n"
                                        f"Language '{lang}' not available.\n"
                                        "Install language packs:\n"
                                        "  macOS: brew install tesseract-lang\n"
                                        "  Ubuntu: sudo apt-get install tesseract-ocr-{lang}\n"
                                        "  Windows: Download from github.com/tesseract-ocr/tessdata"
                                    )
                                else:
                                    raise Exception(f"Tesseract OCR failed: {e2}")
                        
                        # Sauvegarder dans le cache
                        self.ocr_cache[cache_key] = ocr_data
                    
                    # Transformer les coordonnées image → PDF (si nécessaire)
                    ocr_data_transformed = self._transform_coords_to_pdf(ocr_data, img.width, img.height, page_num)
                    
                    # Méthode SANDWICH manuelle
                    self._render_page_sandwich_method(doc, img, ocr_data_transformed, img.width, img.height)
            
            # Sauvegarde du cache OCR
            self._save_ocr_cache()
            
            # Sauvegarde
            doc.save(str(self.docx_path))
            logger.info(f"✅ Conversion completed: {self.docx_path.name}")
        
        except Exception as e:
            logger.error(f"OCR conversion failed: {e}")
            if self.verbose:
                import traceback
                traceback.print_exc()
            raise
    
    def _enhance_image_for_ocr(self, img):
        """
        Améliore l'image pour meilleur résultat OCR (OPTIMISÉ - rapide et efficace).
        """
        from PIL import ImageEnhance, ImageFilter
        
        # Convertir en niveaux de gris si nécessaire
        if img.mode != 'L':
            img = img.convert('L')
        
        # 1. DENOISING simple (rapide)
        try:
            img = img.filter(ImageFilter.MedianFilter(size=3))
        except Exception:
            pass
        
        # 2. AMÉLIORATION DU CONTRASTE
        try:
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.2)
        except Exception:
            pass
        
        # 3. AMÉLIORATION DE LA NETTETÉ
        try:
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.1)
        except Exception:
            pass
        
        return img
    
    # ============================================================================
    # _reconstruct_page_faithfully() - SUPPRIMÉE/COMMENTÉE
    # Cette fonction appelait des méthodes inexistantes (_render_text_line_ocr, etc.)
    # Utiliser _render_page_sandwich_method() à la place (version simple et fonctionnelle)
    # ============================================================================
    # def _reconstruct_page_faithfully(self, doc, ocr_data, img_width, img_height):
    #     """FONCTION DÉSACTIVÉE - Utiliser _render_page_sandwich_method() à la place"""
    #     pass
    
    def _convert_with_paddleocr_ppstructure(self, doc, images, start_page, end_page, ocr_language):
        """
        Méthode OCR PROFESSIONNELLE : Utilise PaddleOCR avec PPStructure.
        Meilleure qualité avec détection de structure (text, table, title, image).
        """
        import numpy as np
        
        # Initialiser PaddleOCR avec détection de structure
        try:
            # Essayer PPStructure d'abord (meilleure structure)
            from paddleocr import PPStructure
            structure_engine = PPStructure(show_log=False, use_gpu=False)
            use_ppstructure = True
        except Exception:
            # Fallback vers PaddleOCR standard
            from paddleocr import PaddleOCR
            ocr_engine = PaddleOCR(use_angle_cls=True, lang='en', det=True, rec=True, show_log=False)
            use_ppstructure = False
        
        lang_map = {'eng': 'en', 'en': 'en', 'fra': 'fr', 'fr': 'fr', 'deu': 'de', 'de': 'de'}
        ocr_lang = lang_map.get(ocr_language.lower(), 'en')
        
        for i, img in enumerate(images):
            page_num = start_page + i + 1
            logger.info(f"  📸 Processing page {page_num}/{end_page} with PaddleOCR...")
            
            if i > 0:
                doc.add_page_break()
            
            # Convertir PIL Image en numpy array
            img_array = np.array(img)
            
            # Utiliser PPStructure si disponible
            if use_ppstructure:
                try:
                    result = structure_engine(img_array)
                    # PPStructure retourne une structure hiérarchique
                    self._render_paddleocr_structure(doc, img, result, img.width, img.height)
                except Exception as e:
                    logger.warning(f"PPStructure failed: {e}, using standard PaddleOCR...")
                    use_ppstructure = False
            
            # Fallback vers PaddleOCR standard
            if not use_ppstructure:
                from paddleocr import PaddleOCR
                ocr_engine = PaddleOCR(use_angle_cls=True, lang=ocr_lang, det=True, rec=True, show_log=False)
                result = ocr_engine.ocr(img_array, cls=True)
                
                # Convertir le résultat PaddleOCR en format standardisé
                ocr_data = self._convert_paddleocr_to_standard(result, img.width, img.height)
                self._render_page_sandwich_method(doc, img, ocr_data, img.width, img.height)
    
    def _convert_paddleocr_to_standard(self, paddle_result, img_width, img_height):
        """
        Convertit le résultat PaddleOCR en format standardisé (compatible avec Tesseract).
        PaddleOCR retourne: [[[x1,y1], [x2,y2], [x3,y3], [x4,y4]], (text, confidence)]
        """
        words = []
        
        if not paddle_result or not paddle_result[0]:
            return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'conf': []}
        
        for line in paddle_result[0]:
            if not line:
                continue
            
            # Extraire les coordonnées et le texte
            bbox = line[0]  # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
            text_info = line[1]  # (text, confidence)
            
            if len(bbox) >= 4 and text_info:
                text = text_info[0] if isinstance(text_info, tuple) else str(text_info)
                conf = text_info[1] if isinstance(text_info, tuple) and len(text_info) > 1 else 95
                
                # Calculer bounding box rectangulaire
                x_coords = [point[0] for point in bbox]
                y_coords = [point[1] for point in bbox]
                
                left = int(min(x_coords))
                top = int(min(y_coords))
                right = int(max(x_coords))
                bottom = int(max(y_coords))
                
                width = right - left
                height = bottom - top
                
                if text.strip() and conf > 30:
                    words.append({
                        'text': text.strip(),
                        'left': left,
                        'top': top,
                        'width': width,
                        'height': height,
                        'conf': int(conf * 100) if conf < 1 else int(conf)
                    })
        
        # Convertir en format dict compatible avec Tesseract
        if words:
            return {
                'text': [w['text'] for w in words],
                'left': [w['left'] for w in words],
                'top': [w['top'] for w in words],
                'width': [w['width'] for w in words],
                'height': [w['height'] for w in words],
                'rotation': [w.get('rotation', 0.0) for w in words],  # 5ème paramètre
                'conf': [w['conf'] for w in words]
            }
        else:
            return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'rotation': [], 'conf': []}
    
    def _render_paddleocr_structure(self, doc, original_img, structure_result, img_width, img_height):
        """
        Rend le résultat de PPStructure en préservant la structure hiérarchique.
        PPStructure détecte: text, table, title, image, etc.
        """
        # Reconstruire UNIQUEMENT le texte avec structure préservée (pas d'image)
        # PPStructure fournit déjà la structure, on la rend directement
        if not structure_result:
            return
        
        # PPStructure retourne une liste d'éléments avec type et bbox
        for element in structure_result:
            if not element:
                continue
            
            elem_type = element.get('type', 'text')
            bbox = element.get('bbox', [])
            content = element.get('res', {})
            
            # Extraire le texte selon le type
            if elem_type == 'table':
                # Rendre comme tableau Word
                self._render_table_from_paddleocr(doc, content, bbox)
            elif elem_type == 'title':
                # Rendre comme titre
                text = content.get('text', '') if isinstance(content, dict) else str(content)
                if text:
                    para = doc.add_paragraph(text.strip())
                    para.style = 'Heading 1'
                    for run in para.runs:
                        run.bold = True
            else:
                # Texte normal
                text = content.get('text', '') if isinstance(content, dict) else str(content)
                if text:
                    doc.add_paragraph(text.strip())
    
    def _render_table_from_paddleocr(self, doc, table_content, bbox):
        """Rend un tableau détecté par PPStructure."""
        from docx.shared import Inches
        
        if not table_content or not isinstance(table_content, dict):
            return
        
        # PPStructure peut retourner le tableau en format structuré
        # Essayer d'extraire les cellules
        cells = table_content.get('cells', [])
        
        if cells:
            # Créer un tableau Word
            num_rows = len(cells)
            num_cols = max(len(row) for row in cells) if cells else 1
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            for row_idx, row in enumerate(cells):
                for col_idx, cell in enumerate(row):
                    if col_idx < num_cols:
                        cell_text = cell.get('text', '') if isinstance(cell, dict) else str(cell)
                        table.rows[row_idx].cells[col_idx].paragraphs[0].text = cell_text.strip()
        else:
            # Fallback : texte simple
            text = str(table_content)
            doc.add_paragraph(f"[Table: {text}]")
    
    def _parse_hocr(self, hocr_output, img_width, img_height, img):
        """
        Parse le résultat hOCR de Tesseract pour extraire les coordonnées précises.
        hOCR contient des bounding boxes au niveau du mot avec coordonnées pixel.
        Ajoute aussi le paramètre de rotation (5ème paramètre).
        """
        from xml.etree import ElementTree as ET
        
        # Détecter l'angle de rotation global
        rotation_angle = self._detect_rotation_angle(img)
        
        try:
            # Parser le XML hOCR
            root = ET.fromstring(hocr_output)
            
            # Trouver tous les éléments de mots (classe 'ocrx_word')
            words = []
            for word_elem in root.iter():
                if 'class' in word_elem.attrib and 'ocrx_word' in word_elem.attrib.get('class', ''):
                    # Extraire le titre qui contient les coordonnées
                    title = word_elem.attrib.get('title', '')
                    if 'bbox' in title:
                        # Parser: bbox x1 y1 x2 y2
                        parts = title.split()
                        bbox_idx = parts.index('bbox')
                        x1, y1, x2, y2 = map(int, parts[bbox_idx+1:bbox_idx+5])
                        
                        # Extraire le texte
                        text = ''.join(word_elem.itertext()).strip()
                        
                        # Extraire la rotation si disponible (dans certains formats hOCR)
                        word_rotation = rotation_angle
                        if 'rotate' in title:
                            try:
                                rotate_idx = parts.index('rotate')
                                word_rotation = float(parts[rotate_idx+1])
                            except (ValueError, IndexError):
                                pass
                        
                        if text:
                            words.append({
                                'text': text,
                                'left': x1,
                                'top': y1,
                                'width': x2 - x1,
                                'height': y2 - y1,
                                'rotation': word_rotation,  # 5ème paramètre
                                'conf': 95  # hOCR ne donne pas toujours la confiance
                            })
            
            # Convertir en format dict
            if words:
                return {
                    'text': [w['text'] for w in words],
                    'left': [w['left'] for w in words],
                    'top': [w['top'] for w in words],
                    'width': [w['width'] for w in words],
                    'height': [w['height'] for w in words],
                    'rotation': [w['rotation'] for w in words],  # 5ème paramètre
                    'conf': [w['conf'] for w in words]
                }
        except Exception as e:
            logger.warning(f"hOCR parsing failed: {e}, using standard format")
            return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'rotation': [], 'conf': []}
        
        return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'rotation': [], 'conf': []}
    
    def _convert_with_ocrmypdf_sandwich(self, doc, start_page, end_page, ocr_language):
        """
        Méthode SANDWICH optimale : Utilise ocrmypdf pour créer un PDF searchable,
        puis convertit en Word avec pdf2docx.
        C'est la meilleure approche industry-standard.
        """
        import tempfile
        
        # Créer un PDF temporaire avec ocrmypdf
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            tmp_pdf_path = tmp_pdf.name
        
        try:
            # Obtenir le nombre total de pages
            from PyPDF2 import PdfReader
            reader = PdfReader(str(self.pdf_path))
            total_pages = len(reader.pages)
            
            # Extraire les pages nécessaires du PDF original
            if start_page > 0 or (end_page and end_page < total_pages):
                # Extraire les pages avec PyPDF2
                from PyPDF2 import PdfWriter
                writer = PdfWriter()
                
                actual_end = end_page if end_page else total_pages
                for i in range(start_page, actual_end):
                    writer.add_page(reader.pages[i])
                
                with open(tmp_pdf_path, 'wb') as f:
                    writer.write(f)
                input_pdf = tmp_pdf_path
            else:
                input_pdf = str(self.pdf_path)
            
            # Appliquer OCR avec ocrmypdf (méthode Sandwich automatique)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as ocr_pdf:
                ocr_pdf_path = ocr_pdf.name
            
            try:
                # Mapper la langue pour ocrmypdf
                lang_map = {'eng': 'eng', 'en': 'eng', 'fra': 'fra', 'fr': 'fra', 'deu': 'deu', 'de': 'deu'}
                ocr_lang = lang_map.get(ocr_language.lower(), 'eng')
                
                try:
                    ocrmypdf.ocr(
                        input_pdf,
                        ocr_pdf_path,
                        language=ocr_lang,
                        deskew=True,  # Redressement automatique
                        clean=False,  # Nettoyage désactivé (nécessite unpaper)
                        force_ocr=True,  # Forcer OCR même si texte présent
                        optimize=1    # Optimisation niveau 1
                    )
                except Exception as e:
                    # Si ocrmypdf échoue (dépendances manquantes), fallback vers méthode manuelle
                    error_msg = str(e).lower()
                    if 'unpaper' in error_msg or 'missing' in error_msg:
                        logger.warning("OCRmyPDF dependencies missing, falling back to manual method")
                        raise Exception("Use manual method") from e
                    else:
                        raise
                
                # Extraire le texte avec structure depuis le PDF OCR'd en utilisant pdfplumber
                logger.info("🔄 Extracting text with layout from OCR'd PDF using pdfplumber...")
                if PDFPLUMBER_AVAILABLE:
                    self._extract_text_with_pdfplumber(doc, ocr_pdf_path, start_page, end_page)
                else:
                    # Fallback : méthode manuelle avec OCR direct
                    logger.warning("pdfplumber not available, using manual OCR method...")
                    self._extract_text_with_layout_from_ocr_pdf(doc, ocr_pdf_path, start_page, end_page)
                logger.info("✅ OCR'd PDF converted successfully")
                    
            finally:
                try:
                    os.unlink(ocr_pdf_path)
                except:
                    pass
                    
        finally:
            try:
                if tmp_pdf_path != str(self.pdf_path):
                    os.unlink(tmp_pdf_path)
            except:
                pass
    
    def _extract_text_with_pdfplumber(self, doc, ocr_pdf_path, start_page, end_page):
        """
        Extrait le texte d'un PDF OCR'd avec pdfplumber qui peut lire le texte overlay.
        Préserve la structure et la mise en page.
        """
        import pdfplumber
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from pdf2image import convert_from_path
        
        try:
            logger.info("📖 Opening OCR'd PDF with pdfplumber...")
            with pdfplumber.open(ocr_pdf_path) as pdf:
                total_pages = len(pdf.pages)
                actual_end = end_page if end_page else total_pages
                
                for page_idx in range(start_page, actual_end):
                    pdf_page = pdf.pages[page_idx]
                    
                    if page_idx > start_page:
                        doc.add_page_break()
                    
                    logger.info(f"  📄 Processing page {page_idx + 1}/{actual_end}...")
                    
                    # Extraire le texte avec coordonnées
                    words = pdf_page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False,
                        use_text_flow=True,
                        extra_attrs=["fontname", "size"]
                    )
                    
                    if not words:
                        logger.warning(f"No words extracted from page {page_idx + 1}")
                        doc.add_paragraph(f"[Page {page_idx + 1}: No text extracted]")
                        continue
                    
                    # VERSION SIMPLE : Extraire le texte directement
                    # Trier les mots par position (top, puis left)
                    words_sorted = sorted(words, key=lambda w: (w.get('top', 0), w.get('left', 0)))
                    
                    # Grouper par lignes (simple)
                    lines = []
                    current_line = []
                    current_y = None
                    
                    for word in words_sorted:
                        word_y = word.get('top', 0)
                        word_text = word.get('text', '').strip()
                        
                        if not word_text:
                            continue
                        
                        if current_y is None or abs(word_y - current_y) <= 10:
                            current_line.append(word)
                            if current_y is None:
                                current_y = word_y
                        else:
                            if current_line:
                                lines.append(sorted(current_line, key=lambda w: w.get('left', 0)))
                            current_line = [word]
                            current_y = word_y
                    
                    if current_line:
                        lines.append(sorted(current_line, key=lambda w: w.get('left', 0)))
                    
                    # Rendre les lignes simplement
                    for line in lines:
                        if not line:
                            doc.add_paragraph("")
                            continue
                        
                        line_text = ' '.join([w.get('text', '') for w in line])
                        
                        if line_text.strip():
                            # Détecter si titre (tout en majuscules ET court)
                            is_title = line_text.isupper() and len(line_text) < 100 and len(line_text.split()) < 15
                            
                            if is_title:
                                para = doc.add_paragraph(line_text.strip())
                                para.style = 'Heading 2'
                                for run in para.runs:
                                    run.bold = True
                            else:
                                doc.add_paragraph(line_text.strip())
                        
        except Exception as e:
            logger.error(f"Error extracting with pdfplumber: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def _group_pdfplumber_words_by_lines(self, words, page_height, tolerance=5):
        """Groupe les mots de pdfplumber par lignes."""
        if not words:
            return []
        
        # Trier par position Y (top)
        sorted_words = sorted(words, key=lambda w: (w.get('top', 0), w.get('left', 0)))
        
        lines = []
        current_line = []
        current_y = None
        
        for word in sorted_words:
            word_y = word.get('top', 0)
            word_text = word.get('text', '').strip()
            
            if not word_text:
                continue
            
            if current_y is None:
                current_y = word_y
                current_line = [word]
            elif abs(word_y - current_y) <= tolerance:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(sorted(current_line, key=lambda w: w.get('left', 0)))
                current_line = [word]
                current_y = word_y
        
        if current_line:
            lines.append(sorted(current_line, key=lambda w: w.get('left', 0)))
        
        return lines
    
    def _analyze_page_structure_from_words(self, words, lines, page_width, page_height):
        """Analyse la structure depuis les mots de pdfplumber."""
        structure = {
            'has_columns': False,
            'has_table': False,
            'column_count': 1,
            'table_regions': [],
            'title_lines': []
        }
        
        if not words or not lines:
            return structure
        
        # Analyser la distribution horizontale
        x_positions = [w.get('left', 0) for w in words]
        if x_positions:
            x_min, x_max = min(x_positions), max(x_positions)
            x_range = x_max - x_min
            
            # Si le texte est concentré sur moins de 70% de la largeur, probablement 2 colonnes
            if x_range < page_width * 0.7:
                # Vérifier les gaps pour confirmer les colonnes
                sorted_x = sorted(set(x_positions))
                large_gaps = []
                for i in range(1, len(sorted_x)):
                    gap = sorted_x[i] - sorted_x[i-1]
                    if gap > page_width * 0.2:  # Gap de plus de 20% de la largeur
                        large_gaps.append(gap)
                
                if len(large_gaps) >= 1:
                    structure['has_columns'] = True
                    structure['column_count'] = 2
        
        # Détecter les titres (police plus grande, ou texte en majuscules)
        for line_idx, line in enumerate(lines):
            if not line:
                continue
            
            line_text = ' '.join([w.get('text', '') for w in line])
            avg_font_size = sum(w.get('size', 10) for w in line) / len(line) if line else 10
            
            is_title = (
                (line_text.isupper() and len(line_text) < 100 and len(line_text.split()) < 15) or
                (avg_font_size > 12)  # Police plus grande
            )
            
            if is_title:
                structure['title_lines'].append(line_idx)
        
        return structure
    
    def _render_multicolumn_from_pdfplumber(self, doc, words, lines, structure, page_width):
        """Rend un layout multi-colonnes depuis pdfplumber."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        column_width = page_width / structure['column_count']
        
        for line_idx, line in enumerate(lines):
            if not line:
                doc.add_paragraph("")
                continue
            
            # Déterminer la colonne
            avg_x = sum(w.get('left', 0) for w in line) / len(line) if line else 0
            column_idx = int(avg_x / column_width) if column_width > 0 else 0
            
            # Créer le texte
            line_text = ' '.join([w.get('text', '') for w in line])
            
            if line_idx in structure.get('title_lines', []):
                para = doc.add_paragraph(line_text.strip())
                para.style = 'Heading 2'
                for run in para.runs:
                    run.bold = True
            else:
                para = doc.add_paragraph(line_text.strip())
    
    def _render_table_from_pdfplumber(self, doc, words, lines, structure, page_width):
        """Rend un tableau depuis pdfplumber."""
        # Détecter les colonnes du tableau
        all_x_positions = sorted(set(w.get('left', 0) for w in words))
        
        # Grouper les positions proches
        columns = []
        if all_x_positions:
            columns.append(all_x_positions[0])
            for x in all_x_positions[1:]:
                if x - columns[-1] > 50:  # Nouvelle colonne si gap > 50
                    columns.append(x)
        
        if len(columns) < 2:
            self._render_single_column_from_pdfplumber(doc, words, lines, structure)
            return
        
        # Créer le tableau
        table = doc.add_table(rows=len(lines), cols=len(columns))
        table.style = 'Table Grid'
        
        for row_idx, line in enumerate(lines):
            if not line:
                continue
            
            for w in line:
                x_pos = w.get('left', 0)
                col_idx = min(range(len(columns)), key=lambda i: abs(x_pos - columns[i]))
                if col_idx < len(columns):
                    cell = table.rows[row_idx].cells[col_idx]
                    if not cell.paragraphs[0].text.strip():
                        cell.paragraphs[0].text = w.get('text', '')
                    else:
                        cell.paragraphs[0].text += " " + w.get('text', '')
    
    def _render_single_column_from_pdfplumber(self, doc, words, lines, structure):
        """Rend un layout simple depuis pdfplumber."""
        from docx.shared import Pt
        
        for line_idx, line in enumerate(lines):
            if not line:
                doc.add_paragraph("")
                continue
            
            line_text = ' '.join([w.get('text', '') for w in line])
            
            if line_idx in structure.get('title_lines', []):
                para = doc.add_paragraph(line_text.strip())
                para.style = 'Heading 2'
                for run in para.runs:
                    run.bold = True
            else:
                doc.add_paragraph(line_text.strip())
    
    def _extract_text_with_layout_from_ocr_pdf(self, doc, ocr_pdf_path, start_page, end_page):
        """
        Extrait le texte d'un PDF OCR'd avec préservation de la structure et de la mise en page.
        Utilise pdfplumber pour extraire le texte avec coordonnées, puis reconstruit fidèlement.
        """
        from pdf2image import convert_from_path
        from io import BytesIO
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import pytesseract
        from pytesseract import Output
        
        try:
            # Convertir le PDF OCR'd en images pour OCR avec coordonnées
            logger.info("🔄 Converting OCR'd PDF to images for layout analysis...")
            images = convert_from_path(
                ocr_pdf_path,
                first_page=1,
                last_page=None,
                dpi=300,
                fmt='PNG'
            )
            
            total_pages = len(images)
            actual_end = end_page if end_page else total_pages
            
            # Traiter chaque page avec OCR layout-aware
            for page_idx in range(start_page, actual_end):
                if page_idx > start_page:
                    doc.add_page_break()
                
                img = images[page_idx]
                logger.info(f"  📄 Processing page {page_idx + 1}/{actual_end} with layout-aware OCR...")
                
                # Améliorer l'image pour meilleur OCR
                img_enhanced = self._enhance_image_for_ocr(img)
                
                # OCR avec extraction de coordonnées précises (hOCR)
                try:
                    hocr_output = pytesseract.image_to_pdf_or_hocr(
                        img_enhanced,
                        lang='eng',
                        extension='hocr',
                        config='--psm 1 --oem 3'  # PSM 1 = Automatic page segmentation with OSD
                    )
                    ocr_data = self._parse_hocr(hocr_output, img.width, img.height, img_enhanced)
                except Exception as e:
                    # Fallback vers image_to_data
                    try:
                        ocr_data_raw = pytesseract.image_to_data(
                            img_enhanced,
                            lang='eng',
                            output_type=Output.DICT,
                            config='--psm 1 --oem 3'
                        )
                        ocr_data = self._add_rotation_to_ocr_data(ocr_data_raw, img_enhanced)
                    except Exception as e2:
                        logger.error(f"OCR failed for page {page_idx + 1}: {e2}")
                        # Fallback : image seule
                        img_buffer = BytesIO()
                        img.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        para = doc.add_paragraph()
                        run = para.add_run()
                        run.add_picture(img_buffer, width=Inches(img.width / 300.0))
                        continue
                
                # Reconstruire la page avec structure préservée
                self._reconstruct_page_with_layout(doc, ocr_data, img.width, img.height)
                
        except Exception as e:
            logger.error(f"Error extracting text with layout from OCR'd PDF: {e}")
            raise
    
    def _reconstruct_page_with_layout(self, doc, ocr_data, img_width, img_height):
        """
        Reconstruit le texte avec layout préservé - VERSION SIMPLE.
        Alias pour _render_page_sandwich_method pour compatibilité.
        """
        # Appeler la méthode simple (ignore original_img car on n'ajoute plus d'image)
        self._render_page_sandwich_method(doc, None, ocr_data, img_width, img_height)
    
    def _analyze_page_structure(self, lines, words, img_width, img_height):
        """
        Analyse la structure de la page pour détecter colonnes, tableaux, titres.
        """
        structure = {
            'has_columns': False,
            'has_table': False,
            'column_count': 1,
            'table_regions': [],
            'title_lines': []
        }
        
        if not lines:
            return structure
        
        # Analyser la distribution horizontale des mots pour détecter les colonnes
        x_positions = [w['left'] for w in words if w.get('left', 0) > 0]
        if x_positions:
            x_min, x_max = min(x_positions), max(x_positions)
            x_range = x_max - x_min
            
            # Si les mots sont distribués sur moins de 60% de la largeur, probablement 2 colonnes
            if x_range < img_width * 0.6:
                # Analyser les clusters de positions X
                sorted_x = sorted(set(x_positions))
                if len(sorted_x) > 10:
                    # Calculer les gaps entre positions
                    gaps = []
                    for i in range(1, len(sorted_x)):
                        gap = sorted_x[i] - sorted_x[i-1]
                        if gap > img_width * 0.15:  # Gap significatif
                            gaps.append((sorted_x[i-1], sorted_x[i]))
                    
                    if len(gaps) >= 1:
                        structure['has_columns'] = True
                        structure['column_count'] = 2
        
        # Détecter les tableaux (lignes avec plusieurs mots alignés verticalement)
        if len(lines) > 3:
            # Vérifier si plusieurs lignes ont des mots alignés verticalement
            vertical_alignment_count = 0
            for i in range(len(lines) - 1):
                if i + 1 < len(lines):
                    line1_words = lines[i]
                    line2_words = lines[i + 1]
                    
                    # Vérifier l'alignement vertical
                    for w1 in line1_words:
                        for w2 in line2_words:
                            if abs(w1['left'] - w2['left']) < 20:  # Alignés à 20 pixels près
                                vertical_alignment_count += 1
                                break
                        if vertical_alignment_count > 0:
                            break
                    
                    if vertical_alignment_count > 3:
                        structure['has_table'] = True
                        break
        
        # Détecter les titres (lignes courtes, en majuscules, ou avec grande police)
        for line_idx, line in enumerate(lines):
            if not line:
                continue
            
            line_text = ' '.join([w['text'] for w in line])
            avg_height = sum(w['height'] for w in line) / len(line) if line else 0
            
            # Titre si : tout en majuscules ET court OU grande hauteur de police
            is_title = (
                (line_text.isupper() and len(line_text) < 100 and len(line_text.split()) < 15) or
                (avg_height > 20)  # Police plus grande
            )
            
            if is_title:
                structure['title_lines'].append(line_idx)
        
        return structure
    
    def _render_multicolumn_layout(self, doc, lines, structure, img_width):
        """Rend un layout multi-colonnes."""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Calculer les positions des colonnes
        column_width = img_width / structure['column_count']
        
        for line_idx, line in enumerate(lines):
            if not line:
                doc.add_paragraph("")
                continue
            
            # Déterminer dans quelle colonne se trouve la ligne
            avg_x = sum(w['left'] for w in line) / len(line) if line else 0
            column_idx = int(avg_x / column_width) if column_width > 0 else 0
            
            # Créer le texte de la ligne
            line_text = ' '.join([w['text'] for w in sorted(line, key=lambda w: w['left'])])
            
            # Appliquer le style selon le type
            if line_idx in structure['title_lines']:
                para = doc.add_paragraph(line_text.strip())
                para.style = 'Heading 2'
                for run in para.runs:
                    run.bold = True
            else:
                para = doc.add_paragraph(line_text.strip())
            
            # Alignement selon la colonne
            if column_idx == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif column_idx == structure['column_count'] - 1:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _render_table_layout(self, doc, lines, structure, img_width):
        """Rend un layout de tableau."""
        from docx.shared import Pt
        
        # Détecter les colonnes du tableau
        all_x_positions = set()
        for line in lines:
            for w in line:
                all_x_positions.add(w['left'])
        
        sorted_x = sorted(all_x_positions)
        # Grouper les positions proches (colonnes)
        columns = []
        current_col = [sorted_x[0]] if sorted_x else []
        for x in sorted_x[1:]:
            if x - current_col[-1] < 50:  # Positions proches = même colonne
                current_col.append(x)
            else:
                if current_col:
                    columns.append(sum(current_col) / len(current_col))  # Position moyenne
                current_col = [x]
        if current_col:
            columns.append(sum(current_col) / len(current_col))
        
        if len(columns) < 2:
            # Pas assez de colonnes, rendre comme texte simple
            self._render_single_column_layout(doc, lines, structure)
            return
        
        # Créer un tableau Word
        table = doc.add_table(rows=len(lines), cols=len(columns))
        table.style = 'Table Grid'
        
        for row_idx, line in enumerate(lines):
            if not line:
                continue
            
            # Assigner chaque mot à sa colonne
            for w in line:
                # Trouver la colonne la plus proche
                col_idx = min(range(len(columns)), key=lambda i: abs(w['left'] - columns[i]))
                if col_idx < len(columns):
                    cell = table.rows[row_idx].cells[col_idx]
                    if not cell.paragraphs[0].text.strip():
                        cell.paragraphs[0].text = w['text']
                    else:
                        cell.paragraphs[0].text += " " + w['text']
    
    def _render_single_column_layout(self, doc, lines, structure):
        """Rend un layout simple à une colonne."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for line_idx, line in enumerate(lines):
            if not line:
                doc.add_paragraph("")
                continue
            
            line_sorted = sorted(line, key=lambda w: w['left'])
            line_text = ' '.join([w['text'] for w in line_sorted])
            
            if line_idx in structure.get('title_lines', []):
                para = doc.add_paragraph(line_text.strip())
                para.style = 'Heading 2'
                for run in para.runs:
                    run.bold = True
            else:
                doc.add_paragraph(line_text.strip())
    
    def _render_page_sandwich_method(self, doc, original_img, ocr_data, img_width, img_height):
        """VERSION SIMPLE QUI MARCHE"""
        if not ocr_data or 'text' not in ocr_data:
            return
        
        # Extraire mots
        n_boxes = len(ocr_data['text'])
        words = []
        for i in range(n_boxes):
            try:
                conf = int(ocr_data['conf'][i]) if ocr_data['conf'][i] else 0
                text = str(ocr_data['text'][i]).strip() if ocr_data['text'][i] else ''
                if conf > 30 and text:
                    words.append({
                        'text': text,
                        'left': int(ocr_data['left'][i]),
                        'top': int(ocr_data['top'][i])
                    })
            except:
                continue
        
        # Grouper par lignes
        words.sort(key=lambda w: (w['top'], w['left']))
        lines = []
        current_line = []
        current_y = None
        
        for word in words:
            if current_y is None or abs(word['top'] - current_y) <= 10:
                current_line.append(word)
                current_y = word['top']
            else:
                if current_line:
                    lines.append(current_line)
                current_line = [word]
                current_y = word['top']
        
        if current_line:
            lines.append(current_line)
        
        # Afficher texte
        for line in lines:
            line.sort(key=lambda w: w['left'])
            text = ' '.join([w['text'] for w in line])
            if text.strip():
                doc.add_paragraph(text.strip())
    
    def _group_words_by_lines_sandwich(self, words, img_height):
        """
        Groupe les mots par lignes pour la méthode Sandwich.
        Similaire à _group_words_by_lines_ocr mais optimisé pour le rendu invisible.
        """
        if not words:
            return []
        
        words_sorted = sorted(words, key=lambda w: (w['top'], w['left']))
        lines = []
        current_line = []
        current_y = None
        y_tolerance = 5  # Tolérance pour grouper les mots sur la même ligne
        
        for word in words_sorted:
            word_y = word['top']
            if current_y is None:
                current_y = word_y
                current_line.append(word)
            elif abs(word_y - current_y) <= y_tolerance:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(current_line)
                current_line = [word]
                current_y = word_y
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _group_words_by_lines_ocr(self, words):
        """
        Groupe les mots par lignes basées sur leur position Y.
        Retourne simplement des listes de mots par ligne, sans séparation de colonnes.
        """
        if not words:
            return []
        
        words_sorted = sorted(words, key=lambda w: (w['top'], w['left']))
        lines = []
        current_line = []
        current_y = None
        y_tolerance = 5
        
        for word in words_sorted:
            word_y = word['top']
            if current_y is None:
                current_y = word_y
                current_line.append(word)
            elif abs(word_y - current_y) <= y_tolerance:
                current_line.append(word)
            else:
                if current_line:
                    # Trier par X pour avoir l'ordre de gauche à droite
                    current_line.sort(key=lambda w: w['left'])
                    lines.append(current_line)
                current_line = [word]
                current_y = word_y
        
        if current_line:
            current_line.sort(key=lambda w: w['left'])
            lines.append(current_line)
        
        return lines
    
    def _detect_column_structure(self, words, img_width):
        """
        Détecte la structure de colonnes en analysant TOUS les mots de la page.
        Utilise le clustering pour identifier les colonnes distinctes.
        
        Returns:
            dict avec 'has_columns', 'left_column_x', 'right_column_x', 'split_x'
        """
        if not words or len(words) < 20:
            return {'has_columns': False, 'split_x': img_width / 2}
        
        # Extraire toutes les positions X
        x_positions = [w['left'] for w in words]
        x_positions.sort()
        
        # Analyser la distribution pour trouver deux clusters
        # Approche: trouver le point de séparation optimal
        midpoint = img_width / 2
        tolerance = img_width * 0.15  # 15% de tolérance
        
        # Compter les mots à gauche et à droite du milieu
        left_count = sum(1 for x in x_positions if x < midpoint - tolerance)
        right_count = sum(1 for x in x_positions if x > midpoint + tolerance)
        center_count = len(x_positions) - left_count - right_count
        
        total = len(x_positions)
        
        # Si on a une bonne distribution 30-70% gauche et 30-70% droite, c'est probablement deux colonnes
        left_ratio = left_count / total if total > 0 else 0
        right_ratio = right_count / total if total > 0 else 0
        
        # Détection: au moins 25% des mots à gauche ET 25% à droite, avec moins de 20% au centre
        has_columns = (left_ratio >= 0.25 and right_ratio >= 0.25 and 
                      (center_count / total if total > 0 else 1) < 0.2)
        
        if has_columns:
            # Trouver le point de séparation optimal (entre les deux colonnes)
            # Utiliser la médiane des positions dans la zone centrale
            center_x = [x for x in x_positions if midpoint - tolerance <= x <= midpoint + tolerance]
            if center_x:
                split_x = sum(center_x) / len(center_x)
            else:
                split_x = midpoint
            
            # Calculer les positions moyennes des colonnes
            left_x = [x for x in x_positions if x < split_x]
            right_x = [x for x in x_positions if x >= split_x]
            
            left_column_x = sum(left_x) / len(left_x) if left_x else 0
            right_column_x = sum(right_x) / len(right_x) if right_x else img_width
            
            return {
                'has_columns': True,
                'split_x': split_x,
                'left_column_x': left_column_x,
                'right_column_x': right_column_x
            }
        else:
            return {'has_columns': False, 'split_x': midpoint}
    
    def _separate_lines_by_columns(self, lines, column_structure):
        """
        Sépare les mots de chaque ligne selon leur colonne.
        Retourne des lignes avec structure multi-colonnes.
        """
        split_x = column_structure['split_x']
        separated_lines = []
        
        for line in lines:
            if not line:
                separated_lines.append([])
                continue
            
            # Séparer les mots selon leur position X
            left_words = []
            right_words = []
            
            for word in line:
                word_x = word.get('left', 0)
                if word_x < split_x:
                    left_words.append(word)
                else:
                    right_words.append(word)
            
            # Si on a des mots des deux côtés, créer une structure multi-colonnes
            if left_words and right_words:
                separated_lines.append({
                    'type': 'multi_column',
                    'columns': [left_words, right_words],
                    'all_words': line
                })
            elif left_words:
                # Seulement à gauche (peut être un header ou titre)
                separated_lines.append(left_words)
            elif right_words:
                # Seulement à droite
                separated_lines.append(right_words)
            else:
                separated_lines.append([])
        
        return separated_lines
    
    def _separate_columns_in_line(self, line):
        """
        Détecte et sépare les colonnes dans une ligne.
        Retourne une structure qui préserve l'ordre des colonnes.
        """
        if not line or len(line) < 2:
            return line
        
        # Calculer les écarts entre les mots
        gaps = []
        for i in range(1, len(line)):
            prev_x_end = line[i-1]['left'] + line[i-1]['width']
            curr_x = line[i]['left']
            gap = curr_x - prev_x_end
            gaps.append(gap)
        
        if not gaps:
            return line
        
        # Trouver le gap le plus large (probable séparation de colonnes)
        max_gap = max(gaps)
        gap_threshold = max(50, max_gap * 0.3)  # Au moins 50px ou 30% du max gap
        
        # Si le gap max est significatif, c'est probablement deux colonnes
        if max_gap > gap_threshold:
            # Trouver l'index où se trouve le grand gap
            split_idx = gaps.index(max_gap) + 1
            
            # Séparer en deux colonnes
            left_col = line[:split_idx]
            right_col = line[split_idx:]
            
            # Retourner une structure qui indique les colonnes
            return {
                'type': 'multi_column',
                'columns': [left_col, right_col],
                'all_words': line  # Pour compatibilité
            }
        
        # Pas de colonnes multiples, retourner la ligne normale
        return line
    
    def _detect_page_structure_ocr(self, lines, img_width):
        """Détecte la structure de la page (positions des colonnes, zones de tableaux)."""
        all_x_positions = []
        for line in lines:
            # Gérer les lignes multi-colonnes
            if isinstance(line, dict) and line.get('type') == 'multi_column':
                for col in line.get('columns', []):
                    for word in col:
                        all_x_positions.append(word['left'])
            elif isinstance(line, list):
                for word in line:
                    if isinstance(word, dict):
                        all_x_positions.append(word['left'])
        
        if not all_x_positions:
            return {'column_positions': [], 'table_zones': []}
        
        # Clusteriser les positions X pour détecter les colonnes
        all_x_positions.sort()
        column_positions = []
        if all_x_positions:
            current_cluster = [all_x_positions[0]]
            cluster_tolerance = 20  # 20 pixels de tolérance
            
            for x in all_x_positions[1:]:
                if x - current_cluster[-1] < cluster_tolerance:
                    current_cluster.append(x)
                else:
                    if len(current_cluster) >= 3:
                        column_positions.append(sum(current_cluster) / len(current_cluster))
                    current_cluster = [x]
            
            if len(current_cluster) >= 3:
                column_positions.append(sum(current_cluster) / len(current_cluster))
        
        # Détecter les zones de tableaux
        table_zones = []
        for line_idx, line in enumerate(lines):
            # Extraire les mots réels
            if isinstance(line, dict) and line.get('type') == 'multi_column':
                line_for_analysis = []
                for col in line.get('columns', []):
                    line_for_analysis.extend(col)
            else:
                line_for_analysis = line if isinstance(line, list) else []
            
            if len(line_for_analysis) >= 4:
                gaps = []
                for i in range(1, len(line_for_analysis)):
                    prev_x_end = line_for_analysis[i-1]['left'] + line_for_analysis[i-1]['width']
                    curr_x = line_for_analysis[i]['left']
                    gap = curr_x - prev_x_end
                    gaps.append(gap)
                
                avg_gap = sum(gaps) / len(gaps) if gaps else 0
                if avg_gap > 15:
                    table_zones.append(line_idx)
        
        return {
            'column_positions': sorted(set(column_positions)),
            'table_zones': table_zones,
            'page_width': img_width
        }
    
    def _is_table_line_ocr(self, line, all_lines, line_idx, page_structure):
        """Détermine si une ligne fait partie d'un tableau."""
        # Gérer les lignes multi-colonnes
        if isinstance(line, dict) and line.get('type') == 'multi_column':
            # Pour les colonnes multiples, utiliser les mots de toutes les colonnes
            all_words = []
            for col in line.get('columns', []):
                all_words.extend(col)
            line = all_words
        
        if not isinstance(line, list) or len(line) < 3:
            return False
        
        if line_idx in page_structure['table_zones']:
            return True
        
        gaps = []
        for i in range(1, len(line)):
            prev_x_end = line[i-1]['left'] + line[i-1]['width']
            curr_x = line[i]['left']
            gap = curr_x - prev_x_end
            gaps.append(gap)
        
        large_gaps = [g for g in gaps if g > 20]
        if len(large_gaps) >= 2:
            return True
        
        text_content = " ".join([w['text'] for w in line])
        digit_count = sum(c.isdigit() for c in text_content)
        if digit_count > len(text_content) * 0.3:
            return True
        
        return False
    
    def _render_table_line_ocr(self, doc, line, page_structure, img_width):
        """Rend une ligne comme ligne de tableau en utilisant des tabulations pour aligner."""
        from docx.shared import Pt
        
        # Gérer les lignes multi-colonnes
        if isinstance(line, dict) and line.get('type') == 'multi_column':
            self._render_multicolumn_line(doc, line['columns'])
            return
        
        if not isinstance(line, list):
            return
        
        column_positions = page_structure.get('column_positions', [])
        if not column_positions:
            text = " ".join([w['text'] if isinstance(w, dict) else str(w) for w in line])
            doc.add_paragraph(text)
            return
        
        # Créer un tableau Word si c'est une zone de tableau détectée
        # Sinon, utiliser des tabulations pour aligner
        para = doc.add_paragraph()
        
        current_x = 0
        for word in line:
            word_x = word['left']
            
            # Trouver la colonne la plus proche
            closest_col_idx = min(range(len(column_positions)),
                                 key=lambda i: abs(column_positions[i] - word_x))
            target_col = column_positions[closest_col_idx]
            
            # Ajouter des tabulations pour aligner
            if word_x > current_x + 10:
                tabs_needed = max(1, int((target_col - current_x) / 50))
                for _ in range(min(tabs_needed, 15)):
                    para.add_run("\t")
            
            para.add_run(word['text'])
            current_x = word_x + word['width']
            
            if word != line[-1]:
                para.add_run(" ")
    
    def _render_text_line_ocr(self, doc, line, page_structure):
        """Rend une ligne de texte normal."""
        # Gérer les lignes multi-colonnes
        if isinstance(line, dict) and line.get('type') == 'multi_column':
            self._render_multicolumn_line(doc, line['columns'])
            return
        
        # Ligne normale
        if not isinstance(line, list):
            return
            
        text = " ".join([w['text'] if isinstance(w, dict) else str(w) for w in line])
        text = text.strip()
        
        if not text:
            return
        
        # Détecter si c'est un titre
        if text.isupper() and len(text) < 100 and len(text.split()) < 15:
            para = doc.add_paragraph(text)
            para.style = 'Heading 2'
            for run in para.runs:
                run.bold = True
        else:
            doc.add_paragraph(text)
    
    def _render_multicolumn_line(self, doc, columns):
        """
        Rend une ligne avec colonnes multiples.
        Utilise un tableau Word à 2 colonnes pour préserver la structure.
        """
        from docx.shared import Inches
        
        if not columns or len(columns) < 2:
            # Fallback : rendre comme texte normal
            all_words = []
            for col in columns:
                all_words.extend(col)
            text = " ".join([w['text'] for w in all_words])
            if text.strip():
                doc.add_paragraph(text.strip())
            return
        
        # Créer un tableau avec 2 colonnes pour cette ligne
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Normal Table'
        
        # Largeur des colonnes (50% chacune)
        for col in table.columns:
            col.width = Inches(3.25)
        
        row = table.rows[0]
        
        # Colonne gauche
        left_words = sorted(columns[0], key=lambda w: w['left'])
        left_text = " ".join([w['text'] for w in left_words])
        row.cells[0].paragraphs[0].text = left_text.strip()
        
        # Colonne droite
        right_words = sorted(columns[1], key=lambda w: w['left'])
        right_text = " ".join([w['text'] for w in right_words])
        row.cells[1].paragraphs[0].text = right_text.strip()
    
    def _detect_multicolumn_layout(self, lines, img_width):
        """
        Détecte si la page a un layout à colonnes multiples (ex: table des matières).
        """
        if len(lines) < 5:
            return False
        
        # Analyser les positions X pour détecter deux colonnes distinctes
        x_positions = []
        for line in lines[:20]:  # Analyser les 20 premières lignes
            if not line:
                continue
            
            # Gérer les lignes multi-colonnes (dictionnaires)
            if isinstance(line, dict) and line.get('type') == 'multi_column':
                # Si on a déjà des lignes multi-colonnes, c'est probablement un layout multi-colonnes
                return True
            elif isinstance(line, list):
                # Ligne normale : extraire les positions X
                for word in line:
                    if isinstance(word, dict) and 'left' in word:
                        x_positions.append(word['left'])
        
        if len(x_positions) < 10:
            return False
        
        # Trouver les clusters de positions X
        x_positions.sort()
        
        # Séparer en deux groupes (gauche et droite)
        midpoint = img_width / 2
        
        left_words = sum(1 for x in x_positions if x < midpoint)
        right_words = sum(1 for x in x_positions if x >= midpoint)
        
        # Si environ 40-60% des mots sont à gauche et 40-60% à droite, c'est probablement 2 colonnes
        total_words = len(x_positions)
        left_ratio = left_words / total_words if total_words > 0 else 0
        right_ratio = right_words / total_words if total_words > 0 else 0
        
        # Vérifier aussi si les lignes ont souvent des mots des deux côtés
        lines_with_both_sides = 0
        for line in lines[:20]:
            if not line:
                continue
            has_left = any(w['left'] < midpoint for w in line)
            has_right = any(w['left'] >= midpoint for w in line)
            if has_left and has_right:
                lines_with_both_sides += 1
        
        both_sides_ratio = lines_with_both_sides / min(20, len(lines)) if lines else 0
        
        # Détection : si bonne distribution gauche/droite OU lignes avec deux côtés
        is_multicolumn = (0.3 <= left_ratio <= 0.7 and 0.3 <= right_ratio <= 0.7) or both_sides_ratio > 0.3
        
        return is_multicolumn
    
    def _render_multicolumn_layout(self, doc, lines, img_width, column_structure=None):
        """
        Rend un layout à colonnes multiples en utilisant un tableau Word.
        Chaque ligne de texte devient une ligne du tableau avec 2 cellules.
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Créer un tableau avec une ligne par ligne de texte
        # On ne crée pas toutes les lignes d'un coup, on les ajoute au fur et à mesure
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Normal Table'
        
        # Largeur des colonnes (50% chacune)
        for col in table.columns:
            col.width = Inches(3.25)
        
        for line in lines:
            if not line:
                # Ligne vide : ajouter quand même une ligne vide pour préserver l'espacement
                row = table.add_row()
                row.cells[0].paragraphs[0].text = ""
                row.cells[1].paragraphs[0].text = ""
                continue
            
            # Ajouter une nouvelle ligne
            row = table.add_row()
            left_cell = row.cells[0]
            right_cell = row.cells[1]
            
            # Gérer les lignes multi-colonnes (déjà séparées)
            if isinstance(line, dict) and line.get('type') == 'multi_column':
                left_words = line['columns'][0] if len(line['columns']) > 0 else []
                right_words = line['columns'][1] if len(line['columns']) > 1 else []
            elif isinstance(line, list):
                # Ligne normale : vérifier si on doit la traiter comme multi-colonnes
                if column_structure and column_structure.get('has_columns'):
                    split_x = column_structure.get('split_x', img_width / 2)
                    left_words = [w for w in line if isinstance(w, dict) and w.get('left', 0) < split_x]
                    right_words = [w for w in line if isinstance(w, dict) and w.get('left', 0) >= split_x]
                else:
                    # Pas de colonnes : mettre tout à gauche
                    left_words = line if isinstance(line, list) else []
                    right_words = []
            else:
                continue
            
            # Trier les mots par position X pour avoir l'ordre correct
            left_words = sorted(left_words, key=lambda w: w.get('left', 0)) if left_words else []
            right_words = sorted(right_words, key=lambda w: w.get('left', 0)) if right_words else []
            
            # Remplir la cellule gauche
            if left_words:
                left_text = " ".join([w['text'] for w in left_words])
                left_cell.paragraphs[0].text = left_text.strip()
                
                # Alignement à droite pour les numéros de page uniquement
                left_text_clean = left_text.strip()
                if left_text_clean.isdigit() or (len(left_text_clean) <= 4 and left_text_clean.replace('.', '').isdigit()):
                    left_cell.paragraphs[0].alignment = 2  # Right align
            
            # Remplir la cellule droite
            if right_words:
                right_text = " ".join([w['text'] for w in right_words])
                right_cell.paragraphs[0].text = right_text.strip()
                
                # Alignement à droite pour les numéros de page uniquement
                right_text_clean = right_text.strip()
                if right_text_clean.isdigit() or (len(right_text_clean) <= 4 and right_text_clean.replace('.', '').isdigit()):
                    right_cell.paragraphs[0].alignment = 2  # Right align
                
                # Alignement à droite aussi si c'est juste un numéro en fin de ligne
                if right_text_clean and len(right_text_clean) <= 4:
                    # Vérifier si c'est principalement des chiffres
                    if sum(c.isdigit() for c in right_text_clean) >= len(right_text_clean) * 0.8:
                        right_cell.paragraphs[0].alignment = 2  # Right align
    
    def _detect_table_blocks_precise(self, lines, img_width):
        """
        Détecte les blocs de tableaux avec haute précision.
        Analyse la structure verticale et horizontale pour identifier les vrais tableaux.
        
        Returns:
            Liste de dicts {'start_line': int, 'end_line': int, 'columns': list}
        """
        if len(lines) < 3:
            return []
        
        table_blocks = []
        current_block = None
        
        for line_idx in range(len(lines)):
            line = lines[line_idx]
            if not line:
                continue
            
            # Analyser si cette ligne fait partie d'un tableau
            is_table, columns = self._analyze_line_for_table(line, lines, line_idx, img_width)
            
            if is_table:
                if current_block is None:
                    # Démarrer un nouveau bloc de tableau
                    current_block = {
                        'start_line': line_idx,
                        'end_line': line_idx + 1,
                        'columns': columns if columns else []
                    }
                else:
                    # Continuer le bloc actuel
                    current_block['end_line'] = line_idx + 1
                    # Mettre à jour les colonnes (union)
                    if columns:
                        current_block['columns'] = self._merge_columns(current_block['columns'], columns)
            else:
                # Fin d'un bloc de tableau
                if current_block and current_block['end_line'] - current_block['start_line'] >= 3:
                    # Un tableau doit avoir au moins 3 lignes
                    table_blocks.append(current_block)
                current_block = None
        
        # Ajouter le dernier bloc s'il existe
        if current_block and current_block['end_line'] - current_block['start_line'] >= 3:
            table_blocks.append(current_block)
        
        return table_blocks
    
    def _analyze_line_for_table(self, line, all_lines, line_idx, img_width):
        """
        Analyse une ligne pour déterminer si elle fait partie d'un tableau.
        
        Returns:
            (is_table: bool, columns: list of x positions)
        """
        if len(line) < 3:
            return False, []
        
        # Calculer les écarts entre les mots
        gaps = []
        x_positions = []
        
        for i, word in enumerate(line):
            x_positions.append(word['left'])
            if i > 0:
                prev_x_end = line[i-1]['left'] + line[i-1]['width']
                curr_x = word['left']
                gap = curr_x - prev_x_end
                gaps.append(gap)
        
        # Détecter les grands écarts (signe de colonnes)
        large_gaps = [g for g in gaps if g > 30]  # Seuil de 30 pixels
        
        # Si au moins 2 grands écarts, probablement un tableau
        if len(large_gaps) >= 2:
            # Retourner les positions X des colonnes
            columns = self._detect_column_positions(line, img_width)
            return True, columns
        
        # Vérifier si c'est une ligne avec beaucoup de chiffres (tableau financier)
        text_content = " ".join([w['text'] for w in line])
        digit_count = sum(c.isdigit() for c in text_content)
        digit_ratio = digit_count / len(text_content) if text_content else 0
        
        # Si beaucoup de chiffres et plusieurs mots, probablement un tableau
        if digit_ratio > 0.25 and len(line) >= 4:
            columns = self._detect_column_positions(line, img_width)
            return True, columns
        
        # Vérifier l'alignement avec les lignes voisines (signe de tableau structuré)
        if line_idx > 0 and line_idx < len(all_lines) - 1:
            prev_line = all_lines[line_idx - 1]
            next_line = all_lines[line_idx + 1] if line_idx + 1 < len(all_lines) else None
            
            # Si les lignes précédente et suivante ont des positions X similaires, c'est un tableau
            if prev_line and next_line:
                alignment_score = self._calculate_alignment_score([prev_line, line, next_line])
                if alignment_score > 0.6:  # Seuil de 60% d'alignement
                    columns = self._detect_column_positions(line, img_width)
                    return True, columns
        
        return False, []
    
    def _detect_column_positions(self, line, img_width):
        """
        Détecte les positions X des colonnes à partir d'une ligne.
        """
        x_positions = [w['left'] for w in line]
        
        # Clusteriser les positions X pour trouver les colonnes
        x_positions.sort()
        columns = []
        
        if x_positions:
            current_cluster = [x_positions[0]]
            cluster_tolerance = 25  # Tolérance de 25 pixels
            
            for x in x_positions[1:]:
                if x - current_cluster[-1] < cluster_tolerance:
                    current_cluster.append(x)
                else:
                    if len(current_cluster) >= 1:
                        columns.append(sum(current_cluster) / len(current_cluster))
                    current_cluster = [x]
            
            if len(current_cluster) >= 1:
                columns.append(sum(current_cluster) / len(current_cluster))
        
        return sorted(columns)
    
    def _merge_columns(self, cols1, cols2):
        """
        Fusionne deux listes de colonnes en trouvant les correspondances.
        """
        if not cols1:
            return cols2
        if not cols2:
            return cols1
        
        all_cols = sorted(cols1 + cols2)
        merged = []
        tolerance = 30
        
        for col in all_cols:
            if not merged:
                merged.append(col)
            else:
                # Vérifier si cette colonne est proche d'une colonne déjà mergée
                closest = min(merged, key=lambda x: abs(x - col))
                if abs(closest - col) > tolerance:
                    merged.append(col)
                else:
                    # Fusionner avec la colonne existante
                    idx = merged.index(closest)
                    merged[idx] = (merged[idx] + col) / 2
        
        return sorted(merged)
    
    def _calculate_alignment_score(self, lines):
        """
        Calcule un score d'alignement entre plusieurs lignes.
        Score entre 0 et 1 : 1 = parfait alignement, 0 = aucun alignement.
        """
        if len(lines) < 2:
            return 0.0
        
        # Extraire toutes les positions X de toutes les lignes
        all_x_positions = []
        for line in lines:
            for word in line:
                all_x_positions.append(word['left'])
        
        if not all_x_positions:
            return 0.0
        
        # Clusteriser les positions X
        all_x_positions.sort()
        clusters = []
        current_cluster = [all_x_positions[0]]
        tolerance = 25
        
        for x in all_x_positions[1:]:
            if x - current_cluster[-1] < tolerance:
                current_cluster.append(x)
            else:
                clusters.append(current_cluster)
                current_cluster = [x]
        
        if current_cluster:
            clusters.append(current_cluster)
        
        # Calculer le score : combien de mots s'alignent avec les clusters
        aligned_count = 0
        total_words = sum(len(line) for line in lines)
        
        for line in lines:
            for word in line:
                word_x = word['left']
                # Vérifier si ce mot s'aligne avec un cluster
                for cluster in clusters:
                    cluster_center = sum(cluster) / len(cluster)
                    if abs(word_x - cluster_center) < tolerance:
                        aligned_count += 1
                        break
        
        return aligned_count / total_words if total_words > 0 else 0.0
    
    def _render_table_block_precise(self, doc, table_lines, img_width):
        """
        Rend un bloc de tableau avec précision MAXIMALE (Best Practice Industry Standard).
        Utilise une approche professionnelle avec :
        - Détection précise des colonnes
        - Préservation des cellules vides
        - Formatage financier (nombres alignés à droite)
        - Gestion des cellules fusionnées potentielles
        """
        if not table_lines:
            return
        
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 1. DÉTECTION PRÉCISE DES COLONNES (Best Practice)
        # Analyser toutes les lignes pour trouver les colonnes communes
        all_columns = []
        column_votes = {}  # Compteur de votes pour chaque position de colonne
        
        for line in table_lines:
            if not line:
                continue
            
            cols = self._detect_column_positions(line, img_width)
            for col in cols:
                # Voter pour les colonnes proches (clustering)
                found_existing = False
                for existing_col in column_votes:
                    if abs(existing_col - col) < 25:  # Tolérance de 25 pixels
                        column_votes[existing_col] += 1
                        found_existing = True
                        break
                
                if not found_existing:
                    column_votes[col] = 1
        
        # Prendre les colonnes avec le plus de votes (au moins 2 votes)
        final_columns = sorted([col for col, votes in column_votes.items() if votes >= 2])
        
        # Si pas assez de colonnes détectées, utiliser toutes les colonnes
        if len(final_columns) < 2:
            final_columns = sorted(set(all_columns)) if all_columns else [img_width * 0.2, img_width * 0.6]
        
        # 2. CRÉER LE TABLEAU WORD
        num_cols = max(len(final_columns), 2)
        table = doc.add_table(rows=len(table_lines), cols=num_cols)
        table.style = 'Table Grid'  # Bordures pour visibilité
        
        # Calculer les largeurs de colonnes proportionnelles
        total_width_points = 7 * 72  # 7 inches en points
        if len(final_columns) > 1:
            col_widths = []
            for i in range(len(final_columns)):
                if i == 0:
                    width = final_columns[0]
                elif i < len(final_columns):
                    width = final_columns[i] - final_columns[i-1]
                else:
                    width = img_width - final_columns[-1]
                col_widths.append(width)
            
            # Normaliser les largeurs
            total_width = sum(col_widths)
            for i, col in enumerate(table.columns):
                if i < len(col_widths):
                    relative_width = col_widths[i] / total_width if total_width > 0 else 1.0 / num_cols
                    col.width = Inches(relative_width * 7)
                else:
                    col.width = Inches(7.0 / num_cols)
        else:
            for col in table.columns:
                col.width = Inches(7.0 / num_cols)
        
        # 3. REMPLIR LE TABLEAU AVEC PRÉCISION
        for row_idx, line in enumerate(table_lines):
            if not line:
                continue
            
            row = table.rows[row_idx]
            
            # Créer un mapping mot -> colonne
            word_column_map = {}
            for word in line:
                word_x = word['left']
                
                # Trouver la colonne la plus proche
                if final_columns:
                    closest_col_idx = min(range(len(final_columns)),
                                         key=lambda i: abs(final_columns[i] - word_x))
                    closest_col_idx = min(closest_col_idx, num_cols - 1)
                else:
                    closest_col_idx = 0
                
                # Grouper les mots par colonne
                if closest_col_idx not in word_column_map:
                    word_column_map[closest_col_idx] = []
                word_column_map[closest_col_idx].append(word)
            
            # Remplir chaque cellule
            for col_idx in range(num_cols):
                cell = row.cells[col_idx]
                
                if col_idx in word_column_map:
                    # Joindre les mots de cette colonne (triés par X)
                    words_in_col = sorted(word_column_map[col_idx], key=lambda w: w['left'])
                    cell_text = " ".join([w['text'] for w in words_in_col])
                    cell.paragraphs[0].text = cell_text
                    
                    # FORMATAGE FINANCIER (Best Practice)
                    # Détecter si c'est un nombre financier
                    is_financial_number = self._is_financial_number(cell_text)
                    
                    if is_financial_number:
                        # Alignement à droite pour les nombres
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        # Formatage du texte : nettoyer et formater
                        para = cell.paragraphs[0]
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                else:
                    # Cellule vide : préserver (Best Practice)
                    cell.paragraphs[0].text = ""
        
        # 4. POST-PROCESSING : Détecter et formater les en-têtes de colonnes
        if len(table_lines) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                if cell.text.strip():
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(10)
    
    def _is_financial_number(self, text):
        """
        Détecte si un texte est un nombre financier (Best Practice).
        """
        if not text or not text.strip():
            return False
        
        text_clean = text.strip()
        
        # Supprimer les caractères monétaires et ponctuation
        clean = text_clean.replace(',', '').replace('.', '').replace('(', '').replace(')', '').replace('£', '').replace('$', '').replace('€', '').replace('-', '').replace('m', '').replace('k', '').strip()
        
        # Vérifier si c'est principalement des chiffres
        if not clean:
            return False
        
        digit_ratio = sum(c.isdigit() for c in clean) / len(clean)
        
        # C'est un nombre financier si :
        # - Ratio de chiffres > 50%
        # - OU commence par un chiffre ou signe négatif
        # - OU contient des caractères monétaires
        return (digit_ratio > 0.5) or \
               (text_clean[0].isdigit() if text_clean else False) or \
               (text_clean[0] == '(' and text_clean[-1] == ')') or \
               ('£' in text_clean or '$' in text_clean or '€' in text_clean)
    
    def _reconstruct_page_from_pdfplumber(self, doc, page):
        """
        Reconstruit une page Word à partir des données pdfplumber.
        Préserve les positions spatiales, colonnes, tableaux et alignements.
        
        Args:
            doc: Document Word
            page: Page pdfplumber
        """
        # Extraire les mots avec leurs coordonnées
        words = page.extract_words(x_tolerance=3, y_tolerance=3)
        
        if not words:
            return
        
        # Extraire les tableaux
        tables = page.extract_tables()
        table_bboxes = []
        for table in tables:
            if table:
                # Calculer la bbox du tableau (première et dernière cellule)
                # Note: pdfplumber ne donne pas directement la bbox, on l'estime
                table_bboxes.append(None)  # On utilisera une heuristique différente
        
        # Grouper les mots par ligne (même Y)
        lines = self._group_words_into_lines(words)
        
        # Détecter les blocs (texte, tableaux, titres)
        blocks = self._detect_blocks_from_lines(lines, page.width)
        
        # Rendre chaque bloc dans Word
        for block in blocks:
            if block['type'] == 'table':
                # Créer un tableau Word à partir des lignes
                # Flatten toutes les lignes du bloc en une seule liste de mots
                all_words = []
                for line in block['lines']:
                    all_words.extend(line)
                self._create_table_from_words(doc, all_words, page.width)
            elif block['type'] == 'title':
                # Titre formaté
                for line in block['lines']:
                    text = " ".join([w.get('text', '') for w in line])
                    if text.strip():
                        para = doc.add_paragraph(text.strip())
                        para.style = 'Heading 2'
                        for run in para.runs:
                            run.bold = True
            else:
                # Texte narratif normal
                for line in block['lines']:
                    text = " ".join([w.get('text', '') for w in line])
                    if text.strip():
                        doc.add_paragraph(text.strip())
    
    def _group_words_into_lines(self, words):
        """
        Groupe les mots en lignes basées sur leur position Y.
        
        Args:
            words: Liste de mots de pdfplumber (dicts avec 'text', 'x0', 'top', etc.)
        
        Returns:
            Liste de lignes, chaque ligne = liste de mots triés par X
        """
        if not words:
            return []
        
        # Trier par Y puis X
        words_sorted = sorted(words, key=lambda w: (w.get('top', 0), w.get('x0', 0)))
        
        lines = []
        current_line = []
        current_y = None
        y_tolerance = 5  # Tolérance pour considérer que c'est la même ligne
        
        for word in words_sorted:
            word_y = word.get('top', 0)
            
            if current_y is None:
                current_y = word_y
                current_line.append(word)
            elif abs(word_y - current_y) <= y_tolerance:
                current_line.append(word)
            else:
                # Nouvelle ligne
                if current_line:
                    current_line.sort(key=lambda w: w.get('x0', 0))
                    lines.append(current_line)
                current_line = [word]
                current_y = word_y
        
        if current_line:
            current_line.sort(key=lambda w: w.get('x0', 0))
            lines.append(current_line)
        
        return lines
    
    def _detect_blocks_from_lines(self, lines, page_width):
        """
        Détecte les types de blocs (texte, tableau, titre) à partir des lignes.
        
        Args:
            lines: Liste de lignes (chaque ligne = liste de mots)
            page_width: Largeur de la page en points
        
        Returns:
            Liste de blocs avec 'type' et 'lines' ou 'words'
        """
        blocks = []
        if not lines:
            return blocks
        
        current_block = None
        
        for line in lines:
            if not line:
                continue
            
            # Classifier la ligne
            line_type = self._classify_line_from_words(line, page_width)
            
            # Si même type que le bloc actuel, ajouter à ce bloc
            if current_block and current_block['type'] == line_type:
                if 'lines' in current_block:
                    current_block['lines'].append(line)
                else:
                    current_block['lines'] = [line]
            else:
                # Nouveau bloc
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': line_type, 'lines': [line]}
        
        if current_block:
            blocks.append(current_block)
        
        return blocks
    
    def _classify_line_from_words(self, line, page_width):
        """
        Classifie une ligne comme 'table', 'title', ou 'text'.
        
        Args:
            line: Liste de mots (dicts de pdfplumber)
            page_width: Largeur de la page en points
        
        Returns:
            'table', 'title', ou 'text'
        """
        if not line:
            return 'text'
        
        # Calculer les écarts entre les mots
        gaps = []
        text_content = ""
        for i, word in enumerate(line):
            if i > 0:
                prev_x1 = line[i-1].get('x1', 0)
                curr_x0 = word.get('x0', 0)
                gap = curr_x0 - prev_x1
                gaps.append(gap)
            text_content += word.get('text', '') + " "
        
        # Détecter les grands écarts (colonnes)
        large_gaps = [g for g in gaps if g > 20]  # 20 points = colonne
        
        # Compter les chiffres
        digit_count = sum(c.isdigit() for c in text_content)
        digit_ratio = digit_count / len(text_content) if text_content else 0
        
        # Heuristique pour tableau : grands écarts OU beaucoup de chiffres
        if len(large_gaps) >= 1 or (len(line) > 3 and digit_ratio > 0.3):
            return 'table'
        
        # Heuristique pour titre
        text_stripped = text_content.strip()
        if text_stripped.isupper() and len(text_stripped) < 100 and len(text_stripped) > 5:
            return 'title'
        
        return 'text'
    
    def _create_table_from_words(self, doc, words, page_width):
        """
        Crée un tableau Word à partir des mots.
        
        Args:
            doc: Document Word
            words: Liste de mots (dicts de pdfplumber)
            page_width: Largeur de la page en points
        """
        if not words:
            return
        
        # Grouper en lignes
        lines = self._group_words_into_lines(words)
        
        if not lines:
            return
        
        # Détecter les positions X des colonnes
        x_positions = []
        for line in lines:
            for word in line:
                x_positions.append(word.get('x0', 0))
        
        # Clusteriser les X
        x_positions.sort()
        column_positions = []
        if x_positions:
            current_cluster = [x_positions[0]]
            for x in x_positions[1:]:
                if x - current_cluster[-1] < 30:  # Tolérance 30 points
                    current_cluster.append(x)
                else:
                    column_positions.append(sum(current_cluster) / len(current_cluster))
                    current_cluster = [x]
            column_positions.append(sum(current_cluster) / len(current_cluster))
        
        num_cols = len(column_positions) if column_positions else 1
        
        # Créer le tableau Word
        table = doc.add_table(rows=len(lines), cols=num_cols)
        table.style = 'Normal Table'
        
        # Remplir le tableau
        for i, line in enumerate(lines):
            row = table.rows[i]
            for word in line:
                # Trouver la colonne la plus proche
                word_x = word.get('x0', 0)
                closest_col_idx = min(range(num_cols), key=lambda k: abs(column_positions[k] - word_x)) if column_positions else 0
                
                # Ajouter le texte dans la cellule
                cell = row.cells[closest_col_idx]
                if cell.text:
                    cell.text += " " + word.get('text', '')
                else:
                    cell.text = word.get('text', '')
                
                # Alignement à droite pour les chiffres
                text = word.get('text', '')
                clean_text = text.replace(',', '').replace('.', '').replace('(', '').replace(')', '').replace('£', '').replace('$', '').strip()
                if clean_text and (clean_text.isdigit() or (clean_text.startswith('-') and clean_text[1:].isdigit())):
                    cell.paragraphs[0].alignment = 2  # Right align
    
    def _reconstruct_page_layout_aware(self, doc, data, page_width, page_height):
        """
        Reconstruction intelligente de la page avec préservation du layout.
        Détecte les colonnes, tableaux, titres et paragraphes.
        
        Args:
            doc: Document Word
            data: Données OCR de Tesseract (format DICT)
            page_width: Largeur de la page en pixels
            page_height: Hauteur de la page en pixels
        """
        # Grouper les mots en lignes
        lines = self._group_ocr_data_by_lines(data)
        
        if not lines:
            return
        
        # Détecter les colonnes et les blocs (texte vs tableaux)
        blocks = self._detect_layout_blocks(lines, page_width)
        
        # Rendre chaque bloc dans Word
        for block in blocks:
            if block['type'] == 'table':
                # Créer un tableau Word
                self._create_word_table_from_lines(doc, block['lines'], page_width)
            elif block['type'] == 'title':
                # Titre formaté
                for line in block['lines']:
                    text = " ".join([w['text'] for w in line])
                    if text.strip():
                        para = doc.add_paragraph(text.strip())
                        para.style = 'Heading 2'
                        for run in para.runs:
                            run.bold = True
            else:
                # Texte narratif normal
                for line in block['lines']:
                    text = " ".join([w['text'] for w in line])
                    if text.strip():
                        doc.add_paragraph(text.strip())
    
    def _group_ocr_data_by_lines(self, data):
        """
        Groupe les données OCR en lignes visuelles.
        
        Args:
            data: Données OCR de Tesseract (format DICT)
        
        Returns:
            Liste de lignes, chaque ligne étant une liste de mots (dicts)
        """
        n_boxes = len(data['text'])
        words = []
        for i in range(n_boxes):
            if int(data['conf'][i]) > 30 and data['text'][i].strip():
                words.append({
                    'text': data['text'][i],
                    'left': data['left'][i],
                    'top': data['top'][i],
                    'width': data['width'][i],
                    'height': data['height'][i]
                })
        
        # Tri Y puis X
        words.sort(key=lambda x: (x['top'], x['left']))
        
        lines = []
        current_line = []
        current_y = -1
        y_tolerance = 15
        
        for word in words:
            if current_y == -1:
                current_y = word['top']
                current_line.append(word)
            elif abs(word['top'] - current_y) <= y_tolerance:
                current_line.append(word)
            else:
                current_line.sort(key=lambda x: x['left'])
                lines.append(current_line)
                current_line = [word]
                current_y = word['top']
        
        if current_line:
            current_line.sort(key=lambda x: x['left'])
            lines.append(current_line)
        
        return lines
    
    def _detect_layout_blocks(self, lines, page_width):
        """
        Détecte les types de blocs (texte, tableau, titre) dans les lignes.
        
        Args:
            lines: Liste de lignes (chaque ligne = liste de mots)
            page_width: Largeur de la page en pixels
        
        Returns:
            Liste de blocs avec 'type' et 'lines'
        """
        blocks = []
        if not lines:
            return blocks
        
        current_block = None
        
        for line in lines:
            if not line:
                continue
            
            # Analyser la ligne pour déterminer son type
            line_type = self._classify_line(line, page_width)
            
            # Si même type que le bloc actuel, ajouter à ce bloc
            if current_block and current_block['type'] == line_type:
                current_block['lines'].append(line)
            else:
                # Nouveau bloc
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': line_type, 'lines': [line]}
        
        if current_block:
            blocks.append(current_block)
        
        return blocks
    
    def _classify_line(self, line, page_width):
        """
        Classifie une ligne comme 'table', 'title', ou 'text'.
        
        Args:
            line: Liste de mots (dicts)
            page_width: Largeur de la page en pixels
        
        Returns:
            'table', 'title', ou 'text'
        """
        if not line:
            return 'text'
        
        # Calculer les écarts entre les mots
        gaps = []
        text_content = ""
        for i, word in enumerate(line):
            if i > 0:
                gap = word['left'] - (line[i-1]['left'] + line[i-1]['width'])
                gaps.append(gap)
            text_content += word['text'] + " "
        
        # Détecter les grands écarts (colonnes)
        large_gaps = [g for g in gaps if g > 30]
        
        # Compter les chiffres
        digit_count = sum(c.isdigit() for c in text_content)
        digit_ratio = digit_count / len(text_content) if text_content else 0
        
        # Heuristique pour tableau : grands écarts OU beaucoup de chiffres
        if len(large_gaps) >= 1 or (len(line) > 3 and digit_ratio > 0.3):
            return 'table'
        
        # Heuristique pour titre : court, majuscules, ou centré
        text_stripped = text_content.strip()
        if (text_stripped.isupper() and len(text_stripped) < 100 and len(text_stripped) > 5) or \
           (len(text_stripped) < 80 and len(line) <= 5):
            # Vérifier si centré (position X proche du centre)
            first_word_x = line[0]['left']
            last_word_x = line[-1]['left'] + line[-1]['width']
            line_center = (first_word_x + last_word_x) / 2
            page_center = page_width / 2
            if abs(line_center - page_center) < page_width * 0.15:  # Tolérance 15%
                return 'title'
            elif text_stripped.isupper():
                return 'title'
        
        return 'text'
    
    def _create_word_table_from_lines(self, doc, lines, page_width):
        """
        Crée un tableau Word pour un bloc de lignes détectées comme tabulaires.
        
        Args:
            doc: Document Word
            lines: Liste de lignes, chaque ligne étant une liste de mots (dicts)
            page_width: Largeur de la page en pixels
        """
        if not lines:
            return
        
        # Détecter les positions X des colonnes
        x_starts = []
        for line in lines:
            for word in line:
                x_starts.append(word['left'])
        
        # Clusteriser les X proches (même colonne)
        x_starts.sort()
        columns_clusters = []
        if x_starts:
            current_cluster = [x_starts[0]]
            for x in x_starts[1:]:
                if x - current_cluster[-1] < 50:  # Tolérance de 50px
                    current_cluster.append(x)
                else:
                    columns_clusters.append(sum(current_cluster) / len(current_cluster))
                    current_cluster = [x]
            columns_clusters.append(sum(current_cluster) / len(current_cluster))
        
        num_cols = len(columns_clusters)
        if num_cols == 0:
            num_cols = 1
        
        # Créer le tableau Word
        table = doc.add_table(rows=len(lines), cols=num_cols)
        table.style = 'Normal Table'
        
        # Remplir le tableau
        for i, line in enumerate(lines):
            row = table.rows[i]
            for word in line:
                # Trouver la colonne la plus proche
                word_center = word['left']
                closest_col_idx = min(range(num_cols), key=lambda k: abs(columns_clusters[k] - word_center))
                
                # Ajouter le texte dans la cellule
                cell = row.cells[closest_col_idx]
                if cell.text:
                    cell.text += " " + word['text']
                else:
                    cell.text = word['text']
                
                # Alignement à droite pour les chiffres
                clean_text = word['text'].replace(',', '').replace('.', '').replace('(', '').replace(')', '').replace('£', '').replace('$', '').strip()
                if clean_text and (clean_text.isdigit() or (clean_text.startswith('-') and clean_text[1:].isdigit())):
                    cell.paragraphs[0].alignment = 2  # Right align

    def _add_text_blocks_to_doc(self, doc, text_blocks):
        """
        Groupe les données brutes de Tesseract en lignes visuelles.
        
        Args:
            data: Données OCR de Tesseract (format DICT)
        
        Returns:
            Liste de lignes, chaque ligne étant une liste de mots (dicts avec 'text', 'left', 'top', etc.)
        """
        n_boxes = len(data['text'])
        words = []
        for i in range(n_boxes):
            if int(data['conf'][i]) > 30 and data['text'][i].strip():
                words.append({
                    'text': data['text'][i],
                    'left': data['left'][i],
                    'top': data['top'][i],
                    'width': data['width'][i],
                    'height': data['height'][i]
                })
        
        # Tri Y puis X
        words.sort(key=lambda x: (x['top'], x['left']))
        
        lines = []
        current_line = []
        current_y = -1
        y_tolerance = 15
        
        for word in words:
            if current_y == -1:
                current_y = word['top']
                current_line.append(word)
            elif abs(word['top'] - current_y) <= y_tolerance:
                current_line.append(word)
            else:
                current_line.sort(key=lambda x: x['left'])
                lines.append(current_line)
                current_line = [word]
                current_y = word['top']
        
        if current_line:
            current_line.sort(key=lambda x: x['left'])
            lines.append(current_line)
        
        return lines
    
    def _add_text_blocks_to_doc(self, doc, text_blocks):
        """
        Ajoute des blocs de texte au document Word avec détection intelligente de structure.
        
        Args:
            doc: Document Word
            text_blocks: Liste de dictionnaires avec 'text', 'y', etc.
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Grouper en lignes
        lines = []
        current_line = []
        current_y = 0
        y_tolerance = 15
        
        for block in text_blocks:
            if abs(block['y'] - current_y) > y_tolerance:
                if current_line:
                    current_line.sort(key=lambda b: b['x'])
                    lines.append(current_line)
                current_line = [block]
                current_y = block['y']
            else:
                current_line.append(block)
        
        if current_line:
            current_line.sort(key=lambda b: b['x'])
            lines.append(current_line)
        
        # Ajouter au document
        for line_blocks in lines:
            line_text = ' '.join([b['text'] for b in line_blocks])
            line_text = line_text.strip()
            
            if not line_text:
                continue
            
            # Détection améliorée de titres
            if line_text.isupper() and len(line_text) < 100 and len(line_text) > 3:
                para = doc.add_paragraph(line_text)
                para.style = 'Heading 1' if len(line_text) < 50 else 'Heading 2'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
            elif (line_text[0].isupper() and len(line_text) < 80 and 
                  not line_text.endswith('.') and not line_text.endswith(',') and
                  len(line_text.split()) < 15):
                para = doc.add_paragraph(line_text)
                para.style = 'Heading 2'
            else:
                doc.add_paragraph(line_text)



    def convert(self, start_page=0, end_page=None, use_parallel=None, pages_per_chunk=20, 
                use_ocr=False, auto_detect_scanned=False, **kwargs):
        """
        Convertit le PDF en DOCX en préservant la structure originale.
        Utilise la parallélisation automatiquement pour les gros PDFs (>100 pages).
        
        Args:
            start_page: Page de départ (0-indexed)
            end_page: Page de fin (None = toutes les pages)
            use_parallel: Force l'utilisation (True) ou non (False) de la parallélisation.
                         Si None, détection automatique basée sur le nombre de pages.
            pages_per_chunk: Nombre de pages par chunk pour la parallélisation (défaut: 20)
            use_ocr: Force l'utilisation de l'OCR (pour PDFs scannés)
            auto_detect_scanned: Détecte automatiquement si le PDF est scanné
            **kwargs: Arguments supplémentaires pour pdf2docx
        
        Raises:
            Exception: Si la conversion échoue
        """
        self.docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        # ⭐ DÉTECTION AUTOMATIQUE DES PDFs SCANNÉS
        if auto_detect_scanned and not use_ocr:
            is_scanned = self.is_scanned_pdf()
            if is_scanned:
                use_ocr = True
                logger.info("📷 Scanned PDF detected → Switching to OCR mode")
        
        # ⭐ MODE OCR (pour PDFs scannés)
        if use_ocr:
            if not OCRMYPDF_AVAILABLE:
                raise ImportError(
                    "\n" + "="*70 + "\n"
                    "❌ OCRmyPDF NOT AVAILABLE\n"
                    "="*70 + "\n\n"
                    "This PDF appears to be scanned (image-based).\n"
                    "OCRmyPDF is required for OCR conversion.\n\n"
                    "📦 INSTALLATION:\n\n"
                    "1. Install OCRmyPDF:\n"
                    "   pip install ocrmypdf\n\n"
                    "2. Install system dependencies:\n"
                    "   macOS: brew install tesseract ghostscript\n"
                    "   Ubuntu: sudo apt-get install tesseract-ocr ghostscript\n"
                    "   Windows: Download from tesseract-ocr.github.io and ghostscript.com\n\n"
                    "3. (Optional) Install language packs:\n"
                    "   macOS: brew install tesseract-lang\n"
                    "   Ubuntu: sudo apt-get install tesseract-ocr-fra (for French, etc.)\n"
                    "="*70
                )
            
            logger.info("🔄 Using OCR mode (OCRmyPDF + pdf2docx)")
            self.convert_with_ocr(start_page, end_page)
            self._apply_post_processing()
            logger.info("✅ OCR conversion completed, returning now")
            return  # CRUCIAL: Empêche l'exécution du code de conversion normale
        
        # ⭐ MODE NORMAL (reste inchangé)
        if not PDF2DOCX_AVAILABLE:
            raise ImportError("pdf2docx is not available. Install with: pip install pdf2docx")
        
        # Détecter le nombre de pages pour décider de la parallélisation
        total_pages = self._get_pdf_page_count()
        if end_page is not None:
            pages_to_convert = end_page - start_page
        elif total_pages:
            pages_to_convert = total_pages - start_page
        else:
            pages_to_convert = None
        
        # Décision de parallélisation
        should_parallelize = False
        if use_parallel is True:
            should_parallelize = True
        elif use_parallel is False:
            should_parallelize = False
        elif pages_to_convert and pages_to_convert > 100:
            should_parallelize = True
            logger.info(f"Large PDF detected ({pages_to_convert} pages), using parallel conversion")
        
        # Parallélisation pour gros PDFs
        if should_parallelize and pages_to_convert and pages_to_convert > 50:
            try:
                self._convert_parallel(start_page, end_page, pages_per_chunk, **kwargs)
                logger.info("PDF conversion completed successfully (parallel)")
                # Post-traitement après conversion parallèle
                self._apply_post_processing()
                return
            except Exception as e:
                logger.warning(f"Parallel conversion failed, falling back to sequential: {e}")
                # Continue avec la conversion normale
        
        # Conversion normale (séquentielle)
        try:
            self.converter = Converter(str(self.pdf_path))
            
            # Paramètres de conversion optimisés pour préserver la structure
            convert_params = {
                'start': start_page,
                'end': end_page,
            }
            convert_params.update(kwargs)
            
            # Supprimer les logs de pdf2docx en mode simple
            with suppress_output(verbose=self.verbose):
                # Conversion avec paramètres par défaut de pdf2docx
                # (la bibliothèque gère mieux la structure multi-colonnes sans intervention)
                self.converter.convert(str(self.docx_path), **convert_params)
            
            if self.verbose:
                logger.info("PDF conversion completed successfully")
            
        except Exception as e:
            # Nettoyer le convertisseur si ouvert
            if self.converter:
                try:
                    self.converter.close()
                except:
                    pass
            
            error_msg = str(e)
            if "IndexError" in str(type(e).__name__) or "vertically_align" in error_msg:
                raise Exception(
                    f"PDF conversion failed. This PDF has a complex structure that pdf2docx cannot process.\n"
                    f"Error details: {error_msg}\n"
                    f"Suggestions:\n"
                    f"  - Try converting with Adobe Acrobat\n"
                    f"  - The PDF may be corrupted or have non-standard formatting\n"
                    f"  - You can still analyze the PDF for sensitive information (option 3)"
                )
            else:
                raise Exception(f"PDF conversion failed: {error_msg}")
        finally:
            if self.converter:
                try:
                    self.converter.close()
                except:
                    pass
            
            # Post-traitement pour améliorer la conversion (si pas déjà fait en parallèle)
            self._apply_post_processing()
    
    def _convert_parallel(self, start_page=0, end_page=None, pages_per_chunk=20, **kwargs):
        """
        Convertit le PDF en utilisant la parallélisation par chunks de pages.
        
        Args:
            start_page: Page de départ (0-indexed)
            end_page: Page de fin (None = toutes les pages)
            pages_per_chunk: Nombre de pages par chunk
            **kwargs: Arguments supplémentaires (non utilisés en mode parallèle)
        """
        # Obtenir le nombre total de pages
        total_pages = self._get_pdf_page_count()
        if not total_pages:
            raise ValueError("Could not determine PDF page count")
        
        # Calculer les limites
        actual_start = start_page
        actual_end = end_page if end_page is not None else total_pages
        
        # Créer les chunks
        chunks = []
        temp_dir = Path(tempfile.mkdtemp(prefix='pdf_conversion_'))
        
        try:
            chunk_idx = 0
            for chunk_start in range(actual_start, actual_end, pages_per_chunk):
                chunk_end = min(chunk_start + pages_per_chunk, actual_end)
                chunk_output = temp_dir / f"chunk_{chunk_idx}.docx"
                chunks.append((str(self.pdf_path), str(chunk_output), chunk_start, chunk_end))
                chunk_idx += 1
            
            logger.info(f"Converting {len(chunks)} chunks in parallel (pages {actual_start}-{actual_end})")
            
            # Déterminer le nombre de processus (max CPU cores, min chunks)
            num_processes = min(cpu_count(), len(chunks), 4)  # Max 4 processus pour éviter surcharge
            
            # Convertir en parallèle
            with Pool(processes=num_processes) as pool:
                results = pool.map(_convert_chunk, chunks)
            
            # Filtrer les résultats (enlever None)
            successful_chunks = [r for r in results if r and Path(r).exists()]
            
            if not successful_chunks:
                raise Exception("All chunks failed to convert")
            
            if len(successful_chunks) < len(chunks):
                logger.warning(f"Only {len(successful_chunks)}/{len(chunks)} chunks succeeded")
            
            # Fusionner les chunks
            logger.info("Merging chunks...")
            if not self._merge_docx_files(successful_chunks):
                raise Exception("Failed to merge DOCX chunks")
            
        finally:
            # Nettoyer les fichiers temporaires
            try:
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Could not clean up temp directory: {e}")
    
    def _merge_docx_files(self, docx_files):
        """
        Fusionne plusieurs fichiers DOCX en un seul.
        
        Args:
            docx_files: Liste des chemins de fichiers DOCX à fusionner
        
        Returns:
            bool: True si succès, False sinon
        """
        if not PYTHON_DOCX_AVAILABLE:
            return False
        
        try:
            from docx import Document
            import copy
            
            # Créer un nouveau document vide
            merged_doc = Document()
            valid_files = [f for f in docx_files if f and Path(f).exists()]
            
            if not valid_files:
                logger.warning("No valid DOCX files to merge")
                return False
            
            # Fusionner chaque document
            for idx, docx_file in enumerate(valid_files):
                doc = Document(docx_file)
                
                # Copier tous les éléments (paragraphes et tableaux) en utilisant copy.deepcopy
                for element in doc.element.body:
                    # Utiliser copy.deepcopy pour créer une copie profonde de l'élément XML
                    merged_doc.element.body.append(copy.deepcopy(element))
                
                # Ajouter un saut de page entre les chunks (sauf pour le dernier)
                if idx < len(valid_files) - 1:
                    merged_doc.add_page_break()
            
            # Sauvegarder le document fusionné
            merged_doc.save(str(self.docx_path))
            logger.info(f"✓ Merged {len(valid_files)} chunk(s) into {self.docx_path.name}")
            return True
        
        except Exception as e:
            logger.error(f"DOCX merge failed: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _apply_post_processing(self):
        """Applique le post-traitement après conversion."""
        try:
            if PYTHON_DOCX_AVAILABLE:
                logger.info("Post-processing document...")
                
                # 1. Formatage Q&A
                try:
                    self.fix_qa_formatting()
                except Exception as e:
                    logger.warning(f"Q&A formatting skipped: {e}")
                
                # 2. Restaurer les hyperlinks
                try:
                    self.restore_hyperlinks()
                except Exception as e:
                    logger.warning(f"Hyperlink restoration skipped: {e}")
                
                # 3. Reconstruire Table of Contents
                try:
                    self.rebuild_table_of_contents()
                except Exception as e:
                    logger.warning(f"TOC rebuilding skipped: {e}")
                
                # 4. Améliorer formatage des nombres financiers
                try:
                    self.format_financial_numbers()
                except Exception as e:
                    logger.warning(f"Financial number formatting skipped: {e}")
                
                # 5. Évaluer la qualité
                try:
                    metrics = self.assess_conversion_quality()
                    logger.info(f"Conversion quality score: {metrics['quality_score']}/100")
                except Exception as e:
                    logger.warning(f"Quality assessment skipped: {e}")
                    
        except Exception as e:
            logger.warning(f"Post-processing failed: {e}")
    
    def enhance_tables_in_docx(self):
        """
        Post-traite le DOCX pour reconstruire les tableaux avec structure propre.
        CORRIGE les cellules fusionnées/manquantes pour permettre un copier-coller Excel parfait.
        """
        if not PYTHON_DOCX_AVAILABLE:
            logger.warning("python-docx not available, skipping table enhancement")
            return
        
        try:
            # Ouvrir le DOCX généré
            doc = Document(str(self.docx_path))
            
            tables_fixed = 0
            
            for table_idx, table in enumerate(doc.tables):
                try:
                    # ÉTAPE 1 : Analyser la structure du tableau
                    # Compter les colonnes de chaque ligne
                    row_col_counts = []
                    for row in table.rows:
                        row_col_counts.append(len(row.cells))
                    
                    # Trouver le nombre MAXIMUM de colonnes (ligne la plus longue)
                    max_cols = max(row_col_counts) if row_col_counts else 0
                    
                    if max_cols == 0:
                        continue
                    
                    # ÉTAPE 2 : Vérifier si le tableau a des problèmes
                    has_inconsistent_cols = len(set(row_col_counts)) > 1
                    
                    if not has_inconsistent_cols:
                        # Tableau déjà cohérent, juste formater les nombres
                        self._format_table_cells(table)
                        tables_fixed += 1
                        continue
                    
                    # ÉTAPE 3 : Tableau problématique → Le reconstruire
                    logger.info(f"Fixing table {table_idx + 1}: inconsistent columns {set(row_col_counts)}")
                    
                    # Extraire tout le contenu du tableau
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            # Extraire le texte de toutes les cellules
                            cell_text = cell.text.strip()
                            row_data.append(cell_text)
                        
                        # Compléter avec des cellules vides si nécessaire
                        while len(row_data) < max_cols:
                            row_data.append("")
                        
                        table_data.append(row_data)
                    
                    # ÉTAPE 4 : Supprimer l'ancien tableau
                    tbl_element = table._element
                    tbl_element.getparent().remove(tbl_element)
                    
                    # ÉTAPE 5 : Créer un NOUVEAU tableau propre
                    new_table = doc.add_table(rows=len(table_data), cols=max_cols)
                    new_table.style = 'Light Grid Accent 1'  # Style par défaut
                    
                    # Remplir le nouveau tableau
                    for i, row_data in enumerate(table_data):
                        row = new_table.rows[i]
                        for j, cell_value in enumerate(row_data):
                            cell = row.cells[j]
                            cell.text = cell_value
                            
                            # Formater si c'est un nombre
                            if cell_value:
                                clean = cell_value.replace(',', '').replace('$', '').replace('%', '').replace('(', '').replace(')', '').strip()
                                if clean.replace('.', '').replace('-', '').isdigit():
                                    # Nombre → aligner à droite
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = 2  # Right
                    
                    tables_fixed += 1
                    logger.info(f"✓ Table {table_idx + 1} reconstructed with {len(table_data)} rows × {max_cols} cols")
                
                except Exception as e:
                    logger.warning(f"Could not fix table {table_idx + 1}: {e}")
                    continue
            
            # Sauvegarder les modifications
            doc.save(str(self.docx_path))
            
            if tables_fixed > 0:
                logger.info(f"✓ Fixed {tables_fixed} tables with proper cell structure")
        
        except Exception as e:
            logger.warning(f"Table enhancement failed: {e}")
    
    def _format_table_cells(self, table):
        """Formate les cellules d'un tableau (alignement des nombres)."""
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    # Vérifier si c'est un nombre
                    clean = cell_text.replace(',', '').replace('$', '').replace('%', '').replace('(', '').replace(')', '').strip()
                    if clean.replace('.', '').replace('-', '').isdigit():
                        # Aligner à droite
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
    
    def fix_qa_formatting(self):
        """
        Reconstruit la mise en forme Q&A après conversion.
        Détecte les patterns 'Q:' et 'A:' et les formate correctement avec sauts de ligne.
        """
        if not PYTHON_DOCX_AVAILABLE:
            logger.warning("python-docx not available, skipping Q&A formatting")
            return
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            
            doc = Document(str(self.docx_path))
            
            logger.info("Processing Q&A formatting...")
            
            # Parcourir tous les paragraphes
            i = 0
            qa_count = 0
            
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                text = para.text.strip()
                
                # Chercher des patterns 'Q:' ou 'A:' dans le texte
                if not text or len(text) < 3:
                    i += 1
                    continue
                
                # Pattern 1: "Q: [...] A: [...]" sur la même ligne -> Séparer en deux paragraphes
                if text.startswith('Q:') and 'A:' in text and text.find('A:') > 2:
                    # Séparer Q et A
                    a_index = text.find('A:')
                    q_text = text[:a_index].strip()
                    a_text = text[a_index:].strip()
                    
                    # Modifier le paragraphe actuel (Q)
                    para.clear()
                    run = para.add_run(q_text)
                    run.bold = True
                    para.style = 'Normal'
                    para.paragraph_format.space_before = Pt(12)
                    
                    # Pour A, utiliser le paragraphe suivant s'il est vide
                    if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                        next_para = doc.paragraphs[i + 1]
                        next_para.clear()
                        next_para.add_run(a_text)
                        next_para.style = 'Normal'
                        next_para.paragraph_format.left_indent = Inches(0.25)
                        next_para.paragraph_format.space_after = Pt(12)
                        qa_count += 1
                        i += 2
                    else:
                        # Paragraphe suivant occupé, on met juste le texte A après un saut de ligne dans le même para
                        # (Solution de contournement : on garde Q et A séparés visuellement avec formatage)
                        para.add_run('\n' + a_text)
                        qa_count += 1
                        i += 1
                    continue
                
                # Pattern 2: Texte commençant par "Q:" (question seule)
                if re.match(r'^Q:\s+', text):
                    para.clear()
                    run = para.add_run(text)
                    run.bold = True
                    para.style = 'Normal'
                    para.paragraph_format.space_before = Pt(12)
                    qa_count += 1
                
                # Pattern 3: Texte commençant par "A:" (réponse seule)
                elif re.match(r'^A:\s+', text):
                    para.clear()
                    run = para.add_run(text)
                    para.style = 'Normal'
                    para.paragraph_format.left_indent = Inches(0.25)
                    para.paragraph_format.space_after = Pt(12)
                    qa_count += 1
                
                i += 1
            
            # Sauvegarder
            doc.save(str(self.docx_path))
            
            if qa_count > 0:
                logger.info(f"✓ Q&A formatting complete ({qa_count} Q/A pairs/formats found)")
            else:
                logger.info("No Q&A patterns found")
        
        except Exception as e:
            logger.warning(f"Q&A formatting failed: {e}")
            import traceback
            traceback.print_exc()
    
    def extract_and_rebuild_tables(self):
        """
        Extrait les tableaux en utilisant les COORDONNÉES X pour placer correctement
        les cellules dans les bonnes colonnes.
        """
        if not PDFPLUMBER_AVAILABLE or not PYTHON_DOCX_AVAILABLE:
            logger.warning("pdfplumber or python-docx not available")
            return
        
        try:
            import pdfplumber
            from docx import Document
            from docx.shared import Pt
            
            doc = Document(str(self.docx_path))
            
            with pdfplumber.open(str(self.pdf_path)) as pdf:
                
                for page_num, page in enumerate(pdf.pages):
                    logger.info(f"📄 Processing page {page_num + 1}...")
                    
                    # 🔧 MÉTHODE ALTERNATIVE : Extraire avec coordonnées
                    # On extrait le texte avec positions
                    words = page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False
                    )
                    
                    # Détecter les tableaux avec leurs bounding boxes
                    tables = page.find_tables()
                    
                    if not tables:
                        continue
                    
                    logger.info(f"Found {len(tables)} table(s) on page {page_num + 1}")
                    
                    # Pour chaque tableau
                    for table_idx, table in enumerate(tables):
                        bbox = table.bbox  # Bounding box du tableau
                        
                        # Extraire le tableau avec settings précis
                        extracted = table.extract()
                        
                        if not extracted or len(extracted) < 1:
                            continue
                        
                        # 🔧 ÉTAPE 1 : Analyser les coordonnées X de TOUTES les colonnes
                        # On prend la première ligne (header) comme référence
                        header_row = extracted[0]
                        
                        # Trouver les positions X de chaque cellule du header
                        # Pour ça, on utilise les mots extraits
                        words_in_bbox = [w for w in words if 
                                        w['x0'] >= bbox[0] and w['x1'] <= bbox[2] and
                                        w['top'] >= bbox[1] and w['bottom'] <= bbox[3]]
                        
                        # Détecter les positions X des colonnes (centres)
                        # On utilise la première ligne pour ça
                        first_row_y = min(w['top'] for w in words_in_bbox) if words_in_bbox else 0
                        first_row_words = [w for w in words_in_bbox if abs(w['top'] - first_row_y) < 5]
                        
                        # Trier par X
                        first_row_words.sort(key=lambda w: w['x0'])
                        
                        # Créer les "bins" de colonnes basés sur les positions X
                        column_x_positions = []
                        for word in first_row_words:
                            x_center = (word['x0'] + word['x1']) / 2
                            column_x_positions.append(x_center)
                        
                        # Supprimer les doublons proches (tolérance 10 pixels)
                        unique_x_positions = []
                        for x in column_x_positions:
                            if not unique_x_positions or abs(x - unique_x_positions[-1]) > 10:
                                unique_x_positions.append(x)
                        
                        num_cols = len(unique_x_positions)
                        logger.info(f"📊 Table {table_idx + 1}: Detected {num_cols} column positions")
                        
                        # 🔧 ÉTAPE 2 : Mapper chaque ligne aux bonnes colonnes
                        normalized_table = []
                        
                        for row_idx, row in enumerate(extracted):
                            if not row:
                                continue
                            
                            # Créer une ligne vide
                            normalized_row = [""] * num_cols
                            
                            # Pour chaque cellule de la ligne brute
                            for cell in row:
                                if not cell or cell.strip() == "":
                                    continue
                                
                                # Trouver la position X de cette cellule dans le PDF
                                # Chercher le mot correspondant dans words_in_bbox
                                cell_text = cell.strip()
                                
                                # Trouver le mot qui correspond
                                matching_words = [w for w in words_in_bbox if cell_text in w['text']]
                                
                                if matching_words:
                                    # Prendre le premier match
                                    word = matching_words[0]
                                    x_center = (word['x0'] + word['x1']) / 2
                                    
                                    # Trouver la colonne la plus proche
                                    closest_col = min(range(num_cols), 
                                                    key=lambda i: abs(unique_x_positions[i] - x_center))
                                    
                                    # Placer la cellule dans la bonne colonne
                                    normalized_row[closest_col] = cell_text
                            
                            normalized_table.append(normalized_row)
                        
                        # 🔧 ÉTAPE 3 : Créer le tableau Word
                        if table_idx == 0:
                            doc.add_paragraph(f"Page {page_num + 1}", style='Heading 2')
                        
                        num_rows = len(normalized_table)
                        
                        word_table = doc.add_table(rows=num_rows, cols=num_cols)
                        word_table.style = 'Light Grid Accent 1'
                        
                        # Remplir
                        for i, row_data in enumerate(normalized_table):
                            row = word_table.rows[i]
                            
                            for j, cell_value in enumerate(row_data):
                                cell = row.cells[j]
                                cell.text = cell_value if cell_value else ""
                                
                                # Formater nombres
                                if cell_value and cell_value.strip():
                                    clean = cell_value.replace(',', '').replace('$', '').replace('%', '')
                                    clean = clean.replace('(', '').replace(')', '').strip()
                                    
                                    if clean.startswith('-'):
                                        clean = clean[1:]
                                    
                                    if clean.replace('.', '', 1).isdigit():
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = 2  # Right
                                            for run in paragraph.runs:
                                                run.font.size = Pt(9)
                        
                        doc.add_paragraph()
                        logger.info(f"  ✅ {num_rows} rows × {num_cols} cols (position-aligned)")
            
            # Sauvegarder
            doc.save(str(self.docx_path))
            logger.info("✅ All tables aligned by X coordinates")
        
        except Exception as e:
            logger.error(f"Table extraction failed: {e}")
            import traceback
            traceback.print_exc()
    
    def convert_with_structure_preservation(self):
        """
        Convertit le PDF en DOCX en préservant au maximum la structure.
        """
        self.convert()
    
    def add_hyperlink(self, paragraph, url, text):
        """
        Ajoute un hyperlien à un paragraphe.
        
        Args:
            paragraph: Paragraphe python-docx
            url: URL du lien
            text: Texte à afficher
        """
        if not PYTHON_DOCX_AVAILABLE:
            return
        
        try:
            # Créer l'élément hyperlink
            part = paragraph.part
            r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            
            # Créer le run avec le lien
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Style bleu et souligné
            c = OxmlElement('w:color')
            c.set(qn('w:val'), '0563C1')
            rPr.append(c)
            
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            new_run.append(rPr)
            new_t = OxmlElement('w:t')
            new_t.text = text
            new_run.append(new_t)
            
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception as e:
            # Fallback: ajouter le texte normal
            paragraph.add_run(text)
    
    def restore_hyperlinks(self):
        """
        Détecte les URLs et les convertit en hyperliens cliquables.
        """
        if not PYTHON_DOCX_AVAILABLE:
            return
        
        try:
            doc = Document(str(self.docx_path))
            url_pattern = re.compile(r'(https?://[^\s\)]+)')
            url_count = 0
            
            for para in doc.paragraphs:
                text = para.text
                urls = url_pattern.findall(text)
                
                if urls:
                    # Reconstruire le paragraphe avec hyperliens
                    para_text = para.text
                    para.clear()
                    
                    # Séparer le texte en parties (avant URL, URL, après URL)
                    parts = url_pattern.split(para_text)
                    
                    for i, part in enumerate(parts):
                        if url_pattern.match(part):
                            # C'est une URL
                            self.add_hyperlink(para, part, part)
                            url_count += 1
                        else:
                            # Texte normal
                            if part:
                                para.add_run(part)
            
            if url_count > 0:
                doc.save(str(self.docx_path))
                logger.info(f"✓ Restored {url_count} hyperlink(s)")
        except Exception as e:
            logger.warning(f"Hyperlink restoration failed: {e}")
    
    def rebuild_table_of_contents(self):
        """
        Reconstruit la Table of Contents en détectant les patterns de titres.
        """
        if not PYTHON_DOCX_AVAILABLE:
            return
        
        try:
            doc = Document(str(self.docx_path))
            
            # Pattern pour détecter les entrées TOC: "Titre ... 123"
            # Pattern amélioré : Titre avec points de suspension + numéro de page
            toc_pattern = re.compile(r'^([A-Z][A-Za-z\s&,\-\']{3,60}?)\.{2,}\s+(\d{1,3})$')
            
            toc_entries = []
            toc_start_idx = None
            
            # Chercher les paragraphes qui correspondent au pattern TOC
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                match = toc_pattern.match(text)
                
                if match:
                    if toc_start_idx is None:
                        toc_start_idx = i
                    title = match.group(1).strip()
                    page_num = match.group(2)
                    toc_entries.append((title, page_num))
            
            if toc_entries and toc_start_idx is not None:
                logger.info(f"Found {len(toc_entries)} TOC entries")
                
                # Formater les entrées TOC avec indentation
                for idx, (title, page_num) in enumerate(toc_entries):
                    if toc_start_idx + idx < len(doc.paragraphs):
                        para = doc.paragraphs[toc_start_idx + idx]
                        para.clear()
                        
                        # Ajouter le titre
                        run = para.add_run(title)
                        run.font.size = Pt(11)
                        
                        # Ajouter les points de suite
                        dots_run = para.add_run(' ' + '.' * (50 - len(title) - len(page_num)) + ' ')
                        dots_run.font.size = Pt(9)
                        
                        # Ajouter le numéro de page
                        page_run = para.add_run(page_num)
                        page_run.font.size = Pt(11)
                        page_run.bold = True
                
                doc.save(str(self.docx_path))
                logger.info(f"✓ Rebuilt Table of Contents with {len(toc_entries)} entries")
        except Exception as e:
            logger.warning(f"TOC rebuilding failed: {e}")
    
    def format_financial_numbers(self):
        """
        Formate les nombres financiers avec style comptable (rouge si négatif).
        """
        if not PYTHON_DOCX_AVAILABLE:
            return
        
        try:
            doc = Document(str(self.docx_path))
            formatted_count = 0
            
            # Parcourir tous les tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if not text:
                            continue
                        
                        # Détecter si négatif (parenthèses)
                        is_negative = text.startswith('(') and text.endswith(')')
                        
                        # Extraire le nombre
                        clean = text.replace('(', '').replace(')', '').replace(',', '').replace('$', '').replace('%', '').strip()
                        
                        # Vérifier si c'est un nombre
                        if clean.replace('.', '').replace('-', '').isdigit():
                            # Formater
                            for para in cell.paragraphs:
                                para.alignment = 2  # Right align
                                for run in para.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                                    if is_negative:
                                        try:
                                            run.font.color.rgb = RGBColor(192, 0, 0)  # Rouge foncé
                                        except:
                                            pass
                            formatted_count += 1
            
            if formatted_count > 0:
                doc.save(str(self.docx_path))
                logger.info(f"✓ Formatted {formatted_count} financial number(s)")
        except Exception as e:
            logger.warning(f"Financial number formatting failed: {e}")
    
    def assess_conversion_quality(self):
        """
        Évalue la qualité de la conversion et retourne des métriques.
        
        Returns:
            dict: Métriques de qualité (quality_score, broken_tables, etc.)
        """
        if not PYTHON_DOCX_AVAILABLE:
            return {'quality_score': 0, 'error': 'python-docx not available'}
        
        try:
            doc = Document(str(self.docx_path))
            
            metrics = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'broken_tables': 0,
                'empty_paragraphs': 0,
                'quality_score': 0
            }
            
            # Compter les tableaux cassés (colonnes incohérentes)
            for table in doc.tables:
                if not table.rows:
                    continue
                row_counts = [len(row.cells) for row in table.rows]
                if len(set(row_counts)) > 1:
                    metrics['broken_tables'] += 1
            
            # Compter les paragraphes vides
            for para in doc.paragraphs:
                if not para.text.strip():
                    metrics['empty_paragraphs'] += 1
            
            # Calculer le score de qualité (0 à 100)
            # Score : 100 - (5 points par table cassée + 0.5 point par 10 paragraphes vides)
            total_issues = metrics['broken_tables'] * 5 + (metrics['empty_paragraphs'] / 10) * 0.5
            metrics['quality_score'] = max(0, min(100, 100 - total_issues))
            
            logger.info(f"Quality metrics: {metrics['broken_tables']} broken tables, "
                       f"{metrics['empty_paragraphs']} empty paragraphs")
            
            return metrics
            
        except Exception as e:
            logger.warning(f"Quality assessment failed: {e}")
            return {'quality_score': 0, 'error': str(e)}
    
    def _detect_images_in_pdf(self):
        """
        Détecte les images dans le PDF.
        Retourne un dictionnaire {numéro_page: nombre_images}
        """
        images_by_page = {}
        
        if not PYPDF2_AVAILABLE:
            logger.warning("PyPDF2 is not available. Install with: pip install PyPDF2")
            return images_by_page
        
        try:
            with open(self.pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    
                    # Compter les images (méthode basique)
                    try:
                        if '/XObject' in page['/Resources']:
                            xobjects = page['/Resources']['/XObject'].get_object()
                            image_count = sum(1 for obj in xobjects if xobjects[obj]['/Subtype'] == '/Image')
                            if image_count > 0:
                                images_by_page[page_num + 1] = image_count
                    except (KeyError, AttributeError):
                        # Pas d'images sur cette page ou structure différente
                        pass
        except Exception as e:
            logger.warning(f"Error detecting images: {e}")
        
        return images_by_page
    
    def _detect_sensitive_information(self, verbose=False):
        """
        Nouvelle stratégie de détection: approche pragmatique et robuste.

        Args:
            verbose: Si True, affiche les logs détaillés. Si False, utilise tqdm pour progress bar.
        """
        sensitive_info_by_page = {}

        # ========== COMPREHENSIVE FALSE POSITIVE LISTS ==========
        # Document structure terms (multi-word phrases that aren't real entities)
        DOCUMENT_STRUCTURE_TERMS = {
            'balance sheet', 'income statement', 'cash flow', 'table of contents',
            'statement of operations', 'notes to financial', 'management discussion',
            'risk factors', 'executive summary', 'business overview', 'legal proceedings',
            'market for', 'selected financial', 'quantitative and qualitative',
            'controls and procedures', 'financial statements', 'exhibits index',
            'consolidated balance', 'consolidated statements', 'stockholders equity',
            'comprehensive income', 'changes in', 'years ended', 'months ended',
            'unaudited consolidated', 'audited consolidated', 'condensed consolidated',
            'management discussion and analysis', 'form 10-k', 'form 10-q', 'form 8-k',
            'part i', 'part ii', 'part iii', 'part iv', 'item 1', 'item 2', 'item 3'
        }

        # Generic business terms that spaCy often flags as organizations
        GENERIC_BUSINESS_TERMS = {
            'operations', 'management', 'board', 'committee', 'team', 'division',
            'department', 'group', 'unit', 'panel', 'council', 'assembly',
            'business combination', 'financial condition', 'operating results',
            'the business', 'the company', 'the registrant', 'the issuer',
            'securities', 'commission', 'exchange', 'trading', 'market',
            'segment', 'services', 'products', 'business', 'enterprise',
            'consolidated', 'combined', 'merged', 'acquired', 'subsidiary',
            'parent company', 'holding company', 'affiliated', 'related party',
            'operating segment', 'reportable segment', 'business segment',
            'common stock', 'preferred stock', 'class a', 'class b',
            'shareholders', 'stockholders', 'beneficial owners', 'insiders'
        }

        # Invalid first words for company names (articles, prepositions, conjunctions)
        INVALID_COMPANY_FIRST_WORDS = {
            'the', 'a', 'an', 'and', 'or', 'but', 'if', 'then', 'such', 'all',
            'any', 'each', 'every', 'some', 'no', 'this', 'that', 'these', 'those',
            'operations', 'management', 'financial', 'business', 'combined',
            'our', 'we', 'us', 'their', 'its', 'his', 'her', 'your',
            'securities', 'commission', 'exchange', 'consolidated', 'total'
        }

        # Common person name false positives (document terms that look like names)
        PERSON_NAME_FALSE_POSITIVES = {
            'balance sheet', 'income statement', 'cash flow statement',
            'table of contents', 'united states', 'new york', 'los angeles',
            'washington dc', 'the company', 'the registrant', 'the issuer',
            'securities exchange', 'exchange commission', 'internal revenue',
            'generally accepted', 'accounting principles', 'fair value',
            'stock option', 'stock compensation', 'stock award',
            'restricted stock', 'performance share', 'equity award',
            'executive compensation', 'board compensation', 'audit committee',
            'compensation committee', 'nominating committee', 'risk committee',
            'operations committee', 'investment committee', 'credit committee'
        }

        # Invalid words that should NEVER appear in a person name
        INVALID_NAME_WORDS = {
            'and', 'or', 'the', 'a', 'an', 'of', 'in', 'on', 'at', 'to', 'for',
            'with', 'by', 'from', 'as', 'is', 'are', 'was', 'were', 'be', 'been',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'shall',
            'should', 'may', 'might', 'must', 'can', 'could',
            'require', 'required', 'requirements', 'pursuant', 'accordance',
            'including', 'include', 'includes', 'such', 'other', 'any', 'all',
            'each', 'every', 'both', 'either', 'neither', 'none', 'not',
            'only', 'same', 'than', 'then', 'there', 'these', 'those', 'this',
            'that', 'which', 'who', 'whom', 'whose', 'what', 'when', 'where',
            'whether', 'applicable', 'provided', 'however', 'therefore',
            'furthermore', 'moreover', 'otherwise', 'unless', 'until', 'upon',
            'within', 'without', 'respect', 'regard', 'pursuant', 'subject'
        }

        if not PYPDF2_AVAILABLE:
            logger.warning("PyPDF2 is not available. Install with: pip install PyPDF2")
            return sensitive_info_by_page
        
        try:
            with open(self.pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                # Progress indicator: tqdm en mode simple, print en mode verbose
                if verbose:
                    # Mode verbose : afficher le message
                    if total_pages > 10:
                        print(f"   📄 Analyzing {total_pages} pages...", end='', flush=True)
                    page_iterator = range(total_pages)
                else:
                    # Mode simple : utiliser tqdm si disponible
                    if TQDM_AVAILABLE and total_pages > 5:
                        page_iterator = tqdm(range(total_pages), desc="   Analyzing pages", unit="page", 
                                            ncols=80, leave=False, disable=False)
                    else:
                        page_iterator = range(total_pages)
                
                for page_num in page_iterator:
                    page = pdf_reader.pages[page_num]
                    page_findings = []
                    
                    try:
                        text = page.extract_text()
                        if not text or len(text.strip()) < 10:
                            continue
                        
                        # Normalize text
                        text = re.sub(r'\s+', ' ', text)
                        text = text.strip()
                        
                        # ===== EMAILS - IMPROVED WITH VALIDATION =====
                        email_pattern = re.compile(r'\b[a-zA-Z0-9][a-zA-Z0-9._%+-]*@[a-zA-Z0-9][a-zA-Z0-9.-]*\.[a-zA-Z]{2,}\b')

                        # Valid TLDs (top-level domains) - most common ones
                        valid_tlds = {
                            'com', 'org', 'net', 'edu', 'gov', 'mil', 'int',
                            'co', 'io', 'ai', 'biz', 'info', 'name', 'pro',
                            'uk', 'us', 'ca', 'au', 'de', 'fr', 'jp', 'cn', 'in', 'br',
                            'ru', 'nl', 'it', 'es', 'pl', 'se', 'ch', 'be', 'at'
                        }

                        # Fake/test email patterns to reject
                        fake_email_patterns = {
                            'test@', 'example@', 'sample@', 'demo@', 'fake@',
                            'noreply@', 'no-reply@', 'donotreply@',
                            '@test.', '@example.', '@sample.', '@demo.', '@fake.',
                            '@localhost', '@127.0.0.1', '@0.0.0.0',
                            'admin@admin', 'user@user', 'email@email'
                        }

                        for match in email_pattern.finditer(text):
                            email = match.group().strip().lower()

                            # Validation 1: Check TLD
                            tld = email.split('.')[-1] if '.' in email else ''
                            if tld not in valid_tlds:
                                continue  # Skip invalid TLD

                            # Validation 2: Reject fake/test emails
                            if any(fake in email for fake in fake_email_patterns):
                                continue

                            # Validation 3: Basic structure validation
                            if email.count('@') != 1:
                                continue  # Must have exactly one @

                            local, domain = email.split('@')
                            if len(local) < 1 or len(domain) < 3:
                                continue  # Too short

                            if domain.count('.') < 1:
                                continue  # Domain must have at least one dot

                            # Validation 4: Reject suspicious patterns
                            if '..' in email or email.startswith('.') or email.endswith('.'):
                                continue  # Invalid dot placement

                            if '--' in email or '__' in email:
                                continue  # Suspicious patterns

                            # Validation 5: Check length constraints
                            if len(email) > 254 or len(local) > 64:
                                continue  # Exceeds RFC 5321 limits

                            # Valid email - add to findings
                            page_findings.append({'type': 'email', 'value': email, 'page': page_num + 1})
                        
                        # ===== PHONES - VERSION AMÉLIORÉE (US + International) =====
                        phones_found = set()
                        
                        # Pattern 1: Numéros US (format standard)
                        phone_us_pattern = re.compile(r'\b\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b')
                        for match in phone_us_pattern.finditer(text):
                            phone = match.group().strip()
                            cleaned = re.sub(r'[^\d]', '', phone)
                            if len(cleaned) == 10:
                                if PHONENUMBERS_AVAILABLE:
                                    try:
                                        parsed = phonenumbers.parse(phone, "US")
                                        if phonenumbers.is_valid_number(parsed):
                                            formatted = phonenumbers.format_number(parsed, PhoneNumberFormat.NATIONAL)
                                            phones_found.add(formatted)
                                    except:
                                        phones_found.add(phone)
                                else:
                                    phones_found.add(phone)
                        
                        # Pattern 2: Numéros internationaux avec indicatif pays (+XX ou 00XX)
                        # Format: +44 (0)20 7638 0129 ou +44 20 7638 0129 ou +1 555 123 4567
                        phone_international_patterns = [
                            # Format: +XX (0)XX XXXX XXXX (UK, etc.)
                            re.compile(r'\+\d{1,3}\s*\(?0\)?\s*\d{1,4}\s*\d{3,4}\s*\d{3,4}\b'),
                            # Format: +XX XX XXXX XXXX (standard international)
                            re.compile(r'\+\d{1,3}\s+\d{1,4}\s+\d{3,4}\s+\d{3,4}\b'),
                            # Format: +XX XXXXXXXXXX (sans espaces)
                            re.compile(r'\+\d{1,3}\s*\d{6,14}\b'),
                            # Format: 00XX XX XXXX XXXX (format européen avec 00)
                            re.compile(r'00\d{1,3}\s+\d{1,4}\s+\d{3,4}\s+\d{3,4}\b'),
                            # Format: (0)XX XXXX XXXX (UK local avec 0 optionnel)
                            re.compile(r'\(?0\)?\s*\d{2,4}\s+\d{3,4}\s+\d{3,4}\b'),
                        ]
                        
                        for pattern in phone_international_patterns:
                            for match in pattern.finditer(text):
                                phone = match.group().strip()
                                phone = re.sub(r'\s+', ' ', phone)  # Normaliser les espaces
                                
                                # Validation basique: doit avoir entre 7 et 15 chiffres (sans le +)
                                digits_only = re.sub(r'[^\d]', '', phone)
                                if 7 <= len(digits_only) <= 15:
                                    # Vérifier le contexte pour éviter les faux positifs
                                    context_start = max(0, match.start() - 50)
                                    context_end = min(len(text), match.end() + 50)
                                    context = text[context_start:context_end].lower()
                                    
                                    # Indicateurs positifs (suggèrent que c'est un numéro de téléphone)
                                    positive_indicators = [
                                        'phone', 'tel', 'telephone', 'call', 'contact', 'mobile',
                                        'fax', 'telefax', 't:', 'p:', 'f:', 'm:'
                                    ]
                                    
                                    # Indicateurs négatifs (exclure)
                                    negative_indicators = [
                                        'file number', 'commission file', 'cik', 'ein',
                                        'employer identification', 'tax id', 'ssn'
                                    ]
                                    
                                    has_positive = any(ind in context for ind in positive_indicators)
                                    has_negative = any(ind in context for ind in negative_indicators)
                                    
                                    # Si phonenumbers est disponible, valider le numéro
                                    if PHONENUMBERS_AVAILABLE:
                                        try:
                                            # Essayer de parser le numéro (peut être n'importe quel pays)
                                            parsed = phonenumbers.parse(phone, None)  # None = auto-detect country
                                            if phonenumbers.is_valid_number(parsed):
                                                # Formater en format international
                                                formatted = phonenumbers.format_number(parsed, PhoneNumberFormat.INTERNATIONAL)
                                                phones_found.add(formatted)
                                            elif has_positive and not has_negative:
                                                # Si le contexte est positif, accepter même si la validation échoue
                                                phones_found.add(phone)
                                        except:
                                            # Si le parsing échoue mais que le contexte est positif, accepter
                                            if has_positive and not has_negative:
                                                phones_found.add(phone)
                                    else:
                                        # Sans phonenumbers, accepter si contexte positif
                                        if has_positive and not has_negative:
                                            phones_found.add(phone)
                        
                        # Pattern 3: Numéros avec extensions (ex: +1 555 123 4567 ext. 123)
                        phone_with_ext_pattern = re.compile(
                            r'(\+?\d{1,3}[\s\-\(\)]?\d{1,4}[\s\-]?\d{3,4}[\s\-]?\d{3,4})\s*(?:ext|extension|ext\.|x\.?|#)\s*(\d{1,6})',
                            re.IGNORECASE
                        )
                        
                        for match in phone_with_ext_pattern.finditer(text):
                            phone_base = match.group(1).strip()
                            extension = match.group(2).strip()
                            phone_with_ext = f"{phone_base} ext. {extension}"
                            
                            # Valider le numéro de base
                            digits_only = re.sub(r'[^\d]', '', phone_base)
                            if 7 <= len(digits_only) <= 15:
                                if PHONENUMBERS_AVAILABLE:
                                    try:
                                        parsed = phonenumbers.parse(phone_base, None)
                                        if phonenumbers.is_valid_number(parsed):
                                            formatted_base = phonenumbers.format_number(parsed, PhoneNumberFormat.INTERNATIONAL)
                                            phones_found.add(f"{formatted_base} ext. {extension}")
                                    except:
                                        phones_found.add(phone_with_ext)
                                else:
                                    phones_found.add(phone_with_ext)
                        
                        # Ajouter tous les numéros trouvés
                        for phone in phones_found:
                                    page_findings.append({'type': 'phone', 'value': phone, 'page': page_num + 1})
                        
                        # ===== ADDRESSES COMPLÈTES - VERSION ROBUSTE POUR PDF =====
                        
                        addresses_found = set()
                        
                        valid_states = [
                            'AL', 'AK', 'AZ', 'AR', 'CA', 'CO', 'CT', 'DE', 'FL', 'GA',
                            'HI', 'ID', 'IL', 'IN', 'IA', 'KS', 'KY', 'LA', 'ME', 'MD',
                            'MA', 'MI', 'MN', 'MS', 'MO', 'MT', 'NE', 'NV', 'NH', 'NJ',
                            'NM', 'NY', 'NC', 'ND', 'OH', 'OK', 'OR', 'PA', 'RI', 'SC',
                            'SD', 'TN', 'TX', 'UT', 'VT', 'VA', 'WA', 'WV', 'WI', 'WY', 'DC'
                        ]
                        
                        # ÉTAPE 1: Normaliser AGRESSIVEMENT le texte
                        # Remplacer TOUS les types d'espaces (nbsp, tabs, newlines) par un seul espace
                        text_clean = text.replace('\n', ' ')
                        text_clean = text_clean.replace('\r', ' ')
                        text_clean = text_clean.replace('\t', ' ')
                        text_clean = re.sub(r'\s+', ' ', text_clean)  # Multiple espaces -> 1 espace
                        text_clean = text_clean.strip()
                        
                        # ÉTAPE 2: Pattern d'adresse complète DIRECT (plus simple et robuste)
                        # Format: Numéro + Nom + Type + Ville + État + ZIP
                        address_pattern = re.compile(
                            r'\b(\d{1,5})\s+'  # Numéro de rue
                            r'([A-Za-z][A-Za-z\s]{2,40}?)\s+'  # Nom de rue (simplifié)
                            r'(Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Place|Pl|Way|Circle|Cir|Parkway|Pkwy|Highway|Hwy)'  # Type
                            r'\s+'
                            r'([A-Za-z][A-Za-z\s]{2,40}?)'  # Ville/Township
                            r'[,\s]+'
                            r'([A-Z]{2})\s+'  # État
                            r'(\d{5}(?:-\d{4})?)\b',  # ZIP
                            re.IGNORECASE
                        )
                        
                        for match in address_pattern.finditer(text_clean):
                            street_num = match.group(1)
                            street_name = match.group(2).strip()
                            street_type = match.group(3)
                            city = match.group(4).strip()
                            state = match.group(5)
                            zip_code = match.group(6)
                            
                            # Validation de l'état
                            if state not in valid_states:
                                continue
                            
                            # Nettoyer le nom de rue
                            street_name = re.sub(r'\s+', ' ', street_name).strip()
                            
                            # Nettoyer la ville - garder seulement les mots capitalisés
                            city_words = city.split()
                            city_clean_parts = []
                            for word in city_words:
                                if word and len(word) > 1 and word[0].isupper():
                                    city_clean_parts.append(word)
                            
                            if not city_clean_parts:
                                continue
                            
                            city_clean = ' '.join(city_clean_parts)
                            
                            # Construire l'adresse
                            address = f"{street_num} {street_name} {street_type} {city_clean} {state} {zip_code}"
                            
                            # Validation de longueur
                            if not (25 <= len(address) <= 150):
                                continue
                            
                            # Filtrer les faux positifs
                            address_upper = address.upper()
                            bad_keywords = [
                                'COMMISSION FILE', 'SECURITIES', 'EXCHANGE ACT',
                                'TABLE OF CONTENTS', 'WALL STREET', 'BALANCE SHEET',
                                'WASHINGTON'  # Adresse SEC
                            ]
                            
                            if any(bad in address_upper for bad in bad_keywords):
                                continue
                            
                            addresses_found.add(address)
                        
                        # Si le pattern direct ne trouve rien, essayer la méthode inverse (ZIP d'abord)
                        if not addresses_found:
                            zip_pattern = re.compile(r'\b(\d{5}(?:-\d{4})?)\b')
                            
                            for zip_match in zip_pattern.finditer(text_clean):
                                zip_code = zip_match.group(1)
                                zip_pos = zip_match.start()
                                
                                # Prendre 250 caractères AVANT le ZIP
                                context_before = text_clean[max(0, zip_pos - 250):zip_pos]
                                
                                # Chercher : "Numéro Nom Type Ville, État"
                                # Pattern inverse plus flexible
                                inverse_pattern = re.compile(
                                    r'(\d{1,5})\s+'
                                    r'([A-Za-z\s]{3,50}?)\s+'
                                    r'(Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Place|Pl|Way|Circle|Cir|Parkway|Pkwy|Highway|Hwy)'
                                    r'\s+([A-Za-z\s]{3,40}?)[,\s]+'
                                    r'([A-Z]{2})\s*$',
                                    re.IGNORECASE
                                )
                                
                                addr_match = inverse_pattern.search(context_before)
                                
                                if addr_match and addr_match.group(5) in valid_states:
                                    street_num = addr_match.group(1)
                                    street_name = addr_match.group(2).strip()
                                    street_type = addr_match.group(3)
                                    city = addr_match.group(4).strip()
                                    state = addr_match.group(5)
                                    
                                    # Nettoyer
                                    street_name = re.sub(r'\s+', ' ', street_name).strip()
                                    city_words = city.split()
                                    city_clean = ' '.join([w for w in city_words if w and w[0].isupper()])
                                    
                                    if city_clean:
                                        address = f"{street_num} {street_name} {street_type} {city_clean} {state} {zip_code}"
                                        
                                        if 25 <= len(address) <= 150:
                                            if 'WASHINGTON' not in address.upper():
                                                addresses_found.add(address)
                        
                        for address in addresses_found:
                            page_findings.append({'type': 'address', 'value': address, 'page': page_num + 1})
                        
                        # ===== SSN =====
                        ssn_pattern = re.compile(r'\b\d{3}-\d{2}-\d{4}\b')
                        for match in ssn_pattern.finditer(text):
                            ssn = match.group().strip()
                            page_findings.append({'type': 'ssn', 'value': ssn, 'page': page_num + 1})
                        
                        # ===== CREDIT CARDS =====
                        # DÉSACTIVÉ - Ne plus détecter les cartes de crédit
                        # credit_card_pattern = re.compile(r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b')
                        # for match in credit_card_pattern.finditer(text):
                        #     cc = match.group().strip()
                        #     page_findings.append({'type': 'credit_card', 'value': cc, 'page': page_num + 1})
                        
                        # ===== COMPANY IDENTIFICATION =====
                        # 🔧 Chercher d'abord les Commission File Numbers de manière explicite
                        commission_numbers_found = set()
                        cfn_context_pattern = re.compile(r'Commission\s+File\s+Number[:\s]+(\d{3}-\d{5})', re.IGNORECASE)
                        for match in cfn_context_pattern.finditer(text):
                            commission_numbers_found.add(match.group(1))
                            page_findings.append({'type': 'commission_file_number', 'value': match.group(1), 'page': page_num + 1})
                        
                        # 🔧 Puis chercher les EINs (en excluant les CFN déjà trouvés)
                        ein_pattern = re.compile(r'\b(\d{2}-\d{7,8})\b')  # 7 OU 8 chiffres
                        for match in ein_pattern.finditer(text):
                            ein = match.group(1)
                            # Vérifier le contexte pour confirmer que c'est un EIN
                            context = text[max(0, match.start() - 60):match.end() + 60]
                            is_ein_context = bool(re.search(r'(?:I\.?R\.?S\.?|Employer\s+Identification|EIN|Tax\s+ID)', context, re.IGNORECASE))
                            
                            if is_ein_context:
                                page_findings.append({'type': 'irs_ein', 'value': ein, 'page': page_num + 1})
                        
                        # ===== EXECUTIVE NAMES - VERSION AMÉLIORÉE =====
                        # DÉSACTIVÉ - Ne plus détecter les noms d'exécutifs
                        executives_found = set()
                        
                        # Liste étendue de titres exécutifs
                        executive_titles = [
                            r'Chief\s+(?:Executive|Financial|Operating|Technology|Information|Marketing|Revenue|Accounting|Legal|Compliance|Risk|Investment)\s+Officer',
                            r'CEO|CFO|CTO|COO|CIO|CMO|CRO|CAO|CLO|CCO|CRO',
                            r'President',
                            r'Vice\s+President',
                            r'VP\s+(?:of|Finance|Operations|Sales|Marketing|Engineering)',
                            r'Executive\s+Vice\s+President',
                            r'EVP',
                            r'Senior\s+Vice\s+President',
                            r'SVP',
                            r'Chairman',
                            r'Chairman\s+of\s+the\s+Board',
                            r'Vice\s+Chairman',
                            r'Secretary',
                            r'Treasurer',
                            r'Controller',
                            r'Chief\s+Controller',
                            r'Managing\s+Director',
                            r'General\s+Manager',
                            r'Director',
                            r'Executive\s+Director',
                            r'Founder',
                            r'Co-Founder',
                            r'Chief\s+of\s+Staff',
                            r'Head\s+of\s+(?:Finance|Operations|Sales|Marketing|Engineering|Legal)',
                        ]
                        
                        # Pattern 1: Signatures "/s/" ou "/s" ou "s/"
                        signature_patterns = [
                            re.compile(r'/s[/\s]+\s*([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', re.IGNORECASE),
                            re.compile(r'Signature:\s*([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', re.IGNORECASE),
                            re.compile(r'Signed:\s*([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', re.IGNORECASE),
                        ]
                        
                        for pattern in signature_patterns:
                            for match in pattern.finditer(text):
                                name = match.group(1).strip()
                                name = re.sub(r'\s+', ' ', name)
                                if 5 <= len(name) <= 50:
                                    executives_found.add(name)
                        
                        # Pattern 2: Format "Nom, Titre" (amélioré)
                        exec_title_pattern = '|'.join(executive_titles)
                        exec_pattern = re.compile(
                            r'\b([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*[,:]\s*(' + exec_title_pattern + r')\b',
                            re.IGNORECASE
                        )
                        
                        for match in exec_pattern.finditer(text):
                            name = match.group(1).strip()
                            name = re.sub(r'\s+', ' ', name)
                            
                            if 5 <= len(name) <= 50:
                                name_upper = name.upper()
                                false_positives = [
                                    'UNITED STATES', 'NEW YORK', 'LOS ANGELES', 'THE COMPANY',
                                    'BALANCE SHEET', 'CASH FLOW', 'INCOME TAX', 'STOCK OPTION',
                                    'SECURITIES AND EXCHANGE', 'COMMISSION', 'WASHINGTON'
                                ]
                                if name_upper not in false_positives:
                                    executives_found.add(name)
                        
                        # Pattern 3: Format "Titre: Nom" ou "Titre Nom"
                        exec_pattern_reverse = re.compile(
                            r'\b(' + exec_title_pattern + r')\s*[:]\s*([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b',
                            re.IGNORECASE
                        )
                        
                        for match in exec_pattern_reverse.finditer(text):
                            name = match.group(2).strip()
                            name = re.sub(r'\s+', ' ', name)
                            
                            if 5 <= len(name) <= 50:
                                name_upper = name.upper()
                                false_positives = [
                                    'UNITED STATES', 'NEW YORK', 'LOS ANGELES', 'THE COMPANY',
                                    'BALANCE SHEET', 'CASH FLOW', 'INCOME TAX', 'STOCK OPTION'
                                ]
                                if name_upper not in false_positives:
                                    executives_found.add(name)
                        
                        # Pattern 4: Sections de signatures avec contexte étendu
                        name_context_patterns = [
                            re.compile(r'(?:By|Name|Signed\s+By|Signature\s+of|Name\s+of\s+Officer|Name\s+of\s+Signatory)[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', re.IGNORECASE),
                            re.compile(r'(?:Officer|Executive|Director)[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', re.IGNORECASE),
                        ]
                        
                        for pattern in name_context_patterns:
                            for match in pattern.finditer(text):
                                name = match.group(1).strip()
                                name = re.sub(r'\s+', ' ', name)

                                if 5 <= len(name) <= 50:
                                    # Vérifier le contexte pour éviter les faux positifs
                                    context_start = max(0, match.start() - 100)
                                    context_end = min(len(text), match.end() + 100)
                                    context = text[context_start:context_end].lower()

                                    # Éviter les contextes non pertinents
                                    bad_contexts = ['table of contents', 'index', 'appendix', 'reference']
                                    if not any(bad in context for bad in bad_contexts):
                                        executives_found.add(name)
                        
                        # Pattern 5: Noms dans les tableaux/listes (format "Nom | Titre" ou "Nom - Titre")
                        table_patterns = [
                            re.compile(r'([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*[|\-]\s*(' + exec_title_pattern + r')\b', re.IGNORECASE),
                            re.compile(r'([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s+(' + exec_title_pattern + r')\b', re.IGNORECASE),
                        ]
                        
                        for pattern in table_patterns:
                            for match in pattern.finditer(text):
                                name = match.group(1).strip()
                                name = re.sub(r'\s+', ' ', name)
                                
                                if 5 <= len(name) <= 50:
                                    name_upper = name.upper()
                                    false_positives = [
                                        'UNITED STATES', 'NEW YORK', 'LOS ANGELES', 'THE COMPANY',
                                        'BALANCE SHEET', 'CASH FLOW', 'INCOME TAX', 'STOCK OPTION'
                                    ]
                                    if name_upper not in false_positives:
                                        executives_found.add(name)
                        
                        # Pattern 6: Noms dans les sections "Officers" ou "Management"
                        officers_section_pattern = re.compile(
                            r'(?:Officers?|Management|Executive\s+Officers?|Key\s+Personnel)[\s\S]{0,500}?'
                            r'([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*[,:]\s*(' + exec_title_pattern + r')\b',
                            re.IGNORECASE | re.MULTILINE
                        )
                        
                        for match in officers_section_pattern.finditer(text):
                            name = match.group(1).strip()
                            name = re.sub(r'\s+', ' ', name)
                            
                            if 5 <= len(name) <= 50:
                                name_upper = name.upper()
                                false_positives = [
                                    'UNITED STATES', 'NEW YORK', 'LOS ANGELES', 'THE COMPANY',
                                    'BALANCE SHEET', 'CASH FLOW', 'INCOME TAX', 'STOCK OPTION'
                                ]
                                if name_upper not in false_positives:
                                    executives_found.add(name)
                        
                        # Nettoyer et valider les noms trouvés avec filtres stricts
                        cleaned_executives = set()
                        
                        # Mots à exclure des noms d'executives
                        excluded_exec_words = {
                            'and', 'or', 'the', 'a', 'an', 'of', 'in', 'on', 'at', 'to', 'for',
                            'with', 'by', 'from', 'as', 'is', 'are', 'was', 'were', 'be', 'been',
                            'vantiv', 'in', 'to', 'with', 'the', 'an', 'independent', 'and',
                            'secretary', 'be', 'the', 'executive', 'lead', 'check', 'mark',
                            'whether', 'purport', 'exempt', 'principal', 'traders', 'merchants',
                            'later', 'than', 'registrant', 'combined', 'filings', 'security',
                            'holdings', 'team', 'reporting', 'high', 'court', 'incremental',
                            'amendment', 'affirmative', 'vote', 'fact', 'offeree', 'company',
                            'undersigned', 'hereunto', 'title', 'officer', 'director', 'chairman',
                            'co', 'services', 'officer', 'and', 'nelson', 'f', 'greene', 'title',
                            'philip', 'jansen', 'and', 'stephanie', 'ferris'
                        }
                        
                        for exec_name in executives_found:
                            # Nettoyer les espaces multiples
                            exec_name = re.sub(r'\s+', ' ', exec_name).strip()
                            
                            # Validation: doit avoir au moins 2 mots (prénom + nom)
                            words = exec_name.split()
                            if len(words) >= 2 and len(words) <= 5:
                                # Vérifier que chaque mot fait sens (pas trop court, pas trop long)
                                if all(2 <= len(w) <= 20 for w in words):
                                    # FILTRE STRICT: Exclure si contient des mots interdits
                                    words_lower = [w.lower().rstrip('.,;:') for w in words]
                                    
                                    # Ne pas accepter si un mot est dans la liste des mots exclus
                                    if any(w in excluded_exec_words for w in words_lower):
                                        continue
                                    
                                    # Ne pas accepter si le nom contient des mots de liaison
                                    if any(w in ['and', 'or', 'the', 'of', 'in', 'to', 'with', 'an', 'a'] for w in words_lower):
                                        continue
                                    
                                    # Vérifier qu'il n'y a pas trop de majuscules (éviter les acronymes)
                                    if sum(1 for c in exec_name if c.isupper()) <= len(exec_name) * 0.4:
                                        # Éviter les faux positifs spécifiques
                                        exec_name_upper = exec_name.upper()
                                        false_positives_exec = [
                                            'VANTIV IN', 'VANTIV TO', 'VANTIV WITH THE',
                                            'AN INDEPENDENT', 'AND SECRETARY', 'AND AMONG VANTIV',
                                            'ANY OFFEROR AND', 'ANY PERSONS ACTING', 'BE THE EXECUTIVE',
                                            'BE THE LEAD', 'CHECK MARK IF', 'CHECK MARK WHETHER',
                                            'DO NOT PURPORT', 'EXEMPT PRINCIPAL TRADERS', 'MEANS OF',
                                            'MERCHANTS OR', 'NO LATER THAN', 'OF REGISTRANT AS',
                                            'OF THE COMBINED', 'ON THE COMBINED', 'OTHER VANTIV FILINGS',
                                            'SECURITY HOLDINGS OR', 'TEAM REPORTING TO', 'THE HIGH COURT',
                                            'THE INCREMENTAL AMENDMENT', 'THE AFFIRMATIVE VOTE',
                                            'THE FACT THAT', 'THE OFFEREE COMPANY', 'THE UNDERSIGNED HEREUNTO',
                                            'CHAIRMAN AND CO', 'NELSON F. GREENE TITLE', 'OFFICER PHILIP JANSEN',
                                            'OFFICER OF', 'PHILIP JANSEN AND', 'SERVICES OFFICER AND'
                                        ]
                                        if not any(fp in exec_name_upper for fp in false_positives_exec):
                                            # Vérifier que le nom ne se termine pas par des mots suspects
                                            last_word = words_lower[-1]
                                            if last_word not in ['and', 'or', 'the', 'of', 'in', 'to', 'with', 'co', 'title', 'officer', 'director', 'chairman']:
                                                # Validation finale: le nom doit ressembler à un vrai nom
                                                if len(set(words_lower)) == len(words_lower):  # Pas de mots dupliqués
                                                    cleaned_executives.add(exec_name)
                        
                        # Validation finale avec spaCy NER pour executives (OPTIONNELLE - pas obligatoire)
                        # Ne valide QUE si spaCy trouve des noms - sinon garde tous les noms détectés
                        if SPACY_AVAILABLE and SPACY_NLP and cleaned_executives:
                            validated_executives = set()

                            for exec_name in cleaned_executives:
                                # Chercher le nom dans le texte original avec contexte
                                name_pos = text.find(exec_name)
                                if name_pos != -1:
                                    context_start = max(0, name_pos - 150)
                                    context_end = min(len(text), name_pos + len(exec_name) + 150)
                                    context_text = text[context_start:context_end]

                                    # Analyser avec spaCy
                                    try:
                                        doc = SPACY_NLP(context_text)

                                        # Vérifier si spaCy détecte ce nom comme une personne
                                        for ent in doc.ents:
                                            if ent.label_ == "PERSON":
                                                detected_name = ent.text.strip()
                                                detected_name = re.sub(r'\s+', ' ', detected_name)

                                                # Vérifier si notre nom correspond
                                                if exec_name.lower() in detected_name.lower() or detected_name.lower() in exec_name.lower():
                                                    validated_executives.add(exec_name)
                                                    break
                                    except:
                                        # Si spaCy échoue, accepter le nom quand même
                                        validated_executives.add(exec_name)

                            # Si spaCy a validé des noms, utiliser ceux-là
                            # Sinon, GARDER TOUS les noms détectés (pas de filtrage)
                            if validated_executives:
                                cleaned_executives = validated_executives
                            # Si rien validé par spaCy, on garde cleaned_executives tel quel

                        # RÉACTIVÉ - Ajouter les noms d'exécutifs aux résultats
                        for exec_name in cleaned_executives:
                            # === ULTRA-STRICT VALIDATION (same as person names) ===
                            exec_words = exec_name.split()
                            exec_words_lower = [w.lower().strip('.,;:') for w in exec_words]
                            exec_lower = exec_name.lower()

                            # REJECT: Gerund phrases
                            if exec_words_lower and exec_words_lower[0].endswith('ing'):
                                continue

                            # REJECT: Possessive phrases
                            if any(w in {'our', 'his', 'her', 'its', 'their', 'your', 'my'} for w in exec_words_lower):
                                continue

                            # REJECT: Job titles without names
                            if exec_lower in {'senior vice', 'vice president', 'chief executive', 'chief financial',
                                             'chief operating', 'executive chairman', 'lead director',
                                             'audit compensation', 'audit compensation nominating'}:
                                continue

                            # REJECT: Business/document terms
                            if exec_lower in {'grant date fair', 'grant date number', 'grant date threshold',
                                             'total award grant', 'stock ownership guidelines', 'severance plan',
                                             'period total less', 'name age position', 'year tra liability',
                                             'exchange act rule', 'delaware law', 'compensation decisions',
                                             'compensation plans', 'compensation program', 'compensation programs',
                                             'director compensation', 'outstanding awards under'}:
                                continue

                            # REJECT: Activity phrases
                            if exec_lower in {'financing activities', 'investing activities', 'operating activities',
                                             'entering into', 'exposing us', 'enabling us', 'converting floating',
                                             'utilizing direct sales', 'using long', 'exercising its put',
                                             'dividing net income', 'reducing our ability', 'processing electronic payment',
                                             'adding cardtronics', 'advise our corporate', 'appoints charles drucker',
                                             'parties claiming ownership', 'persons offering consumer'}:
                                continue

                            # REJECT: Company references
                            if exec_lower in {'fifth third', 'fifth third bancorp', 'fifth third bank',
                                             'advent international corporation', 'blackrock', 'jpmorgan chase',
                                             'matthew taylor group'}:
                                continue

                            # REJECT: Common verbs at start
                            if exec_words_lower and exec_words_lower[0] in {'adding', 'advise', 'appoints', 'comparing',
                                                                            'contacting', 'converting', 'discussed', 'dividing',
                                                                            'enabling', 'entering', 'exercising', 'exposing',
                                                                            'financing', 'investing', 'operating', 'processing',
                                                                            'reducing', 'reviewing', 'since', 'stockholders',
                                                                            'using', 'utilizing'}:
                                continue

                            # REJECT: Generic business phrases
                            if exec_lower in {'business segment', 'customer incentives', 'merchant category',
                                             'payment networks', 'transaction volume', 'meetings per year',
                                             'merchant acquiring entities', 'numerous laws', 'one stockholder',
                                             'two different persons', 'written consent', 'virtue hereof',
                                             'your bank', 'since november', 'discussed above', 'formal policy regarding',
                                             'general economic conditions', 'rapid technological change',
                                             'registration statement number', 'regulatory guidelines'}:
                                continue

                            # Original validation (as backup)
                            has_invalid_word = any(w in INVALID_NAME_WORDS for w in exec_words_lower)

                            # DÉSACTIVÉ - Executive name detection disabled (too many false positives)
                            # if (exec_lower not in PERSON_NAME_FALSE_POSITIVES and
                            #     exec_lower not in DOCUMENT_STRUCTURE_TERMS and
                            #     not has_invalid_word):
                            #     page_findings.append({'type': 'executive_name', 'value': exec_name, 'page': page_num + 1})
                            pass  # Executive name detection disabled
                        
                        # ===== TOUS LES NOMS DE PERSONNES (pas seulement executives) =====
                        all_person_names_found = set()
                        
                        # Pattern 1: Noms dans les contextes de personnes (directors, officers, employees, etc.)
                        person_context_patterns = [
                            re.compile(
                                r'(?:Director|Officer|Employee|Trustee|Beneficial\s+Owner|Shareholder|Stockholder|Member|Partner|Principal|Agent|Representative|Signatory|Authorized\s+Signatory)[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
                                re.IGNORECASE
                            ),
                            re.compile(
                                r'(?:Name|Person|Individual|Contact)[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
                                re.IGNORECASE
                            ),
                        ]
                        
                        for pattern in person_context_patterns:
                            for match in pattern.finditer(text):
                                name = match.group(1).strip()
                                name = re.sub(r'\s+', ' ', name)
                                
                                if 5 <= len(name) <= 50:
                                    words = name.split()
                                    if len(words) >= 2 and len(words) <= 5:
                                        # Vérifier que ce n'est pas un faux positif
                                        name_upper = name.upper()
                                        false_positives = [
                                            'UNITED STATES', 'NEW YORK', 'LOS ANGELES', 'THE COMPANY',
                                            'BALANCE SHEET', 'CASH FLOW', 'INCOME TAX', 'STOCK OPTION',
                                            'SECURITIES AND EXCHANGE', 'COMMISSION', 'WASHINGTON',
                                            'TABLE OF CONTENTS', 'APPENDIX', 'EXHIBIT'
                                        ]
                                        if name_upper not in false_positives:
                                            all_person_names_found.add(name)
                        
                        # Pattern 2: Noms dans les listes/tableaux (format "Nom, Prénom" ou "Prénom Nom")
                        # Détecter les patterns de noms propres typiques
                        name_list_patterns = [
                            # Format "Nom, Prénom" (nom de famille en premier)
                            re.compile(
                                r'\b([A-Z][a-z]+),\s+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+)\b'
                            ),
                            # Format "Prénom Nom" (prénom en premier)
                            re.compile(
                                r'\b([A-Z][a-z]+)\s+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+)\b'
                            ),
                        ]
                        
                        for pattern in name_list_patterns:
                            for match in pattern.finditer(text):
                                if len(match.groups()) == 2:
                                    part1, part2 = match.groups()
                                    # Construire le nom complet
                                    if ',' in match.group(0):
                                        # Format "Nom, Prénom" -> "Prénom Nom"
                                        full_name = f"{part2} {part1}"
                                    else:
                                        # Format "Prénom Nom"
                                        full_name = f"{part1} {part2}"
                                    
                                    full_name = re.sub(r'\s+', ' ', full_name).strip()
                                    
                                    # Validation
                                    if 5 <= len(full_name) <= 50:
                                        words = full_name.split()
                                        if len(words) >= 2 and len(words) <= 5:
                                            # Vérifier que chaque mot commence par une majuscule
                                            if all(w[0].isupper() for w in words if len(w) > 1):
                                                # Vérifier le contexte pour éviter les faux positifs
                                                context_start = max(0, match.start() - 100)
                                                context_end = min(len(text), match.end() + 100)
                                                context = text[context_start:context_end].lower()
                                                
                                                # Éviter les contextes non pertinents
                                                bad_contexts = [
                                                    'table of contents', 'index', 'appendix', 'reference',
                                                    'balance sheet', 'income statement', 'cash flow',
                                                    'note', 'footnote', 'exhibit', 'schedule',
                                                    'form 10-k', 'form 10-q', 'form 8-k',
                                                    'part i', 'part ii', 'part iii', 'part iv'
                                                ]

                                                # Chercher des indicateurs positifs (REQUIRED - not optional)
                                                good_indicators = [
                                                    'director', 'officer', 'employee', 'trustee',
                                                    'shareholder', 'beneficial owner', 'signatory',
                                                    'authorized', 'representative', 'agent',
                                                    'management', 'board', 'committee', 'team',
                                                    'signed', 'by:', 'name:', '/s/'
                                                ]

                                                has_good_indicator = any(ind in context for ind in good_indicators)
                                                has_bad_context = any(bad in context for bad in bad_contexts)

                                                # Check against false positives (EXACT MATCH ONLY)
                                                full_name_lower = full_name.lower()
                                                is_false_positive = (
                                                    full_name_lower in PERSON_NAME_FALSE_POSITIVES or
                                                    full_name_lower in DOCUMENT_STRUCTURE_TERMS
                                                    # Don't check partial matches - too aggressive
                                                )

                                                # RELAXED: Accept if (has indicator OR no bad context) AND not false positive
                                                # This allows names in signature sections even without explicit indicators
                                                if not is_false_positive and not has_bad_context:
                                                    # Check that no word in the name is an invalid word
                                                    words_lower = [w.lower().strip('.,;:') for w in words]
                                                    has_invalid_word = any(w in INVALID_NAME_WORDS for w in words_lower)

                                                    if not has_invalid_word:
                                                        # Additional validation: check capitalization pattern
                                                        # Reject all-caps or names with too many capitals
                                                        cap_ratio = sum(1 for c in full_name if c.isupper()) / len(full_name)
                                                        if cap_ratio <= 0.5:  # Max 50% capitals (allows for initials)
                                                            # If no good indicator, require at least 2 words (first + last name)
                                                            if has_good_indicator or len(words) >= 2:
                                                                all_person_names_found.add(full_name)
                        
                        # Pattern 3: Noms dans les sections de signatures et certifications
                        signature_section_patterns = [
                            re.compile(
                                r'(?:Signed|Certified|Attested|Witnessed|Notarized)[\s\S]{0,200}?([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
                                re.IGNORECASE | re.MULTILINE
                            ),
                            re.compile(
                                r'(?:By|Name\s+of|Signature\s+of)[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
                                re.IGNORECASE
                            ),
                        ]
                        
                        for pattern in signature_section_patterns:
                            for match in pattern.finditer(text):
                                name = match.group(1).strip()
                                name = re.sub(r'\s+', ' ', name)
                                
                                if 5 <= len(name) <= 50:
                                    words = name.split()
                                    if len(words) >= 2 and len(words) <= 5:
                                        # Vérifier que ce n'est pas déjà dans les executives
                                        if name not in cleaned_executives:
                                            all_person_names_found.add(name)
                        
                        # Pattern 4: Noms dans les tableaux de personnes (format tabulaire)
                        # Chercher des lignes qui ressemblent à des noms dans des contextes de personnes
                        # Format: "Nom | Titre | Autre info" ou "Nom - Titre"
                        table_name_patterns = [
                            re.compile(
                                r'([A-Z][a-z]+(?:\s+[A-Z]\.?\s*)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*[|\-]\s*[A-Z]',
                                re.MULTILINE
                            ),
                        ]
                        
                        for pattern in table_name_patterns:
                            for match in pattern.finditer(text):
                                name = match.group(1).strip()
                                name = re.sub(r'\s+', ' ', name)
                                
                                if 5 <= len(name) <= 50:
                                    words = name.split()
                                    if len(words) >= 2 and len(words) <= 5:
                                        # Vérifier le contexte
                                        context_start = max(0, match.start() - 150)
                                        context_end = min(len(text), match.end() + 150)
                                        context = text[context_start:context_end].lower()
                                        
                                        # Chercher des indicateurs de tableau de personnes
                                        person_table_indicators = [
                                            'director', 'officer', 'executive', 'management',
                                            'board', 'committee', 'trustee', 'beneficial owner',
                                            'shareholder', 'stockholder', 'employee', 'personnel'
                                        ]
                                        
                                        if any(ind in context for ind in person_table_indicators):
                                            if name not in cleaned_executives:
                                                all_person_names_found.add(name)
                        
                        # Pattern 5: Noms dans les emails (avant @)
                        # Les emails contiennent souvent des noms
                        email_name_pattern = re.compile(r'\b([a-z]+\.?[a-z]+)\.([a-z]+)@', re.IGNORECASE)
                        for match in email_name_pattern.finditer(text):
                            first_part = match.group(1).capitalize()
                            second_part = match.group(2).capitalize()
                            potential_name = f"{first_part} {second_part}"
                            
                            if 5 <= len(potential_name) <= 30:
                                # Vérifier que ça ressemble à un nom (pas un mot technique)
                                if len(first_part) >= 2 and len(second_part) >= 2:
                                    if potential_name not in cleaned_executives:
                                        all_person_names_found.add(potential_name)
                        
                        # Pattern 6: Détection générale de noms propres (prénom + nom) - VERSION TRÈS STRICTE
                        # DÉSACTIVÉ par défaut car trop de faux positifs
                        # Seulement activé si contexte très clair
                        # Format: "Prénom Nom" ou "Prénom Initial Nom"
                        general_name_pattern = re.compile(
                            r'\b([A-Z][a-z]{2,15})\s+([A-Z]\.?\s*)?([A-Z][a-z]{2,20})\b'
                        )
                        
                        for match in general_name_pattern.finditer(text):
                            first_name = match.group(1)
                            middle_initial = match.group(2) if match.group(2) else ''
                            last_name = match.group(3)
                            
                            # Construire le nom complet
                            if middle_initial:
                                full_name = f"{first_name} {middle_initial.strip()} {last_name}"
                            else:
                                full_name = f"{first_name} {last_name}"
                            
                            full_name = re.sub(r'\s+', ' ', full_name).strip()
                            
                            # Validation TRÈS stricte pour éviter les faux positifs
                            if 5 <= len(full_name) <= 50:
                                words = full_name.split()
                                if len(words) >= 2 and len(words) <= 4:
                                    # Vérifier le contexte pour s'assurer que c'est bien un nom de personne
                                    context_start = max(0, match.start() - 200)
                                    context_end = min(len(text), match.end() + 200)
                                    context = text[context_start:context_end]
                                    context_lower = context.lower()
                                    
                                    # Indicateurs positifs FORTS (contexte suggérant un nom de personne)
                                    strong_positive_indicators = [
                                        'director:', 'officer:', 'executive:', 'employee:', 'trustee:',
                                        'shareholder:', 'stockholder:', 'beneficial owner:', 'owner:',
                                        'signatory:', 'authorized signatory:', 'representative:', 'agent:',
                                        'by:', 'name:', 'person:', 'individual:', 'contact:',
                                        'signed by', 'certified by', 'attested by', 'witnessed by',
                                        'notarized by', 'email:', 'phone:', 'address:', 'residence:'
                                    ]
                                    
                                    # Indicateurs négatifs (contexte suggérant que ce n'est PAS un nom)
                                    negative_indicators = [
                                        'table of contents', 'index', 'appendix', 'exhibit', 'schedule',
                                        'balance sheet', 'income statement', 'cash flow', 'statement',
                                        'note', 'footnote', 'page', 'section', 'chapter',
                                        'united states', 'new york', 'los angeles', 'washington',
                                        'commission', 'securities', 'exchange', 'filing',
                                        'the company', 'our company', 'such company', 'a company',
                                        'provides', 'entered', 'completed', 'acquired', 'merged',
                                        'income', 'taxes', 'financial', 'segment', 'revenue',
                                        'street', 'avenue', 'road', 'drive', 'boulevard', 'lane',
                                        'inc.', 'llc', 'corp', 'corporation', 'limited',
                                        'vantiv', 'and', 'or', 'the', 'of', 'in', 'to', 'with',
                                        'is', 'are', 'was', 'were', 'be', 'been', 'will', 'would'
                                    ]
                                    
                                    has_strong_positive = any(ind in context_lower for ind in strong_positive_indicators)
                                    has_negative = any(ind in context_lower for ind in negative_indicators)
                                    
                                    # Vérifier la position (début de ligne ou après ponctuation)
                                    line_start = text.rfind('\n', 0, match.start())
                                    if line_start == -1:
                                        line_start = 0
                                    before_name = text[line_start:match.start()].strip()
                                    is_at_line_start = len(before_name) < 5 or before_name.endswith(('.', ',', ':', ';', '-', '|'))
                                    
                                    # FILTRE TRÈS STRICT: Accepter SEULEMENT si:
                                    # 1. Il y a un indicateur positif FORT (avec ':')
                                    # 2. ET pas d'indicateur négatif
                                    # 3. ET ce n'est pas déjà dans les executives
                                    # 4. ET le nom ne contient pas de mots interdits
                                    words_lower = [w.lower().rstrip('.,;:') for w in words]
                                    excluded_words_check = {'and', 'or', 'the', 'a', 'an', 'of', 'in', 'to', 'with', 'is', 'are', 'was', 'were', 'be', 'been', 'will', 'would', 'vantiv', 'company', 'co', 'title', 'officer', 'director', 'chairman'}
                                    
                                    if full_name not in cleaned_executives:
                                        if has_strong_positive and not has_negative and not any(w in excluded_words_check for w in words_lower):
                                            # Validation finale: vérifier que les mots ressemblent à des noms
                                            if all(3 <= len(w) <= 15 for w in words_lower):
                                                all_person_names_found.add(full_name)
                        
                        # Nettoyer et valider tous les noms trouvés avec filtres stricts
                        cleaned_person_names = set()
                        
                        # Liste étendue de mots à exclure (mots de liaison, verbes, articles, etc.)
                        excluded_words = {
                            'and', 'or', 'the', 'a', 'an', 'of', 'in', 'on', 'at', 'to', 'for',
                            'with', 'by', 'from', 'as', 'is', 'are', 'was', 'were', 'be', 'been',
                            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                            'should', 'may', 'might', 'must', 'can', 'shall', 'if', 'that', 'this',
                            'these', 'those', 'which', 'who', 'whom', 'whose', 'where', 'when',
                            'why', 'how', 'what', 'all', 'any', 'each', 'every', 'some', 'no',
                            'not', 'but', 'than', 'then', 'there', 'here', 'other', 'another',
                            'such', 'same', 'more', 'most', 'less', 'least', 'very', 'too', 'so',
                            'only', 'just', 'also', 'even', 'still', 'yet', 'already', 'again',
                            'about', 'above', 'below', 'under', 'over', 'through', 'during',
                            'before', 'after', 'since', 'until', 'while', 'because', 'although',
                            'however', 'therefore', 'thus', 'hence', 'moreover', 'furthermore',
                            'additionally', 'further', 'indeed', 'rather', 'quite', 'rather',
                            'vantiv', 'company', 'companies', 'corporation', 'inc', 'llc', 'corp',
                            'ltd', 'limited', 'incorporated', 'group', 'holdings', 'enterprises',
                            'services', 'business', 'financial', 'securities', 'exchange',
                            'commission', 'registrant', 'issuer', 'filing', 'report', 'statement',
                            'balance', 'sheet', 'income', 'cash', 'flow', 'tax', 'taxes',
                            'director', 'officer', 'executive', 'employee', 'trustee', 'shareholder',
                            'stockholder', 'beneficial', 'owner', 'signatory', 'authorized',
                            'representative', 'agent', 'management', 'board', 'committee', 'team',
                            'chairman', 'president', 'secretary', 'treasurer', 'controller',
                            'chief', 'vice', 'senior', 'junior', 'lead', 'head', 'general',
                            'additional', 'information', 'combination', 'act', 'court', 'high',
                            'mastercard', 'visa', 'stock', 'combined', 'uk', 'companies', 'worldpay'
                        }
                        
                        # Mots qui ne peuvent PAS être des noms de personnes (trop génériques)
                        invalid_name_words = {
                            'additional', 'information', 'business', 'combination', 'companies',
                            'act', 'court', 'high', 'mastercard', 'visa', 'stock', 'combined',
                            'worldpay', 'vantiv', 'services', 'executive', 'chairman', 'director',
                            'officer', 'secretary', 'treasurer', 'lead', 'title', 'and', 'co',
                            'the', 'of', 'in', 'to', 'with', 'an', 'a', 'or', 'be', 'is', 'are',
                            'was', 'were', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
                            'may', 'might', 'must', 'can', 'shall', 'if', 'that', 'this', 'these',
                            'those', 'which', 'who', 'whom', 'whose', 'where', 'when', 'why',
                            'how', 'what', 'all', 'any', 'each', 'every', 'some', 'no', 'not',
                            'but', 'than', 'then', 'there', 'here', 'other', 'another', 'such',
                            'same', 'more', 'most', 'less', 'least', 'very', 'too', 'so', 'only',
                            'just', 'also', 'even', 'still', 'yet', 'already', 'again', 'about',
                            'above', 'below', 'under', 'over', 'through', 'during', 'before',
                            'after', 'since', 'until', 'while', 'because', 'although', 'however',
                            'therefore', 'thus', 'hence', 'moreover', 'furthermore', 'additionally',
                            'further', 'indeed', 'rather', 'quite', 'means', 'check', 'mark', 'whether',
                            'purport', 'exempt', 'principal', 'traders', 'merchants', 'later',
                            'registrant', 'combined', 'filings', 'security', 'holdings', 'team',
                            'reporting', 'incremental', 'amendment', 'affirmative', 'vote', 'fact',
                            'offeree', 'undersigned', 'hereunto'
                        }
                        
                        # Termes financiers à exclure
                        financial_terms = [
                            'accounting', 'adjustment', 'amount', 'shares', 'fair', 'value',
                            'restricted', 'net', 'activity', 'attributable', 'pretax', 'tax',
                            'related', 'transactions', 'secondary', 'offering', 'purchase',
                            'plan', 'time', 'awards', 'weighted', 'merchant', 'financial',
                            'require', 'services'
                        ]
                        
                        # Phrases financières complètes à exclure
                        financial_phrases = [
                            'accounting adjustment', 'amount shares', 'fair value', 'restricted',
                            'net activity', 'attributable', 'pretax activity', 'tax',
                            'related transactions', 'secondary offering', 'purchase plan',
                            'share purchase', 'time awards', 'weighted', 'merchant services',
                            'financial', 'and/or require', 'shares amount', 'adjustment change',
                            'value restricted', 'activity attributable', 'activity tax',
                            'offering purchase', 'purchase plan', 'awards weighted'
                        ]
                        
                        # Noms d'entreprises connus à exclure
                        company_names_to_exclude = [
                            'vantiv', 'jpmorgan', 'morgan stanley', 'credit suisse',
                            'worldpay', 'tokyo-mitsubishi', 'mitsubishi ufj',
                            'fleetcor', 'broadridge', 'paymetric'
                        ]
                        
                        for person_name in all_person_names_found:
                            # Nettoyer les espaces multiples
                            person_name = re.sub(r'\s+', ' ', person_name).strip()

                            # === NETTOYAGE DES CHIFFRES/DATES/CARACTÈRES PARASITES (AVANT VALIDATION) ===
                            # Enlever les dates/nombres au début (ex: "2/17/2026 S. Ferris" → "S. Ferris")
                            person_name = re.sub(r'^[\d/,\s]+(?=[A-Z])', '', person_name)

                            # Enlever les nombres/dates à la fin (ex: "Jeffrey Stieﬂer 138,750" → "Jeffrey Stieﬂer")
                            person_name = re.sub(r'\s+[\d/,]+$', '', person_name)

                            # Enlever les parenthèses avec nombres à la fin (ex: "Charles Drucker(7)(8" → "Charles Drucker")
                            person_name = re.sub(r'\([^\)]*\d[^\)]*\)*$', '', person_name)
                            person_name = re.sub(r'\(\d+$', '', person_name)  # Cas "(7" sans fermeture

                            # Enlever les tirets multiples à la fin (ex: "David Karnstedt - - -" → "David Karnstedt")
                            person_name = re.sub(r'[\s\-]+$', '', person_name)

                            # Enlever les marqueurs de signature (tous les patterns)
                            # Ex: "Lars Anderson /s/", "Lisa Hook /S/", "Name (s)", "Name signed:", etc.
                            person_name = re.sub(r'\s*/[sS]/\s*', ' ', person_name, flags=re.IGNORECASE)
                            person_name = re.sub(r'\s*\([sS]\)\s*', ' ', person_name, flags=re.IGNORECASE)
                            person_name = re.sub(r'\s+signed:?\s*', ' ', person_name, flags=re.IGNORECASE)
                            person_name = re.sub(r'\s+/signed/?\s*', ' ', person_name, flags=re.IGNORECASE)

                            # Nettoyer à nouveau les espaces
                            person_name = re.sub(r'\s+', ' ', person_name).strip()

                            # Si vide après nettoyage, skip
                            if not person_name or len(person_name) < 2:
                                continue

                            # === ULTRA-STRICT VALIDATION ===
                            words = person_name.split()
                            words_lower = [w.lower().strip('.,;:') for w in words]
                            person_name_lower = person_name.lower()  # Define this early for all checks

                            # REJECT: Gerund phrases (verb-ing + words)
                            if words_lower and words_lower[0].endswith('ing'):
                                continue  # Skip: "enabling us", "reducing our ability", etc.

                            # REJECT: Possessive phrases (our/his/her/its/their/your/my + words)
                            possessives = {'our', 'his', 'her', 'its', 'their', 'your', 'my'}
                            if any(w in possessives for w in words_lower):
                                continue  # Skip: "our clients", "his prior employer", etc.

                            # REJECT: Job titles without actual names
                            job_title_phrases = {
                                'senior vice', 'vice president', 'chief executive', 'chief financial',
                                'chief operating', 'executive chairman', 'lead director',
                                'audit compensation', 'audit compensation nominating'
                            }
                            if person_name_lower in job_title_phrases:
                                continue

                            # REJECT: Mots-clés qui indiquent un faux positif (vérification partielle)
                            # Si le nom contient ces mots, c'est probablement pas une personne
                            false_positive_keywords = [
                                'accountant', 'transaction', 'addendum', 'letter', 'notice',
                                'award', 'repurchase', 'salary', 'severance', 'performance',
                                'intellectual', 'investor', 'relations', 'property', 'processing',
                                'principal', 'related', 'restated', 'offer', 'grant'
                            ]
                            if any(keyword in person_name_lower for keyword in false_positive_keywords):
                                continue

                            # REJECT: Document/business structure terms
                            business_structure_terms = {
                                'grant date fair', 'grant date number', 'grant date threshold',
                                'total award grant', 'stock ownership guidelines', 'severance plan',
                                'period total less', 'name age position', 'year tra liability',
                                'exchange act rule', 'delaware law', 'formal policy regarding',
                                'business segment', 'payment networks', 'merchant category',
                                'customer incentives', 'transaction volume', 'general economic conditions',
                                'rapid technological change', 'numerous laws', 'written consent',
                                'virtue hereof', 'since november', 'discussed above', 'one stockholder',
                                'two different persons', 'registration statement number', 'regulatory guidelines',
                                # Document sections/headings
                                'principal accountant fees', 'related transactions', 'performance awards',
                                'performance graph', 'leadership structur', 'financial ofﬁcer',
                                'award notice', 'offer letter', 'restated offer letter', 'grant notice',
                                'repurchase addendum', 'position year salary', 'bridge lenders',
                                'stockholder outr', 'withhold allall', 'performance shar',
                                'year salary', 'severance beneﬁts', 'martin cash severance',
                                # Table headers and form instructions
                                'check', 'non', 'smaller', 'emer', 'meetin g',
                                # Legal/financial terms
                                'pro forma', 'non-de minimis', 'de minimis', 'compris ed',
                                'frank act', 'sarbanes oxley', 'secur ed cr',
                                # Department/function names
                                'investor relations', 'intellectual property', 'card processing',
                                'capital research global',
                                # Company/service names
                                'worldwide fleetcor', 'fleetcor', 'broadridge investor communication',
                                'broadridge', 'paymetric', 'moneris usa', 'jpmorgan', 'matthew taylor group',
                                # Broken/partial words
                                'jeffr ey', 'nelson f',
                                'nelson f. greene title', 's. ferris cash', 's. ferris vc'
                            }
                            if person_name_lower in business_structure_terms:
                                continue

                            # REJECT: Patterns suspects (références de documents)
                            # Rejeter les noms suivis directement de ".10-K" ou ".10-Q" (références de documents)
                            if re.search(r'\.\d+[KQ-]', person_name):
                                continue

                            # REJECT: Very short fragments (single words < 3 characters, or obvious non-names)
                            if len(words) == 1:
                                word = words[0].strip('.,;:')
                                word_clean = words_lower[0].strip('.,;:')
                                # Reject single words that are < 3 chars OR common non-name words
                                if len(word_clean) < 3 or word_clean in {'check', 'non', 'emer', 'all', 'yes', 'no'}:
                                    continue
                                # REJECT: Mauvaise capitalisation pour mot unique
                                # Accepter: "Smith" (première maj, reste min) ou initiales courtes
                                # Rejeter: "SMITH" (tout maj si > 2 chars), "smith" (commence par min)
                                if not word[0].isupper():
                                    continue  # Rejeter si commence par minuscule
                                if len(word) > 2 and word.isupper():
                                    continue  # Rejeter si tout majuscule (sauf initiales)

                            # REJECT: Broken/truncated words (contain spaces in odd places or end with incomplete patterns)
                            # Example: "compris ed", "Performance Shar", "Leadership Structur"
                            if any(' ' in word for word in words if len(word) > 2):  # Space inside a word
                                continue
                            # Reject if ends with common truncation patterns (incomplete words)
                            if person_name_lower.endswith(('structur', ' shar', ' outr', ' ed', ' ing')):
                                continue

                            # REJECT: Document references with numbers/dots (e.g., "Rules 8.1")
                            if re.search(r'\b\d+\.\d+\b', person_name):  # Contains numbered references like "8.1"
                                continue

                            # REJECT: Activity/action phrases
                            activity_phrases = {
                                'financing activities', 'investing activities', 'operating activities',
                                'financing activities during', 'entering into', 'exposing us',
                                'enabling us', 'converting floating', 'utilizing direct sales',
                                'utilizing our integrated', 'using long', 'exercising its put',
                                'dividing net income', 'reducing our ability', 'comparing our actual',
                                'contacting investor relations', 'processing electronic payment',
                                'reviewing overhang levels', 'adding cardtronics', 'advise our corporate',
                                'appoints charles drucker', 'parties claiming ownership',
                                'persons offering consumer', 'stockholders using substantially',
                                'variance between our', 'various payment networks'
                            }
                            if person_name_lower in activity_phrases:
                                continue

                            # REJECT: Compensation/awards terms
                            compensation_terms = {
                                'compensation decisions', 'compensation plans', 'compensation program',
                                'compensation program follows', 'compensation programs',
                                'director compensation', 'outstanding awards under'
                            }
                            if person_name_lower in compensation_terms:
                                continue

                            # REJECT: Company/bank references
                            company_refs = {
                                'fifth third', 'fifth third bancorp', 'fifth third bank',
                                'fifth third represents', 'advent international corporation',
                                'blackrock', 'jpmorgan chase', 'matthew taylor group'
                            }
                            if person_name_lower in company_refs:
                                continue

                            # REJECT: Waiver phrases (partial names we don't want)
                            if any(phrase in person_name_lower for phrase in ['waived his', 'waived her']):
                                continue

                            # REJECT: Names that are actually sentence fragments
                            # Check if first word is a common verb or preposition
                            common_verbs = {
                                'adding', 'advise', 'appoints', 'comparing', 'contacting',
                                'converting', 'discussed', 'dividing', 'enabling', 'entering',
                                'exercising', 'exposing', 'financing', 'investing', 'operating',
                                'parties', 'payme', 'persons', 'processing', 'reducing',
                                'registration', 'reviewing', 'since', 'stockholders', 'using',
                                'utilizing', 'variance', 'various', 'virtue'
                            }
                            if words_lower and words_lower[0] in common_verbs:
                                continue

                            # EXCLURE immédiatement si c'est un nom d'entreprise connu
                            # (person_name_lower already defined at start of validation)
                            if any(company in person_name_lower for company in company_names_to_exclude):
                                continue
                            
                            # EXCLURE si contient des phrases financières
                            if any(phrase in person_name_lower for phrase in financial_phrases):
                                continue
                            
                            # Validation: doit avoir au moins 2 mots (prénom + nom)
                            words = person_name.split()
                            if len(words) >= 2 and len(words) <= 5:
                                # Vérifier que chaque mot fait sens (pas trop court, pas trop long)
                                if all(2 <= len(w) <= 20 for w in words):
                                    # Vérifier que chaque mot commence par une majuscule
                                    if all(w[0].isupper() for w in words if len(w) > 1):
                                        # FILTRE STRICT: Exclure si contient des mots interdits
                                        words_lower = [w.lower().rstrip('.,;:') for w in words]
                                        
                                        # Ne pas accepter si un mot est dans la liste des mots exclus
                                        if any(w in excluded_words for w in words_lower):
                                            continue
                                        
                                        # Ne pas accepter si un mot est dans la liste des mots invalides
                                        if any(w in invalid_name_words for w in words_lower):
                                            continue
                                        
                                        # EXCLURE si contient des termes financiers dans les mots
                                        if any(term in words_lower for term in financial_terms):
                                            continue
                                        
                                        # EXCLURE si le nom contient des mots collés (ex: "AdjustmentChange", "SharesAmount")
                                        # Vérifier si un mot contient une majuscule au milieu (suggère deux mots collés)
                                        has_collapsed_words = False
                                        for word in words:
                                            if len(word) > 8:  # Mots longs suspects
                                                # Compter les majuscules au milieu du mot (pas au début)
                                                mid_caps = sum(1 for i, c in enumerate(word[1:], 1) if c.isupper() and word[i-1].isalpha())
                                                if mid_caps > 0:
                                                    # C'est probablement deux mots collés (ex: "AdjustmentChange")
                                                    has_collapsed_words = True
                                                    break
                                        
                                        if has_collapsed_words:
                                            continue
                                        
                                        # Ne pas accepter si le nom contient des mots de liaison communs
                                        if any(w in ['and', 'or', 'the', 'of', 'in', 'to', 'with', 'an', 'a', 'and/or'] for w in words_lower):
                                            continue
                                        
                                        # Vérifier qu'il n'y a pas trop de majuscules (éviter les acronymes)
                                        uppercase_ratio = sum(1 for c in person_name if c.isupper()) / len(person_name)
                                        if uppercase_ratio <= 0.4:
                                            # Éviter les faux positifs finaux
                                            person_name_upper = person_name.upper()
                                            final_false_positives = [
                                                'UNITED STATES', 'NEW YORK', 'LOS ANGELES',
                                                'BALANCE SHEET', 'CASH FLOW', 'INCOME TAX',
                                                'STOCK OPTION', 'SECURITIES', 'EXCHANGE',
                                                'COMMISSION', 'WASHINGTON', 'TABLE OF',
                                                'CONTENTS', 'APPENDIX', 'EXHIBIT', 'SCHEDULE',
                                                'NOTE', 'FOOTNOTE', 'PAGE', 'SECTION',
                                                'VANTIV IN', 'VANTIV TO', 'VANTIV WITH',
                                                'AN INDEPENDENT', 'AND SECRETARY', 'BE THE',
                                                'THE EXECUTIVE', 'THE LEAD', 'CHECK MARK',
                                                'NO LATER', 'OF REGISTRANT', 'THE COMBINED',
                                                'ON THE COMBINED', 'OTHER VANTIV', 'TEAM REPORTING',
                                                'THE HIGH COURT', 'THE INCREMENTAL', 'THE AFFIRMATIVE',
                                                'THE FACT', 'THE OFFEREE', 'THE UNDERSIGNED',
                                                'ADDITIONAL INFORMATION', 'BUSINESS COMBINATION',
                                                'COMPANIES ACT', 'EXECUTIVE CHAIRMAN', 'HIGH COURT',
                                                'LEAD DIRECTOR', 'MASTERCARD VISA', 'THE BUSINESS',
                                                'VANTIV STOCK', 'IS AN EMERGING', 'WILL CO',
                                                # Ajouter les faux positifs financiers
                                                'ACCOUNTING ADJUSTMENT', 'AMOUNT SHARES', 'FAIR VALUE',
                                                'RESTRICTED', 'NET ACTIVITY', 'ATTRIBUTABLE', 'PRETAX ACTIVITY',
                                                'RELATED TRANSACTIONS', 'SECONDARY OFFERING', 'PURCHASE PLAN',
                                                'SHARE PURCHASE', 'TIME AWARDS', 'WEIGHTED', 'MERCHANT SERVICES',
                                                'FINANCIAL', 'AND/OR REQUIRE', 'SHARES AMOUNT', 'ADJUSTMENT CHANGE',
                                                'VALUE RESTRICTED', 'ACTIVITY ATTRIBUTABLE', 'ACTIVITY TAX',
                                                'OFFERING PURCHASE', 'PURCHASE PLAN', 'AWARDS WEIGHTED',
                                                'ACCOUNTING ADJUSTMENTCHANGE', 'AMOUNT SHARESAMOUNT',
                                                'FAIR VALUERESTRICTED', 'MERCHANT SERVICESFINANCIAL',
                                                'NET ACTIVITYATTRIBUTABLE', 'PRETAX ACTIVITYTAX',
                                                'RELATED TRANSACTIONS', 'SECONDARY OFFERINGPURCHASE',
                                                'SHARE PURCHASE PLAN', 'SHARES AMOUNTSHARES',
                                                'TIME AWARDSWEIGHTED', 'VANTIV'
                                            ]
                                            if not any(fp in person_name_upper for fp in final_false_positives):
                                                # Vérifier que le nom ne se termine pas par des mots suspects
                                                last_word = words_lower[-1]
                                                invalid_endings = ['and', 'or', 'the', 'of', 'in', 'to', 'with', 'co', 
                                                                  'title', 'officer', 'director', 'chairman', 'name',
                                                                  'financial', 'services', 'plan', 'tax', 'amount',
                                                                  'shares', 'awards', 'transactions', 'require']
                                                if last_word not in invalid_endings:
                                                    # S'assurer que ce n'est pas déjà dans les executives
                                                    if person_name not in cleaned_executives:
                                                        # Validation finale: le nom doit ressembler à un vrai nom
                                                        # (pas de mots trop courts, pas de répétitions)
                                                        if len(set(words_lower)) == len(words_lower):  # Pas de mots dupliqués
                                                            # VALIDATION CAPITALISATION: Chaque mot doit commencer par majuscule
                                                            # Format attendu: "John Smith", "M. Taylor", "Smith"
                                                            valid_capitalization = True
                                                            for word in words:
                                                                # Retirer ponctuation éventuelle
                                                                word_clean = word.rstrip('.,;:')
                                                                if len(word_clean) == 0:
                                                                    valid_capitalization = False
                                                                    break

                                                                # Initiales acceptées: 1-2 chars majuscules (ex: "M.", "S.", "Jr")
                                                                if len(word_clean) <= 2:
                                                                    if not word_clean[0].isupper():
                                                                        valid_capitalization = False
                                                                        break
                                                                else:
                                                                    # Mots normaux: première lettre majuscule, reste minuscule
                                                                    if not word_clean[0].isupper():
                                                                        valid_capitalization = False
                                                                        break
                                                                    # Rejeter si tout en majuscules (sauf si <= 2 chars)
                                                                    if word_clean.isupper() and len(word_clean) > 2:
                                                                        valid_capitalization = False
                                                                        break

                                                            if valid_capitalization:
                                                                cleaned_person_names.add(person_name)
                        
                        # Validation finale avec spaCy NER (si disponible)
                        if SPACY_AVAILABLE and SPACY_NLP and cleaned_person_names:
                            # Utiliser spaCy pour valider les noms détectés
                            validated_person_names = set()
                            
                            # Créer un texte avec tous les noms pour analyse
                            # Analyser chaque nom individuellement avec son contexte
                            for person_name in cleaned_person_names:
                                # Chercher le nom dans le texte original avec contexte
                                name_pos = text.find(person_name)
                                if name_pos != -1:
                                    # Prendre un contexte autour du nom
                                    context_start = max(0, name_pos - 100)
                                    context_end = min(len(text), name_pos + len(person_name) + 100)
                                    context_text = text[context_start:context_end]
                                    
                                    # Analyser avec spaCy
                                    doc = SPACY_NLP(context_text)
                                    
                                    # Vérifier si spaCy détecte ce nom comme une personne
                                    for ent in doc.ents:
                                        if ent.label_ == "PERSON":
                                            # Normaliser le nom détecté par spaCy
                                            detected_name = ent.text.strip()
                                            detected_name = re.sub(r'\s+', ' ', detected_name)
                                            
                                            # Vérifier si notre nom correspond (exact ou partiel)
                                            if person_name.lower() in detected_name.lower() or detected_name.lower() in person_name.lower():
                                                validated_person_names.add(person_name)
                                                break
                                    
                                    # Si spaCy n'a pas trouvé, mais que le nom est très probable (déjà validé par nos filtres),
                                    # on peut quand même l'accepter si le contexte est bon
                                    if person_name not in validated_person_names:
                                        # Vérifier si le contexte contient des indicateurs positifs
                                        context_lower = context_text.lower()
                                        strong_indicators = ['director', 'officer', 'executive', 'employee', 'trustee',
                                                           'shareholder', 'by:', 'name:', 'signed', 'certified']
                                        if any(ind in context_lower for ind in strong_indicators):
                                            validated_person_names.add(person_name)
                            
                            # Utiliser les noms validés par spaCy, ou tous si spaCy n'a rien trouvé
                            if validated_person_names:
                                cleaned_person_names = validated_person_names
                        
                        # Ajouter tous les noms de personnes trouvés avec déduplication
                        # Dédupliquer les noms qui sont les mêmes mais avec différentes capitalisations
                        deduplicated_names = {}
                        
                        for person_name in cleaned_person_names:
                            # Normaliser le nom pour la comparaison (minuscules, sans espaces multiples)
                            normalized = re.sub(r'\s+', ' ', person_name.lower().strip())

                            # Enlever les initiales du milieu pour la comparaison
                            # Ex: "stephanie l. ferris" → "stephanie ferris" pour matcher avec "stephanie ferris"
                            normalized = re.sub(r'\s+[a-z]\.\s+', ' ', normalized)  # "X. " au milieu
                            normalized = re.sub(r'\s+[a-z]\s+', ' ', normalized)    # "X " au milieu (sans point)

                            # Vérifier si ce nom est contenu dans un nom existant ou vice-versa
                            # Ex: "Lisa Hook" et "Lisa Hook BOON" doivent être dédupliqués
                            matching_key = None
                            for existing_key in deduplicated_names.keys():
                                # Si le nouveau nom est un préfixe d'un nom existant
                                if existing_key.startswith(normalized + ' '):
                                    matching_key = existing_key
                                    break
                                # Si un nom existant est un préfixe du nouveau nom
                                elif normalized.startswith(existing_key + ' '):
                                    matching_key = existing_key
                                    break

                            # Si on a trouvé une correspondance partielle, utiliser la clé du nom le plus court (le plus propre)
                            if matching_key:
                                # Garder le nom le plus court (probablement le plus propre)
                                if len(normalized) < len(matching_key):
                                    # Le nouveau nom est plus court, l'utiliser comme nouvelle clé
                                    deduplicated_names[normalized] = person_name
                                    # Supprimer l'ancienne clé plus longue
                                    del deduplicated_names[matching_key]
                                else:
                                    # Le nom existant est plus court, ne rien faire (garder l'existant)
                                    pass
                                continue

                            # Si on a déjà vu ce nom exact (normalisé), garder la meilleure version
                            if normalized in deduplicated_names:
                                existing_name = deduplicated_names[normalized]
                                # Préférer la version avec capitalisation standard (première lettre de chaque mot en majuscule)
                                # plutôt que tout en majuscules
                                if person_name.isupper() and not existing_name.isupper():
                                    # Garder la version existante (meilleure capitalisation)
                                    continue
                                elif not person_name.isupper() and existing_name.isupper():
                                    # Remplacer par la version avec meilleure capitalisation
                                    deduplicated_names[normalized] = person_name
                                # Sinon, garder la première version trouvée
                            else:
                                # Nouveau nom, l'ajouter
                                deduplicated_names[normalized] = person_name
                        
                        # Ajouter seulement les noms dédupliqués
                        # (Le nettoyage des chiffres a déjà été fait avant la validation)
                        for person_name in deduplicated_names.values():
                            page_findings.append({'type': 'person_name', 'value': person_name, 'page': page_num + 1})
                        
                        # ===== TICKER SYMBOLS (pour documents SEC) =====
                        # Format: (NYSE: XXXX) ou (NASDAQ: XXXX)
                        ticker_pattern = re.compile(r'\((?:NYSE|NASDAQ|AMEX|OTC):\s*([A-Z]{1,5})\)', re.IGNORECASE)
                        for match in ticker_pattern.finditer(text):
                            ticker = match.group(1)
                            page_findings.append({'type': 'ticker_symbol', 'value': f"{match.group(0)}", 'page': page_num + 1})
                        
                        # ===== DATES IMPORTANTES - VERSION AMÉLIORÉE =====
                        # DÉSACTIVÉ - Ne plus détecter les dates importantes
                        # important_dates_found = set()
                        # 
                        # # Pattern 1: Dates en format texte (December 31, 2017)
                        # date_pattern1 = re.compile(r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', re.IGNORECASE)
                        # for match in date_pattern1.finditer(text):
                        #     date_str = match.group(0)
                        #     # Contexte étendu pour détecter les dates importantes
                        #     context_start = max(0, match.start() - 80)
                        #     context_end = min(len(text), match.end() + 80)
                        #     context = text[context_start:context_end].lower()
                        #     
                        #     important_keywords = [
                        #         'period ended', 'as of', 'dated', 'effective', 'fiscal year',
                        #         'quarter ended', 'year ended', 'date', 'filing date',
                        #         'report date', 'balance sheet', 'statement date', 'closing date',
                        #         'maturity date', 'expiration date', 'commencement', 'termination',
                        #         'agreement date', 'contract date', 'execution date', 'signature date'
                        #     ]
                        #     
                        #     if any(keyword in context for keyword in important_keywords):
                        #         important_dates_found.add(date_str)
                        # 
                        # # Pattern 2: Dates en format numérique (12/31/2017, 12-31-2017, 12.31.2017)
                        # date_pattern2 = re.compile(r'\b(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})\b')
                        # for match in date_pattern2.finditer(text):
                        #     month, day, year = match.groups()
                        #     # Validation basique (mois entre 1-12, jour entre 1-31)
                        #     if 1 <= int(month) <= 12 and 1 <= int(day) <= 31 and 1900 <= int(year) <= 2100:
                        #         date_str = f"{month}/{day}/{year}"
                        #         context_start = max(0, match.start() - 80)
                        #         context_end = min(len(text), match.end() + 80)
                        #         context = text[context_start:context_end].lower()
                        #         
                        #         important_keywords = [
                        #             'period ended', 'as of', 'dated', 'effective', 'fiscal year',
                        #             'quarter ended', 'year ended', 'date', 'filing date',
                        #             'report date', 'balance sheet', 'statement date', 'closing date',
                        #             'maturity date', 'expiration date', 'commencement', 'termination',
                        #             'agreement date', 'contract date', 'execution date', 'signature date'
                        #         ]
                        #         
                        #         if any(keyword in context for keyword in important_keywords):
                        #             important_dates_found.add(date_str)
                        # 
                        # # Pattern 3: Dates en format ISO ou autre (2017-12-31)
                        # date_pattern3 = re.compile(r'\b(\d{4})[/\-\.](\d{1,2})[/\-\.](\d{1,2})\b')
                        # for match in date_pattern3.finditer(text):
                        #     year, month, day = match.groups()
                        #     if 1900 <= int(year) <= 2100 and 1 <= int(month) <= 12 and 1 <= int(day) <= 31:
                        #         date_str = f"{year}-{month}-{day}"
                        #         context_start = max(0, match.start() - 80)
                        #         context_end = min(len(text), match.end() + 80)
                        #         context = text[context_start:context_end].lower()
                        #         
                        #         important_keywords = [
                        #             'period ended', 'as of', 'dated', 'effective', 'fiscal year',
                        #             'quarter ended', 'year ended', 'date', 'filing date',
                        #             'report date', 'balance sheet', 'statement date', 'closing date'
                        #         ]
                        #         
                        #         if any(keyword in context for keyword in important_keywords):
                        #             important_dates_found.add(date_str)
                        # 
                        # for date_str in important_dates_found:
                        #     page_findings.append({'type': 'important_date', 'value': date_str, 'page': page_num + 1})
                        
                        # ===== COMPANY NAMES - VERSION AMÉLIORÉE =====
                        
                        company_names_found = set()
                        
                        # Liste étendue de suffixes d'entreprise (incluant AG, SA, etc.)
                        company_suffixes = r'(?:Inc\.?|LLC|LLP|L\.L\.C\.|L\.L\.P\.|Corporation|Corp\.?|Corp|Incorporated|Inc|Ltd\.?|Limited|LP|L\.P\.|PC|P\.C\.|PLLC|PLC|Co\.?|Company|Companies|Group|Holdings|Holdings?|Enterprises|Partners|Partnership|Bank|Banks|Trust|Capital|Securities|Financial|Services|AG|S\.A\.|SA|GmbH|BV|NV|SpA|S\.p\.A\.)'
                        
                        # Mots qui ne peuvent PAS être le premier mot d'un nom de compagnie (défini tôt pour être utilisé partout)
                        invalid_first_words = {
                            'amended', 'and', 'backstop', 'bridge', 'business', 'combined', 'commission',
                            'dealing', 'delaware', 'disclosure', 'executive', 'financial', 'incremental',
                            'irs', 'loan', 'london', 'offer', 'operations', 'original', 'panel', 'press',
                            'registrant', 'rule', 'scheme', 'securities', 'stock', 'u.s', 'us',
                            'a', 'an', 'is', 'will', 'be', 'are', 'was', 'were',
                            # Termes financiers/comptables
                            'accounting', 'accounts', 'accumulated', 'amount', 'average', 'based',
                            'capital', 'common', 'comprehensive', 'consolidated', 'consolidation',
                            'contents', 'controller', 'customer', 'earnings', 'equity', 'income',
                            'interest', 'internal', 'investor', 'management', 'marketing', 'merchant',
                            'net', 'non-operating', 'oci', 'office', 'parent', 'performance', 'period',
                            'principal', 'pro', 'public', 'restricted', 'retained', 'secondary',
                            'shares', 'significant', 'subsidiaries', 'tax', 'unit', 'visa',
                            # Termes génériques
                            'company', 'state', 'the', 'if', 'as', 'on', 'of', 'in', 'to', 'for',
                            'with', 'by', 'from', 'at', 'or', 'and', 'an', 'a',
                            # Comités et fonctions
                            'committee', 'committees', 'community', 'compliance', 'secretary'
                        }
                        
                        # D'abord, utiliser spaCy pour détecter TOUTES les organisations (si disponible)
                        # Cela permet de détecter des entreprises comme JPMorgan, WorldPay même sans suffixe
                        if SPACY_AVAILABLE and SPACY_NLP:
                            try:
                                # Analyser le texte complet de la page avec spaCy
                                text_to_analyze = text[:1000000] if len(text) > 1000000 else text
                                doc = SPACY_NLP(text_to_analyze)
                                
                                # Collecter toutes les organisations détectées par spaCy
                                for ent in doc.ents:
                                    if ent.label_ == "ORG":
                                        org_name = ent.text.strip()
                                        org_name = re.sub(r'\s+', ' ', org_name)
                                        
                                        # Validation basique
                                        if 3 <= len(org_name) <= 80:
                                            # Vérifier le contexte pour éviter les faux positifs
                                            start_char = ent.start_char
                                            end_char = ent.end_char
                                            context_start = max(0, start_char - 150)
                                            context_end = min(len(text_to_analyze), end_char + 150)
                                            context = text_to_analyze[context_start:context_end].lower()
                                            
                                            # Indicateurs négatifs (exclure les titres de documents)
                                            negative_indicators = [
                                                'agreement', 'letter', 'release', 'document', 'amendment',
                                                'commitment', 'combination', 'condition', 'operation',
                                                'disclosure', 'scheme', 'rule', 'act', 'commission',
                                                'file number', 'employer identification', 'delaware',
                                                'london stock exchange', 'securities exchange act',
                                                'press release', 'annual meeting', 'stockholder',
                                                'the business combination', 'the original loan',
                                                'the backstop commitment', 'the bridge commitment',
                                                'amended and restated', 'backstop commitment letter',
                                                'bridge commitment letter', 'loan agreement',
                                                'financial condition', 'operations', 'panel',
                                                'registrant', 'executive', 'the disclosure table',
                                                'the offer period', 'the incremental amendment'
                                            ]
                                            
                                            # Indicateurs positifs (suggèrent que c'est une vraie entreprise)
                                            positive_indicators = [
                                                'company', 'corporation', 'inc', 'llc', 'corp', 'ltd',
                                                'limited', 'bank', 'trust', 'capital', 'securities',
                                                'financial', 'services', 'group', 'holdings',
                                                'name of registrant', 'name of issuer', 'name of company',
                                                'registrant\'s name', 'issuer\'s name', 'company\'s name',
                                                'exact name', 'entity name', 'jpmorgan', 'worldpay',
                                                'vantiv', 'morgan stanley', 'credit suisse', 'bank of',
                                                'tokyo-mitsubishi', 'mitsubishi ufj', 'ag', 'sa', 'gmbh'
                                            ]
                                            
                                            has_negative = any(ind in context for ind in negative_indicators)
                                            has_positive = any(ind in context for ind in positive_indicators) or \
                                                         any(ind in org_name.lower() for ind in ['bank', 'trust', 'capital', 'securities', 'financial', 'services', 'group', 'holdings'])
                                            
                                            # Accepter si:
                                            # 1. Pas d'indicateur négatif ET (indicateur positif OU ressemble à une entreprise)
                                            # 2. OU le nom contient un suffixe d'entreprise
                                            has_suffix = bool(re.search(company_suffixes, org_name, re.IGNORECASE))
                                            
                                            if (not has_negative and (has_positive or has_suffix)) or has_suffix:
                                                # Vérifier que ce n'est pas un faux positif spécifique
                                                org_lower = org_name.lower()
                                                org_upper = org_name.upper()
                                                
                                                excluded_phrases_spacy = [
                                                    'amended and restated', 'backstop commitment', 'backstop fee',
                                                    'bridge commitment', 'bridge documents', 'bridge lenders',
                                                    'business combination', 'combined company', 'commission file number',
                                                    'dealing disclosures', 'disclosure table', 'executive',
                                                    'financial condition', 'incremental amendment', 'irs employer',
                                                    'loan agreement', 'london stock', 'offer period', 'original loan',
                                                    'press release', 'rule 8.3', 'scheme', 'securities exchange',
                                                    'the business', 'the disclosure', 'the incremental', 'the london',
                                                    'the offer', 'the original', 'the securities', 'the u.s',
                                                    'the us', 'the u.s. securities', 'vantiv\'s 2017', 'worldpay 10.1',
                                                    'the backstop', 'the bridge', 'the business combination',
                                                    'the "business', 'the "original', 'operations', 'panel',
                                                    'registrant', 'executive', 'financial condition'
                                                ]
                                                
                                                if not any(phrase in org_lower for phrase in excluded_phrases_spacy):
                                                    # Vérifier que le nom commence par une majuscule
                                                    if org_name[0].isupper():
                                                        company_names_found.add(org_name)
                            except Exception as e:
                                if verbose:
                                    logger.debug(f"spaCy NER failed on page {page_num + 1}: {e}")
                        
                        # Méthode 1: Contexte "registrant" ou "company name" (amélioré)
                        context_patterns = [
                            re.compile(
                            r'(?:Exact\s+name\s+of\s+registrant|'
                                r'Name\s+of\s+(?:the\s+)?(?:registrant|issuer|company|corporation)|'
                                r'(?:Registrant|Issuer|Company|Corporation)(?:\'s)?\s+name|'
                                r'Company\s+Name|'
                                r'Entity\s+Name)'
                                r'[:\s]+([A-Z][A-Za-z0-9\s&,\.\-]+?(?:,\s*)?' + company_suffixes + r')',
                            re.IGNORECASE
                            ),
                            re.compile(
                                r'(?:Registrant|Issuer|Company)[:\s]+([A-Z][A-Za-z0-9\s&,\.\-]+?(?:,\s*)?' + company_suffixes + r')',
                                re.IGNORECASE
                            ),
                        ]
                        
                        for pattern in context_patterns:
                            for match in pattern.finditer(text):
                                company_name = match.group(1).strip()
                                company_name = re.sub(r'\s+', ' ', company_name)

                                if 5 <= len(company_name) <= 80:
                                    if company_name.count(' ') <= 8:  # Max 9 mots
                                        company_names_found.add(company_name)
                        
                        # Méthode 2: Format standalone "XXX, Inc." ou "XXX AG" (AMÉLIORÉ)
                        standalone_patterns = [
                            re.compile(
                                r'(?<=[\.\n\(\s])([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){0,6}),?\s+' + company_suffixes + r'(?=[\.\n\)\s,])',
                            re.MULTILINE
                            ),
                            re.compile(
                                r'\b([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){0,6})\s+' + company_suffixes + r'\b',
                                re.MULTILINE
                            ),
                            # Pattern spécial pour "Credit Suisse Securities AG" (suffixe AG séparé)
                            re.compile(
                                r'\b([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){1,4})\s+(AG|S\.A\.|SA|GmbH|BV|NV|SpA|S\.p\.A\.)\b',
                                re.MULTILINE
                            ),
                        ]
                        
                        for pattern_idx, pattern in enumerate(standalone_patterns):
                            for match in pattern.finditer(text):
                                # Construire le nom selon le pattern
                                if pattern_idx == 2:  # Pattern avec AG/SA/etc.
                                    company_name = f"{match.group(1)} {match.group(2)}"
                                else:
                                    company_name = match.group(0).strip()

                                company_name = re.sub(r'\s+', ' ', company_name)

                                # Vérifier qu'il n'y a pas de verbes ou mots suspects autour
                                ctx_before = text[max(0, match.start() - 100):match.start()].lower()
                                ctx_after = text[match.end():min(len(text), match.end() + 100)].lower()
                                context_full = ctx_before + ' ' + ctx_after

                                # Liste étendue de mots suspects
                                suspect_words = [
                                    'provides', 'entered', 'completed', 'acquired', 'merged',
                                    'income', 'taxes', 'financial', 'statements', 'segment',
                                    'the company', 'our company', 'such company', 'a company',
                                    'any company', 'each company', 'this company',
                                    'agreement', 'letter', 'release', 'document', 'amendment',
                                    'commitment', 'combination', 'condition', 'operation',
                                    'disclosure', 'scheme', 'rule', 'act', 'commission',
                                    'backstop', 'bridge', 'loan', 'credit', 'business',
                                    'amended', 'restated', 'original', 'incremental'
                                ]

                                # Vérifier aussi que ce n'est pas dans une phrase générique
                                if not any(word in context_full for word in suspect_words):
                                    # Vérifier que le nom commence bien par une majuscule et contient des lettres
                                    if company_name[0].isupper() and any(c.isalpha() for c in company_name):
                                        # Vérifier que le premier mot n'est pas un terme invalide
                                        first_word = company_name.split()[0].lower().rstrip('.,;:')
                                        if first_word not in INVALID_COMPANY_FIRST_WORDS:
                                            if 5 <= len(company_name) <= 80:
                                                company_names_found.add(company_name)
                        
                        # Méthode 3: Noms de compagnies dans les en-têtes ou sections spéciales
                        header_patterns = [
                            re.compile(
                                r'(?:^|\n)\s*([A-Z][A-Za-z0-9\s&,\.\-]+?(?:,\s*)?' + company_suffixes + r')\s*(?:\n|$)',
                                re.MULTILINE
                            ),
                        ]
                        
                        for pattern in header_patterns:
                            for match in pattern.finditer(text):
                                company_name = match.group(1).strip()
                                company_name = re.sub(r'\s+', ' ', company_name)
                                
                                # Vérifier que c'est bien en début de ligne ou après un saut de ligne
                                line_start = text.rfind('\n', 0, match.start())
                                if line_start == -1:
                                    line_start = 0
                                line_text = text[line_start:match.end()].strip()
                                
                                # Si la ligne est courte (probablement un en-tête), c'est probablement un nom de compagnie
                                if len(line_text) < 100 and 5 <= len(company_name) <= 80:
                                    company_names_found.add(company_name)
                        
                        # Méthode 4: Détecter les noms de compagnies SANS suffixe (comme JPMorgan, WorldPay, Credit Suisse)
                        # Pattern pour détecter les noms propres composés qui ressemblent à des entreprises
                        # Format: "NomPropre NomPropre" ou "NomPropre NomPropre Bank/Trust/etc." ou "The NomPropre..."
                        company_name_patterns_no_suffix = [
                            # Pattern 1: Noms composés avec mots-clés d'entreprise (Bank, Trust, Capital, Securities, etc.)
                            re.compile(
                                r'\b([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){1,4})\s+(Bank|Trust|Capital|Securities|Financial|Services|Group|Holdings|Partners)\b',
                                re.MULTILINE
                            ),
                            # Pattern 2: "The" + nom d'entreprise (The Bank of..., The Company...)
                            re.compile(
                                r'\bThe\s+([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){1,5})\s+(?:of\s+)?([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){0,4})?\s*(?:Bank|Trust|Capital|Securities|Financial|Services|Group|Holdings|Company|Corporation)?\b',
                                re.MULTILINE
                            ),
                            # Pattern 3: Noms composés connus (Credit Suisse, Morgan Stanley, etc.)
                            re.compile(
                                r'\b([A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9]+){1,3})\b'
                            ),
                        ]
                        
                        # Acronymes d'entreprises connus (détection intelligente)
                        # Pattern pour détecter les acronymes en majuscules (3-6 lettres) dans des contextes d'entreprise
                        acronym_pattern = re.compile(
                            r'\b([A-Z]{3,6})\b'
                        )
                        
                        # Mots-clés qui suggèrent qu'un acronyme est une entreprise
                        company_acronym_indicators = [
                            'bank', 'trust', 'capital', 'securities', 'financial', 'services',
                            'group', 'holdings', 'funding', 'lender', 'creditor', 'underwriter',
                            'company', 'corporation', 'inc', 'llc', 'corp', 'ltd'
                        ]
                        
                        for pattern_idx, pattern in enumerate(company_name_patterns_no_suffix):
                            for match in pattern.finditer(text):
                                # Construire le nom de l'entreprise selon le pattern
                                if pattern_idx == 1:  # Pattern "The ..."
                                    if match.group(2):  # Si "of" est présent
                                        potential_company = f"The {match.group(1)} of {match.group(2)}"
                                    else:
                                        potential_company = f"The {match.group(1)}"
                                else:
                                    potential_company = match.group(0).strip()
                                
                                potential_company = re.sub(r'\s+', ' ', potential_company)
                                
                                # Vérifier le contexte
                                match_start = match.start()
                                match_end = match.end()
                                context_start = max(0, match_start - 250)
                                context_end = min(len(text), match_end + 250)
                                context = text[context_start:context_end].lower()
                                
                                # Indicateurs positifs (suggèrent que c'est une entreprise)
                                positive_indicators = [
                                    'company', 'corporation', 'bank', 'trust', 'capital', 'securities',
                                    'financial', 'services', 'group', 'holdings', 'partners',
                                    'name of registrant', 'name of issuer', 'name of company',
                                    'registrant\'s name', 'issuer\'s name', 'company\'s name',
                                    'exact name', 'entity name', 'underwriter', 'lender',
                                    'creditor', 'debtor', 'party', 'counterparty', 'funding',
                                    'credit suisse', 'morgan stanley', 'jpmorgan', 'worldpay',
                                    'vantiv', 'tokyo-mitsubishi', 'mitsubishi ufj'
                                ]
                                
                                # Indicateurs négatifs (exclure)
                                negative_indicators = [
                                    'agreement', 'letter', 'release', 'document', 'amendment',
                                    'commitment', 'combination', 'condition', 'operation',
                                    'disclosure', 'scheme', 'rule', 'act', 'commission',
                                    'file number', 'employer identification', 'delaware',
                                    'london stock exchange', 'securities exchange act',
                                    'press release', 'annual meeting', 'stockholder',
                                    'the business combination', 'the original loan',
                                    'amended and restated', 'backstop commitment letter',
                                    'bridge commitment letter', 'loan agreement',
                                    'financial condition', 'operations', 'panel',
                                    'registrant', 'executive', 'the disclosure table'
                                ]
                                
                                has_positive = any(ind in context for ind in positive_indicators)
                                has_negative = any(ind in context for ind in negative_indicators)
                                
                                # Validation: doit avoir au moins 1 mot (ou 2 pour certains patterns)
                                words = potential_company.split()
                                
                                # Noms d'entreprises connus (même avec 1-2 mots)
                                known_companies = ['credit suisse', 'morgan stanley', 'jpmorgan', 'worldpay',
                                                  'vantiv', 'tokyo-mitsubishi', 'mitsubishi', 'ufj']
                                
                                is_known_company = any(known in potential_company.lower() for known in known_companies)
                                
                                if len(words) >= 1 and len(words) <= 8:
                                    # Vérifier que tous les mots commencent par majuscule (nom propre) ou sont "The", "of"
                                    valid_words = [w for w in words if w.lower() not in ['the', 'of', 'and', 'or']]
                                    if all(w[0].isupper() for w in valid_words if len(w) > 1):
                                        # Vérifier que ce n'est pas déjà détecté
                                        if potential_company not in company_names_found:
                                            # Accepter si:
                                            # 1. Pattern 1 (avec Bank/Trust/etc.) - toujours accepter
                                            # 2. Pattern 2 (The...) - accepter si contexte positif
                                            # 3. Pattern 3 - accepter si connu OU contexte positif
                                            # 4. Noms connus - toujours accepter
                                            if 'Bank' in potential_company or 'Trust' in potential_company or \
                                               'Capital' in potential_company or 'Securities' in potential_company or \
                                               'Financial' in potential_company or 'Services' in potential_company or \
                                               'Group' in potential_company or 'Holdings' in potential_company or \
                                               'Partners' in potential_company:
                                                # Toujours accepter les noms avec ces mots-clés
                                                if not has_negative:
                                                    company_names_found.add(potential_company)
                                            elif pattern_idx == 1:  # Pattern "The ..."
                                                # Pour "The Bank of...", accepter si contexte positif
                                                if has_positive and not has_negative:
                                                    company_names_found.add(potential_company)
                                            elif is_known_company:
                                                # Toujours accepter les entreprises connues
                                                if not has_negative:
                                                    company_names_found.add(potential_company)
                                            elif has_positive and not has_negative:
                                                # Pour les autres, être plus strict
                                                # Vérifier que le nom ne contient pas de mots interdits
                                                words_lower = [w.lower() for w in words]
                                                invalid_words = ['a', 'an', 'and', 'or', 'in', 'to', 'for', 'with', 'by']
                                                if not any(w in invalid_words for w in words_lower):
                                                    # Vérifier que ce n'est pas un faux positif spécifique
                                                    potential_lower = potential_company.lower()
                                                    excluded = ['combined company', 'business combination', 'executive', 'operations', 'panel', 'registrant']
                                                    if not any(ex in potential_lower for ex in excluded):
                                                        company_names_found.add(potential_company)
                        
                        # Méthode 5: Détecter les acronymes d'entreprises (MSSF, etc.)
                        for match in acronym_pattern.finditer(text):
                            acronym = match.group(1)
                            
                            # Vérifier le contexte
                            match_start = match.start()
                            match_end = match.end()
                            context_start = max(0, match_start - 150)
                            context_end = min(len(text), match_end + 150)
                            context = text[context_start:context_end].lower()
                            
                            # Vérifier si le contexte suggère que c'est une entreprise
                            has_company_indicator = any(ind in context for ind in company_acronym_indicators)
                            
                            # Vérifier si l'acronyme est suivi ou précédé d'un nom d'entreprise connu
                            context_upper = text[context_start:context_end].upper()
                            known_company_acronyms = ['MSSF', 'JPM', 'CS', 'MS', 'BOT', 'UFJ']
                            
                            if has_company_indicator or acronym in known_company_acronyms:
                                # Vérifier que ce n'est pas dans un contexte négatif
                                negative_context = [
                                    'agreement', 'letter', 'release', 'document', 'amendment',
                                    'file number', 'employer identification', 'rule', 'act'
                                ]
                                if not any(neg in context for neg in negative_context):
                                    # Construire le nom complet si possible
                                    # Chercher le nom complet avant ou après l'acronyme
                                    before_text = text[max(0, match_start - 100):match_start].strip()
                                    after_text = text[match_end:min(len(text), match_end + 100)].strip()
                                    
                                    # Si on trouve un nom d'entreprise connu, utiliser celui-ci
                                    # Sinon, utiliser l'acronyme seul
                                    full_name = None
                                    for known in known_companies:
                                        if known in before_text.lower() or known in after_text.lower():
                                            # Extraire le nom complet
                                            name_match = re.search(r'\b([A-Z][A-Za-z\s-]+' + known.replace('-', r'[\s-]') + r'[A-Za-z\s-]*)\b', 
                                                                   before_text + ' ' + after_text, re.IGNORECASE)
                                            if name_match:
                                                full_name = name_match.group(1).strip()
                                                break
                                    
                                    if full_name:
                                        company_names_found.add(full_name)
                                    else:
                                        # Pour les acronymes connus, les ajouter tels quels
                                        if acronym in known_company_acronyms:
                                            company_names_found.add(acronym)
                        
                        # Nettoyer et valider les noms de compagnies avec filtres ULTRA stricts
                        cleaned_companies = set()
                        
                        # Liste massive de termes à exclure (titres de documents, termes génériques, etc.)
                        excluded_company_terms = {
                            # Articles et verbes
                            'is', 'an', 'a', 'the', 'and', 'or', 'of', 'in', 'on', 'at', 'to', 'for', 'with', 'by',
                            'from', 'as', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does',
                            'did', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'shall', 'will',
                            # Titres de documents et termes légaux
                            'agreement', 'letter', 'release', 'document', 'disclosure', 'disclosures', 'scheme',
                            'amendment', 'amendments', 'commitment', 'commitments', 'combination', 'combinations',
                            'condition', 'conditions', 'operation', 'operations', 'executive', 'executives',
                            'registrant', 'registrants', 'issuer', 'issuers', 'panel', 'panels', 'rule', 'rules',
                            'act', 'acts', 'commission', 'commissions', 'exchange', 'exchanges', 'securities',
                            'stockholder', 'stockholders', 'meeting', 'meetings', 'annual', 'stock',
                            # Termes génériques
                            'combined', 'original', 'amended', 'restated', 'backstop', 'bridge', 'loan', 'credit',
                            'business', 'financial', 'disclosure', 'dealing', 'offer', 'period', 'table',
                            'press', 'fee', 'lenders', 'documents', 'lender', 'incremental',
                            # Termes spécifiques à exclure
                            'emerging', 'growth', 'co', 'file', 'number', 'employer', 'identification',
                            'no', 'delaware', 'london', 'stock', 'u.s', 'us', 'usa'
                        }
                        
                        # Phrases complètes à exclure (faux positifs communs)
                        excluded_phrases = [
                            'amended and restated', 'backstop commitment', 'backstop fee', 'backstop credit',
                            'bridge commitment', 'bridge documents', 'bridge lenders', 'business combination',
                            'combined company', 'commission file', 'dealing disclosures', 'disclosure table',
                            'executive', 'financial condition', 'incremental amendment', 'irs employer',
                            'loan agreement', 'london stock', 'offer period', 'original loan',
                            'press release', 'rule 8.3', 'scheme', 'securities exchange', 'the business',
                            'the disclosure', 'the incremental', 'the london', 'the offer', 'the original',
                            'the securities', 'the u.s', 'the us', 'the u.s. securities', 'vantiv\'s 2017',
                            'worldpay 10.1', 'the backstop', 'the bridge', 'the business combination',
                            'the "business', 'the "original',
                            # Comités et fonctions
                            'committee going', 'committee how', 'committee separately', 'committee the board',
                            'committeegroup', 'committeeremuneration', 'committee the', 'committees',
                            'committees (audit', 'committees link', 'community', 'community diversity',
                            'company secretary', 'company secretary derek', 'company secretary joined',
                            'compliance functions', 'group risk committee', 'remuneration committee',
                            'nomination committee', 'audit committee'
                            # Termes financiers/comptables
                            'accounting adjustment', 'accounting standards', 'accounts receivable',
                            'accumulated common stock', 'amount capital', 'amount shares', 'average price',
                            'based compensation', 'capital stock', 'common stock', 'comprehensive income',
                            'consolidated statement', 'consolidated statements', 'consolidation',
                            'contents bank', 'contents item', 'contents merchant', 'contents net',
                            'contents part', 'contents sales', 'contents segment', 'contents the',
                            'contents tra', 'contents we', 'customer incentives', 'earnings income',
                            'equity incentive plan', 'equity shares', 'income tax', 'income taxes',
                            'interest coverage', 'interest rate', 'internal control', 'internal revenue',
                            'investor relations', 'management\'s report', 'marketing sales',
                            'merchant services total', 'merchant services year', 'merchant servicesfinancial',
                            'net income per share', 'non-operating income', 'oci component',
                            'office services', 'parent company', 'performance share', 'period leverage',
                            'principal agent', 'pro forma adjustments', 'public company accounting',
                            'restricted stock', 'retained comprehensive', 'secondary offering',
                            'shares amount', 'significant accounting', 'subsidiaries report',
                            'tax receivable', 'unit incentive plan', 'visa debit processing',
                            # Sections de documents
                            'consolidated statement of equity', 'consolidated statements of comprehensive',
                            'consolidated statements of financial', 'consolidated statements of income',
                            'contingencies', 'mine safety disclosures', 'net income per share'
                        ]
                        
                        # NOTE: invalid_first_words est maintenant défini plus tôt (ligne ~5165) pour être utilisé partout
                        
                        # Mots qui ne peuvent PAS être dans un nom de compagnie (sauf si c'est vraiment une compagnie)
                        invalid_company_words = {
                            'agreement', 'letter', 'release', 'document', 'disclosure', 'scheme', 'amendment',
                            'commitment', 'combination', 'condition', 'operation', 'executive', 'registrant',
                            'issuer', 'panel', 'rule', 'act', 'commission', 'exchange', 'securities', 'stockholder',
                            'meeting', 'annual', 'stock', 'combined', 'original', 'amended', 'restated', 'backstop',
                            'bridge', 'loan', 'credit', 'business', 'financial', 'disclosure', 'dealing', 'offer',
                            'period', 'table', 'press', 'fee', 'lenders', 'documents', 'lender', 'incremental',
                            'file', 'number', 'employer', 'identification', 'no', 'delaware', 'london', 'u.s', 'us'
                        }
                        
                        for company in company_names_found:
                            # Nettoyer les espaces multiples
                            company = re.sub(r'\s+', ' ', company).strip()
                            
                            # Validation: doit avoir au moins 3 caractères et un suffixe valide
                            if len(company) >= 3 and len(company) <= 80:
                                # Vérifier qu'il contient un suffixe d'entreprise
                                if re.search(company_suffixes, company, re.IGNORECASE):
                                    # FILTRE ULTRA STRICT: Validation en plusieurs étapes
                                    words = company.split()
                                    words_lower = [w.lower().rstrip('.,;:') for w in words]
                                    company_lower = company.lower()
                                    company_upper = company.upper()
                                    
                                    # ÉTAPE 1: Exclure si contient des phrases interdites
                                    if any(phrase in company_lower for phrase in excluded_phrases):
                                        continue
                                    
                                    # ÉTAPE 1.5: Exclure les sections de documents
                                    document_section_indicators = [
                                        'consolidated statement', 'consolidated statements', 'contents',
                                        'mine safety', 'management\'s report', 'subsidiaries report',
                                        'independent registered', 'public accounting', 'accounting firm',
                                        # Comités et fonctions
                                        'committee going', 'committee how', 'committee separately',
                                        'committee the board', 'committee the', 'committees (audit',
                                        'committees link', 'community diversity', 'company secretary',
                                        'compliance functions', 'group risk committee', 'remuneration committee',
                                        'nomination committee', 'audit committee'
                                    ]
                                    if any(ind in company_lower for ind in document_section_indicators):
                                        continue
                                    
                                    # ÉTAPE 1.5.5: Exclure les termes commençant par "Committee" (sauf si vraiment une compagnie)
                                    if company_lower.startswith('committee') and not any(suffix in company_lower for suffix in ['inc', 'llc', 'corp', 'ltd', 'limited', 'company']):
                                        continue
                                    
                                    # ÉTAPE 1.5.6: Exclure "Community" seul ou avec descriptions
                                    if company_lower.strip() == 'community' or company_lower.startswith('community '):
                                        continue
                                    
                                    # ÉTAPE 1.5.7: Exclure "Company Secretary" et variations
                                    if 'company secretary' in company_lower:
                                        continue
                                    
                                    # ÉTAPE 1.5.8: Exclure "Compliance Functions"
                                    if 'compliance functions' in company_lower:
                                        continue
                                    
                                    # ÉTAPE 1.6: Exclure les phrases mal formées (mots collés, répétitions)
                                    has_collapsed_words = False
                                    for word in words:
                                        if len(word) > 10:
                                            mid_caps = sum(1 for i, c in enumerate(word[1:], 1) if c.isupper() and word[i-1].isalpha())
                                            if mid_caps > 1:  # Plus d'une majuscule au milieu = mots collés
                                                has_collapsed_words = True
                                                break
                                    if has_collapsed_words:
                                        continue
                                    
                                    # ÉTAPE 1.6.5: Exclure les patterns avec "Committee" collé (ex: "CommitteeGroup", "CommitteeRemuneration")
                                    if re.search(r'committeegroup|committeeremuneration|committeethe', company_lower):
                                        continue
                                    
                                    # Vérifier les répétitions de mots
                                    if len(words_lower) != len(set(words_lower)):  # Mots dupliqués
                                        continue
                                    
                                    # ÉTAPE 1.7: Exclure les patterns suspects
                                    suspicious_patterns = [
                                        r'sharesamount', r'amountshares', r'amount\s+shares\s+amount',
                                        r'equity\s+shares\s+amount', r'shares\s+amount\s+shares',
                                        r'amount\s+capital\s+earnings', r'earnings\s+income',
                                        r'comprehensive\s+controlling', r'retained\s+comprehensive'
                                    ]
                                    if any(re.search(pattern, company_lower) for pattern in suspicious_patterns):
                                        continue
                                    
                                    # ÉTAPE 1.8: Exclure les termes génériques seuls
                                    generic_terms = ['company', 'controller', 'state', 'the', 'principal', 'public company']
                                    if company_lower.strip() in generic_terms:
                                        continue
                                    
                                    # ÉTAPE 2: Exclure si le premier mot est invalide
                                    if len(words_lower) > 0:
                                        first_word = words_lower[0]
                                        if first_word in invalid_first_words:
                                            continue
                                    
                                    # ÉTAPE 3: Exclure si contient trop de mots interdits
                                    invalid_word_count = sum(1 for w in words_lower if w in invalid_company_words)
                                    if invalid_word_count > 0 and len(words_lower) <= 3:
                                        # Si c'est un nom court avec des mots interdits, c'est probablement un faux positif
                                        continue
                                    
                                    # ÉTAPE 4: Exclure les faux positifs spécifiques
                                    false_positives = [
                                        'SECURITIES AND EXCHANGE COMMISSION', 'COMMISSION FILE',
                                        'UNITED STATES', 'NEW YORK', 'LOS ANGELES',
                                        'IS AN EMERGING GROWTH CO', 'WILL CO', 'THE CO',
                                        'AN EMERGING', 'GROWTH CO', 'IS AN',
                                        'AMENDED AND RESTATED', 'BACKSTOP COMMITMENT', 'BACKSTOP FEE',
                                        'BRIDGE COMMITMENT', 'BRIDGE DOCUMENTS', 'BRIDGE LENDERS',
                                        'BUSINESS COMBINATION', 'COMBINED COMPANY', 'COMMISSION FILE NUMBER',
                                        'DEALING DISCLOSURES', 'DISCLOSURE TABLE', 'EXECUTIVE',
                                        'FINANCIAL CONDITION', 'INCREMENTAL AMENDMENT', 'IRS EMPLOYER',
                                        'LOAN AGREEMENT', 'LONDON STOCK', 'OFFER PERIOD', 'ORIGINAL LOAN',
                                        'PRESS RELEASE', 'RULE 8.3', 'SCHEME', 'SECURITIES EXCHANGE',
                                        'THE BUSINESS', 'THE DISCLOSURE', 'THE INCREMENTAL', 'THE LONDON',
                                        'THE OFFER', 'THE ORIGINAL', 'THE SECURITIES', 'THE U.S',
                                        'THE US', 'THE U.S. SECURITIES', 'VANTIV\'S 2017', 'WORLDPAY 10.1',
                                        'THE BACKSTOP', 'THE BRIDGE', 'THE BUSINESS COMBINATION',
                                        'THE "BUSINESS', 'THE "ORIGINAL', 'OPERATIONS', 'PANEL',
                                        'REGISTRANT', 'EXECUTIVE', 'FINANCIAL CONDITION',
                                        # Ajouter les faux positifs de comités et fonctions
                                        'COMMITTEE GOING', 'COMMITTEE HOW', 'COMMITTEE SEPARATELY',
                                        'COMMITTEE THE BOARD', 'COMMITTEEGROUP', 'COMMITTEEREMUNERATION',
                                        'COMMITTEE THE', 'COMMITTEES', 'COMMITTEES (AUDIT', 'COMMITTEES LINK',
                                        'COMMUNITY', 'COMMUNITY DIVERSITY', 'COMPANY SECRETARY',
                                        'COMPANY SECRETARY DEREK', 'COMPANY SECRETARY JOINED',
                                        'COMPLIANCE FUNCTIONS', 'GROUP RISK COMMITTEE', 'REMUNERATION COMMITTEE',
                                        'NOMINATION COMMITTEE', 'AUDIT COMMITTEE',
                                        # Ajouter les faux positifs financiers
                                        'ACCOUNTING ADJUSTMENT', 'ACCOUNTING STANDARDS', 'ACCOUNTS RECEIVABLE',
                                        'AMOUNT CAPITAL', 'AMOUNT SHARES', 'AVERAGE PRICE', 'BASED COMPENSATION',
                                        'CAPITAL STOCK', 'COMMON STOCK', 'COMPREHENSIVE INCOME', 'CONSOLIDATED STATEMENT',
                                        'CONSOLIDATED STATEMENTS', 'CONSOLIDATION', 'CONTENTS BANK', 'CONTENTS ITEM',
                                        'CONTENTS MERCHANT', 'CONTENTS NET', 'CONTENTS PART', 'CONTENTS SALES',
                                        'CONTENTS SEGMENT', 'CONTENTS THE', 'CONTENTS TRA', 'CONTENTS WE',
                                        'CUSTOMER INCENTIVES', 'EARNINGS INCOME', 'EQUITY INCENTIVE PLAN',
                                        'EQUITY SHARES', 'INCOME TAX', 'INCOME TAXES', 'INTEREST COVERAGE',
                                        'INTEREST RATE', 'INTERNAL CONTROL', 'INTERNAL REVENUE', 'INVESTOR RELATIONS',
                                        'MANAGEMENT\'S REPORT', 'MARKETING SALES', 'MERCHANT SERVICES TOTAL',
                                        'MERCHANT SERVICES YEAR', 'MERCHANT SERVICESFINANCIAL', 'NET INCOME PER SHARE',
                                        'NON-OPERATING INCOME', 'OCI COMPONENT', 'OFFICE SERVICES', 'PARENT COMPANY',
                                        'PERFORMANCE SHARE', 'PERIOD LEVERAGE', 'PRINCIPAL AGENT', 'PRO FORMA ADJUSTMENTS',
                                        'PUBLIC COMPANY ACCOUNTING', 'RESTRICTED STOCK', 'RETAINED COMPREHENSIVE',
                                        'SECONDARY OFFERING', 'SHARES AMOUNT', 'SIGNIFICANT ACCOUNTING',
                                        'SUBSIDIARIES REPORT', 'TAX RECEIVABLE', 'UNIT INCENTIVE PLAN',
                                        'VISA DEBIT PROCESSING', 'MINE SAFETY DISCLOSURES', 'CONTROLLER', 'STATE',
                                        'COMPANY', 'PRINCIPAL'
                                    ]
                                    if any(fp in company_upper for fp in false_positives):
                                        continue
                                    
                                    # ÉTAPE 5: Vérifier que le nom ne commence pas par "the" suivi d'un terme générique
                                    # MAIS autoriser "The Bank of...", "The Company...", etc.
                                    if len(words_lower) >= 2 and words_lower[0] == 'the':
                                        second_word = words_lower[1]
                                        # Termes génériques à exclure
                                        invalid_after_the = ['business', 'combined', 'disclosure', 'incremental',
                                                             'london', 'offer', 'original', 'securities', 'u.s', 'us',
                                                             'backstop', 'bridge', 'backstop commitment', 'bridge commitment',
                                                             'executive', 'operations', 'panel', 'registrant']
                                        # Termes valides après "The" (noms d'entreprises)
                                        valid_after_the = ['bank', 'company', 'corporation', 'trust', 'capital',
                                                          'securities', 'financial', 'services', 'group', 'holdings']
                                        
                                        if second_word in invalid_after_the:
                                            continue
                                        # Si c'est un terme valide (Bank, Company, etc.), continuer la validation
                                    
                                    # ÉTAPE 6: Extraire les mots avant le suffixe et valider
                                    words_before_suffix = []
                                    for word in words:
                                        word_lower = word.lower().rstrip('.,;:')
                                        if word_lower not in ['inc', 'llc', 'llp', 'corp', 'corporation', 'ltd', 'limited', 
                                                              'co', 'company', 'companies', 'group', 'holdings', 'enterprises',
                                                              'partners', 'partnership', 'incorporated']:
                                            words_before_suffix.append(word)
                                    
                                    if len(words_before_suffix) >= 1:
                                        # Vérifier que le premier mot réel n'est pas invalide
                                        first_real_word = words_before_suffix[0].lower().rstrip('.,;:')
                                        
                                        # Exclure si le premier mot est un terme invalide
                                        if first_real_word in invalid_first_words:
                                            continue
                                        
                                        # Exclure si le premier mot est un article/verbe suivi d'un terme générique
                                        if first_real_word in ['the', 'a', 'an'] and len(words_before_suffix) > 1:
                                            second_real_word = words_before_suffix[1].lower().rstrip('.,;:')
                                            if second_real_word in invalid_company_words:
                                                continue
                                        
                                        # ÉTAPE 7: Validation contextuelle - vérifier le contexte dans le texte
                                        company_pos = text.find(company)
                                        if company_pos != -1:
                                            context_start = max(0, company_pos - 200)
                                            context_end = min(len(text), company_pos + len(company) + 200)
                                            context = text[context_start:context_end].lower()
                                            
                                            # Indicateurs négatifs (suggèrent que ce n'est PAS un nom de compagnie)
                                            negative_context_indicators = [
                                                'agreement', 'letter', 'release', 'document', 'amendment',
                                                'commitment', 'combination', 'condition', 'operation',
                                                'disclosure', 'scheme', 'rule', 'act', 'commission',
                                                'file number', 'employer identification', 'delaware',
                                                'london stock exchange', 'securities exchange act',
                                                'press release', 'annual meeting', 'stockholder',
                                                'the business combination', 'the original loan',
                                                'the backstop commitment', 'the bridge commitment',
                                                'the disclosure table', 'the offer period',
                                                'the incremental amendment', 'the securities exchange',
                                                'amended and restated', 'backstop commitment letter',
                                                'bridge commitment letter', 'loan agreement',
                                                'financial condition', 'operations'
                                            ]
                                            
                                            # Si le contexte contient des indicateurs négatifs, exclure
                                            if any(ind in context for ind in negative_context_indicators):
                                                # Sauf si c'est clairement dans un contexte de nom de compagnie
                                                positive_context_indicators = [
                                                    'name of registrant', 'name of issuer', 'name of company',
                                                    'company name', 'registrant name', 'issuer name',
                                                    'exact name', 'entity name'
                                                ]
                                                if not any(pos_ind in context for pos_ind in positive_context_indicators):
                                                    continue
                                        
                                        # ÉTAPE 8: Validation finale - le nom doit ressembler à un vrai nom de compagnie
                                        # Un vrai nom de compagnie devrait avoir au moins un mot substantif (nom propre)
                                        # qui n'est pas dans la liste des mots interdits
                                        valid_substantive_words = [w for w in words_before_suffix 
                                                                  if w.lower().rstrip('.,;:') not in excluded_company_terms 
                                                                  and len(w) >= 2]
                                        
                                        if len(valid_substantive_words) >= 1:
                                            # Vérifier que le premier mot substantif commence par une majuscule
                                            if valid_substantive_words[0][0].isupper():
                                                cleaned_companies.add(company)
                        
                        # Validation finale avec spaCy NER pour company names (si disponible)
                        if SPACY_AVAILABLE and SPACY_NLP and cleaned_companies:
                            validated_companies = set()
                            
                            for company in cleaned_companies:
                                # Chercher le nom dans le texte original avec contexte
                                company_pos = text.find(company)
                                if company_pos != -1:
                                    context_start = max(0, company_pos - 200)
                                    context_end = min(len(text), company_pos + len(company) + 200)
                                    context_text = text[context_start:context_end]
                                    context_lower = context_text.lower()
                                    
                                    # Vérifier d'abord le contexte pour éviter les faux positifs
                                    negative_context = [
                                        'agreement', 'letter', 'release', 'document', 'amendment',
                                        'commitment', 'combination', 'condition', 'operation',
                                        'disclosure', 'scheme', 'rule', 'act', 'commission',
                                        'file number', 'employer identification', 'delaware',
                                        'london stock exchange', 'securities exchange act',
                                        'press release', 'annual meeting', 'stockholder',
                                        'the business combination', 'the original loan',
                                        'the backstop commitment', 'the bridge commitment'
                                    ]
                                    
                                    # Si le contexte est négatif, exclure même si spaCy le détecte
                                    if any(neg in context_lower for neg in negative_context):
                                        # Sauf si c'est clairement dans un contexte de nom de compagnie
                                        positive_context = [
                                            'name of registrant', 'name of issuer', 'name of company',
                                            'company name', 'registrant name', 'issuer name',
                                            'exact name', 'entity name'
                                        ]
                                        if not any(pos in context_lower for pos in positive_context):
                                            continue
                                    
                                    # Analyser avec spaCy
                                    doc = SPACY_NLP(context_text)
                                    
                                    # Vérifier si spaCy détecte ce nom comme une organisation
                                    found_in_spacy = False
                                    for ent in doc.ents:
                                        if ent.label_ == "ORG":
                                            detected_org = ent.text.strip()
                                            detected_org = re.sub(r'\s+', ' ', detected_org)
                                            
                                            # Vérifier si notre nom correspond
                                            if company.lower() in detected_org.lower() or detected_org.lower() in company.lower():
                                                # Validation supplémentaire: vérifier que spaCy n'a pas détecté un faux positif
                                                detected_lower = detected_org.lower()
                                                if not any(phrase in detected_lower for phrase in excluded_phrases):
                                                    if detected_org[0].isupper():  # Doit commencer par majuscule
                                                        found_in_spacy = True
                                                        validated_companies.add(company)
                                                        break
                                    
                                    # Si spaCy ne l'a pas trouvé mais que le contexte est très positif, on peut quand même l'accepter
                                    if not found_in_spacy:
                                        # Vérifier le contexte pour des indicateurs très positifs
                                        very_positive_context = [
                                            'exact name of registrant', 'name of the registrant',
                                            'name of the issuer', 'name of the company',
                                            'registrant\'s name', 'issuer\'s name', 'company\'s name'
                                        ]
                                        if any(pos in context_lower for pos in very_positive_context):
                                            # Vérifier que le nom ne contient pas de termes interdits
                                            company_lower = company.lower()
                                            company_words = company.split()
                                            company_words_lower = [w.lower().rstrip('.,;:') for w in company_words]
                                            if not any(phrase in company_lower for phrase in excluded_phrases):
                                                if len(company_words_lower) > 0 and company_words_lower[0] not in invalid_first_words:
                                                    validated_companies.add(company)
                            
                            # Utiliser les companies validées par spaCy, ou garder celles qui ont passé les filtres
                            if validated_companies:
                                # Ajouter les companies validées par spaCy
                                cleaned_companies.update(validated_companies)
                            # Note: On garde aussi les companies qui ont passé les filtres stricts précédents
                            # car elles peuvent être valides même si spaCy ne les a pas détectées
                        
                        for company in cleaned_companies:
                            page_findings.append({'type': 'company_name', 'value': company, 'page': page_num + 1})
                        
                        # ===== CIK NUMBERS (SEC Central Index Key) =====
                        # Format: 7-10 chiffres (ex: 1065280)
                        # Chercher dans un contexte spécifique pour éviter faux positifs
                        cik_context_pattern = re.compile(
                            r'(?:CIK|Central\s+Index\s+Key|File\s+Number)[:\s]+(\d{7,10})',
                            re.IGNORECASE
                        )
                        cik_found = False
                        for match in cik_context_pattern.finditer(text):
                            cik = match.group(1)
                            page_findings.append({'type': 'cik_number', 'value': cik, 'page': page_num + 1})
                            cik_found = True
                        
                        # Si pas de contexte, chercher les nombres de 7-10 chiffres près de "Commission"
                        if not cik_found:
                            cik_generic_pattern = re.compile(r'\b(\d{7,10})\b')
                            for match in cik_generic_pattern.finditer(text):
                                context = text[max(0, match.start() - 100):match.end() + 100]
                                if 'Commission' in context or 'SEC' in context or 'File' in context:
                                    cik = match.group(1)
                                    page_findings.append({'type': 'cik_number', 'value': cik, 'page': page_num + 1})
                                    break  # Prendre seulement le premier
                        
                        # ===== CUSIP/ISIN CODES =====
                        # CUSIP: 9 caractères alphanumériques (ex: 64110LAH9)
                        # ISIN: 2 lettres + 9 alphanumériques + 1 chiffre (ex: US64110LAH96)
                        cusip_pattern = re.compile(r'\b([0-9A-Z]{6,9}[A-Z0-9]{2}[0-9])\b')
                        isin_pattern = re.compile(r'\b(US[0-9A-Z]{9}[0-9])\b')
                        
                        for match in cusip_pattern.finditer(text):
                            cusip = match.group(1)
                            # Vérifier que c'est bien un code financier (contexte)
                            context = text[max(0, match.start() - 50):match.end() + 50]
                            if any(kw in context.upper() for kw in ['CUSIP', 'SECURITY', 'BOND', 'NOTE', 'DEBT']):
                                page_findings.append({'type': 'cusip_code', 'value': cusip, 'page': page_num + 1})
                        
                        for match in isin_pattern.finditer(text):
                            isin = match.group(1)
                            page_findings.append({'type': 'isin_code', 'value': isin, 'page': page_num + 1})
                        
                        # ===== CITY NAMES - SUPPRESSION COMPLÈTE =====
                        # NE PAS détecter les villes seules - trop de faux positifs
                        # Les villes sont capturées dans les adresses complètes seulement
                        
                        # ===== PARTIAL ADDRESSES - SEULEMENT si pas dans adresse complète =====
                        
                        partial_addresses_found = set()
                        
                        # Pattern strict: "Ville, État ZIP" seulement
                        city_state_zip_pattern = re.compile(
                            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2}),\s+([A-Z]{2})\s+(\d{5}(?:-\d{4})?)\b'
                        )
                        
                        for match in city_state_zip_pattern.finditer(text):
                            city = match.group(1)
                            state = match.group(2)
                            zip_code = match.group(3)
                            
                            partial = f"{city}, {state} {zip_code}"
                            
                            # Vérifier que ce n'est pas déjà dans une adresse complète
                            already_in_full = any(partial in addr for addr in addresses_found)
                            
                            if not already_in_full and state in valid_states:
                                # Vérifier qu'il n'y a pas un numéro de rue juste avant
                                ctx_before = text[max(0, match.start() - 50):match.start()]
                                has_street_number = bool(re.search(r'\b\d{1,5}\s+[A-Za-z]+\s+(Street|Avenue|Drive|Road|Lane|Boulevard)\s*$', ctx_before, re.IGNORECASE))
                                
                                if not has_street_number:
                                    partial_addresses_found.add(partial)
                        
                        for partial in partial_addresses_found:
                            page_findings.append({'type': 'partial_address', 'value': partial, 'page': page_num + 1})
                        
                        # ===== PROFESSIONAL FIRMS - VERSION ULTRA STRICTE =====
                        
                        firms_found = set()
                        
                        # Pattern: Chercher SEULEMENT les noms propres se terminant par LLP/LLC/etc.
                        # ET qui sont précédés/suivis de ponctuation ou début/fin de phrase
                        firm_pattern = re.compile(
                            r'(?:^|[.\n\(])\s*([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){0,4})[,\s]+(LLP|LLC|L\.L\.P\.|L\.L\.C\.|PLLC)\b',
                            re.MULTILINE
                        )
                        
                        for match in firm_pattern.finditer(text):
                            firm_name = f"{match.group(1)} {match.group(2)}"
                            firm_name = re.sub(r'\s+', ' ', firm_name).strip()
                            
                            # Validation stricte
                            if 10 <= len(firm_name) <= 60:
                                # Vérifier que c'est bien un nom propre (tous les mots commencent par majuscule)
                                words = firm_name.replace(',', '').split()
                                if all(w[0].isupper() for w in words if len(w) > 2):
                                    # Vérifier le contexte (50 caractères avant/après)
                                    ctx_start = max(0, match.start() - 50)
                                    ctx_end = min(len(text), match.end() + 50)
                                    context = text[ctx_start:ctx_end].upper()
                                    
                                    # Déterminer le type
                                    firm_type = 'professional_firm'
                                    
                                    if any(kw in context for kw in ['AUDIT', 'ACCOUNTANT', 'KPMG', 'ERNST', 'DELOITTE', 'PWC', 'PRICEWATERHOUSE']):
                                        firm_type = 'auditor'
                                    elif any(kw in context for kw in ['UNDERWRITER', 'UNDERWRITING', 'BOOK-RUNNING', 'LEAD MANAGER']):
                                        firm_type = 'underwriter'
                                    elif any(kw in context for kw in ['COUNSEL', 'ATTORNEY', 'LAW FIRM', 'LEGAL']):
                                        firm_type = 'law_firm'
                                    elif any(kw in context for kw in ['BANK', 'TRUST', 'TRUSTEE', 'PAYING AGENT']):
                                        firm_type = 'financial_institution'
                                    
                                    # Filtrer les faux positifs
                                    false_positives = ['Vantiv LLC', 'Vantiv Holding LLC', 'Mercury Payment Systems LLC']
                                    if firm_name not in false_positives:
                                        if firm_name not in firms_found:
                                            firms_found.add(firm_name)
                                            page_findings.append({'type': firm_type, 'value': firm_name, 'page': page_num + 1})
                        
                        # ===== SHOW/PRODUCT TITLES - DÉSACTIVÉ =====
                        # DÉSACTIVÉ - Trop de faux positifs
                        # show_pattern = re.compile(r'[""]([A-Z][A-Za-z\s]+(?:of|is|the|and)[A-Za-z\s]+)[""]')
                        # for match in show_pattern.finditer(text):
                        #     title = match.group(1).strip()
                        #     if 10 <= len(title) <= 50:
                        #         page_findings.append({'type': 'quoted_title', 'value': title, 'page': page_num + 1})
                        
                        # ===== VALIDATION FINALE AVEC SPACY NER (si disponible) =====
                        # Utiliser spaCy pour détecter des entités supplémentaires et valider celles trouvées
                        if SPACY_AVAILABLE and SPACY_NLP and len(text) > 50:
                            try:
                                # Analyser le texte complet de la page avec spaCy
                                # Limiter à 1M de caractères pour éviter les problèmes de mémoire
                                text_to_analyze = text[:1000000] if len(text) > 1000000 else text
                                doc = SPACY_NLP(text_to_analyze)
                                
                                # Collecter les entités détectées par spaCy
                                spacy_persons = set()
                                spacy_orgs = set()
                                
                                for ent in doc.ents:
                                    if ent.label_ == "PERSON":
                                        person_name = ent.text.strip()
                                        person_name = re.sub(r'\s+', ' ', person_name)
                                        # Validation basique
                                        words = person_name.split()
                                        if len(words) >= 2 and len(words) <= 5:
                                            # Vérifier que ce n'est pas déjà détecté
                                            already_found = any(
                                                person_name.lower() in f.get('value', '').lower() or 
                                                f.get('value', '').lower() in person_name.lower()
                                                for f in page_findings if f.get('type') in ['person_name']
                                            )
                                            if not already_found:
                                                # Vérifier le contexte pour éviter les faux positifs
                                                start_char = ent.start_char
                                                end_char = ent.end_char
                                                context_start = max(0, start_char - 100)
                                                context_end = min(len(text_to_analyze), end_char + 100)
                                                context = text_to_analyze[context_start:context_end].lower()
                                                
                                                # Indicateurs positifs (MANDATORY)
                                                positive_indicators = [
                                                    'director', 'officer', 'executive', 'employee',
                                                    'trustee', 'shareholder', 'by:', 'name:', 'signed',
                                                    'beneficial owner', 'signatory', 'authorized',
                                                    'representative', 'agent', '/s/'
                                                ]
                                                # Indicateurs négatifs (expanded)
                                                negative_indicators = [
                                                    'table of contents', 'balance sheet', 'income statement',
                                                    'cash flow', 'note', 'footnote', 'page', 'section',
                                                    'form 10-k', 'form 10-q', 'form 8-k',
                                                    'part i', 'part ii', 'part iii', 'part iv',
                                                    'exhibit', 'schedule', 'appendix', 'index',
                                                    'consolidated', 'financial statements', 'statement of operations'
                                                ]

                                                has_positive = any(ind in context for ind in positive_indicators)
                                                has_negative = any(ind in context for ind in negative_indicators)

                                                # Check against false positive lists (EXACT MATCH ONLY)
                                                person_name_lower = person_name.lower()
                                                is_false_positive = (
                                                    person_name_lower in PERSON_NAME_FALSE_POSITIVES or
                                                    person_name_lower in DOCUMENT_STRUCTURE_TERMS or
                                                    person_name_lower in GENERIC_BUSINESS_TERMS
                                                    # Don't check partial matches - too aggressive
                                                )

                                                # Validate capitalization (reject all-caps or excessive capitals)
                                                cap_ratio = sum(1 for c in person_name if c.isupper()) / len(person_name) if person_name else 0
                                                valid_capitalization = cap_ratio <= 0.5

                                                # Check that no word in the name is an invalid word
                                                words = person_name.split()
                                                words_lower = [w.lower().strip('.,;:') for w in words]
                                                has_invalid_word = any(w in INVALID_NAME_WORDS for w in words_lower)

                                                # RELAXED: Accept if no negative AND not false positive AND valid capitalization AND no invalid words
                                                # Positive indicator is helpful but not mandatory for spaCy-detected persons
                                                if not has_negative and not is_false_positive and valid_capitalization and not has_invalid_word:
                                                    spacy_persons.add(person_name)
                                    
                                    elif ent.label_ == "ORG":
                                        org_name = ent.text.strip()
                                        org_name = re.sub(r'\s+', ' ', org_name)
                                        # Validation basique
                                        if 5 <= len(org_name) <= 80:
                                            # Vérifier que ce n'est pas déjà détecté
                                            already_found = any(
                                                org_name.lower() in f.get('value', '').lower() or 
                                                f.get('value', '').lower() in org_name.lower()
                                                for f in page_findings if f.get('type') == 'company_name'
                                            )
                                            if not already_found:
                                                # Vérifier que c'est une organisation valide (contient des mots substantifs)
                                                words = org_name.split()
                                                if len(words) >= 1:
                                                    # Vérifier le contexte
                                                    start_char = ent.start_char
                                                    end_char = ent.end_char
                                                    context_start = max(0, start_char - 100)
                                                    context_end = min(len(text_to_analyze), end_char + 100)
                                                    context = text_to_analyze[context_start:context_end].lower()
                                                    
                                                    # Check against false positive lists FIRST (EXACT MATCH ONLY)
                                                    org_name_lower = org_name.lower()
                                                    is_false_positive = (
                                                        org_name_lower in DOCUMENT_STRUCTURE_TERMS or
                                                        org_name_lower in GENERIC_BUSINESS_TERMS
                                                        # Don't check partial matches - too aggressive
                                                    )

                                                    # Check if first word is invalid
                                                    first_word = org_name.split()[0].lower().rstrip('.,;:') if org_name.split() else ''
                                                    has_invalid_first_word = first_word in INVALID_COMPANY_FIRST_WORDS

                                                    # Only proceed if not a false positive and valid first word
                                                    if not is_false_positive and not has_invalid_first_word:
                                                        # Indicateurs négatifs (expanded)
                                                        negative_indicators = [
                                                            'table of contents', 'balance sheet', 'income statement',
                                                            'cash flow', 'note', 'footnote', 'form 10-k', 'form 10-q',
                                                            'exhibit', 'schedule', 'appendix', 'part i', 'part ii'
                                                        ]

                                                        has_negative = any(ind in context for ind in negative_indicators)

                                                        # RELAXED: Accept if no negative context
                                                        # Don't require positive indicators - just filter out bad contexts
                                                        if not has_negative:
                                                            spacy_orgs.add(org_name)
                                
                                # Ajouter les entités détectées par spaCy qui n'ont pas été trouvées par nos patterns
                                for person_name in spacy_persons:
                                    page_findings.append({'type': 'person_name', 'value': person_name, 'page': page_num + 1})
                                
                                # Filtrer les organisations détectées par spaCy avec les mêmes règles strictes
                                for org_name in spacy_orgs:
                                    org_lower = org_name.lower()
                                    org_upper = org_name.upper()

                                    # Check against comprehensive false positive lists (EXACT MATCH ONLY)
                                    is_false_positive = (
                                        org_lower in DOCUMENT_STRUCTURE_TERMS or
                                        org_lower in GENERIC_BUSINESS_TERMS
                                        # Don't check partial matches - filters out legitimate companies
                                    )

                                    if is_false_positive:
                                        continue

                                    # Additional specific false positives (highly specific terms)
                                    specific_false_positives = [
                                        'SECURITIES AND EXCHANGE COMMISSION', 'COMMISSION FILE',
                                        'UNITED STATES', 'NEW YORK', 'LOS ANGELES',
                                        'IS AN EMERGING GROWTH CO', 'WILL CO', 'THE CO',
                                        'AN EMERGING', 'GROWTH CO', 'IS AN',
                                        'AMENDED AND RESTATED', 'BACKSTOP COMMITMENT', 'BACKSTOP FEE',
                                        'BRIDGE COMMITMENT', 'BRIDGE DOCUMENTS', 'BRIDGE LENDERS',
                                        'BUSINESS COMBINATION', 'COMBINED COMPANY', 'COMMISSION FILE NUMBER',
                                        'DEALING DISCLOSURES', 'DISCLOSURE TABLE',
                                        'INCREMENTAL AMENDMENT', 'IRS EMPLOYER',
                                        'LOAN AGREEMENT', 'LONDON STOCK', 'OFFER PERIOD', 'ORIGINAL LOAN',
                                        'PRESS RELEASE', 'RULE 8.3', 'SCHEME', 'SECURITIES EXCHANGE',
                                        'THE DISCLOSURE', 'THE INCREMENTAL', 'THE LONDON',
                                        'THE OFFER', 'THE ORIGINAL', 'THE SECURITIES', 'THE U.S',
                                        'THE US', 'THE U.S. SECURITIES',
                                        'THE BACKSTOP', 'THE BRIDGE', 'THE BUSINESS COMBINATION',
                                        'THE "BUSINESS', 'THE "ORIGINAL'
                                    ]
                                    if any(fp in org_upper for fp in specific_false_positives):
                                        continue

                                    # Vérifier que le nom commence par une majuscule
                                    if not org_name or not org_name[0].isupper():
                                        continue

                                    # Vérifier que le premier mot n'est pas invalide
                                    org_words = org_name.split()
                                    if len(org_words) > 0:
                                        first_word = org_words[0].lower().rstrip('.,;:')
                                        if first_word in INVALID_COMPANY_FIRST_WORDS:
                                            continue

                                    # Vérifier qu'il contient un suffixe d'entreprise ou est dans un contexte TRÈS positif
                                    # Common company suffixes pattern
                                    company_suffixes = r'\b(Inc|LLC|Corp|Corporation|Company|Co|Ltd|Limited|LP|LLP|PLLC|PC)\b'
                                    has_suffix = bool(re.search(company_suffixes, org_name, re.IGNORECASE))

                                    if not has_suffix:
                                        # Sans suffixe, être TRÈS strict - REQUIRE explicit registrant/issuer context
                                        org_pos = text_to_analyze.find(org_name)
                                        if org_pos != -1:
                                            context_start = max(0, org_pos - 200)
                                            context_end = min(len(text_to_analyze), org_pos + len(org_name) + 200)
                                            context_lower = text_to_analyze[context_start:context_end].lower()

                                            # VERY strict - only accept if explicitly identified as registrant/issuer
                                            very_positive = [
                                                'exact name of registrant', 'name of the registrant',
                                                'name of the issuer', 'name of the company',
                                                'registrant\'s name', 'issuer\'s name',
                                                'registrant as specified', 'issuer as specified'
                                            ]
                                            if not any(pos in context_lower for pos in very_positive):
                                                continue

                                    page_findings.append({'type': 'company_name', 'value': org_name, 'page': page_num + 1})
                                    
                            except Exception as e:
                                # Si spaCy échoue, continuer sans cette validation
                                if verbose:
                                    logger.debug(f"spaCy NER failed on page {page_num + 1}: {e}")
                        
                        if page_findings:
                            sensitive_info_by_page[page_num + 1] = page_findings
                        
                        # Progress indicator (every 10 pages or last page)
                        if verbose and total_pages > 10 and ((page_num + 1) % 10 == 0 or page_num + 1 == total_pages):
                            progress = (page_num + 1) / total_pages * 100
                            print(f"\r   📄 Analyzing {total_pages} pages... {page_num + 1}/{total_pages} ({progress:.1f}%)", end='', flush=True)
                            
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        if verbose and total_pages > 10:
                            progress = (page_num + 1) / total_pages * 100
                            print(f"\r   📄 Analyzing {total_pages} pages... {page_num + 1}/{total_pages} ({progress:.1f}%) [error]", end='', flush=True)
                        continue
                
                # Complete progress indicator (seulement en mode verbose)
                if verbose and total_pages > 10:
                    print(f"\r   📄 Analyzing {total_pages} pages... {total_pages}/{total_pages} (100%) ✓", flush=True)
                        
        except Exception as e:
            logger.error(f"Error during sensitive information detection: {e}")
        
        return sensitive_info_by_page

    def analyze_and_create_report(self, report_path=None, grouping_keywords=None, verbose=False):
        """
        Analyse le PDF et crée un rapport.

        Args:
            report_path: Chemin du fichier de rapport (optionnel)
            grouping_keywords: Liste de mots-clés pour regrouper les occurrences
            verbose: Si True, affiche les logs détaillés. Si False, mode simple.
        """
        if verbose:
            logger.info("Starting PDF analysis...")
        
        # Détecter les images
        self.images_by_page = self._detect_images_in_pdf()
        if verbose:
            logger.info(f"Images detected on {len(self.images_by_page)} page(s)")
        
        # Détecter les informations sensibles
        self.sensitive_info_by_page = self._detect_sensitive_information(verbose=verbose)
        
        total_sensitive_pages = len(self.sensitive_info_by_page)
        logger.info(f"Sensitive information detected on {total_sensitive_pages} page(s)")
        
        # Créer le chemin du rapport
        if report_path is None:
            report_path = self.pdf_path.parent / f"{self.pdf_path.stem}_ANALYSE.xlsx"
        else:
            report_path = Path(report_path)
        
        # Créer le rapport
        if OPENPYXL_AVAILABLE:
            self._create_excel_report(report_path, grouping_keywords)
        else:
            self._create_text_report(report_path)
        
        logger.info(f"Report created: {report_path}")
        return report_path
    
    def _create_excel_report(self, report_path, grouping_keywords=None):
        """
        Crée un rapport Excel avec la structure demandée.

        Args:
            report_path: Chemin du fichier Excel à créer
            grouping_keywords: Liste de mots-clés pour regrouper les occurrences
                              (ex: ['vantiv', 'netflix'] pour regrouper toutes les mentions)
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Analysis Report"
        
        # Headers
        headers = ["Information Type", "Information", "", "Document", "Pages", "Anonymized Information"]
        ws.append(headers)

        # Dictionnaire pour maintenir la cohérence des anonymisations
        anonymization_map = {}
        
        # Style headers
        header_font = Font(bold=True)
        for cell in ws[1]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Collecter toutes les informations avec leurs pages
        info_dict = {}  # (type, value) -> [pages]
        
        # Images
        for page_num, image_count in self.images_by_page.items():
            key = ('image', f"{image_count} image(s)")
            if key not in info_dict:
                info_dict[key] = []
            info_dict[key].append(page_num)
        
        # Informations sensibles
        type_labels = {
            'email': 'Email',
            'phone': 'Phone',
            'address': 'Address',
            'partial_address': 'Partial Address',
            'city_name': 'City Name',
            'ssn': 'SSN',
            'credit_card': 'Credit Card',
            'irs_ein': 'IRS EIN',
            'commission_file_number': 'Commission File Number',
            'cik_number': 'CIK Number',
            'cusip_code': 'CUSIP Code',
            'isin_code': 'ISIN Code',
            'executive_name': 'Executive Name',
            'person_name': 'Person Name',
            'company_name': 'Company Name',
            'ticker_symbol': 'Ticker Symbol',
            'important_date': 'Important Date',
            'auditor': 'Auditor',
            'underwriter': 'Underwriter',
            'law_firm': 'Law Firm',
            'financial_institution': 'Financial Institution',
            'professional_firm': 'Professional Firm',
            'quoted_title': 'Quoted Title',
            'image': 'Image'
        }
        
        # Dictionnaires pour dédupliquer (normalisés)
        person_names_dedup = {}    # normalized_name -> (best_version, pages)
        company_names_dedup = {}   # normalized_name -> (best_version, pages)
        executive_names_dedup = {} # normalized_name -> (best_version, pages)
        addresses_dedup = {}       # normalized_address -> (best_version, pages)

        for page_num, findings in self.sensitive_info_by_page.items():
            for finding in findings:
                info_type = finding['type']
                info_value = finding['value']

                # Déduplication spéciale pour noms et adresses (avec déduplication par préfixe)
                if info_type in ['person_name', 'executive_name', 'company_name', 'address', 'partial_address']:
                    # Normaliser le nom pour la comparaison
                    normalized = re.sub(r'\s+', ' ', str(info_value).lower().strip())

                    # Choisir le bon dictionnaire
                    if info_type == 'person_name':
                        dedup_dict = person_names_dedup
                    elif info_type == 'executive_name':
                        dedup_dict = executive_names_dedup
                    elif info_type == 'company_name':
                        dedup_dict = company_names_dedup
                    else:  # address or partial_address
                        dedup_dict = addresses_dedup

                    # Vérifier si ce nom est contenu dans un nom existant ou vice-versa (déduplication par préfixe)
                    matching_key = None
                    for existing_key in dedup_dict.keys():
                        # Si le nouveau nom est un préfixe d'un nom existant
                        if existing_key.startswith(normalized + ' '):
                            matching_key = existing_key
                            break
                        # Si un nom existant est un préfixe du nouveau nom
                        elif normalized.startswith(existing_key + ' '):
                            matching_key = existing_key
                            break

                    # Si correspondance partielle trouvée, fusionner avec le nom le plus court
                    if matching_key:
                        best_version, pages = dedup_dict[matching_key]
                        pages.append(page_num)

                        # Garder le nom le plus court (le plus propre)
                        if len(normalized) < len(matching_key):
                            # Le nouveau nom est plus court, le préférer
                            del dedup_dict[matching_key]
                            dedup_dict[normalized] = (str(info_value), pages)
                        # Sinon garder l'existant (ne rien faire)
                        continue

                    if normalized in dedup_dict:
                        # Nom déjà vu, mettre à jour les pages et garder la meilleure version
                        best_version, pages = dedup_dict[normalized]
                        pages.append(page_num)

                        # Préférer la version avec capitalisation standard plutôt que tout en majuscules
                        if str(info_value).isupper() and not best_version.isupper():
                            # Garder la version existante (meilleure capitalisation)
                            dedup_dict[normalized] = (best_version, pages)
                        elif not str(info_value).isupper() and best_version.isupper():
                            # Remplacer par la version avec meilleure capitalisation
                            dedup_dict[normalized] = (str(info_value), pages)
                        elif not str(info_value).isupper() and not best_version.isupper():
                            # Les deux ont une bonne capitalisation - préférer la plus longue (plus de détails)
                            if len(str(info_value)) > len(best_version):
                                dedup_dict[normalized] = (str(info_value), pages)
                            else:
                                dedup_dict[normalized] = (best_version, pages)
                        # Sinon, garder la version existante
                    else:
                        # Nouveau nom
                        dedup_dict[normalized] = (str(info_value), [page_num])
                else:
                    # Pour les autres types, comportement normal avec déduplication simple
                    # Normaliser la valeur pour éviter les doublons exacts
                    normalized_value = str(info_value).strip()
                    key = (info_type, normalized_value)
                    if key not in info_dict:
                        info_dict[key] = []
                    info_dict[key].append(page_num)

        # Ajouter les noms dédupliqués
        for normalized, (best_version, pages) in person_names_dedup.items():
            key = ('person_name', best_version)
            info_dict[key] = pages

        for normalized, (best_version, pages) in executive_names_dedup.items():
            key = ('executive_name', best_version)
            info_dict[key] = pages

        for normalized, (best_version, pages) in company_names_dedup.items():
            key = ('company_name', best_version)
            info_dict[key] = pages

        for normalized, (best_version, pages) in addresses_dedup.items():
            # Déterminer le type exact (address ou partial_address)
            # On utilise 'address' par défaut car la distinction n'est pas importante après déduplication
            key = ('address', best_version)
            info_dict[key] = pages

        # Regrouper par mots-clés si spécifiés
        if grouping_keywords:
            grouped_dict = {}
            keywords_lower = [kw.lower() for kw in grouping_keywords]
            
            for (info_type, info_value), pages in info_dict.items():
                info_value_lower = str(info_value).lower()
                
                # Chercher si cette valeur contient un mot-clé
                matched_keyword = None
                for keyword in keywords_lower:
                    if keyword in info_value_lower:
                        matched_keyword = keyword
                        break
                
                if matched_keyword:
                    # Regrouper sous le mot-clé
                    grouped_key = (info_type, matched_keyword.title())
                    if grouped_key not in grouped_dict:
                        grouped_dict[grouped_key] = []
                    grouped_dict[grouped_key].extend(pages)
                else:
                    # Garder tel quel
                    if (info_type, info_value) not in grouped_dict:
                        grouped_dict[(info_type, info_value)] = []
                    grouped_dict[(info_type, info_value)].extend(pages)
            
            info_dict = grouped_dict
        
        # Trier et ajouter les lignes par ordre alphabétique par type d'information
        # Créer une liste avec les labels pour le tri
        items_with_labels = []
        
        for (info_type, info_value), pages in info_dict.items():
            type_label = type_labels.get(info_type, info_type.capitalize())
            items_with_labels.append((type_label, info_type, info_value, pages))
        
        # Trier par ordre alphabétique du type d'information (type_label), puis par valeur
        sorted_items = sorted(items_with_labels, key=lambda x: (x[0].lower(), str(x[2]).lower()))
        
        # Code couleur pour différents types d'informations
        from openpyxl.styles import PatternFill

        color_map = {
            'email': 'E3F2FD',           # Bleu clair
            'phone': 'E8F5E9',           # Vert clair
            'address': 'FFF9C4',         # Jaune clair
            'person_name': 'F3E5F5',     # Violet clair
            'executive_name': 'E1BEE7',  # Violet moyen
            'company_name': 'FFE0B2',    # Orange clair
            'ssn': 'FFCDD2',             # Rouge clair
            'irs_ein': 'D7CCC8',         # Marron clair
            'commission_file_number': 'CFD8DC',  # Bleu-gris clair
            'cik_number': 'B2DFDB',      # Teal clair
            'cusip_code': 'C5CAE9',      # Indigo clair
            'isin_code': 'C5CAE9',       # Indigo clair
            'auditor': 'DCEDC8',         # Vert olive clair
            'law_firm': 'F0F4C3',        # Citron clair
        }

        for type_label, info_type, info_value, pages in sorted_items:
            pages_sorted = sorted(set(pages))
            pages_str = ', '.join(f"p.{p}" for p in pages_sorted)
            document_name = self.pdf_path.stem

            # Column D: Just document name with .docx extension
            document_column = f"{document_name}.docx"

            # Column E: Document name + page numbers
            pages_column = f"{document_name} {pages_str}"

            # Column F: Anonymized value
            anonymized_value = _anonymize_value(str(info_value), info_type, anonymization_map)

            row = [type_label, info_value, "", document_column, pages_column, anonymized_value]
            ws.append(row)

            # Appliquer la couleur de fond si définie
            if info_type in color_map:
                fill = PatternFill(start_color=color_map[info_type],
                                 end_color=color_map[info_type],
                                 fill_type='solid')
                current_row = ws.max_row
                for cell in ws[current_row]:
                    cell.fill = fill
        
        # Ajuster la largeur des colonnes
        ws.column_dimensions['A'].width = 25  # Information Type
        ws.column_dimensions['B'].width = 50  # Information
        ws.column_dimensions['C'].width = 5   # Empty column
        ws.column_dimensions['D'].width = 30  # Document (just filename.docx)
        ws.column_dimensions['E'].width = 40  # Pages (document + page numbers)
        ws.column_dimensions['F'].width = 50  # Anonymized Information
        
        # Freeze header row
        ws.freeze_panes = 'A2'
        
        # Sauvegarder
        wb.save(report_path)
    
    def _create_text_report(self, report_path):
        """Crée un rapport texte si openpyxl n'est pas disponible."""
        type_labels = {
            'email': 'Email',
            'phone': 'Phone',
            'address': 'Address',
            'ssn': 'SSN',
            'credit_card': 'Credit Card',
            'irs_ein': 'IRS EIN',
            'commission_file_number': 'Commission File Number',
            'zip_code': 'ZIP Code',
            'executive_name': 'Executive Name',
            'person_name': 'Person Name',
            'image': 'Image'
        }
        
        info_regrouped = {}
        
        # Images
        for page_num, image_count in self.images_by_page.items():
            key = ('image', f"{image_count} image(s)")
            if key not in info_regrouped:
                info_regrouped[key] = []
            info_regrouped[key].append(page_num)
        
        # Informations sensibles
        for page_num, findings in self.sensitive_info_by_page.items():
            for finding in findings:
                info_type = finding['type']
                info_value = finding['value']
                key = (info_type, info_value)
                if key not in info_regrouped:
                    info_regrouped[key] = []
                info_regrouped[key].append(page_num)
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("RAPPORT D'ANALYSE DU PDF\n")
            f.write("=" * 80 + "\n\n")
            f.write(f"Fichier analysé: {self.pdf_path.name}\n")
            f.write(f"Date d'analyse: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            if info_regrouped:
                f.write(f"{'Information Type':<30} {'Information':<40} {'Document and Pages':<30}\n")
                f.write("-" * 80 + "\n")
                
                # Créer une liste avec les labels pour le tri alphabétique
                items_with_labels = []
                for (info_type, info_value), pages in info_regrouped.items():
                    type_str = type_labels.get(info_type, info_type)
                    items_with_labels.append((type_str, info_type, info_value, pages))
                
                # Trier par ordre alphabétique du type d'information, puis par valeur
                sorted_items = sorted(items_with_labels, key=lambda x: (x[0].lower(), str(x[2]).lower()))
                
                for type_str, info_type, info_value, pages in sorted_items:
                    value_str = str(info_value)[:38]
                    pages_sorted = sorted(set(pages))
                    pages_str = ', '.join(f"p.{p}" for p in pages_sorted)
                    document_name = self.pdf_path.stem
                    pages_info = f"{document_name} {pages_str}"
                    
                    f.write(f"{type_str:<30} {value_str:<40} {pages_info:<30}\n")
            else:
                f.write("No images or sensitive information detected.\n")


# ==============================================================================
# BATCH PROCESSING HELPER FUNCTIONS
# ==============================================================================

def _anonymize_value(original_value, info_type, anonymization_map):
    """
    Génère un alias anonymisé pour une valeur donnée.

    Args:
        original_value: Valeur originale à anonymiser
        info_type: Type d'information (email, phone, company_name, etc.)
        anonymization_map: Dictionnaire pour maintenir la cohérence des mappings

    Returns:
        Valeur anonymisée
    """
    # Si déjà anonymisé, retourner la valeur existante
    if original_value in anonymization_map:
        return anonymization_map[original_value]

    # Extraire les suffixes à préserver (Inc, LLC, Ltd, etc.)
    suffix_patterns = r'\s+(Inc\.?|LLC\.?|Ltd\.?|Corporation|Corp\.?|Company|Co\.?|LP|LLP|SA|SAS|SARL|GmbH|AG)(\s*,?\s*)$'
    suffix_match = re.search(suffix_patterns, original_value, re.IGNORECASE)
    suffix = suffix_match.group(0) if suffix_match else ''
    base_value = original_value[:len(original_value) - len(suffix)] if suffix else original_value

    anonymized = ''

    if info_type == 'company_name':
        # IMPORTANT: Vérifier si cette valeur est liée à un nom d'entreprise déjà anonymisé
        # Ex: Si "Fidelity" → "Apex", alors "Fidelity Inc" → "Apex Inc"
        found_base = None
        found_base_anon = None

        # Chercher la correspondance la plus longue dans le mapping existant (bidirectionnel)
        for existing_original, existing_anon in anonymization_map.items():
            # Extraire la base de l'original existant (sans suffixes)
            existing_base = re.sub(r'\s+(Inc\.?|LLC\.?|Ltd\.?|Corporation|Corp\.?|Company|Co\.?|LP|LLP|SA|SAS|SARL|GmbH|AG)(\s*,?\s*)$', '', existing_original, flags=re.IGNORECASE).strip()

            # CAS 1: L'original existant est contenu dans la nouvelle valeur
            # Ex: existing="Fifth Third" in new="Fifth Third Bank"
            if existing_base in base_value and len(existing_base) > len(found_base or ''):
                found_base = existing_base
                # Extraire l'alias de base (sans suffixes Inc, LLC, etc.)
                found_base_anon = re.sub(r'\s+(Inc\.?|LLC\.?|Ltd\.?|Corporation|Corp\.?|Company|Co\.?|LP|LLP|SA|SAS|SARL|GmbH|AG)(\s*,?\s*)$', '', existing_anon, flags=re.IGNORECASE).strip()

            # CAS 2: La nouvelle valeur est contenue dans l'original existant
            # Ex: new="Fifth Third" in existing="Fifth Third Bancorp"
            elif base_value in existing_base and len(base_value) > len(found_base or ''):
                found_base = base_value
                # Utiliser le même alias de base
                found_base_anon = re.sub(r'\s+(Inc\.?|LLC\.?|Ltd\.?|Corporation|Corp\.?|Company|Co\.?|LP|LLP|SA|SAS|SARL|GmbH|AG)(\s*,?\s*)$', '', existing_anon, flags=re.IGNORECASE).strip()

        if found_base and found_base_anon:
            # Utiliser le même alias de base
            # Remplacer la partie commune et garder le reste
            remaining_part = base_value.replace(found_base, '').strip()

            if remaining_part:
                # Il y a des mots supplémentaires après le nom de base
                # Ex: "Fidelity Global Payments" → "Apex Global Payments"
                anonymized = f"{found_base_anon} {remaining_part}{suffix}"
            else:
                # Juste le nom de base avec suffix
                anonymized = f"{found_base_anon}{suffix}"
        else:
            # Nouveau nom d'entreprise, créer un nouvel alias
            # Générateur de noms d'entreprise aléatoires
            prefixes = ['Alpha', 'Beta', 'Gamma', 'Delta', 'Epsilon', 'Zeta', 'Theta', 'Lambda',
                       'Sigma', 'Omega', 'Nova', 'Stellar', 'Quantum', 'Nexus', 'Apex', 'Vertex',
                       'Zenith', 'Prime', 'Core', 'Phoenix', 'Atlas', 'Titan', 'Orion', 'Helix']
            suffixes_name = ['Tech', 'Systems', 'Solutions', 'Group', 'Enterprises', 'Industries',
                            'Labs', 'Dynamics', 'Global', 'International', 'Partners', 'Ventures',
                            'Capital', 'Holdings', 'Services', 'Networks', 'Digital', 'Innovation']

            import random
            random.seed(hash(base_value))  # Seed pour cohérence
            prefix = random.choice(prefixes)
            suffix_name = random.choice(suffixes_name)
            anonymized = f"{prefix} {suffix_name}{suffix}"

    elif info_type == 'person_name':
        # Générateur de noms de personnes
        first_names = ['James', 'John', 'Robert', 'Michael', 'William', 'David', 'Richard', 'Joseph',
                      'Thomas', 'Charles', 'Mary', 'Patricia', 'Jennifer', 'Linda', 'Elizabeth',
                      'Barbara', 'Susan', 'Jessica', 'Sarah', 'Karen', 'Alex', 'Chris', 'Jordan',
                      'Taylor', 'Morgan', 'Casey', 'Riley', 'Avery', 'Quinn', 'Blake']
        last_names = ['Smith', 'Johnson', 'Williams', 'Brown', 'Jones', 'Garcia', 'Miller', 'Davis',
                     'Rodriguez', 'Martinez', 'Hernandez', 'Lopez', 'Wilson', 'Anderson', 'Thomas',
                     'Taylor', 'Moore', 'Jackson', 'Martin', 'Lee', 'White', 'Harris', 'Clark',
                     'Lewis', 'Robinson', 'Walker', 'Young', 'Allen', 'King', 'Wright']

        import random
        random.seed(hash(base_value))
        # Détecter la structure (Prénom Nom, Initiale. Nom, etc.)
        parts = base_value.split()
        if len(parts) >= 2:
            # Check si première partie est une initiale
            if len(parts[0].rstrip('.')) <= 2:
                # Initiale + Nom
                anonymized = f"{random.choice(first_names)[0]}. {random.choice(last_names)}"
            else:
                # Prénom Nom ou Prénom I. Nom
                first = random.choice(first_names)
                last = random.choice(last_names)
                if len(parts) == 3 and len(parts[1].rstrip('.')) <= 2:
                    # Prénom I. Nom
                    anonymized = f"{first} {random.choice(first_names)[0]}. {last}"
                else:
                    # Prénom Nom
                    anonymized = f"{first} {last}"
        else:
            # Juste un nom
            anonymized = random.choice(last_names)

        # Préserver la capitalisation
        if base_value.isupper():
            anonymized = anonymized.upper()

    elif info_type == 'email':
        # Générer email anonymisé
        parts = original_value.split('@')
        if len(parts) == 2:
            local_part = parts[0]
            domain = parts[1]

            # Anonymiser le domaine si c'est un domaine d'entreprise
            if domain in anonymization_map:
                anon_domain = anonymization_map[domain].replace('http://', '').replace('https://', '').replace('www.', '')
            else:
                import random
                random.seed(hash(domain))
                domain_names = ['alpaca', 'beta', 'gamma', 'delta', 'epsilon', 'zeta', 'theta']
                anon_domain = f"{random.choice(domain_names)}.com"
                anonymization_map[domain] = anon_domain

            # Anonymiser la partie locale
            import random
            random.seed(hash(local_part))
            names = ['john.smith', 'jane.doe', 'alice.wonder', 'bob.builder', 'charlie.brown',
                    'david.jones', 'emma.watson', 'frank.miller', 'grace.kelly', 'harry.potter']
            anon_local = random.choice(names)

            anonymized = f"{anon_local}@{anon_domain}"

    elif info_type == 'phone':
        # Randomiser les chiffres en gardant le format
        import random
        random.seed(hash(original_value))
        anonymized = re.sub(r'\d', lambda m: str(random.randint(0, 9)), original_value)

    elif info_type in ['address', 'partial_address']:
        # Générer adresse aléatoire ou partielle
        import random
        random.seed(hash(base_value))
        street_nums = [str(random.randint(100, 999))]
        street_names = ['Main Street', 'Oak Avenue', 'Maple Drive', 'Cedar Lane', 'Pine Road',
                      'Elm Street', 'Park Avenue', 'Washington Boulevard', 'Lincoln Way', 'Market Street']
        cities = ['New York', 'Los Angeles', 'Chicago', 'Houston', 'Phoenix', 'Philadelphia',
                 'San Antonio', 'San Diego', 'Dallas', 'San Jose', 'Boston', 'Seattle', 'Denver']
        states = ['NY', 'CA', 'TX', 'FL', 'IL', 'PA', 'OH', 'GA', 'NC', 'MI', 'NJ', 'VA', 'WA', 'MA']
        zip_code = ''.join([str(random.randint(0, 9)) for _ in range(5)])

        # Détecter le format de l'adresse partielle
        if ',' in base_value:
            # Format avec virgule (ex: "Edgewood, NY 11717" ou "Washington, DC 20549")
            parts = base_value.split(',')
            if len(parts) >= 2:
                # Générer format similaire
                city = random.choice(cities)
                state = random.choice(states)
                anonymized = f"{city}, {state} {zip_code}"
            else:
                anonymized = f"{random.choice(street_nums)} {random.choice(street_names)}, {random.choice(cities)} {zip_code}"
        else:
            # Adresse complète ou autre format
            anonymized = f"{random.choice(street_nums)} {random.choice(street_names)}, {random.choice(cities)} {zip_code}"

    elif info_type == 'city_name':
        # Générer nom de ville aléatoire
        import random
        random.seed(hash(base_value))
        cities = ['Springfield', 'Riverside', 'Centerville', 'Georgetown', 'Franklin', 'Clinton',
                 'Madison', 'Washington', 'Arlington', 'Manchester', 'Oxford', 'Cambridge', 'Salem']
        anonymized = random.choice(cities)

    elif info_type == 'ssn':
        # Randomiser SSN en gardant le format XXX-XX-XXXX
        import random
        random.seed(hash(original_value))
        anonymized = f"{random.randint(100, 999):03d}-{random.randint(10, 99):02d}-{random.randint(1000, 9999):04d}"

    elif info_type == 'credit_card':
        # Randomiser numéro de carte en gardant le format
        import random
        random.seed(hash(original_value))
        anonymized = re.sub(r'\d', lambda m: str(random.randint(0, 9)), original_value)

    elif info_type in ['cusip_code', 'isin_code']:
        # Randomiser codes en gardant le format
        import random
        random.seed(hash(original_value))
        anonymized = re.sub(r'[A-Z0-9]', lambda m: random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789') if m.group().isalnum() else m.group(), original_value)

    elif info_type == 'quoted_title':
        # Garder les guillemets, anonymiser le contenu
        import random
        random.seed(hash(base_value))
        titles = ['Strategic Initiative', 'Annual Report', 'Market Analysis', 'Financial Overview',
                 'Corporate Strategy', 'Business Plan', 'Performance Review', 'Investment Proposal']
        if '"' in original_value or '"' in original_value:
            anonymized = f'"{random.choice(titles)}"'
        else:
            anonymized = random.choice(titles)

    elif info_type in ['cik_number', 'commission_file_number', 'irs_ein']:
        # Randomiser les numéros en gardant le format
        import random
        random.seed(hash(original_value))
        anonymized = re.sub(r'\d', lambda m: str(random.randint(0, 9)), original_value)

    elif info_type == 'ticker_symbol':
        # Générer ticker aléatoire
        import random
        random.seed(hash(base_value))
        letters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        ticker_len = len(re.sub(r'[^A-Z]', '', base_value.upper()))
        anonymized = ''.join(random.choice(letters) for _ in range(ticker_len))
        # Préserver le format (ex: (NYSE: XXXX))
        if '(' in original_value:
            match = re.match(r'\(([^:]+):\s*', original_value)
            if match:
                exchange = match.group(1)
                anonymized = f"({exchange}: {anonymized})"

    elif info_type == 'important_date':
        # Randomiser les dates en gardant le format
        import random
        random.seed(hash(original_value))
        anonymized = re.sub(r'\d', lambda m: str(random.randint(0, 9)), original_value)

    elif info_type in ['auditor', 'underwriter', 'law_firm', 'financial_institution', 'professional_firm']:
        # Générateur de noms de firmes
        import random
        random.seed(hash(base_value))
        firm_names = ['Morgan & Associates', 'Smith Partners LLP', 'Johnson Group', 'Williams & Co.',
                     'Brown Associates', 'Davis & Partners', 'Miller Group LLP', 'Wilson & Associates',
                     'Anderson Partners', 'Taylor & Co.', 'Thomas Associates', 'Moore & Partners']
        anonymized = random.choice(firm_names)

    else:
        # Par défaut, masquer avec des X
        anonymized = 'X' * min(len(base_value), 10)

    # Sauvegarder le mapping
    anonymization_map[original_value] = anonymized

    return anonymized


def _consolidate_batch_findings(all_findings_by_pdf, all_images_by_pdf):
    """
    Consolidate findings from multiple PDFs, deduplicating by (type, value).

    Args:
        all_findings_by_pdf: {
            'file1.pdf': {page_num: [{type, value, page}, ...]},
            'file2.pdf': {page_num: [{type, value, page}, ...]},
            ...
        }
        all_images_by_pdf: {
            'file1.pdf': {page_num: count},
            'file2.pdf': {page_num: count},
            ...
        }

    Returns:
        consolidated: {
            ('email', 'user@example.com'): {
                'type': 'email',
                'value': 'user@example.com',
                'occurrences': [
                    {'document': 'file1.docx', 'pages': [5, 12]},
                    {'document': 'file2.docx', 'pages': [3]}
                ]
            },
            ...
        }
    """
    consolidated = {}

    # Process sensitive information findings
    for pdf_filename, findings_by_page in all_findings_by_pdf.items():
        # Convert pdf filename to docx for display
        doc_name = pdf_filename.replace('.pdf', '.docx')

        for page_num, findings_list in findings_by_page.items():
            for finding in findings_list:
                info_type = finding.get('type', '')
                info_value = finding.get('value', '')
                page = finding.get('page', page_num)

                # Normalize value for better deduplication
                normalized_value = info_value
                if info_type == 'email':
                    normalized_value = info_value.lower()  # Emails: lowercase
                elif info_type in ['person_name', 'executive_name', 'company_name']:
                    # Names: prefer proper case over ALL CAPS
                    if info_value.isupper() and len(info_value) > 3:
                        # If ALL CAPS, check if we have a better version
                        pass  # Will be handled in merge logic
                    else:
                        normalized_value = info_value

                # Create unique key (type, normalized_value)
                key = (info_type, normalized_value)

                if key not in consolidated:
                    consolidated[key] = {
                        'type': info_type,
                        'value': info_value,  # Keep original formatting
                        'occurrences': []
                    }

                # Find if this document already has an entry
                existing_doc = None
                for occ in consolidated[key]['occurrences']:
                    if occ['document'] == doc_name:
                        existing_doc = occ
                        break

                if existing_doc:
                    # Add page to existing document entry
                    if page not in existing_doc['pages']:
                        existing_doc['pages'].append(page)
                        existing_doc['pages'].sort()
                else:
                    # Create new document entry
                    consolidated[key]['occurrences'].append({
                        'document': doc_name,
                        'pages': [page]
                    })

                # Update value to prefer better formatting
                if info_type in ['person_name', 'executive_name', 'company_name']:
                    current_value = consolidated[key]['value']
                    # Prefer proper case over ALL CAPS
                    if current_value.isupper() and not info_value.isupper():
                        consolidated[key]['value'] = info_value
                    # Prefer longer version if same case
                    elif len(info_value) > len(current_value):
                        if info_value.lower() == current_value.lower():
                            consolidated[key]['value'] = info_value

    # Process images
    for pdf_filename, images_by_page in all_images_by_pdf.items():
        doc_name = pdf_filename.replace('.pdf', '.docx')

        for page_num, image_count in images_by_page.items():
            if image_count > 0:
                info_type = 'image'
                info_value = f"{image_count} image(s)"

                key = (info_type, info_value)

                if key not in consolidated:
                    consolidated[key] = {
                        'type': info_type,
                        'value': info_value,
                        'occurrences': []
                    }

                # Find if this document already has an entry
                existing_doc = None
                for occ in consolidated[key]['occurrences']:
                    if occ['document'] == doc_name:
                        existing_doc = occ
                        break

                if existing_doc:
                    if page_num not in existing_doc['pages']:
                        existing_doc['pages'].append(page_num)
                        existing_doc['pages'].sort()
                else:
                    consolidated[key]['occurrences'].append({
                        'document': doc_name,
                        'pages': [page_num]
                    })

    return consolidated


def _create_consolidated_excel_report(consolidated_findings, report_path, grouping_keywords=None):
    """
    Create consolidated Excel report from findings across multiple PDFs.

    Args:
        consolidated_findings: Dict from _consolidate_batch_findings()
        report_path: Path to save the Excel file
        grouping_keywords: Optional list of keywords for grouping
    """
    if not OPENPYXL_AVAILABLE:
        logger.warning("openpyxl not available, cannot create Excel report")
        return None

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Consolidated Analysis"

    # Headers
    headers = ["Information Type", "Information", "", "Documents", "Pages", "Anonymized Information"]
    ws.append(headers)

    # Dictionnaire pour maintenir la cohérence des anonymisations
    anonymization_map = {}

    # Style headers
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Type labels (same as in _create_excel_report)
    type_labels = {
        'email': 'Email',
        'phone': 'Phone',
        'address': 'Address',
        'partial_address': 'Partial Address',
        'city_name': 'City Name',
        'ssn': 'SSN',
        'credit_card': 'Credit Card',
        'irs_ein': 'IRS EIN',
        'commission_file_number': 'Commission File Number',
        'cik_number': 'CIK Number',
        'cusip_code': 'CUSIP Code',
        'isin_code': 'ISIN Code',
        'executive_name': 'Executive Name',
        'person_name': 'Person Name',
        'company_name': 'Company Name',
        'ticker_symbol': 'Ticker Symbol',
        'important_date': 'Important Date',
        'auditor': 'Auditor',
        'underwriter': 'Underwriter',
        'law_firm': 'Law Firm',
        'financial_institution': 'Financial Institution',
        'professional_firm': 'Professional Firm',
        'quoted_title': 'Quoted Title',
        'image': 'Image'
    }

    # Color mapping (same as in _create_excel_report)
    color_map = {
        'email': 'E3F2FD',
        'phone': 'E8F5E9',
        'address': 'FFF9C4',
        'partial_address': 'FFF9C4',
        'person_name': 'F3E5F5',
        'executive_name': 'F3E5F5',
        'company_name': 'FFE0B2',
        'ssn': 'FFCDD2',
        'credit_card': 'FFCDD2',
        'irs_ein': 'E1F5FE',
        'commission_file_number': 'E1F5FE',
        'cik_number': 'E1F5FE',
        'cusip_code': 'E1F5FE',
        'isin_code': 'E1F5FE',
        'ticker_symbol': 'E1F5FE',
        'image': 'F5F5F5'
    }

    # Prepare data for Excel
    rows_data = []

    # GROUPING BY KEYWORDS (if provided)
    # Group findings that contain the same keyword
    if grouping_keywords:
        grouped_findings = {}  # {(type, keyword): {value, occurrences}}

        for (info_type, _), finding_data in consolidated_findings.items():
            info_value = finding_data['value']
            info_value_lower = info_value.lower()

            # Check if this value contains any of the keywords
            matched_keyword = None
            for keyword in grouping_keywords:
                if keyword.lower() in info_value_lower:
                    matched_keyword = keyword
                    break

            if matched_keyword:
                # Group under the keyword
                key = (info_type, matched_keyword.lower())
                if key not in grouped_findings:
                    grouped_findings[key] = {
                        'type': info_type,
                        'value': matched_keyword,  # Display keyword as the value
                        'occurrences': []
                    }
                # Merge occurrences
                for occ in finding_data['occurrences']:
                    # Check if document already exists
                    existing = None
                    for existing_occ in grouped_findings[key]['occurrences']:
                        if existing_occ['document'] == occ['document']:
                            existing = existing_occ
                            break

                    if existing:
                        # Merge pages
                        for page in occ['pages']:
                            if page not in existing['pages']:
                                existing['pages'].append(page)
                        existing['pages'].sort()
                    else:
                        # Add new document entry
                        grouped_findings[key]['occurrences'].append({
                            'document': occ['document'],
                            'pages': occ['pages'][:]
                        })
            else:
                # Not grouped - keep original
                key = (info_type, info_value.lower())
                grouped_findings[key] = finding_data

        # Use grouped findings
        consolidated_findings = grouped_findings

    for (info_type, _), finding_data in consolidated_findings.items():
        info_value = finding_data['value']
        occurrences = finding_data['occurrences']

        # Column A: Type label
        type_label = type_labels.get(info_type, info_type.replace('_', ' ').title())

        # Column B: Information value
        value_str = str(info_value)

        # Column D: Documents (comma-separated)
        documents = [occ['document'] for occ in occurrences]
        documents_str = ', '.join(documents)

        # Column E: Pages (grouped by document)
        pages_parts = []
        for occ in occurrences:
            doc_name = occ['document']
            pages = occ['pages']
            pages_str = ', '.join(f"p.{p}" for p in sorted(pages))
            pages_parts.append(f"{doc_name} {pages_str}")

        pages_str = '; '.join(pages_parts)

        rows_data.append({
            'type': info_type,
            'type_label': type_label,
            'value': value_str,
            'documents': documents_str,
            'pages': pages_str,
            'anonymized': ''  # Will be filled when creating rows
        })

    # Sort by type, then by value
    rows_data.sort(key=lambda x: (x['type_label'], x['value']))

    # Write rows to Excel
    for row_data in rows_data:
        # Generate anonymized value
        anonymized_value = _anonymize_value(row_data['value'], row_data['type'], anonymization_map)

        row = [
            row_data['type_label'],  # Column A
            row_data['value'],        # Column B
            '',                        # Column C (empty spacer)
            row_data['documents'],     # Column D
            row_data['pages'],         # Column E
            anonymized_value           # Column F
        ]
        ws.append(row)

        # Apply color coding
        row_num = ws.max_row
        info_type = row_data['type']
        if info_type in color_map:
            fill_color = color_map[info_type]
            fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
            for col in range(1, 7):  # Columns A-F
                ws.cell(row=row_num, column=col).fill = fill

    # Set column widths
    ws.column_dimensions['A'].width = 25  # Information Type
    ws.column_dimensions['B'].width = 50  # Information
    ws.column_dimensions['C'].width = 5   # Empty spacer
    ws.column_dimensions['D'].width = 40  # Documents
    ws.column_dimensions['E'].width = 60  # Pages (longer for multi-document)
    ws.column_dimensions['F'].width = 50  # Anonymized Information

    # Freeze header row
    ws.freeze_panes = 'A2'

    # Save
    wb.save(str(report_path))
    logger.info(f"Consolidated report saved to {report_path}")

    return report_path


def get_pdf_directory():
    """Get or create the PDF input directory."""
    pdf_dir = Path.cwd() / "pdf"
    pdf_dir.mkdir(exist_ok=True)
    return pdf_dir


def get_output_directory():
    """Get or create the output directory."""
    output_dir = Path.cwd() / "output"
    output_dir.mkdir(exist_ok=True)
    return output_dir


def list_pdf_files(directory=None):
    """List all PDF files in the specified directory (or current directory if None)."""
    if directory is None:
        directory = Path.cwd()
    else:
        directory = Path(directory)
    
    if not directory.exists():
        return []
    
    pdf_files = sorted([f for f in directory.glob("*.pdf") if f.is_file()])
    return pdf_files


def select_pdf_file():
    """Display list of PDFs and allow selection by number."""
    pdf_files = list_pdf_files()
    
    if not pdf_files:
        print("\n" + "=" * 70)
        print("⚠️  No PDF files found in the current directory.")
        print("=" * 70)
        return None
    
    print("\n" + "=" * 70)
    print(" " * 25 + "📄 AVAILABLE PDF FILES")
    print("=" * 70)
    
    for idx, pdf_file in enumerate(pdf_files, 1):
        file_size = pdf_file.stat().st_size / 1024  # Size in KB
        if file_size < 1024:
            size_str = f"{file_size:.1f} KB"
        else:
            size_str = f"{file_size / 1024:.2f} MB"
        print(f"  [{idx}] {pdf_file.name:<50} {size_str:>10}")
    
    print("=" * 70)
    
    while True:
        try:
            choice = input(f"\n👉 Select a file (1-{len(pdf_files)}) or press Enter to cancel: ").strip()
            if not choice:
                return None
            
            file_idx = int(choice) - 1
            if 0 <= file_idx < len(pdf_files):
                selected = pdf_files[file_idx]
                print(f"✓ Selected: {selected.name}")
                return str(selected)
            else:
                print(f"❌ Invalid number. Please choose between 1 and {len(pdf_files)}")
        except ValueError:
            print("❌ Please enter a valid number.")
        except KeyboardInterrupt:
            print("\n\n⚠️  Operation cancelled.")
            return None


def show_main_menu():
    """
    Menu principal moderne avec catégories claires et informations sur les dépendances.
    """
    # Vérifier les capacités disponibles
    ocr_status = "✅ Available" if OCR_AVAILABLE else "❌ Not installed"
    conversion_status = "✅ Available" if PDF2DOCX_AVAILABLE else "❌ Not installed"
    analysis_status = "✅ Available" if (PYPDF2_AVAILABLE and OPENPYXL_AVAILABLE) else "❌ Not installed"
    
    print("\n" + "="*80)
    print(" " * 25 + "🚀 xAI PDF CONVERTER 🚀")
    print("="*80)
    
    # Section 1: Single File Operations
    print("\n" + "─" * 80)
    print("  📄 SINGLE FILE OPERATIONS")
    print("─" * 80)
    print(f"  1. Convert PDF → Word              [{conversion_status}]")
    print(f"  2. Convert PDF → Word + Analysis   [{conversion_status}]")
    print(f"  3. Analyze PDF Only (Excel Report) [{analysis_status}]")
    print(f"  4. OCR: Scanned PDF → Word         [{ocr_status}]")
    
    # Section 2: Batch Processing
    print("\n" + "─" * 80)
    print("  📦 BATCH PROCESSING (Multiple PDFs)")
    print("─" * 80)
    print("  5. Batch: Convert All PDFs")
    print("  6. Batch: Convert + Analyze All")
    print("  7. Batch: Analyze All PDFs")
    
    # Section 3: Settings & Info
    print("\n" + "─" * 80)
    print("  ⚙️  SETTINGS & INFO")
    print("─" * 80)
    print("  8. View System Status")
    print("  9. Configure Settings")
    print("  0. Exit")
    
    print("\n" + "="*80)
    
    choice = input("➤ Your choice (0-9): ").strip()
    return choice


def show_system_status():
    """
    Affiche le statut des dépendances et des capacités du système.
    """
    print("\n" + "="*80)
    print(" " * 28 + "SYSTEM STATUS")
    print("="*80)
    
    # Core dependencies
    print("\n🔧 CORE DEPENDENCIES:")
    deps = [
        ("PyPDF2", PYPDF2_AVAILABLE, "PDF text extraction"),
        ("pdf2docx", PDF2DOCX_AVAILABLE, "PDF to Word conversion"),
        ("python-docx", PYTHON_DOCX_AVAILABLE, "Word document processing"),
        ("openpyxl", OPENPYXL_AVAILABLE, "Excel report generation"),
        ("pdfplumber", PDFPLUMBER_AVAILABLE, "Advanced table extraction"),
    ]
    
    # NLP dependencies for entity recognition
    print("\n🧠 NLP DEPENDENCIES (for improved entity detection):")
    nlp_deps = [
        ("spaCy", SPACY_AVAILABLE, "Named Entity Recognition (NER) for person/company detection"),
        ("transformers", TRANSFORMERS_AVAILABLE, "Advanced NLP models (optional)"),
    ]
    
    for name, available, description in nlp_deps:
        status = "✅" if available else "❌"
        print(f"  {status} {name:20} - {description}")
    
    if SPACY_AVAILABLE and SPACY_NLP:
        model_name = SPACY_NLP.meta.get('name', 'unknown')
        print(f"     └─ Model loaded: {model_name}")
    elif not SPACY_AVAILABLE:
        print("     └─ Install: pip install spacy")
        print("     └─ Then: python -m spacy download en_core_web_sm")
    
    for name, available, description in deps:
        status = "✅" if available else "❌"
        print(f"  {status} {name:20} - {description}")
    
    # Optional dependencies
    print("\n🎨 OPTIONAL FEATURES:")
    opt_deps = [
        ("tqdm", TQDM_AVAILABLE, "Progress bars"),
        ("phonenumbers", PHONENUMBERS_AVAILABLE, "Phone number validation"),
    ]
    
    for name, available, description in opt_deps:
        status = "✅" if available else "⚠️ "
        print(f"  {status} {name:20} - {description}")
    
    # OCR capabilities
    print("\n🔍 OCR ENGINES (for scanned PDFs):")
    ocr_engines = [
        ("PaddleOCR", PADDLEOCR_AVAILABLE, "🏆 Best quality + table detection"),
        ("EasyOCR", EASYOCR_AVAILABLE, "⭐ Good quality, easy setup"),
        ("Tesseract", PYTESSERACT_AVAILABLE, "✓  Basic OCR"),
        ("pdf2image", PDF2IMAGE_AVAILABLE, "Required for all OCR methods"),
    ]
    
    for name, available, description in ocr_engines:
        status = "✅" if available else "❌"
        print(f"  {status} {name:20} - {description}")
    
    # Summary
    print("\n" + "─"*80)
    total_features = len(deps) + len(opt_deps) + len(ocr_engines) + len(nlp_deps)
    available_features = sum([
        sum([1 for _, av, _ in deps if av]),
        sum([1 for _, av, _ in opt_deps if av]),
        sum([1 for _, av, _ in ocr_engines if av]),
        sum([1 for _, av, _ in nlp_deps if av]),
    ])
    
    print(f"📊 SUMMARY: {available_features}/{total_features} features available")
    
    if available_features < total_features:
        print("\n💡 To install missing dependencies:")
        if not PDF2DOCX_AVAILABLE:
            print("   pip install pdf2docx")
        if not OPENPYXL_AVAILABLE:
            print("   pip install openpyxl")
        if not OCR_AVAILABLE:
            print("   pip install paddlepaddle paddleocr  (recommended)")
            print("   pip install pdf2image pillow")
        if not SPACY_AVAILABLE:
            print("   pip install spacy")
            print("   python -m spacy download en_core_web_sm  (for better entity detection)")
    
    print("="*80)
    input("\n Press Enter to continue...")


def show_settings_menu():
    """
    Menu de configuration avec options persistantes.
    """
    print("\n" + "="*80)
    print(" " * 30 + "SETTINGS")
    print("="*80)
    
    print("\n⚙️  CONVERSION SETTINGS:")
    print("  1. Set default output directory")
    print("  2. Enable/disable verbose logging")
    print("  3. Configure parallelization (pages per chunk)")
    print("  4. Set OCR language")
    
    print("\n📋 ANALYSIS SETTINGS:")
    print("  5. Configure keyword grouping")
    print("  6. Set detection thresholds")
    
    print("\n🎨 DISPLAY SETTINGS:")
    print("  7. Toggle progress bars")
    print("  8. Set report format (Excel/Text)")
    
    print("\n  0. Back to main menu")
    print("="*80)
    
    choice = input("➤ Your choice (0-8): ").strip()
    return choice


def handle_single_convert(use_ocr=False):
    """
    Gère la conversion d'un seul fichier PDF avec UI améliorée.
    """
    pdf_file = select_pdf_file()
    if not pdf_file:
        return
    
    # Options de sortie
    print("\n" + "─"*80)
    print("📝 OUTPUT OPTIONS:")
    print("─"*80)
    docx_file = input("  Output Word file (press Enter for auto): ").strip()
    if not docx_file:
        docx_file = None
    
    if use_ocr:
        if not OCR_AVAILABLE:
            print("\n❌ OCR is not installed.")
            print("\nInstall ONE of these (in order of quality):")
            print("  🏆 pip install paddlepaddle paddleocr  (BEST)")
            print("  ⭐ pip install easyocr                 (GOOD)")
            print("  ✓  pip install pytesseract             (OK)")
            print("\nAlso required:")
            print("  pip install pdf2image pillow")
            input("\nPress Enter to continue...")
            return
    
    try:
        print("\n" + "="*80)
        print(" " * 28 + "🔄 CONVERTING...")
        print("="*80)
        
        converter = SECPDFConverter(pdf_file, docx_file, verbose=True)
        
        if use_ocr:
            converter.convert(use_ocr=True, auto_detect_scanned=False)
        else:
            converter.convert()
        
        print("\n" + "="*80)
        print(" " * 28 + "✅ SUCCESS!")
        print("="*80)
        print(f"  📄 Input:  {converter.pdf_path.name}")
        print(f"  📝 Output: {converter.docx_path.name}")
        print(f"  📁 Location: {converter.docx_path.parent}")
        print("="*80)
        
    except Exception as e:
        print("\n" + "="*80)
        print(" " * 28 + "❌ ERROR")
        print("="*80)
        print(f"  {e}")
        print("="*80)
        import traceback
        traceback.print_exc()
    
    input("\nPress Enter to continue...")


def handle_single_convert_analyze():
    """
    Gère la conversion + analyse d'un seul fichier.
    """
    pdf_file = select_pdf_file()
    if not pdf_file:
        return
    
    print("\n" + "─"*80)
    print("📝 OUTPUT OPTIONS:")
    print("─"*80)
    docx_file = input("  Output Word file (press Enter for auto): ").strip()
    if not docx_file:
        docx_file = None
    
    # Keyword grouping
    grouping_keywords = get_grouping_keywords()
    
    try:
        print("\n" + "="*80)
        print(" " * 25 + "🔄 CONVERTING + ANALYZING...")
        print("="*80)
        
        converter = SECPDFConverter(pdf_file, docx_file, verbose=True)
        
        # Conversion
        print("\n📄 Step 1/2: Converting to Word...")
        converter.convert()
        
        # Analysis
        print("\n🔍 Step 2/2: Analyzing document...")
        report_path = converter.analyze_and_create_report(
            grouping_keywords=grouping_keywords,
            verbose=True
        )
        
        total_images = sum(converter.images_by_page.values())
        total_sensitive = sum(len(findings) for findings in converter.sensitive_info_by_page.values())
        
        print("\n" + "="*80)
        print(" " * 28 + "✅ SUCCESS!")
        print("="*80)
        print(f"  📄 Input PDF:  {converter.pdf_path.name}")
        print(f"  📝 Word file:  {converter.docx_path.name}")
        print(f"  📊 Report:     {report_path.name}")
        print(f"  📁 Location:   {converter.docx_path.parent}")
        print("\n📈 ANALYSIS SUMMARY:")
        print(f"  • Images detected:        {total_images}")
        print(f"  • Sensitive items:        {total_sensitive}")
        print(f"  • Pages with images:      {len(converter.images_by_page)}")
        print(f"  • Pages with info:        {len(converter.sensitive_info_by_page)}")
        print("="*80)
        
    except Exception as e:
        print("\n" + "="*80)
        print(" " * 28 + "❌ ERROR")
        print("="*80)
        print(f"  {e}")
        print("="*80)
        import traceback
        traceback.print_exc()
    
    input("\nPress Enter to continue...")


def handle_analyze_only():
    """
    Gère l'analyse seule d'un fichier PDF.
    """
    pdf_file = select_pdf_file()
    if not pdf_file:
        return
    
    grouping_keywords = get_grouping_keywords()
    
    try:
        print("\n" + "="*80)
        print(" " * 28 + "🔍 ANALYZING...")
        print("="*80)
        
        converter = SECPDFConverter(pdf_file, verbose=True)
        report_path = converter.analyze_and_create_report(
            grouping_keywords=grouping_keywords,
            verbose=True
        )
        
        total_images = sum(converter.images_by_page.values())
        total_sensitive = sum(len(findings) for findings in converter.sensitive_info_by_page.values())
        
        print("\n" + "="*80)
        print(" " * 25 + "✅ ANALYSIS COMPLETE!")
        print("="*80)
        print(f"  📄 Input PDF:  {converter.pdf_path.name}")
        print(f"  📊 Report:     {report_path.name}")
        print(f"  📁 Location:   {report_path.parent}")
        print("\n📈 SUMMARY:")
        print(f"  • Images detected:        {total_images}")
        print(f"  • Sensitive items:        {total_sensitive}")
        print(f"  • Pages with images:      {len(converter.images_by_page)}")
        print(f"  • Pages with info:        {len(converter.sensitive_info_by_page)}")
        print("="*80)
        
    except Exception as e:
        print("\n" + "="*80)
        print(" " * 28 + "❌ ERROR")
        print("="*80)
        print(f"  {e}")
        print("="*80)
        import traceback
        traceback.print_exc()
    
    input("\nPress Enter to continue...")


def get_grouping_keywords():
    """Ask user for keywords to group occurrences."""
    print("\n" + "─" * 70)
    print(" " * 15 + "🔑 KEYWORD GROUPING (Optional)")
    print("─" * 70)
    print("\n   Group occurrences containing the same keyword together.")
    print("   Example: 'vantiv' will group 'Vantiv Inc', 'Vantiv LLC', etc.")
    print("   All pages where these occur will be listed together.")
    print()
    print("─" * 70)
    
    keywords_input = input("   Keywords (comma-separated, or press Enter to skip): ").strip()
    
    if not keywords_input:
        return None
    
    keywords = [kw.strip() for kw in keywords_input.split(',') if kw.strip()]
    if keywords:
        print(f"   ✓ Grouping enabled for: {', '.join(keywords)}")
    return keywords if keywords else None


def process_batch(do_conversion=True, do_analysis=True, grouping_keywords=None, verbose=False, create_consolidated_report=False):
    """
    Process all PDFs in the pdf/ directory and organize outputs in output/.

    Args:
        do_conversion: Whether to convert PDFs to Word documents
        do_analysis: Whether to analyze PDFs for sensitive information
        grouping_keywords: Optional list of keywords for grouping analysis results
        verbose: If True, show detailed logs. If False (default), show simple progress only.
        create_consolidated_report: If True, create a consolidated report across all PDFs

    Returns:
        dict: Summary of processed files with success/failure counts
    """
    # Configure logging based on verbose mode
    configure_logging(verbose=verbose)
    
    pdf_dir = get_pdf_directory()
    output_dir = get_output_directory()
    
    pdf_files = list_pdf_files(pdf_dir)
    
    if not pdf_files:
        print("\n" + "=" * 70)
        print("⚠️  No PDF files found in the 'pdf/' directory.")
        print("=" * 70)
        print(f"\n   📁 Please place your PDF files in: {pdf_dir.absolute()}")
        return {
            'total': 0,
            'success': 0,
            'failed': 0,
            'results': []
        }
    
    # Header - toujours affiché, mais moins détaillé en mode simple
    print("\n" + "╔" + "═" * 68 + "╗")
    print("║" + " " * 20 + "📦 BATCH PROCESSING" + " " * 28 + "║")
    print("╚" + "═" * 68 + "╝")
    print(f"\n   📂 Input directory:  {pdf_dir.absolute()}")
    print(f"   📂 Output directory: {output_dir.absolute()}")
    print(f"   📄 PDF files found:  {len(pdf_files)}")
    print(f"   🔄 Conversion:       {'✓ Enabled' if do_conversion else '✗ Disabled'}")
    print(f"   🔍 Analysis:         {'✓ Enabled' if do_analysis else '✗ Disabled'}")
    if grouping_keywords:
        print(f"   🔑 Keywords:         {', '.join(grouping_keywords)}")
    print("\n" + "─" * 70)
    
    results = {
        'total': len(pdf_files),
        'success': 0,
        'failed': 0,
        'results': []
    }

    # Collections for consolidated report
    all_findings = {}  # {pdf_filename: sensitive_info_by_page}
    all_images = {}    # {pdf_filename: images_by_page}

    # Progress bar pour mode simple
    if not verbose and TQDM_AVAILABLE:
        file_iterator = tqdm(pdf_files, desc="Processing PDFs", unit="file", ncols=80, leave=True)
    else:
        file_iterator = pdf_files
    
    for idx, pdf_file in enumerate(file_iterator, 1):
        pdf_name = pdf_file.stem
        
        # Toujours afficher le header de fichier
        if verbose:
            print(f"\n[{idx}/{len(pdf_files)}] Processing: {pdf_file.name}")
            print("─" * 70)
        else:
            # Mode simple : affichage minimal
            if TQDM_AVAILABLE:
                file_iterator.set_description(f"Processing: {pdf_file.name[:50]}")
            else:
                print(f"\n[{idx}/{len(pdf_files)}] {pdf_file.name}")
        
        # Create output subdirectory for this PDF
        pdf_output_dir = output_dir / pdf_name
        pdf_output_dir.mkdir(exist_ok=True)
        
        # Define output paths
        docx_path = pdf_output_dir / f"{pdf_name}.docx" if do_conversion else None
        analysis_path = pdf_output_dir / f"{pdf_name}_analysis.xlsx" if do_analysis else None
        
        result = {
            'pdf': pdf_file.name,
            'conversion': None,
            'analysis': None,
            'error': None
        }
        
        try:
            converter = SECPDFConverter(pdf_file, docx_path, verbose=verbose)
            
            # Convert if requested
            if do_conversion:
                try:
                    if verbose:
                        print(f"   ⚙️  Converting to Word...")
                    converter.convert()
                    result['conversion'] = 'success'
                    if verbose:
                        print(f"   ✓ Conversion complete: {docx_path.name}")
                except Exception as e:
                    result['conversion'] = 'failed'
                    result['error'] = str(e)
                    logger.error(f"Conversion failed for {pdf_file.name}: {e}")
                    if verbose:
                        print(f"   ❌ Conversion failed: {e}")
            
            # Analyze if requested
            if do_analysis:
                try:
                    if verbose:
                        print(f"   🔍 Analyzing document...")
                    report_path = converter.analyze_and_create_report(
                        report_path=analysis_path,
                        grouping_keywords=grouping_keywords,
                        verbose=verbose
                    )
                    result['analysis'] = 'success'
                    
                    total_images = sum(converter.images_by_page.values())
                    total_sensitive = sum(len(findings) for findings in converter.sensitive_info_by_page.values())
                    
                    if verbose:
                        print(f"   ✓ Analysis complete: {report_path.name}")
                        print(f"      • Images: {total_images}, Sensitive items: {total_sensitive}")

                    # Collect findings for consolidated report
                    if create_consolidated_report:
                        all_findings[pdf_file.name] = converter.sensitive_info_by_page.copy()
                        all_images[pdf_file.name] = converter.images_by_page.copy()

                except Exception as e:
                    result['analysis'] = 'failed'
                    if not result['error']:
                        result['error'] = str(e)
                    logger.error(f"Analysis failed for {pdf_file.name}: {e}")
                    if verbose:
                        print(f"   ❌ Analysis failed: {e}")
            
            # Determine overall success
            if do_conversion and do_analysis:
                overall_success = result['conversion'] == 'success' and result['analysis'] == 'success'
            elif do_conversion:
                overall_success = result['conversion'] == 'success'
            elif do_analysis:
                overall_success = result['analysis'] == 'success'
            else:
                overall_success = False
            
            if overall_success:
                results['success'] += 1
                if not verbose:
                    # Mode simple : afficher succès
                    status_msg = "✓"
                    if TQDM_AVAILABLE:
                        file_iterator.set_postfix(status=status_msg)
                    else:
                        print(f"   ✓ Done")
            else:
                results['failed'] += 1
                if not verbose:
                    status_msg = "✗"
                    if TQDM_AVAILABLE:
                        file_iterator.set_postfix(status=status_msg)
                    else:
                        print(f"   ✗ Failed")
            
        except Exception as e:
            result['error'] = str(e)
            results['failed'] += 1
            logger.error(f"Failed to process {pdf_file.name}: {e}")
            if verbose:
                print(f"   ❌ Error: {e}")
            elif not TQDM_AVAILABLE:
                print(f"   ✗ Error: {e}")
        
        results['results'].append(result)

    # Generate consolidated report if requested
    if create_consolidated_report and all_findings:
        print("\n" + "─" * 70)
        print("   📊 Creating consolidated analysis report...")

        try:
            # Consolidate findings from all PDFs
            consolidated = _consolidate_batch_findings(all_findings, all_images)

            # Create consolidated Excel file
            consolidated_report_path = output_dir / "consolidated statements.xlsx"
            _create_consolidated_excel_report(
                consolidated,
                consolidated_report_path,
                grouping_keywords=grouping_keywords
            )

            print(f"   ✓ Consolidated report saved: {consolidated_report_path.name}")
        except Exception as e:
            logger.error(f"Failed to create consolidated report: {e}")
            print(f"   ❌ Consolidated report failed: {e}")

    # Print summary - toujours affiché de la même manière
    print("\n" + "╔" + "═" * 68 + "╗")
    print("║" + " " * 20 + "📊 BATCH SUMMARY" + " " * 33 + "║")
    print("╚" + "═" * 68 + "╝")
    print(f"\n   ✅ Successful: {results['success']}/{results['total']}")
    print(f"   ❌ Failed:     {results['failed']}/{results['total']}")
    print(f"\n   📁 All outputs saved to: {output_dir.absolute()}")
    if results['failed'] > 0:
        print("\n   ⚠️  Failed files:")
        for res in results['results']:
            if res['error']:
                print(f"      • {res['pdf']}: {res['error']}")
    
    return results


def main():
    """
    Main function avec menu interactif amélioré.
    """
    # Check for command-line arguments (batch mode)
    if len(sys.argv) > 1:
        # Legacy support for command-line usage
        batchmode = True
        doconversion = True
        doanalysis = True
        verbose = False
        grouping_keywords = None
        
        # Parse arguments
        if sys.argv[1].endswith('.pdf') and Path(sys.argv[1]).exists():
            batchmode = False
        elif '--batch' in sys.argv or '--no-batch' in sys.argv:
            batchmode = '--batch' in sys.argv
        elif '--conversion-only' in sys.argv:
            doanalysis = False
        elif '--analysis-only' in sys.argv:
            doconversion = False
        elif '--keywords' in ' '.join(sys.argv):
            for arg in sys.argv:
                if arg.startswith('--keywords='):
                    keywords_str = arg.split('=', 1)[1]
                    grouping_keywords = [kw.strip() for kw in keywords_str.split(',') if kw.strip()]
        
        if batchmode:
            configure_logging(verbose=verbose)
            results = process_batch(
                do_conversion=doconversion,
                do_analysis=doanalysis,
                grouping_keywords=grouping_keywords,
                verbose=verbose
            )
            return
    
    # Interactive menu mode
    while True:
        try:
            choice = show_main_menu()
            
            if choice == '1':
                handle_single_convert(use_ocr=False)
            
            elif choice == '2':
                handle_single_convert_analyze()
            
            elif choice == '3':
                handle_analyze_only()
            
            elif choice == '4':
                handle_single_convert(use_ocr=True)
            
            elif choice == '5':
                # Batch convert
                configure_logging(verbose=False)
                results = process_batch(do_conversion=True, do_analysis=False, verbose=False)
            
            elif choice == '6':
                # Batch convert + analyze
                grouping_keywords = get_grouping_keywords()
                configure_logging(verbose=False)
                results = process_batch(
                    do_conversion=True,
                    do_analysis=True,
                    grouping_keywords=grouping_keywords,
                    verbose=False,
                    create_consolidated_report=True
                )

            elif choice == '7':
                # Batch analyze
                grouping_keywords = get_grouping_keywords()
                configure_logging(verbose=False)
                results = process_batch(
                    do_conversion=False,
                    do_analysis=True,
                    grouping_keywords=grouping_keywords,
                    verbose=False,
                    create_consolidated_report=True
                )
            
            elif choice == '8':
                show_system_status()
            
            elif choice == '9':
                settings_choice = show_settings_menu()
                # TODO: Implement settings handlers
                print("\n⚠️  Settings persistence coming soon!")
                input("Press Enter to continue...")
            
            elif choice == '0':
                print("\n" + "="*80)
                print(" " * 30 + "👋 Goodbye!")
                print("="*80 + "\n")
                break
            
            else:
                print("\n❌ Invalid choice. Please select 0-9.")
                input("Press Enter to continue...")
        
        except KeyboardInterrupt:
            print("\n\n" + "="*80)
            print(" " * 28 + "⚠️  Interrupted")
            print("="*80)
            confirm = input("\nDo you want to exit? (y/n): ").strip().lower()
            if confirm == 'y':
                break
        
        except Exception as e:
            print("\n" + "="*80)
            print(" " * 28 + "❌ ERROR")
            print("="*80)
            print(f"  {e}")
            print("="*80)
            import traceback
            traceback.print_exc()
            input("\nPress Enter to continue...")


if __name__ == "__main__":
    main()
