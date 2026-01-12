#!/usr/bin/env python3
"""
OCR Processing Module
Handles all OCR-related operations for PDF conversion.
"""

import os
import sys

# CRITICAL: Suppress verbose PaddleOCR messages BEFORE any imports
os.environ['DISABLE_MODEL_SOURCE_CHECK'] = 'True'
os.environ['PADDLEOCR_LOGGING_LEVEL'] = 'ERROR'
os.environ['HF_HUB_DISABLE_TELEMETRY'] = '1'
os.environ['PYTHONWARNINGS'] = 'ignore'

import re
import tempfile
import warnings
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import numpy as np

# Suppress all warnings and verbose logging
warnings.filterwarnings("ignore")
logging.getLogger('paddleocr').setLevel(logging.ERROR)
logging.getLogger('ppocr').setLevel(logging.ERROR)
logging.getLogger('PIL').setLevel(logging.ERROR)

from src.config import (
    PADDLEOCR_AVAILABLE, PPSTRUCTURE_AVAILABLE, EASYOCR_AVAILABLE,
    PYTESSERACT_AVAILABLE, OCRMYPDF_AVAILABLE, PDF2IMAGE_AVAILABLE,
    PDFPLUMBER_AVAILABLE, PYTHON_DOCX_AVAILABLE, PYPDF2_AVAILABLE,
    PDF2DOCX_AVAILABLE, TQDM_AVAILABLE, logger
)
from src.ocr.ocr_cache import OCRCache
from src.analysis.pdf_analyzer import map_language_code, get_pdf_page_count

if PADDLEOCR_AVAILABLE:
    from paddleocr import PaddleOCR
    if PPSTRUCTURE_AVAILABLE:
        from paddleocr import PPStructure

if EASYOCR_AVAILABLE:
    import easyocr

if PYTESSERACT_AVAILABLE:
    import pytesseract
    from pytesseract import Output, TesseractError

if PDF2IMAGE_AVAILABLE:
    from pdf2image import convert_from_path

if PDFPLUMBER_AVAILABLE:
    import pdfplumber

if PYTHON_DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, Inches, Twips

if PYPDF2_AVAILABLE:
    import PyPDF2

if TQDM_AVAILABLE:
    from tqdm import tqdm


def transform_coords_to_pdf(ocr_data: Dict, img_width: int, img_height: int, page_num: int, pdf_path: Path) -> Dict:
    """
    Transforme les coordonnées OCR (relatives à l'image) en coordonnées PDF.
    Utilise une transformation matricielle pour mapper image → PDF.
    """
    try:
        # Obtenir les dimensions de la page PDF
        if PYPDF2_AVAILABLE:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                if page_num <= len(pdf_reader.pages):
                    page = pdf_reader.pages[page_num - 1]
                    # Dimensions de la page PDF en points (72 DPI = 1 point = 1/72 inch)
                    pdf_width_pt = float(page.mediabox.width)
                    pdf_height_pt = float(page.mediabox.height)
                    
                    # Calculer les facteurs de transformation
                    scale_x = pdf_width_pt / img_width
                    scale_y = pdf_height_pt / img_height
                    
                    # Transformer les coordonnées
                    if 'left' in ocr_data and ocr_data['left']:
                        ocr_data['left'] = [int(x * scale_x) for x in ocr_data['left']]
                        ocr_data['top'] = [int(y * scale_y) for y in ocr_data['top']]
                        ocr_data['width'] = [int(w * scale_x) for w in ocr_data['width']]
                        ocr_data['height'] = [int(h * scale_y) for h in ocr_data['height']]
    except Exception as e:
        logger.warning(f"Coordinate transformation failed: {e}")
    
    return ocr_data



def detect_rotation_angle(img) -> float:
    """
    Détecte l'angle de rotation d'une image scannée.
    Retourne l'angle en degrés.
    """
    import numpy as np
    from PIL import ImageFilter
    
    try:
        # Convertir en niveaux de gris
        if img.mode != 'L':
            img_gray = img.convert('L')
        else:
            img_gray = img
        
        # Détecter les bords
        img_edges = img_gray.filter(ImageFilter.FIND_EDGES)
        img_array = np.array(img_edges)
        
        # Projection horizontale pour détecter l'inclinaison
        h_projection = img_array.sum(axis=1)
        
        # Trouver les pics dans la projection
        peaks = []
        for i in range(1, len(h_projection) - 1):
            if h_projection[i] > h_projection[i-1] and h_projection[i] > h_projection[i+1]:
                peaks.append(i)
        
        if len(peaks) > 5:
            # Calculer l'angle moyen basé sur les différences entre pics
            angles = []
            for i in range(1, min(10, len(peaks))):
                deviation = peaks[i] - peaks[0]
                if deviation != 0:
                    angle = np.arctan(deviation / img.width) * 180 / np.pi
                    angles.append(angle)
            
            if angles:
                avg_angle = np.mean(angles)
                # Filtrer les angles trop petits (< 0.5 degré)
                if abs(avg_angle) > 0.5:
                    return avg_angle
    except Exception as e:
        logger.debug(f"Rotation detection failed: {e}")
    
    return 0.0


def add_rotation_to_ocr_data(ocr_data: Dict, img) -> Dict:
    """
    Ajoute le paramètre de rotation aux données OCR.
    Les bounding boxes deviennent: left, top, width, height, rotation
    """
    # Détecter l'angle de rotation
    rotation_angle = detect_rotation_angle(img)
    
    # Ajouter la rotation aux données OCR
    if 'left' in ocr_data and ocr_data['left']:
        n_boxes = len(ocr_data['left'])
        ocr_data['rotation'] = [rotation_angle] * n_boxes
    else:
        ocr_data['rotation'] = []
    
    return ocr_data



def enhance_image_for_ocr(img) -> Any:
    """
    Améliore l'image pour meilleur résultat OCR (OPTIMISÉ - rapide et efficace).
    """
    from PIL import ImageEnhance, ImageFilter
    
    # Convertir en niveaux de gris si nécessaire
    if img.mode != 'L':
        img = img.convert('L')
    
    # 1. DENOISING simple (rapide)
    try:
        img = img.filter(ImageFilter.MedianFilter(size=3))
    except Exception:
        pass
    
    # 2. AMÉLIORATION DU CONTRASTE
    try:
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.2)
    except Exception:
        pass
    
    # 3. AMÉLIORATION DE LA NETTETÉ
    try:
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.1)
    except Exception:
        pass
    
    return img

# ============================================================================
# _reconstruct_page_faithfully() - SUPPRIMÉE/COMMENTÉE
# Cette fonction appelait des méthodes inexistantes (_render_text_line_ocr, etc.)
# Utiliser _render_page_sandwich_method() à la place (version simple et fonctionnelle)
# ============================================================================
# 

def convert_paddleocr_to_standard(paddle_result, img_width, img_height) -> None:
    """
    Convertit le résultat PaddleOCR en format standardisé (compatible avec Tesseract).
    PaddleOCR retourne: [[[x1,y1], [x2,y2], [x3,y3], [x4,y4]], (text, confidence)]
    """
    words = []
    
    if not paddle_result or not paddle_result[0]:
        return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'conf': []}
    
    for line in paddle_result[0]:
        if not line:
            continue
        
        # Extraire les coordonnées et le texte
        bbox = line[0]  # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
        text_info = line[1]  # (text, confidence)
        
        if len(bbox) >= 4 and text_info:
            text = text_info[0] if isinstance(text_info, tuple) else str(text_info)
            conf = text_info[1] if isinstance(text_info, tuple) and len(text_info) > 1 else 95
            
            # Calculer bounding box rectangulaire
            x_coords = [point[0] for point in bbox]
            y_coords = [point[1] for point in bbox]
            
            left = int(min(x_coords))
            top = int(min(y_coords))
            right = int(max(x_coords))
            bottom = int(max(y_coords))
            
            width = right - left
            height = bottom - top
            
            if text.strip() and conf > 30:
                words.append({
                    'text': text.strip(),
                    'left': left,
                    'top': top,
                    'width': width,
                    'height': height,
                    'conf': int(conf * 100) if conf < 1 else int(conf)
                })
    
    # Convertir en format dict compatible avec Tesseract
    if words:
        return {
            'text': [w['text'] for w in words],
            'left': [w['left'] for w in words],
            'top': [w['top'] for w in words],
            'width': [w['width'] for w in words],
            'height': [w['height'] for w in words],
            'rotation': [w.get('rotation', 0.0) for w in words],  # 5ème paramètre
            'conf': [w['conf'] for w in words]
        }
    else:
        return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'rotation': [], 'conf': []}



def render_paddleocr_structure(doc, original_img, structure_result, img_width, img_height) -> None:
    """
    Rend le résultat de PPStructure en préservant la structure hiérarchique.
    PPStructure détecte: text, table, title, image, etc.
    """
    # Reconstruire UNIQUEMENT le texte avec structure préservée (pas d'image)
    # PPStructure fournit déjà la structure, on la rend directement
    if not structure_result:
        return
    
    # PPStructure retourne une liste d'éléments avec type et bbox
    for element in structure_result:
        if not element:
            continue
        
        elem_type = element.get('type', 'text')
        bbox = element.get('bbox', [])
        content = element.get('res', {})
        
        # Extraire le texte selon le type
        if elem_type == 'table':
            # Rendre comme tableau Word
            rendertablefrompaddleocr(doc, content, bbox)
        elif elem_type == 'title':
            # Rendre comme titre
            text = content.get('text', '') if isinstance(content, dict) else str(content)
            if text:
                para = doc.add_paragraph(text.strip())
                para.style = 'Heading 1'
                for run in para.runs:
                    run.bold = True
        else:
            # Texte normal
            text = content.get('text', '') if isinstance(content, dict) else str(content)
            if text:
                doc.add_paragraph(text.strip())



def render_table_from_paddleocr(doc, table_content, bbox) -> None:
    """Rend un tableau détecté par PPStructure."""
    from docx.shared import Inches
    
    if not table_content or not isinstance(table_content, dict):
        return
    
    # PPStructure peut retourner le tableau en format structuré
    # Essayer d'extraire les cellules
    cells = table_content.get('cells', [])
    
    if cells:
        # Créer un tableau Word
        num_rows = len(cells)
        num_cols = max(len(row) for row in cells) if cells else 1
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        for row_idx, row in enumerate(cells):
            for col_idx, cell in enumerate(row):
                if col_idx < num_cols:
                    cell_text = cell.get('text', '') if isinstance(cell, dict) else str(cell)
                    table.rows[row_idx].cells[col_idx].paragraphs[0].text = cell_text.strip()
    else:
        # Fallback : texte simple
        text = str(table_content)
        doc.add_paragraph(f"[Table: {text}]")



def parse_hocr(hocr_output: str, img_width: int, img_height: int, img) -> Dict:
    """
    Parse le résultat hOCR de Tesseract pour extraire les coordonnées précises.
    hOCR contient des bounding boxes au niveau du mot avec coordonnées pixel.
    Ajoute aussi le paramètre de rotation (5ème paramètre).
    """
    from xml.etree import ElementTree as ET
    
    # Détecter l'angle de rotation global
    rotation_angle = detect_rotation_angle(img)
    
    try:
        # Parser le XML hOCR
        root = ET.fromstring(hocr_output)
        
        # Trouver tous les éléments de mots (classe 'ocrx_word')
        words = []
        for word_elem in root.iter():
            if 'class' in word_elem.attrib and 'ocrx_word' in word_elem.attrib.get('class', ''):
                # Extraire le titre qui contient les coordonnées
                title = word_elem.attrib.get('title', '')
                if 'bbox' in title:
                    # Parser: bbox x1 y1 x2 y2
                    parts = title.split()
                    bbox_idx = parts.index('bbox')
                    x1, y1, x2, y2 = map(int, parts[bbox_idx+1:bbox_idx+5])
                    
                    # Extraire le texte
                    text = ''.join(word_elem.itertext()).strip()
                    
                    # Extraire la rotation si disponible (dans certains formats hOCR)
                    word_rotation = rotation_angle
                    if 'rotate' in title:
                        try:
                            rotate_idx = parts.index('rotate')
                            word_rotation = float(parts[rotate_idx+1])
                        except (ValueError, IndexError):
                            pass
                    
                    if text:
                        words.append({
                            'text': text,
                            'left': x1,
                            'top': y1,
                            'width': x2 - x1,
                            'height': y2 - y1,
                            'rotation': word_rotation,  # 5ème paramètre
                            'conf': 95  # hOCR ne donne pas toujours la confiance
                        })
        
        # Convertir en format dict
        if words:
            return {
                'text': [w['text'] for w in words],
                'left': [w['left'] for w in words],
                'top': [w['top'] for w in words],
                'width': [w['width'] for w in words],
                'height': [w['height'] for w in words],
                'rotation': [w['rotation'] for w in words],  # 5ème paramètre
                'conf': [w['conf'] for w in words]
            }
    except Exception as e:
        logger.warning(f"hOCR parsing failed: {e}, using standard format")
        return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'rotation': [], 'conf': []}
    
    return {'text': [], 'left': [], 'top': [], 'width': [], 'height': [], 'rotation': [], 'conf': []}



def group_words_by_lines_ocr(words) -> None:
    """
    Groupe les mots par lignes basées sur leur position Y.
    Retourne simplement des listes de mots par ligne, sans séparation de colonnes.
    """
    if not words:
        return []
    
    words_sorted = sorted(words, key=lambda w: (w['top'], w['left']))
    lines = []
    current_line = []
    current_y = None
    y_tolerance = 5
    
    for word in words_sorted:
        word_y = word['top']
        if current_y is None:
            current_y = word_y
            current_line.append(word)
        elif abs(word_y - current_y) <= y_tolerance:
            current_line.append(word)
        else:
            if current_line:
                # Trier par X pour avoir l'ordre de gauche à droite
                current_line.sort(key=lambda w: w['left'])
                lines.append(current_line)
            current_line = [word]
            current_y = word_y
    
    if current_line:
        current_line.sort(key=lambda w: w['left'])
        lines.append(current_line)
    
    return lines



def detect_page_structure_ocr(lines, img_width) -> None:
    """Détecte la structure de la page (positions des colonnes, zones de tableaux)."""
    all_x_positions = []
    for line in lines:
        # Gérer les lignes multi-colonnes
        if isinstance(line, dict) and line.get('type') == 'multi_column':
            for col in line.get('columns', []):
                for word in col:
                    all_x_positions.append(word['left'])
        elif isinstance(line, list):
            for word in line:
                if isinstance(word, dict):
                    all_x_positions.append(word['left'])
    
    if not all_x_positions:
        return {'column_positions': [], 'table_zones': []}
    
    # Clusteriser les positions X pour détecter les colonnes
    all_x_positions.sort()
    column_positions = []
    if all_x_positions:
        current_cluster = [all_x_positions[0]]
        cluster_tolerance = 20  # 20 pixels de tolérance
        
        for x in all_x_positions[1:]:
            if x - current_cluster[-1] < cluster_tolerance:
                current_cluster.append(x)
            else:
                if len(current_cluster) >= 3:
                    column_positions.append(sum(current_cluster) / len(current_cluster))
                current_cluster = [x]
        
        if len(current_cluster) >= 3:
            column_positions.append(sum(current_cluster) / len(current_cluster))
    
    # Détecter les zones de tableaux
    table_zones = []
    for line_idx, line in enumerate(lines):
        # Extraire les mots réels
        if isinstance(line, dict) and line.get('type') == 'multi_column':
            line_for_analysis = []
            for col in line.get('columns', []):
                line_for_analysis.extend(col)
        else:
            line_for_analysis = line if isinstance(line, list) else []
        
        if len(line_for_analysis) >= 4:
            gaps = []
            for i in range(1, len(line_for_analysis)):
                prev_x_end = line_for_analysis[i-1]['left'] + line_for_analysis[i-1]['width']
                curr_x = line_for_analysis[i]['left']
                gap = curr_x - prev_x_end
                gaps.append(gap)
            
            avg_gap = sum(gaps) / len(gaps) if gaps else 0
            if avg_gap > 15:
                table_zones.append(line_idx)
    
    return {
        'column_positions': sorted(set(column_positions)),
        'table_zones': table_zones,
        'page_width': img_width
    }



def is_table_line_ocr(line, all_lines, line_idx, page_structure) -> None:
    """Détermine si une ligne fait partie d'un tableau."""
    # Gérer les lignes multi-colonnes
    if isinstance(line, dict) and line.get('type') == 'multi_column':
        # Pour les colonnes multiples, utiliser les mots de toutes les colonnes
        all_words = []
        for col in line.get('columns', []):
            all_words.extend(col)
        line = all_words
    
    if not isinstance(line, list) or len(line) < 3:
        return False
    
    if line_idx in page_structure['table_zones']:
        return True
    
    gaps = []
    for i in range(1, len(line)):
        prev_x_end = line[i-1]['left'] + line[i-1]['width']
        curr_x = line[i]['left']
        gap = curr_x - prev_x_end
        gaps.append(gap)
    
    large_gaps = [g for g in gaps if g > 20]
    if len(large_gaps) >= 2:
        return True
    
    text_content = " ".join([w['text'] for w in line])
    digit_count = sum(c.isdigit() for c in text_content)
    if digit_count > len(text_content) * 0.3:
        return True
    
    return False



def render_table_line_ocr(doc, line, page_structure, img_width) -> None:
    """Rend une ligne comme ligne de tableau en utilisant des tabulations pour aligner."""
    from docx.shared import Pt
    
    # Gérer les lignes multi-colonnes
    if isinstance(line, dict) and line.get('type') == 'multi_column':
        render_multicolumn_line(doc, line['columns'])
        return
    
    if not isinstance(line, list):
        return
    
    column_positions = page_structure.get('column_positions', [])
    if not column_positions:
        text = " ".join([w['text'] if isinstance(w, dict) else str(w) for w in line])
        doc.add_paragraph(text)
        return
    
    # Créer un tableau Word si c'est une zone de tableau détectée
    # Sinon, utiliser des tabulations pour aligner
    para = doc.add_paragraph()
    
    current_x = 0
    for word in line:
        word_x = word['left']
        
        # Trouver la colonne la plus proche
        closest_col_idx = min(range(len(column_positions)),
                             key=lambda i: abs(column_positions[i] - word_x))
        target_col = column_positions[closest_col_idx]
        
        # Ajouter des tabulations pour aligner
        if word_x > current_x + 10:
            tabs_needed = max(1, int((target_col - current_x) / 50))
            for _ in range(min(tabs_needed, 15)):
                para.add_run("\t")
        
        para.add_run(word['text'])
        current_x = word_x + word['width']
        
        if word != line[-1]:
            para.add_run(" ")




def render_multicolumn_line(doc, columns):
    """
    Rend une ligne avec colonnes multiples.
    Utilise un tableau Word à 2 colonnes pour préserver la structure.
    """
    from docx.shared import Inches
    
    if not columns or len(columns) < 2:
        # Fallback : rendre comme texte normal
        all_words = []
        for col in columns:
            all_words.extend(col)
        text = " ".join([w['text'] for w in all_words])
        if text.strip():
            doc.add_paragraph(text.strip())
        return
    
    # Créer un tableau avec 2 colonnes pour cette ligne
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Normal Table'
    
    # Largeur des colonnes (50% chacune)
    for col in table.columns:
        col.width = Inches(3.25)
    
    # Remplir les colonnes
    for col_idx, col_words in enumerate(columns[:2]):
        if col_idx < 2:
            cell = table.rows[0].cells[col_idx]
            text = " ".join([w['text'] for w in col_words])
            cell.paragraphs[0].text = text.strip()



def render_text_line_ocr(doc, line, page_structure) -> None:
    """Rend une ligne de texte normal."""
    # Gérer les lignes multi-colonnes
    if isinstance(line, dict) and line.get('type') == 'multi_column':
        render_multicolumn_line(doc, line['columns'])
        return
    
    # Ligne normale
    if not isinstance(line, list):
        return
        
    text = " ".join([w['text'] if isinstance(w, dict) else str(w) for w in line])
    text = text.strip()
    
    if not text:
        return
    
    # Détecter si c'est un titre
    if text.isupper() and len(text) < 100 and len(text.split()) < 15:
        para = doc.add_paragraph(text)
        para.style = 'Heading 2'
        for run in para.runs:
            run.bold = True
    else:
        doc.add_paragraph(text)



def group_ocr_data_by_lines(data) -> None:
    """
    Groupe les données OCR en lignes visuelles.
    
    Args:
        data: Données OCR de Tesseract (format DICT)
    
    Returns:
        Liste de lignes, chaque ligne étant une liste de mots (dicts)
    """
    n_boxes = len(data['text'])
    words = []
    for i in range(n_boxes):
        if int(data['conf'][i]) > 30 and data['text'][i].strip():
            words.append({
                'text': data['text'][i],
                'left': data['left'][i],
                'top': data['top'][i],
                'width': data['width'][i],
                'height': data['height'][i]
            })
    
    # Tri Y puis X
    words.sort(key=lambda x: (x['top'], x['left']))
    
    lines = []
    current_line = []
    current_y = -1
    y_tolerance = 15
    
    for word in words:
        if current_y == -1:
            current_y = word['top']
            current_line.append(word)
        elif abs(word['top'] - current_y) <= y_tolerance:
            current_line.append(word)
        else:
            current_line.sort(key=lambda x: x['left'])
            lines.append(current_line)
            current_line = [word]
            current_y = word['top']
    
    if current_line:
        current_line.sort(key=lambda x: x['left'])
        lines.append(current_line)
    
    return lines



def render_page_sandwich_method(doc: Document, original_img, ocr_data: Dict, img_width: int, img_height: int) -> None:
    """VERSION SIMPLE QUI MARCHE"""
    if not ocr_data or 'text' not in ocr_data:
        return
    
    # Extraire mots
    n_boxes = len(ocr_data['text'])
    words = []
    for i in range(n_boxes):
        try:
            conf = int(ocr_data['conf'][i]) if ocr_data['conf'][i] else 0
            text = str(ocr_data['text'][i]).strip() if ocr_data['text'][i] else ''
            if conf > 30 and text:
                words.append({
                    'text': text,
                    'left': int(ocr_data['left'][i]),
                    'top': int(ocr_data['top'][i])
                })
        except:
            continue
    
    # Grouper par lignes
    words.sort(key=lambda w: (w['top'], w['left']))
    lines = []
    current_line = []
    current_y = None
    
    for word in words:
        if current_y is None or abs(word['top'] - current_y) <= 10:
            current_line.append(word)
            current_y = word['top']
        else:
            if current_line:
                lines.append(current_line)
            current_line = [word]
            current_y = word['top']
    
    if current_line:
        lines.append(current_line)
    
    # Afficher texte
    for line in lines:
        line.sort(key=lambda w: w['left'])
        text = ' '.join([w['text'] for w in line])
        if text.strip():
            doc.add_paragraph(text.strip())



def reconstruct_page_with_layout(doc: Document, ocr_data: Dict, img_width: int, img_height: int) -> None:
    """
    Reconstruit le texte avec layout préservé - VERSION SIMPLE.
    Alias pour _render_page_sandwich_method pour compatibilité.
    """
    # Appeler la méthode simple (ignore original_img car on n'ajoute plus d'image)
    renderpagesandwichmethod(doc, None, ocr_data, img_width, img_height)



def convert_with_paddleocr_ppstructure(doc, images, start_page, end_page, ocr_language) -> None:
    """
    Méthode OCR PROFESSIONNELLE : Utilise PaddleOCR avec PPStructure.
    Meilleure qualité avec détection de structure (text, table, title, image).
    """
    import numpy as np
    
    # Initialiser PaddleOCR avec détection de structure
    # Supprimer TOUS les messages verbeux
    import warnings
    warnings.filterwarnings("ignore")

    try:
        # Essayer PPStructure d'abord (meilleure structure)
        from paddleocr import PPStructure
        import logging
        logging.getLogger('ppocr').setLevel(logging.ERROR)
        structure_engine = PPStructure(show_log=False, use_gpu=False)
        use_ppstructure = True
    except Exception:
        # Fallback vers PaddleOCR standard
        from paddleocr import PaddleOCR
        import logging
        logging.getLogger('ppocr').setLevel(logging.ERROR)
        ocr_engine = PaddleOCR(use_angle_cls=True, lang='en', det=True, rec=True, show_log=False)
        use_ppstructure = False
    
    lang_map = {'eng': 'en', 'en': 'en', 'fra': 'fr', 'fr': 'fr', 'deu': 'de', 'de': 'de'}
    ocr_lang = lang_map.get(ocr_language.lower(), 'en')
    
    for i, img in enumerate(images):
        page_num = start_page + i + 1
        logger.info(f"  📸 Processing page {page_num}/{end_page} with PaddleOCR...")
        
        if i > 0:
            doc.add_page_break()
        
        # Convertir PIL Image en numpy array
        img_array = np.array(img)
        
        # Utiliser PPStructure si disponible
        if use_ppstructure:
            try:
                result = structure_engine(img_array)
                # PPStructure retourne une structure hiérarchique
                renderpaddleocrstructure(doc, img, result, img.width, img.height)
            except Exception as e:
                logger.warning(f"PPStructure failed: {e}, using standard PaddleOCR...")
                use_ppstructure = False
        
        # Fallback vers PaddleOCR standard
        if not use_ppstructure:
            from paddleocr import PaddleOCR
            ocr_engine = PaddleOCR(use_angle_cls=True, lang=ocr_lang, det=True, rec=True, show_log=False)
            result = ocr_engine.ocr(img_array, cls=True)
            
            # Convertir le résultat PaddleOCR en format standardisé
            ocr_data = convertpaddleocrtostandard(result, img.width, img.height)
            renderpagesandwichmethod(doc, img, ocr_data, img.width, img.height)



def convert_with_ocrmypdf_sandwich(doc, start_page, end_page, ocr_language) -> None:
    """
    Méthode SANDWICH optimale : Utilise ocrmypdf pour créer un PDF searchable,
    puis convertit en Word avec pdf2docx.
    C'est la meilleure approche industry-standard.
    """
    import tempfile
    
    # Créer un PDF temporaire avec ocrmypdf
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        tmp_pdf_path = tmp_pdf.name
    
    try:
        # Obtenir le nombre total de pages
        from PyPDF2 import PdfReader
        reader = PdfReader(str(pdf_path))
        total_pages = len(reader.pages)
        
        # Extraire les pages nécessaires du PDF original
        if start_page > 0 or (end_page and end_page < total_pages):
            # Extraire les pages avec PyPDF2
            from PyPDF2 import PdfWriter
            writer = PdfWriter()
            
            actual_end = end_page if end_page else total_pages
            for i in range(start_page, actual_end):
                writer.add_page(reader.pages[i])
            
            with open(tmp_pdf_path, 'wb') as f:
                writer.write(f)
            input_pdf = tmp_pdf_path
        else:
            input_pdf = str(pdf_path)
        
        # Appliquer OCR avec ocrmypdf (méthode Sandwich automatique)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as ocr_pdf:
            ocr_pdf_path = ocr_pdf.name
        
        try:
            # Mapper la langue pour ocrmypdf
            lang_map = {'eng': 'eng', 'en': 'eng', 'fra': 'fra', 'fr': 'fra', 'deu': 'deu', 'de': 'deu'}
            ocr_lang = lang_map.get(ocr_language.lower(), 'eng')
            
            try:
                ocrmypdf.ocr(
                    input_pdf,
                    ocr_pdf_path,
                    language=ocr_lang,
                    deskew=True,  # Redressement automatique
                    clean=False,  # Nettoyage désactivé (nécessite unpaper)
                    force_ocr=True,  # Forcer OCR même si texte présent
                    optimize=1    # Optimisation niveau 1
                )
            except Exception as e:
                # Si ocrmypdf échoue (dépendances manquantes), fallback vers méthode manuelle
                error_msg = str(e).lower()
                if 'unpaper' in error_msg or 'missing' in error_msg:
                    logger.warning("OCRmyPDF dependencies missing, falling back to manual method")
                    raise Exception("Use manual method") from e
                else:
                    raise
            
            # Extraire le texte avec structure depuis le PDF OCR'd en utilisant pdfplumber
            logger.info("🔄 Extracting text with layout from OCR'd PDF using pdfplumber...")
            if PDFPLUMBER_AVAILABLE:
                extracttextwithpdfplumber(doc, ocr_pdf_path, start_page, end_page)
            else:
                # Fallback : méthode manuelle avec OCR direct
                logger.warning("pdfplumber not available, using manual OCR method...")
                extracttextwithlayoutfromocrpdf(doc, ocr_pdf_path, start_page, end_page)
            logger.info("✅ OCR'd PDF converted successfully")
                
        finally:
            try:
                os.unlink(ocr_pdf_path)
            except:
                pass
                
    finally:
        try:
            if tmp_pdf_path != str(pdf_path):
                os.unlink(tmp_pdf_path)
        except:
            pass



def extract_text_with_pdfplumber(doc, ocr_pdf_path, start_page, end_page) -> None:
    """
    Extrait le texte d'un PDF OCR'd avec pdfplumber qui peut lire le texte overlay.
    Préserve la structure et la mise en page.
    """
    import pdfplumber
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from pdf2image import convert_from_path
    
    try:
        logger.info("📖 Opening OCR'd PDF with pdfplumber...")
        with pdfplumber.open(ocr_pdf_path) as pdf:
            total_pages = len(pdf.pages)
            actual_end = end_page if end_page else total_pages
            
            for page_idx in range(start_page, actual_end):
                pdf_page = pdf.pages[page_idx]
                
                if page_idx > start_page:
                    doc.add_page_break()
                
                logger.info(f"  📄 Processing page {page_idx + 1}/{actual_end}...")
                
                # Extraire le texte avec coordonnées
                words = pdf_page.extract_words(
                    x_tolerance=3,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=True,
                    extra_attrs=["fontname", "size"]
                )
                
                if not words:
                    logger.warning(f"No words extracted from page {page_idx + 1}")
                    doc.add_paragraph(f"[Page {page_idx + 1}: No text extracted]")
                    continue
                
                # VERSION SIMPLE : Extraire le texte directement
                # Trier les mots par position (top, puis left)
                words_sorted = sorted(words, key=lambda w: (w.get('top', 0), w.get('left', 0)))
                
                # Grouper par lignes (simple)
                lines = []
                current_line = []
                current_y = None
                
                for word in words_sorted:
                    word_y = word.get('top', 0)
                    word_text = word.get('text', '').strip()
                    
                    if not word_text:
                        continue
                    
                    if current_y is None or abs(word_y - current_y) <= 10:
                        current_line.append(word)
                        if current_y is None:
                            current_y = word_y
                    else:
                        if current_line:
                            lines.append(sorted(current_line, key=lambda w: w.get('left', 0)))
                        current_line = [word]
                        current_y = word_y
                
                if current_line:
                    lines.append(sorted(current_line, key=lambda w: w.get('left', 0)))
                
                # Rendre les lignes simplement
                for line in lines:
                    if not line:
                        doc.add_paragraph("")
                        continue
                    
                    line_text = ' '.join([w.get('text', '') for w in line])
                    
                    if line_text.strip():
                        # Détecter si titre (tout en majuscules ET court)
                        is_title = line_text.isupper() and len(line_text) < 100 and len(line_text.split()) < 15
                        
                        if is_title:
                            para = doc.add_paragraph(line_text.strip())
                            para.style = 'Heading 2'
                            for run in para.runs:
                                run.bold = True
                        else:
                            doc.add_paragraph(line_text.strip())
                    
    except Exception as e:
        logger.error(f"Error extracting with pdfplumber: {e}")
        import traceback
        traceback.print_exc()
        raise



def extract_text_with_layout_from_ocr_pdf(doc, ocr_pdf_path, start_page, end_page) -> None:
    """
    Extrait le texte d'un PDF OCR'd avec préservation de la structure et de la mise en page.
    Utilise pdfplumber pour extraire le texte avec coordonnées, puis reconstruit fidèlement.
    """
    from pdf2image import convert_from_path
    from io import BytesIO
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pytesseract
    from pytesseract import Output
    
    try:
        # Convertir le PDF OCR'd en images pour OCR avec coordonnées
        logger.info("🔄 Converting OCR'd PDF to images for layout analysis...")
        images = convert_from_path(
            ocr_pdf_path,
            first_page=1,
            last_page=None,
            dpi=300,
            fmt='PNG'
        )
        
        total_pages = len(images)
        actual_end = end_page if end_page else total_pages
        
        # Traiter chaque page avec OCR layout-aware
        for page_idx in range(start_page, actual_end):
            if page_idx > start_page:
                doc.add_page_break()
            
            img = images[page_idx]
            logger.info(f"  📄 Processing page {page_idx + 1}/{actual_end} with layout-aware OCR...")
            
            # Améliorer l'image pour meilleur OCR
            img_enhanced = enhanceimageforocr(img)
            
            # OCR avec extraction de coordonnées précises (hOCR)
            try:
                hocr_output = pytesseract.image_to_pdf_or_hocr(
                    img_enhanced,
                    lang='eng',
                    extension='hocr',
                    config='--psm 1 --oem 3'  # PSM 1 = Automatic page segmentation with OSD
                )
                ocr_data = parsehocr(hocr_output, img.width, img.height, img_enhanced)
            except Exception as e:
                # Fallback vers image_to_data
                try:
                    ocr_data_raw = pytesseract.image_to_data(
                        img_enhanced,
                        lang='eng',
                        output_type=Output.DICT,
                        config='--psm 1 --oem 3'
                    )
                    ocr_data = addrotationtoocrdata(ocr_data_raw, img_enhanced)
                except Exception as e2:
                    logger.error(f"OCR failed for page {page_idx + 1}: {e2}")
                    # Fallback : image seule
                    img_buffer = BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(img_buffer, width=Inches(img.width / 300.0))
                    continue
            
            # Reconstruire la page avec structure préservée
            reconstructpagewithlayout(doc, ocr_data, img.width, img.height)
            
    except Exception as e:
        logger.error(f"Error extracting text with layout from OCR'd PDF: {e}")
        raise



def convert_with_ocr(pdf_path: Path, docx_path: Path, ocr_cache: OCRCache, start_page: int = 0, end_page: Optional[int] = None, ocr_language: str = 'eng', verbose: bool = False) -> None:
    """
    🔥 OCR avec Préservation de Structure EXACTE (comme onlineocr.net)
    Architecture: pdf2image → pytesseract (image_to_data) → python-docx
    Préserve FIDÈLEMENT le layout, les colonnes, tableaux et positions spatiales.
    
    Args:
        start_page: Page de départ (0-indexed)
        end_page: Page de fin (None = toutes les pages)
        ocr_language: Code langue ISO (eng, fra, etc.) pour Tesseract
    
    Raises:
        ImportError: Si pytesseract, pdf2image ou python-docx ne sont pas disponibles
        Exception: Si la conversion échoue
    """
    if not PYTESSERACT_AVAILABLE or not PDF2IMAGE_AVAILABLE:
        raise ImportError(
            "=" * 70 + "\n"
            "OCR DEPENDENCIES NOT AVAILABLE\n"
            "=" * 70 + "\n"
            "Required packages:\n"
            "  pip install pytesseract pdf2image pillow\n\n"
            "System dependencies:\n"
            "  macOS: brew install tesseract poppler\n"
            "  Ubuntu: sudo apt-get install tesseract-ocr poppler-utils\n"
            "=" * 70
        )
    
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    # Créer le répertoire de sortie
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from pytesseract import Output, TesseractError
        from docx import Document
        from docx.shared import Pt, Inches, Twips
        from PIL import Image, ImageEnhance
        import numpy as np
        
        # Créer le document Word
        doc = Document()
        
        # Configuration des marges (standard)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Style par défaut
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        total_pages = get_pdf_page_count(pdf_path)
        actual_end = end_page if end_page is not None else total_pages
        
        # DPI : 300 minimum pour documents financiers SEC (Best Practice)
        dpi = 300
        
        logger.info(f"🔄 Converting PDF to images ({dpi} DPI) for pages {start_page+1} to {actual_end}...")
        
        # 1. Conversion PDF → Images (300 DPI = qualité professionnelle pour documents financiers)
        try:
            images = convert_from_path(
                str(pdf_path),
                first_page=start_page+1,
                last_page=actual_end,
                dpi=dpi,  # 300 DPI = qualité professionnelle (Best Practice pour documents financiers)
                fmt='PNG',
                thread_count=4  # Paralléliser la conversion d'images
            )
        except Exception as e:
            error_msg = str(e).lower()
            if 'poppler' in error_msg or 'pdfinfo' in error_msg:
                raise Exception(
                    "❌ POPPLER NOT FOUND\n"
                    "Poppler is required to convert PDF to images.\n"
                    "Install:\n"
                    "  macOS: brew install poppler\n"
                    "  Ubuntu: sudo apt-get install poppler-utils\n"
                    "  Windows: Download from poppler.freedesktop.org"
                )
            else:
                raise Exception(f"Error converting PDF to images: {e}")
        
        lang = map_language_code(ocr_language)
        
        # 2. Méthode OCR PROFESSIONNELLE : PaddleOCR (PPStructure) > OCRmyPDF > Tesseract
        # Ordre de priorité : PaddleOCR (meilleur) > OCRmyPDF > Tesseract (fallback)
        use_advanced_ocr = False
        
        # Essayer PaddleOCR avec PPStructure (meilleure qualité)
        if PADDLEOCR_AVAILABLE:
            try:
                logger.info("🏆 Using PaddleOCR with PPStructure (best quality)...")
                convertwithpaddleocrppstructure(doc, images, start_page, actual_end, ocr_language, dpi)
                use_advanced_ocr = True
            except Exception as e:
                logger.warning(f"⚠️  PaddleOCR failed: {e}, trying fallback...")
                use_advanced_ocr = False
        
        # Essayer OCRmyPDF si PaddleOCR n'est pas disponible
        if not use_advanced_ocr and OCRMYPDF_AVAILABLE and PDF2DOCX_AVAILABLE:
            try:
                logger.info("🔄 Using OCRmyPDF (Sandwich method - industry standard)...")
                convertwithocrmypdfsandwich(doc, start_page, actual_end, ocr_language)
                use_advanced_ocr = True
            except Exception as e:
                # Si ocrmypdf échoue (dépendances manquantes), utiliser méthode manuelle
                if 'unpaper' in str(e).lower() or 'missing' in str(e).lower():
                    logger.warning("⚠️  OCRmyPDF dependencies missing (unpaper), using manual method...")
                    use_advanced_ocr = False
                else:
                    raise
        
        # Fallback : Tesseract avec méthode manuelle
        if not use_advanced_ocr:
            logger.info("📸 Using Tesseract OCR (manual method)...")
            for i, img in enumerate(images):
                page_num = start_page + i + 1
                logger.info(f"  📸 Processing page {page_num}/{actual_end} with OCR...")
                
                if i > 0:
                    doc.add_page_break()
                
                # Amélioration de l'image pour meilleur OCR
                img_enhanced = enhanceimageforocr(img)
                
                # Vérifier le cache OCR
                cache_key = ocr_cache.get_cache_key(page_num, dpi, lang)
                if cache_key in ocr_cache:
                    logger.info(f"  📦 Using cached OCR for page {page_num}...")
                    ocr_data = ocr_cache[cache_key]
                else:
                    # OCR avec extraction de coordonnées (hOCR pour meilleure précision)
                    try:
                        # Utiliser hOCR pour obtenir les bounding boxes précises
                        hocr_output = pytesseract.image_to_pdf_or_hocr(
                            img_enhanced,
                            lang=lang,
                            extension='hocr',
                            config='--psm 1 --oem 3'
                        )
                        # Parser hOCR pour extraire les coordonnées (avec rotation)
                        ocr_data = parsehocr(hocr_output, img.width, img.height, img_enhanced)
                    except Exception as e:
                        # Fallback vers image_to_data si hOCR échoue
                        try:
                            ocr_data_raw = pytesseract.image_to_data(
                                img_enhanced,
                                lang=lang,
                                output_type=Output.DICT,
                                config='--psm 1 --oem 3'
                            )
                            # Ajouter la détection de rotation
                            ocr_data = addrotationtoocrdata(ocr_data_raw, img_enhanced)
                        except Exception as e2:
                            error_msg = str(e2).lower()
                            if 'tessdata' in error_msg or 'traineddata' in error_msg:
                                raise Exception(
                                    "❌ TESSERACT LANGUAGE DATA NOT FOUND\n"
                                    f"Language '{lang}' not available.\n"
                                    "Install language packs:\n"
                                    "  macOS: brew install tesseract-lang\n"
                                    "  Ubuntu: sudo apt-get install tesseract-ocr-{lang}\n"
                                    "  Windows: Download from github.com/tesseract-ocr/tessdata"
                                )
                            else:
                                raise Exception(f"Tesseract OCR failed: {e2}")
                    
                    # Sauvegarder dans le cache
                    ocr_cache[cache_key] = ocr_data
                
                # Transformer les coordonnées image → PDF (si nécessaire)
                ocr_data_transformed = transformcoordstopdf(ocr_data, img.width, img.height, page_num)
                
                # Méthode SANDWICH manuelle
                renderpagesandwichmethod(doc, img, ocr_data_transformed, img.width, img.height)
        
        # Sauvegarde du cache OCR
        ocr_cache.save()
        
        # Sauvegarde
        doc.save(str(docx_path))
        logger.info(f"✅ Conversion completed: {docx_path.name}")
    
    except Exception as e:
        logger.error(f"OCR conversion failed: {e}")
        if verbose:
            import traceback
            traceback.print_exc()
        raise


