#!/usr/bin/env python3
"""
Document Post-Processing Module
"""

import re
from pathlib import Path
from typing import Dict, Any, List

from src.config import PYTHON_DOCX_AVAILABLE, logger

if PYTHON_DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

def is_financial_number(text: str) -> bool:
    """
    Détecte si un texte est un nombre financier (Best Practice).
    
    Args:
        text: Texte à analyser
    
    Returns:
        bool: True si c'est un nombre financier
    """
    if not text or not text.strip():
        return False
    
    text_clean = text.strip()
    
    # Supprimer les caractères monétaires et ponctuation
    clean = text_clean.replace(',', '').replace('.', '').replace('(', '').replace(')', '').replace('£', '').replace('$', '').replace('€', '').replace('-', '').replace('m', '').replace('k', '').strip()
    
    # Vérifier si c'est principalement des chiffres
    if not clean:
        return False
    
    digit_ratio = sum(c.isdigit() for c in clean) / len(clean)
    
    # C'est un nombre financier si :
    # - Ratio de chiffres > 50%
    # - OU commence par un chiffre ou signe négatif
    # - OU contient des caractères monétaires
    return (digit_ratio > 0.5) or \
           (text_clean[0].isdigit() if text_clean else False) or \
           (text_clean[0] == '(' and text_clean[-1] == ')') or \
           ('£' in text_clean or '$' in text_clean or '€' in text_clean)



def apply_post_processing(docx_path: Path) -> None:
    """Applique le post-traitement après conversion."""
    try:
        if PYTHON_DOCX_AVAILABLE:
            logger.info("Post-processing document...")
            
            # 1. Formatage Q&A
            try:
                fix_qa_formatting(docx_path)
            except Exception as e:
                logger.warning(f"Q&A formatting skipped: {e}")
            
            # 2. Restaurer les hyperlinks
            try:
                restore_hyperlinks(docx_path)
            except Exception as e:
                logger.warning(f"Hyperlink restoration skipped: {e}")
            
            # 3. Reconstruire Table of Contents
            try:
                rebuild_table_of_contents(docx_path)
            except Exception as e:
                logger.warning(f"TOC rebuilding skipped: {e}")
            
            # 4. Améliorer formatage des nombres financiers
            try:
                format_financial_numbers(docx_path)
            except Exception as e:
                logger.warning(f"Financial number formatting skipped: {e}")
            
            # 5. Évaluer la qualité
            try:
                metrics = assess_conversion_quality(docx_path)
                logger.info(f"Conversion quality score: {metrics['quality_score']}/100")
            except Exception as e:
                logger.warning(f"Quality assessment skipped: {e}")
                
    except Exception as e:
        logger.warning(f"Post-processing failed: {e}")



def enhance_tables_in_docx(docx_path: Path) -> None:
    """
    Post-traite le DOCX pour reconstruire les tableaux avec structure propre.
    CORRIGE les cellules fusionnées/manquantes pour permettre un copier-coller Excel parfait.
    """
    if not PYTHON_DOCX_AVAILABLE:
        logger.warning("python-docx not available, skipping table enhancement")
        return
    
    try:
        # Ouvrir le DOCX généré
        doc = Document(str(docx_path))
        
        tables_fixed = 0
        
        for table_idx, table in enumerate(doc.tables):
            try:
                # ÉTAPE 1 : Analyser la structure du tableau
                # Compter les colonnes de chaque ligne
                row_col_counts = []
                for row in table.rows:
                    row_col_counts.append(len(row.cells))
                
                # Trouver le nombre MAXIMUM de colonnes (ligne la plus longue)
                max_cols = max(row_col_counts) if row_col_counts else 0
                
                if max_cols == 0:
                    continue
                
                # ÉTAPE 2 : Vérifier si le tableau a des problèmes
                has_inconsistent_cols = len(set(row_col_counts)) > 1
                
                if not has_inconsistent_cols:
                    # Tableau déjà cohérent, juste formater les nombres
                    _format_table_cells(table)
                    tables_fixed += 1
                    continue
                
                # ÉTAPE 3 : Tableau problématique → Le reconstruire
                logger.info(f"Fixing table {table_idx + 1}: inconsistent columns {set(row_col_counts)}")
                
                # Extraire tout le contenu du tableau
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Extraire le texte de toutes les cellules
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    # Compléter avec des cellules vides si nécessaire
                    while len(row_data) < max_cols:
                        row_data.append("")
                    
                    table_data.append(row_data)
                
                # ÉTAPE 4 : Supprimer l'ancien tableau
                tbl_element = table._element
                tbl_element.getparent().remove(tbl_element)
                
                # ÉTAPE 5 : Créer un NOUVEAU tableau propre
                new_table = doc.add_table(rows=len(table_data), cols=max_cols)
                new_table.style = 'Light Grid Accent 1'  # Style par défaut
                
                # Remplir le nouveau tableau
                for i, row_data in enumerate(table_data):
                    row = new_table.rows[i]
                    for j, cell_value in enumerate(row_data):
                        cell = row.cells[j]
                        cell.text = cell_value
                        
                        # Formater si c'est un nombre
                        if cell_value:
                            clean = cell_value.replace(',', '').replace('$', '').replace('%', '').replace('(', '').replace(')', '').strip()
                            if clean.replace('.', '').replace('-', '').isdigit():
                                # Nombre → aligner à droite
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = 2  # Right
                
                tables_fixed += 1
                logger.info(f"✓ Table {table_idx + 1} reconstructed with {len(table_data)} rows × {max_cols} cols")
            
            except Exception as e:
                logger.warning(f"Could not fix table {table_idx + 1}: {e}")
                continue
        
        # Sauvegarder les modifications
        doc.save(str(docx_path))
        
        if tables_fixed > 0:
            logger.info(f"✓ Fixed {tables_fixed} tables with proper cell structure")
    
    except Exception as e:
        logger.warning(f"Table enhancement failed: {e}")



def format_table_cells(table) -> None:
    """Formate les cellules d'un tableau (alignement des nombres)."""
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                # Vérifier si c'est un nombre
                clean = cell_text.replace(',', '').replace('$', '').replace('%', '').replace('(', '').replace(')', '').strip()
                if clean.replace('.', '').replace('-', '').isdigit():
                    # Aligner à droite
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT



def fix_qa_formatting(docx_path: Path) -> None:
    """
    Reconstruit la mise en forme Q&A après conversion.
    Détecte les patterns 'Q:' et 'A:' et les formate correctement avec sauts de ligne.
    """
    if not PYTHON_DOCX_AVAILABLE:
        logger.warning("python-docx not available, skipping Q&A formatting")
        return
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        
        doc = Document(str(docx_path))
        
        logger.info("Processing Q&A formatting...")
        
        # Parcourir tous les paragraphes
        i = 0
        qa_count = 0
        
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Chercher des patterns 'Q:' ou 'A:' dans le texte
            if not text or len(text) < 3:
                i += 1
                continue
            
            # Pattern 1: "Q: [...] A: [...]" sur la même ligne -> Séparer en deux paragraphes
            if text.startswith('Q:') and 'A:' in text and text.find('A:') > 2:
                # Séparer Q et A
                a_index = text.find('A:')
                q_text = text[:a_index].strip()
                a_text = text[a_index:].strip()
                
                # Modifier le paragraphe actuel (Q)
                para.clear()
                run = para.add_run(q_text)
                run.bold = True
                para.style = 'Normal'
                para.paragraph_format.space_before = Pt(12)
                
                # Pour A, utiliser le paragraphe suivant s'il est vide
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                    next_para = doc.paragraphs[i + 1]
                    next_para.clear()
                    next_para.add_run(a_text)
                    next_para.style = 'Normal'
                    next_para.paragraph_format.left_indent = Inches(0.25)
                    next_para.paragraph_format.space_after = Pt(12)
                    qa_count += 1
                    i += 2
                else:
                    # Paragraphe suivant occupé, on met juste le texte A après un saut de ligne dans le même para
                    # (Solution de contournement : on garde Q et A séparés visuellement avec formatage)
                    para.add_run('\n' + a_text)
                    qa_count += 1
                    i += 1
                continue
            
            # Pattern 2: Texte commençant par "Q:" (question seule)
            if re.match(r'^Q:\s+', text):
                para.clear()
                run = para.add_run(text)
                run.bold = True
                para.style = 'Normal'
                para.paragraph_format.space_before = Pt(12)
                qa_count += 1
            
            # Pattern 3: Texte commençant par "A:" (réponse seule)
            elif re.match(r'^A:\s+', text):
                para.clear()
                run = para.add_run(text)
                para.style = 'Normal'
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.space_after = Pt(12)
                qa_count += 1
            
            i += 1
        
        # Sauvegarder
        doc.save(str(docx_path))
        
        if qa_count > 0:
            logger.info(f"✓ Q&A formatting complete ({qa_count} Q/A pairs/formats found)")
        else:
            logger.info("No Q&A patterns found")
    
    except Exception as e:
        logger.warning(f"Q&A formatting failed: {e}")
        import traceback
        traceback.print_exc()


def add_hyperlink(paragraph, url: str, text: str):
    """
    Ajoute un hyperlien à un paragraphe.

    Args:
        paragraph: Paragraphe python-docx
        url: URL du lien
        text: Texte à afficher
    """
    if not PYTHON_DOCX_AVAILABLE:
        return

    try:
        # Créer l'élément hyperlink
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        # Créer le run avec le lien
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Style bleu et souligné
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0563C1')
        rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_t = OxmlElement('w:t')
        new_t.text = text
        new_run.append(new_t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception as e:
        logger.warning(f"Failed to add hyperlink: {e}")


def restore_hyperlinks(docx_path: Path) -> None:
    """
    Détecte les URLs et les convertit en hyperliens cliquables.
    """
    if not PYTHON_DOCX_AVAILABLE:
        return
    
    try:
        doc = Document(str(docx_path))
        url_pattern = re.compile(r'(https?://[^\s\)]+)')
        url_count = 0
        
        for para in doc.paragraphs:
            text = para.text
            urls = url_pattern.findall(text)
            
            if urls:
                # Reconstruire le paragraphe avec hyperliens
                para_text = para.text
                para.clear()
                
                # Séparer le texte en parties (avant URL, URL, après URL)
                parts = url_pattern.split(para_text)
                
                for i, part in enumerate(parts):
                    if url_pattern.match(part):
                        # C'est une URL
                        add_hyperlink(para, part, part)
                        url_count += 1
                    else:
                        # Texte normal
                        if part:
                            para.add_run(part)
        
        if url_count > 0:
            doc.save(str(docx_path))
            logger.info(f"✓ Restored {url_count} hyperlink(s)")
    except Exception as e:
        logger.warning(f"Hyperlink restoration failed: {e}")



def rebuild_table_of_contents(docx_path: Path) -> None:
    """
    Reconstruit la Table of Contents en détectant les patterns de titres.
    """
    if not PYTHON_DOCX_AVAILABLE:
        return
    
    try:
        doc = Document(str(docx_path))
        
        # Pattern pour détecter les entrées TOC: "Titre ... 123"
        # Pattern amélioré : Titre avec points de suspension + numéro de page
        toc_pattern = re.compile(r'^([A-Z][A-Za-z\s&,\-\']{3,60}?)\.{2,}\s+(\d{1,3})$')
        
        toc_entries = []
        toc_start_idx = None
        
        # Chercher les paragraphes qui correspondent au pattern TOC
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            match = toc_pattern.match(text)
            
            if match:
                if toc_start_idx is None:
                    toc_start_idx = i
                title = match.group(1).strip()
                page_num = match.group(2)
                toc_entries.append((title, page_num))
        
        if toc_entries and toc_start_idx is not None:
            logger.info(f"Found {len(toc_entries)} TOC entries")
            
            # Formater les entrées TOC avec indentation
            for idx, (title, page_num) in enumerate(toc_entries):
                if toc_start_idx + idx < len(doc.paragraphs):
                    para = doc.paragraphs[toc_start_idx + idx]
                    para.clear()
                    
                    # Ajouter le titre
                    run = para.add_run(title)
                    run.font.size = Pt(11)
                    
                    # Ajouter les points de suite
                    dots_run = para.add_run(' ' + '.' * (50 - len(title) - len(page_num)) + ' ')
                    dots_run.font.size = Pt(9)
                    
                    # Ajouter le numéro de page
                    page_run = para.add_run(page_num)
                    page_run.font.size = Pt(11)
                    page_run.bold = True
            
            doc.save(str(docx_path))
            logger.info(f"✓ Rebuilt Table of Contents with {len(toc_entries)} entries")
    except Exception as e:
        logger.warning(f"TOC rebuilding failed: {e}")



def format_financial_numbers(docx_path: Path) -> None:
    """
    Formate les nombres financiers avec style comptable (rouge si négatif).
    """
    if not PYTHON_DOCX_AVAILABLE:
        return
    
    try:
        doc = Document(str(docx_path))
        formatted_count = 0
        
        # Parcourir tous les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if not text:
                        continue
                    
                    # Détecter si négatif (parenthèses)
                    is_negative = text.startswith('(') and text.endswith(')')
                    
                    # Extraire le nombre
                    clean = text.replace('(', '').replace(')', '').replace(',', '').replace('$', '').replace('%', '').strip()
                    
                    # Vérifier si c'est un nombre
                    if clean.replace('.', '').replace('-', '').isdigit():
                        # Formater
                        for para in cell.paragraphs:
                            para.alignment = 2  # Right align
                            for run in para.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(10)
                                if is_negative:
                                    try:
                                        run.font.color.rgb = RGBColor(192, 0, 0)  # Rouge foncé
                                    except:
                                        pass
                        formatted_count += 1
        
        if formatted_count > 0:
            doc.save(str(docx_path))
            logger.info(f"✓ Formatted {formatted_count} financial number(s)")
    except Exception as e:
        logger.warning(f"Financial number formatting failed: {e}")



def assess_conversion_quality(docx_path: Path) -> Dict[str, Any]:
    """
    Évalue la qualité de la conversion et retourne des métriques.
    
    Returns:
        dict: Métriques de qualité (quality_score, broken_tables, etc.)
    """
    if not PYTHON_DOCX_AVAILABLE:
        return {'quality_score': 0, 'error': 'python-docx not available'}
    
    try:
        doc = Document(str(docx_path))
        
        metrics = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'broken_tables': 0,
            'empty_paragraphs': 0,
            'quality_score': 0
        }
        
        # Compter les tableaux cassés (colonnes incohérentes)
        for table in doc.tables:
            if not table.rows:
                continue
            row_counts = [len(row.cells) for row in table.rows]
            if len(set(row_counts)) > 1:
                metrics['broken_tables'] += 1
        
        # Compter les paragraphes vides
        for para in doc.paragraphs:
            if not para.text.strip():
                metrics['empty_paragraphs'] += 1
        
        # Calculer le score de qualité (0 à 100)
        # Score : 100 - (5 points par table cassée + 0.5 point par 10 paragraphes vides)
        total_issues = metrics['broken_tables'] * 5 + (metrics['empty_paragraphs'] / 10) * 0.5
        metrics['quality_score'] = max(0, min(100, 100 - total_issues))
        
        logger.info(f"Quality metrics: {metrics['broken_tables']} broken tables, "
                   f"{metrics['empty_paragraphs']} empty paragraphs")
        
        return metrics
        
    except Exception as e:
        logger.warning(f"Quality assessment failed: {e}")
        return {'quality_score': 0, 'error': str(e)}


