#!/usr/bin/env python3
"""
PDF Converter class - Refactored modular version.
Uses all extracted modules for functionality.
"""

from pathlib import Path
from typing import Optional
import logging

from src.config import (
    PDF2DOCX_AVAILABLE, OCR_AVAILABLE, PYPDF2_AVAILABLE,
    PYTHON_DOCX_AVAILABLE, OPENPYXL_AVAILABLE, ADOBE_PDF_AVAILABLE, logger
)

# Import modules
from src.analysis.pdf_analyzer import get_pdf_page_count, is_scanned_pdf, detect_images_in_pdf
from src.ocr.ocr_cache import OCRCache
from src.analysis.sensitive_info_detector import detect_sensitive_information
from src.processing.document_processor import apply_post_processing
from src.analysis.report_generator import analyze_and_create_report

# OCR processor (may not be available)
try:
    from src.ocr.ocr_processor import convert_with_ocr as ocr_convert_with_ocr
except ImportError:
    ocr_convert_with_ocr = None

if PDF2DOCX_AVAILABLE:
    from pdf2docx import Converter

if PYPDF2_AVAILABLE:
    import PyPDF2

if PYTHON_DOCX_AVAILABLE:
    from docx import Document
    import copy


class SECPDFConverter:
    """
    Convertisseur PDF vers DOCX avec préservation de structure et détection d'informations sensibles.
    Version refactorée utilisant des modules modulaires.
    """
    
    def __init__(self, pdf_path, docx_path=None, verbose=False):
        """
        Initialise le convertisseur.
        
        Args:
            pdf_path: Chemin vers le fichier PDF
            docx_path: Chemin de sortie (optionnel, généré automatiquement si non fourni)
            verbose: Mode verbose
        """
        self.pdf_path = Path(pdf_path)
        
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {self.pdf_path}")
        
        if docx_path:
            self.docx_path = Path(docx_path)
        else:
            self.docx_path = self.pdf_path.with_suffix('.docx')
        
        # Créer le dossier parent si nécessaire
        self.docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.converter = None
        self.images_by_page = {}
        self.sensitive_info_by_page = {}
        self.verbose = verbose
        
        # Cache OCR
        self.ocr_cache = OCRCache(self.pdf_path, verbose=verbose)
    
    def convert(self, start_page=0, end_page=None, use_parallel=None, pages_per_chunk=20, 
                use_ocr=False, auto_detect_scanned=False, **kwargs):
        """
        Convertit le PDF en DOCX en préservant la structure originale.
        
        Args:
            start_page: Page de départ (0-indexed)
            end_page: Page de fin (None = toutes les pages)
            use_parallel: Force la parallélisation (True/False/None=auto)
            pages_per_chunk: Nombre de pages par chunk
            use_ocr: Force l'utilisation de l'OCR
            auto_detect_scanned: Détecte automatiquement si le PDF est scanné
            **kwargs: Arguments supplémentaires
        """
        self.docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Détection automatique des PDFs scannés
        if auto_detect_scanned and not use_ocr:
            is_scanned = is_scanned_pdf(self.pdf_path)
            if is_scanned:
                use_ocr = True
                logger.info("📷 Scanned PDF detected → Switching to OCR mode")
        
        # Mode OCR
        if use_ocr:
            if not OCR_AVAILABLE or ocr_convert_with_ocr is None:
                raise ImportError("OCR libraries not available. Install: pip install ocrmypdf")
            
            logger.info("🔄 Using OCR mode")
            ocr_convert_with_ocr(
                self.pdf_path, self.docx_path, self.ocr_cache,
                start_page=start_page, end_page=end_page, verbose=self.verbose
            )
            apply_post_processing(self.docx_path)
            logger.info("✅ OCR conversion completed")
            return
        
        # PRIORITY 1: Adobe PDF Services (industry standard, best quality)
        # Adobe handles everything automatically: complex PDFs, scanned PDFs, OCR, etc.
        # Includes extended timeouts (120s) and automatic retry logic (3 attempts)
        from src.adobe_converter import convert_pdf_to_docx_adobe, is_adobe_available

        if is_adobe_available():
            try:
                logger.info("🚀 Using Adobe PDF Services (industry standard)")
                convert_pdf_to_docx_adobe(self.pdf_path, self.docx_path, verbose=self.verbose)
                logger.info("✅ Adobe conversion completed successfully")
                apply_post_processing(self.docx_path)
                return
            except Exception as e:
                # Adobe conversion failed after retries
                # Re-raise the exception - NO FALLBACK per user request
                logger.error(f"❌ Adobe conversion failed: {e}")
                raise RuntimeError(f"Adobe PDF conversion failed: {e}") from e

        # If Adobe is not available at all, show setup instructions
        if not PDF2DOCX_AVAILABLE:
            from src.adobe_converter import print_adobe_setup_instructions
            logger.error("No conversion library available!")
            print("\n" + "="*80)
            print(" " * 28 + "❌ CONVERSION ERROR")
            print("="*80)
            print("\nNo PDF conversion library is installed.")
            print()
            print("RECOMMENDED (Best Quality):")
            print("  pip install pdfservices-sdk")
            print()
            print("Then configure Adobe credentials (see instructions below)")
            print()
            print_adobe_setup_instructions()
            print("\nALTERNATIVE (Fallback):")
            print("  pip install pdf2docx")
            print("="*80)
            raise ImportError("No PDF conversion library available")

        logger.info("Using pdf2docx (fallback)")

        # Détection du nombre de pages
        total_pages = get_pdf_page_count(self.pdf_path)
        if end_page is not None:
            pages_to_convert = end_page - start_page
        elif total_pages:
            pages_to_convert = total_pages - start_page
        else:
            pages_to_convert = None
        
        # Décision de parallélisation
        should_parallelize = False
        if use_parallel is True:
            should_parallelize = True
        elif use_parallel is False:
            should_parallelize = False
        elif pages_to_convert and pages_to_convert > 100:
            should_parallelize = True
            logger.info(f"Large PDF detected ({pages_to_convert} pages), using parallel conversion")
        
        # Parallélisation pour gros PDFs
        if should_parallelize and pages_to_convert and pages_to_convert > 50:
            try:
                self._convert_parallel(start_page, end_page, pages_per_chunk, **kwargs)
                logger.info("PDF conversion completed successfully (parallel)")
                apply_post_processing(self.docx_path)
                return
            except Exception as e:
                logger.warning(f"Parallel conversion failed, falling back to sequential: {e}")
        
        # Conversion normale (séquentielle)
        try:
            self.converter = Converter(str(self.pdf_path))
            
            convert_params = {
                'start': start_page,
                'end': end_page,
            }
            convert_params.update(kwargs)
            
            from src.config import suppress_output
            with suppress_output(verbose=self.verbose):
                self.converter.convert(str(self.docx_path), **convert_params)
            
            if self.verbose:
                logger.info("PDF conversion completed successfully")
            
        except Exception as e:
            if self.converter:
                try:
                    self.converter.close()
                except:
                    pass
            raise Exception(f"PDF conversion failed: {e}")
        finally:
            if self.converter:
                try:
                    self.converter.close()
                except:
                    pass
            apply_post_processing(self.docx_path)
    
    def convert_with_ocr(self, start_page=0, end_page=None, ocr_language='eng'):
        """Convertit avec OCR (délègue à ocr_processor)."""
        ocr_convert_with_ocr(
            self.pdf_path, self.docx_path, self.ocr_cache,
            start_page=start_page, end_page=end_page, ocr_language=ocr_language, verbose=self.verbose
        )
    
    def analyze_and_create_report(self, report_path=None, grouping_keywords=None, verbose=False):
        """Analyse et crée un rapport (délègue à report_generator)."""
        from src.analysis.pdf_analyzer import detect_images_in_pdf

        # Détecter les images et informations sensibles
        self.images_by_page = detect_images_in_pdf(self.pdf_path)
        self.sensitive_info_by_page = detect_sensitive_information(self.pdf_path, verbose=verbose)

        # Créer le rapport
        return analyze_and_create_report(
            self.pdf_path, report_path, grouping_keywords, verbose=verbose
        )
    
    def _convert_parallel(self, start_page=0, end_page=None, pages_per_chunk=20, **kwargs):
        """Convertit en parallèle (utilise utils._convert_chunk)."""
        from multiprocessing import Pool, cpu_count
        import tempfile
        from src.utils import _convert_chunk
        
        total_pages = get_pdf_page_count(self.pdf_path)
        if not total_pages:
            raise ValueError("Could not determine PDF page count")
        
        actual_start = start_page
        actual_end = end_page if end_page is not None else total_pages
        
        chunks = []
        temp_dir = Path(tempfile.mkdtemp(prefix='pdf_conversion_'))
        
        try:
            chunk_idx = 0
            for chunk_start in range(actual_start, actual_end, pages_per_chunk):
                chunk_end = min(chunk_start + pages_per_chunk, actual_end)
                chunk_output = temp_dir / f"chunk_{chunk_idx}.docx"
                chunks.append((str(self.pdf_path), str(chunk_output), chunk_start, chunk_end))
                chunk_idx += 1
            
            logger.info(f"Converting {len(chunks)} chunks in parallel")
            
            num_processes = min(cpu_count(), len(chunks), 4)
            
            with Pool(processes=num_processes) as pool:
                results = pool.map(_convert_chunk, chunks)
            
            successful_chunks = [r for r in results if r and Path(r).exists()]
            
            if not successful_chunks:
                raise Exception("All chunks failed to convert")
            
            if len(successful_chunks) < len(chunks):
                logger.warning(f"Only {len(successful_chunks)}/{len(chunks)} chunks succeeded")
            
            self._merge_docx_files(successful_chunks)
            
        finally:
            try:
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Could not clean up temp directory: {e}")
    
    def _merge_docx_files(self, docx_files):
        """Fusionne plusieurs fichiers DOCX en un seul."""
        if not PYTHON_DOCX_AVAILABLE:
            return False
        
        try:
            merged_doc = Document()
            valid_files = [f for f in docx_files if f and Path(f).exists()]
            
            if not valid_files:
                logger.warning("No valid DOCX files to merge")
                return False
            
            for idx, docx_file in enumerate(valid_files):
                doc = Document(docx_file)
                
                for element in doc.element.body:
                    merged_doc.element.body.append(copy.deepcopy(element))
                
                if idx < len(valid_files) - 1:
                    merged_doc.add_page_break()
            
            merged_doc.save(str(self.docx_path))
            logger.info(f"✓ Merged {len(valid_files)} chunk(s) into {self.docx_path.name}")
            return True
        
        except Exception as e:
            logger.error(f"DOCX merge failed: {e}")
            return False
    
    def convert_complex(self, start_page=0, end_page=None, **kwargs):
        """
        Conversion avancée pour PDFs complexes avec mise en page désorganisée.
        Utilise pdfplumber + analyse de layout pour mieux organiser le contenu.

        Stratégies:
        1. Analyse de layout avec pdfplumber (détection colonnes, tables)
        2. Extraction du texte avec ordre de lecture optimal
        3. Reconstruction logique du document
        4. Préservation des tables et structures
        5. Fallback sur OCR si nécessaire

        Args:
            start_page: Page de départ
            end_page: Page de fin
            **kwargs: Arguments supplémentaires
        """
        from src.config import PDFPLUMBER_AVAILABLE

        if not PDFPLUMBER_AVAILABLE:
            logger.warning("⚠️  pdfplumber not available, falling back to standard conversion")
            logger.info("   Install with: pip install pdfplumber")
            return self.convert(start_page=start_page, end_page=end_page, **kwargs)

        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for complex conversion")

        import pdfplumber
        from docx import Document
        from docx.shared import Pt

        logger.info("🔧 Using COMPLEX PDF conversion mode")
        logger.info("   ├─ Advanced layout detection")
        logger.info("   ├─ Multi-column handling")
        logger.info("   ├─ Table preservation")
        logger.info("   └─ Logical text ordering")

        try:
            doc = Document()

            with pdfplumber.open(self.pdf_path) as pdf:
                total_pages = len(pdf.pages)
                actual_end = end_page if end_page is not None else total_pages

                if self.verbose:
                    print(f"\n   📄 Processing {actual_end - start_page} pages with advanced layout analysis...")

                for page_num in range(start_page, actual_end):
                    if page_num >= total_pages:
                        break

                    page = pdf.pages[page_num]

                    # Ajouter numéro de page
                    if page_num > start_page:
                        doc.add_page_break()

                    doc.add_paragraph(f"─── Page {page_num + 1} ───")

                    # STRATÉGIE 1: Détecter et extraire les tables
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            self._add_table_to_doc(doc, table_data)

                    # STRATÉGIE 2: Extraire le texte avec layout
                    # Utilise l'ordre de lecture optimal (colonnes, etc.)
                    text = page.extract_text(layout=True)

                    if text and text.strip():
                        # Nettoyer et organiser le texte
                        lines = text.split('\n')
                        current_para = []

                        for line in lines:
                            line = line.strip()
                            if not line:
                                # Ligne vide = nouveau paragraphe
                                if current_para:
                                    para_text = ' '.join(current_para)
                                    if para_text:
                                        p = doc.add_paragraph(para_text)
                                        p.style = 'Normal'
                                        # Police plus petite pour mieux tenir
                                        for run in p.runs:
                                            run.font.size = Pt(10)
                                    current_para = []
                            else:
                                current_para.append(line)

                        # Dernier paragraphe
                        if current_para:
                            para_text = ' '.join(current_para)
                            if para_text:
                                p = doc.add_paragraph(para_text)
                                p.style = 'Normal'
                                for run in p.runs:
                                    run.font.size = Pt(10)

                    if self.verbose and (page_num - start_page + 1) % 10 == 0:
                        print(f"   ├─ Processed {page_num - start_page + 1}/{actual_end - start_page} pages")

            # Sauvegarder
            doc.save(str(self.docx_path))
            logger.info(f"✅ Complex conversion completed: {self.docx_path.name}")

            # Post-processing
            apply_post_processing(self.docx_path)

        except Exception as e:
            logger.error(f"Complex conversion failed: {e}")
            logger.info("Falling back to standard conversion...")
            return self.convert(start_page=start_page, end_page=end_page, **kwargs)

    def _add_table_to_doc(self, doc, table_data):
        """Ajoute une table au document DOCX."""
        if not table_data or not any(table_data):
            return

        # Filtrer les lignes vides
        table_data = [row for row in table_data if row and any(cell for cell in row if cell)]

        if not table_data:
            return

        # Trouver le nombre maximum de colonnes
        max_cols = max(len(row) for row in table_data)

        # Créer la table
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        # Remplir la table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_value in enumerate(row_data):
                if j < max_cols and cell_value:
                    row.cells[j].text = str(cell_value).strip()

        doc.add_paragraph()  # Espace après la table

    # Méthodes de compatibilité (délèguent aux modules)
    def _get_pdf_page_count(self):
        """Retourne le nombre total de pages (délègue à pdf_analyzer)."""
        return get_pdf_page_count(self.pdf_path)
    
    def is_scanned_pdf(self):
        """Détecte si le PDF est scanné (délègue à pdf_analyzer)."""
        return is_scanned_pdf(self.pdf_path)
    
    def _detect_images_in_pdf(self):
        """Détecte les images (délègue à pdf_analyzer)."""
        return detect_images_in_pdf(self.pdf_path)
    
    def _detect_sensitive_information(self, verbose=False):
        """Détecte les informations sensibles (délègue à sensitive_info_detector)."""
        return detect_sensitive_information(self.pdf_path, verbose=verbose)


__all__ = ['SECPDFConverter']
