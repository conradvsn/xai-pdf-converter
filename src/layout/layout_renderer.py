#!/usr/bin/env python3
"""
Layout Renderer Module
Handles layout rendering (multi-column, tables, structure).
"""

from typing import Dict, List, Any, Optional
from src.config import PYTHON_DOCX_AVAILABLE, logger

if PYTHON_DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_multicolumn_layout(doc: Document, lines, structure, img_width) -> None:
    """Rend un layout multi-colonnes."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Calculer les positions des colonnes
    column_width = img_width / structure['column_count']
    
    for line_idx, line in enumerate(lines):
        if not line:
            doc.add_paragraph("")
            continue
        
        # Déterminer dans quelle colonne se trouve la ligne
        avg_x = sum(w['left'] for w in line) / len(line) if line else 0
        column_idx = int(avg_x / column_width) if column_width > 0 else 0
        
        # Créer le texte de la ligne
        line_text = ' '.join([w['text'] for w in sorted(line, key=lambda w: w['left'])])
        
        # Appliquer le style selon le type
        if line_idx in structure['title_lines']:
            para = doc.add_paragraph(line_text.strip())
            para.style = 'Heading 2'
            for run in para.runs:
                run.bold = True
        else:
            para = doc.add_paragraph(line_text.strip())
        
        # Alignement selon la colonne
        if column_idx == 0:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif column_idx == structure['column_count'] - 1:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT



def render_table_layout(doc: Document, lines, structure, img_width) -> None:
    """Rend un layout de tableau."""
    from docx.shared import Pt
    
    # Détecter les colonnes du tableau
    all_x_positions = set()
    for line in lines:
        for w in line:
            all_x_positions.add(w['left'])
    
    sorted_x = sorted(all_x_positions)
    # Grouper les positions proches (colonnes)
    columns = []
    current_col = [sorted_x[0]] if sorted_x else []
    for x in sorted_x[1:]:
        if x - current_col[-1] < 50:  # Positions proches = même colonne
            current_col.append(x)
        else:
            if current_col:
                columns.append(sum(current_col) / len(current_col))  # Position moyenne
            current_col = [x]
    if current_col:
        columns.append(sum(current_col) / len(current_col))
    
    if len(columns) < 2:
        # Pas assez de colonnes, rendre comme texte simple
        render_single_column_layout(doc, lines, structure)
        return
    
    # Créer un tableau Word
    table = doc.add_table(rows=len(lines), cols=len(columns))
    table.style = 'Table Grid'
    
    for row_idx, line in enumerate(lines):
        if not line:
            continue
        
        # Assigner chaque mot à sa colonne
        for w in line:
            # Trouver la colonne la plus proche
            col_idx = min(range(len(columns)), key=lambda i: abs(w['left'] - columns[i]))
            if col_idx < len(columns):
                cell = table.rows[row_idx].cells[col_idx]
                if not cell.paragraphs[0].text.strip():
                    cell.paragraphs[0].text = w['text']
                else:
                    cell.paragraphs[0].text += " " + w['text']



def render_single_column_layout(doc: Document, lines, structure) -> None:
    """Rend un layout simple à une colonne."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    for line_idx, line in enumerate(lines):
        if not line:
            doc.add_paragraph("")
            continue
        
        line_sorted = sorted(line, key=lambda w: w['left'])
        line_text = ' '.join([w['text'] for w in line_sorted])
        
        if line_idx in structure.get('title_lines', []):
            para = doc.add_paragraph(line_text.strip())
            para.style = 'Heading 2'
            for run in para.runs:
                run.bold = True
        else:
            doc.add_paragraph(line_text.strip())



def render_multicolumn_from_pdfplumber(doc: Document, words, lines, structure, page_width) -> None:
    """Rend un layout multi-colonnes depuis pdfplumber."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    column_width = page_width / structure['column_count']
    
    for line_idx, line in enumerate(lines):
        if not line:
            doc.add_paragraph("")
            continue
        
        # Déterminer la colonne
        avg_x = sum(w.get('left', 0) for w in line) / len(line) if line else 0
        column_idx = int(avg_x / column_width) if column_width > 0 else 0
        
        # Créer le texte
        line_text = ' '.join([w.get('text', '') for w in line])
        
        if line_idx in structure.get('title_lines', []):
            para = doc.add_paragraph(line_text.strip())
            para.style = 'Heading 2'
            for run in para.runs:
                run.bold = True
        else:
            para = doc.add_paragraph(line_text.strip())



def render_table_from_pdfplumber(doc: Document, words, lines, structure, page_width) -> None:
    """Rend un tableau depuis pdfplumber."""
    # Détecter les colonnes du tableau
    all_x_positions = sorted(set(w.get('left', 0) for w in words))
    
    # Grouper les positions proches
    columns = []
    if all_x_positions:
        columns.append(all_x_positions[0])
        for x in all_x_positions[1:]:
            if x - columns[-1] > 50:  # Nouvelle colonne si gap > 50
                columns.append(x)
    
    if len(columns) < 2:
        render_single_column_from_pdfplumber(doc, words, lines, structure)
        return
    
    # Créer le tableau
    table = doc.add_table(rows=len(lines), cols=len(columns))
    table.style = 'Table Grid'
    
    for row_idx, line in enumerate(lines):
        if not line:
            continue
        
        for w in line:
            x_pos = w.get('left', 0)
            col_idx = min(range(len(columns)), key=lambda i: abs(x_pos - columns[i]))
            if col_idx < len(columns):
                cell = table.rows[row_idx].cells[col_idx]
                if not cell.paragraphs[0].text.strip():
                    cell.paragraphs[0].text = w.get('text', '')
                else:
                    cell.paragraphs[0].text += " " + w.get('text', '')



def render_single_column_from_pdfplumber(doc: Document, words, lines, structure) -> None:
    """Rend un layout simple depuis pdfplumber."""
    from docx.shared import Pt
    
    for line_idx, line in enumerate(lines):
        if not line:
            doc.add_paragraph("")
            continue
        
        line_text = ' '.join([w.get('text', '') for w in line])
        
        if line_idx in structure.get('title_lines', []):
            para = doc.add_paragraph(line_text.strip())
            para.style = 'Heading 2'
            for run in para.runs:
                run.bold = True
        else:
            doc.add_paragraph(line_text.strip())



def create_word_table_from_lines(doc: Document, lines, page_width) -> None:
    """
    Crée un tableau Word pour un bloc de lignes détectées comme tabulaires.
    
    Args:
        doc: Document Word
        lines: Liste de lignes, chaque ligne étant une liste de mots (dicts)
        page_width: Largeur de la page en pixels
    """
    if not lines:
        return
    
    # Détecter les positions X des colonnes
    x_starts = []
    for line in lines:
        for word in line:
            x_starts.append(word['left'])
    
    # Clusteriser les X proches (même colonne)
    x_starts.sort()
    columns_clusters = []
    if x_starts:
        current_cluster = [x_starts[0]]
        for x in x_starts[1:]:
            if x - current_cluster[-1] < 50:  # Tolérance de 50px
                current_cluster.append(x)
            else:
                columns_clusters.append(sum(current_cluster) / len(current_cluster))
                current_cluster = [x]
        columns_clusters.append(sum(current_cluster) / len(current_cluster))
    
    num_cols = len(columns_clusters)
    if num_cols == 0:
        num_cols = 1
    
    # Créer le tableau Word
    table = doc.add_table(rows=len(lines), cols=num_cols)
    table.style = 'Normal Table'
    
    # Remplir le tableau
    for i, line in enumerate(lines):
        row = table.rows[i]
        for word in line:
            # Trouver la colonne la plus proche
            word_center = word['left']
            closest_col_idx = min(range(num_cols), key=lambda k: abs(columns_clusters[k] - word_center))
            
            # Ajouter le texte dans la cellule
            cell = row.cells[closest_col_idx]
            if cell.text:
                cell.text += " " + word['text']
            else:
                cell.text = word['text']
            
            # Alignement à droite pour les chiffres
            clean_text = word['text'].replace(',', '').replace('.', '').replace('(', '').replace(')', '').replace('£', '').replace('$', '').strip()
            if clean_text and (clean_text.isdigit() or (clean_text.startswith('-') and clean_text[1:].isdigit())):
                cell.paragraphs[0].alignment = 2  # Right align



def add_text_blocks_to_doc(doc: Document, text_blocks) -> None:
    """
    Groupe les données brutes de Tesseract en lignes visuelles.
    
    Args:
        data: Données OCR de Tesseract (format DICT)
    
    Returns:
        Liste de lignes, chaque ligne étant une liste de mots (dicts avec 'text', 'left', 'top', etc.)
    """
    n_boxes = len(data['text'])
    words = []
    for i in range(n_boxes):
        if int(data['conf'][i]) > 30 and data['text'][i].strip():
            words.append({
                'text': data['text'][i],
                'left': data['left'][i],
                'top': data['top'][i],
                'width': data['width'][i],
                'height': data['height'][i]
            })
    
    # Tri Y puis X
    words.sort(key=lambda x: (x['top'], x['left']))
    
    lines = []
    current_line = []
    current_y = -1
    y_tolerance = 15
    
    for word in words:
        if current_y == -1:
            current_y = word['top']
            current_line.append(word)
        elif abs(word['top'] - current_y) <= y_tolerance:
            current_line.append(word)
        else:
            current_line.sort(key=lambda x: x['left'])
            lines.append(current_line)
            current_line = [word]
            current_y = word['top']
    
    if current_line:
        current_line.sort(key=lambda x: x['left'])
        lines.append(current_line)
    
    return lines



def create_table_from_words(doc: Document, words, page_width) -> None:
    """
    Crée un tableau Word à partir des mots.
    
    Args:
        doc: Document Word
        words: Liste de mots (dicts de pdfplumber)
        page_width: Largeur de la page en points
    """
    if not words:
        return
    
    # Grouper en lignes
    lines = group_words_into_lines(words)
    
    if not lines:
        return
    
    # Détecter les positions X des colonnes
    x_positions = []
    for line in lines:
        for word in line:
            x_positions.append(word.get('x0', 0))
    
    # Clusteriser les X
    x_positions.sort()
    column_positions = []
    if x_positions:
        current_cluster = [x_positions[0]]
        for x in x_positions[1:]:
            if x - current_cluster[-1] < 30:  # Tolérance 30 points
                current_cluster.append(x)
            else:
                column_positions.append(sum(current_cluster) / len(current_cluster))
                current_cluster = [x]
        column_positions.append(sum(current_cluster) / len(current_cluster))
    
    num_cols = len(column_positions) if column_positions else 1
    
    # Créer le tableau Word
    table = doc.add_table(rows=len(lines), cols=num_cols)
    table.style = 'Normal Table'
    
    # Remplir le tableau
    for i, line in enumerate(lines):
        row = table.rows[i]
        for word in line:
            # Trouver la colonne la plus proche
            word_x = word.get('x0', 0)
            closest_col_idx = min(range(num_cols), key=lambda k: abs(column_positions[k] - word_x)) if column_positions else 0
            
            # Ajouter le texte dans la cellule
            cell = row.cells[closest_col_idx]
            if cell.text:
                cell.text += " " + word.get('text', '')
            else:
                cell.text = word.get('text', '')
            
            # Alignement à droite pour les chiffres
            text = word.get('text', '')
            clean_text = text.replace(',', '').replace('.', '').replace('(', '').replace(')', '').replace('£', '').replace('$', '').strip()
            if clean_text and (clean_text.isdigit() or (clean_text.startswith('-') and clean_text[1:].isdigit())):
                cell.paragraphs[0].alignment = 2  # Right align



def render_multicolumn_line(doc: Document, columns) -> None:
    """
    Rend une ligne avec colonnes multiples.
    Utilise un tableau Word à 2 colonnes pour préserver la structure.
    """
    from docx.shared import Inches
    
    if not columns or len(columns) < 2:
        # Fallback : rendre comme texte normal
        all_words = []
        for col in columns:
            all_words.extend(col)
        text = " ".join([w['text'] for w in all_words])
        if text.strip():
            doc.add_paragraph(text.strip())
        return
    
    # Créer un tableau avec 2 colonnes pour cette ligne
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Normal Table'
    
    # Largeur des colonnes (50% chacune)
    for col in table.columns:
        col.width = Inches(3.25)
    
    row = table.rows[0]
    
    # Colonne gauche
    left_words = sorted(columns[0], key=lambda w: w['left'])
    left_text = " ".join([w['text'] for w in left_words])
    row.cells[0].paragraphs[0].text = left_text.strip()
    
    # Colonne droite
    right_words = sorted(columns[1], key=lambda w: w['left'])
    right_text = " ".join([w['text'] for w in right_words])
    row.cells[1].paragraphs[0].text = right_text.strip()


